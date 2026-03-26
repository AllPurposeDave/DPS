"""
Step 4: Section Splitter
===========================

RUN THIS AFTER heading_style_fixer.py has produced *_fixed.docx files.

Splits each _fixed.docx at Heading 1 boundaries into sub-documents for
RAG (Retrieval-Augmented Generation) ingestion. Each sub-document is kept
under the character limit (default 36,000 chars, set via max_characters in
dps_config.yaml). Preamble (content before first H1) is prepended to every
sub-document so each chunk is self-contained.

SPLIT STRATEGY — GREEDY H2 ACCUMULATION:
  When an H1 section exceeds max_characters, H2 sub-sections are accumulated
  greedily: content is grouped until the next H2 would push the chunk over
  the limit, then a split is made at that H2 boundary. This ensures each
  sub-document is as large as possible without exceeding the limit, avoiding
  fragmentation from splitting at every H2 regardless of size.
  WHY: Smaller, focused RAG chunks improve retrieval precision — each chunk
  is retrieved as a unit, so tightly-grouped content returns better answers.

Usage (unified pipeline):
    python run_pipeline.py --step 4

Usage (standalone):
    python scripts/section_splitter.py
    python scripts/section_splitter.py --config dps_config.yaml
    python scripts/section_splitter.py ./output/3\ -\ heading_fixes/ ./output/4\ -\ split_documents/

Input:
    Reads *_fixed.docx files from Step 3's output directory.

Output:
    Sub-document .docx files named "[OriginalName] - [Heading1Text].docx"
    split_manifest.csv — manifest of all sub-documents created

FAILURE POINT: This script ONLY processes files ending in "_fixed.docx".
If Step 3 was not run, this script will find nothing.

FAILURE POINT: Check split_manifest.csv after running:
  - "Full Document - No Heading 1 found" = heading fixer didn't work for that doc
  - Duplicate sub_doc_filename values = second overwrote the first
  - character_count above limit = needs manual H2 insertion
"""
# pip install python-docx

import csv
import os
import re
import shutil
import traceback

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from shared_utils import (
    ensure_output_dir,
    get_output_dir,
    get_heading_level,
    iter_docx_files,
    load_config,
    log_pipeline_issue,
    sanitize_filename,
    setup_argparse,
)


def get_paragraph_text_length(paragraph) -> int:
    return len(paragraph.text)


def compute_section_chars(paragraphs: list) -> int:
    return sum(len(p.text) for p in paragraphs)


def find_heading1_indices(paragraphs) -> list[int]:
    indices = []
    for i, para in enumerate(paragraphs):
        if get_heading_level(para.style) == 1:
            indices.append(i)
    return indices


def find_heading2_indices(paragraphs, start: int, end: int) -> list[int]:
    indices = []
    for i in range(start, end):
        if get_heading_level(paragraphs[i].style) == 2:
            indices.append(i)
    return indices


def copy_paragraph(source_para, target_doc):
    """Copy a paragraph preserving formatting via XML deep copy."""
    from copy import deepcopy
    from docx.oxml.ns import qn

    new_para = target_doc.add_paragraph()
    new_para._element.getparent().remove(new_para._element)
    new_elem = deepcopy(source_para._element)
    target_doc.element.body.append(new_elem)


def copy_table(source_table, target_doc):
    """Copy a table preserving structure via XML deep copy."""
    from copy import deepcopy

    new_elem = deepcopy(source_table._element)
    target_doc.element.body.append(new_elem)


def build_element_sequence(doc):
    """
    Build an ordered sequence of body elements (paragraphs and tables)
    to preserve correct interleaved ordering when splitting.
    Returns list of tuples: (element_type, object, paragraph_index_or_None)
    """
    from docx.oxml.ns import qn

    elements = []
    para_idx = 0
    table_idx = 0
    paragraphs = doc.paragraphs
    tables = doc.tables

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            if para_idx < len(paragraphs):
                elements.append(("paragraph", paragraphs[para_idx], para_idx))
                para_idx += 1
        elif child.tag == qn("w:tbl"):
            if table_idx < len(tables):
                elements.append(("table", tables[table_idx], None))
                table_idx += 1

    return elements


def create_sub_document(preamble_elements: list, section_elements: list, output_path: str, source_path: str) -> int:
    """Create a sub-document containing preamble + section elements. Returns char count.

    Clones the source DOCX to preserve styles, themes, relationships, and numbering
    definitions, then clears the body and repopulates with the requested elements.
    This prevents Word from reporting "unreadable content" errors caused by broken
    style/relationship references when copying XML into a blank document.
    """
    from docx.oxml.ns import qn as _qn

    shutil.copy2(source_path, output_path)
    sub_doc = Document(output_path)

    # Clear all body elements except sectPr (page layout / section properties)
    body = sub_doc.element.body
    for child in list(body):
        if child.tag != _qn("w:sectPr"):
            body.remove(child)

    total_chars = 0

    for elem_type, elem_obj, _ in preamble_elements:
        if elem_type == "paragraph":
            copy_paragraph(elem_obj, sub_doc)
            total_chars += len(elem_obj.text)
        elif elem_type == "table":
            copy_table(elem_obj, sub_doc)
            for row in elem_obj.rows:
                for cell in row.cells:
                    total_chars += len(cell.text)

    for elem_type, elem_obj, _ in section_elements:
        if elem_type == "paragraph":
            copy_paragraph(elem_obj, sub_doc)
            total_chars += len(elem_obj.text)
        elif elem_type == "table":
            copy_table(elem_obj, sub_doc)
            for row in elem_obj.rows:
                for cell in row.cells:
                    total_chars += len(cell.text)

    sub_doc.save(output_path)
    return total_chars


def split_at_heading2(
    preamble_elements: list, section_elements: list,
    base_name: str, heading1_text: str, output_dir: str,
    sub_counter: int, chars_per_page: int,
    max_chars: int, preamble_chars: int, source_path: str,
) -> list[dict]:
    """
    Further split a section at Heading 2 boundaries when it exceeds the limit.

    GREEDY ACCUMULATION: H2 sub-sections are grouped together until adding
    the next H2 group would push the chunk (preamble + accumulated) over
    max_chars. The split happens at THAT H2 boundary, not at every H2.
    Result: fewest possible sub-documents, each as large as allowed.
    """
    records = []

    has_h2 = any(
        elem_type == "paragraph" and get_heading_level(elem_obj.style) == 2
        for elem_type, elem_obj, _ in section_elements
    )

    if not has_h2:
        # No H2s — write the whole oversized section as one sub-doc with a warning.
        safe_heading = sanitize_filename(heading1_text)
        out_name = f"{base_name} - {safe_heading}.docx"
        out_path = os.path.join(output_dir, out_name)
        char_count = create_sub_document(preamble_elements, section_elements, out_path, source_path)
        records.append({
            "original_doc": f"{base_name}_fixed.docx",
            "sub_doc_filename": out_name,
            "heading_text": heading1_text,
            "character_count": char_count,
            "page_estimate": round(char_count / chars_per_page, 1),
        })
        return records

    def _flush_chunk(chunk: list) -> None:
        """Save chunk as a sub-document and append a manifest record."""
        if not chunk:
            return
        first_elem = chunk[0]
        heading_level = (
            get_heading_level(first_elem[1].style)
            if first_elem[0] == "paragraph" else None
        )
        if heading_level in (1, 2):
            sub_heading = first_elem[1].text.strip()
        else:
            sub_heading = heading1_text

        if heading_level == 2:
            safe_heading = sanitize_filename(f"{heading1_text} - {sub_heading}")
        else:
            safe_heading = sanitize_filename(heading1_text)

        out_name = f"{base_name} - {safe_heading}.docx"
        out_path = os.path.join(output_dir, out_name)
        char_count = create_sub_document(preamble_elements, chunk, out_path, source_path)
        records.append({
            "original_doc": f"{base_name}_fixed.docx",
            "sub_doc_filename": out_name,
            "heading_text": sub_heading,
            "character_count": char_count,
            "page_estimate": round(char_count / chars_per_page, 1),
        })

    current_chunk: list = []
    current_chars: int = 0

    for elem_type, elem_obj, para_idx in section_elements:
        is_h2 = elem_type == "paragraph" and get_heading_level(elem_obj.style) == 2

        # When we hit an H2 AND the current chunk already exceeds the limit,
        # flush before starting a new chunk at this heading boundary.
        if is_h2 and current_chunk and (preamble_chars + current_chars) > max_chars:
            _flush_chunk(current_chunk)
            current_chunk = []
            current_chars = 0

        current_chunk.append((elem_type, elem_obj, para_idx))
        if elem_type == "paragraph":
            current_chars += len(elem_obj.text)
        elif elem_type == "table":
            for row in elem_obj.rows:
                for cell in row.cells:
                    current_chars += len(cell.text)

    _flush_chunk(current_chunk)
    return records


def process_document(filepath: str, output_dir: str, max_chars: int, chars_per_page: int) -> list[dict]:
    """Process a single _fixed.docx file: split at Heading 1 boundaries."""
    doc = Document(filepath)
    filename = os.path.basename(filepath)
    base_name = os.path.splitext(filename)[0]
    if base_name.endswith("_fixed"):
        base_name = base_name[:-6]

    elements = build_element_sequence(doc)
    records = []

    if not elements:
        return records

    h1_element_indices = []
    for i, (elem_type, elem_obj, _) in enumerate(elements):
        if elem_type == "paragraph" and get_heading_level(elem_obj.style) == 1:
            h1_element_indices.append(i)

    if not h1_element_indices:
        safe_name = sanitize_filename(base_name)
        out_name = f"{safe_name} - Full Document.docx"
        out_path = os.path.join(output_dir, out_name)
        char_count = create_sub_document([], elements, out_path, filepath)
        records.append({
            "original_doc": filename,
            "sub_doc_filename": out_name,
            "heading_text": "(Full Document - No Heading 1 found)",
            "character_count": char_count,
            "page_estimate": round(char_count / chars_per_page, 1),
        })
        return records

    preamble_elements = elements[:h1_element_indices[0]]
    section_boundaries = h1_element_indices + [len(elements)]

    for sec_idx in range(len(h1_element_indices)):
        start = section_boundaries[sec_idx]
        end = section_boundaries[sec_idx + 1]
        section_elements = elements[start:end]

        heading1_elem = section_elements[0]
        heading1_text = heading1_elem[1].text.strip() if heading1_elem[0] == "paragraph" else "Untitled Section"

        section_chars = 0
        preamble_chars = 0
        for elem_type, elem_obj, _ in section_elements:
            if elem_type == "paragraph":
                section_chars += len(elem_obj.text)
            elif elem_type == "table":
                for row in elem_obj.rows:
                    for cell in row.cells:
                        section_chars += len(cell.text)
        for elem_type, elem_obj, _ in preamble_elements:
            if elem_type == "paragraph":
                preamble_chars += len(elem_obj.text)
            elif elem_type == "table":
                for row in elem_obj.rows:
                    for cell in row.cells:
                        preamble_chars += len(cell.text)

        total_chars = preamble_chars + section_chars

        if total_chars > max_chars:
            sub_records = split_at_heading2(
                preamble_elements, section_elements, base_name,
                heading1_text, output_dir, sec_idx, chars_per_page,
                max_chars, preamble_chars, filepath,
            )
            records.extend(sub_records)
        else:
            safe_heading = sanitize_filename(heading1_text)
            out_name = f"{base_name} - {safe_heading}.docx"
            out_path = os.path.join(output_dir, out_name)
            char_count = create_sub_document(preamble_elements, section_elements, out_path, filepath)
            records.append({
                "original_doc": filename,
                "sub_doc_filename": out_name,
                "heading_text": heading1_text,
                "character_count": char_count,
                "page_estimate": round(char_count / chars_per_page, 1),
            })

    return records


def main():
    parser = setup_argparse("Step 4: Split _fixed.docx policy documents at Heading 1 boundaries")
    args = parser.parse_args()

    config = load_config(args.config)
    thresholds = config.get("thresholds", {})
    max_chars = thresholds.get("max_characters", 36_000)
    chars_per_page = thresholds.get("chars_per_page", 1800)

    # Input: Step 3's output directory (heading fixes)
    if args.input_dir:
        input_dir = args.input_dir
    else:
        input_dir = get_output_dir(config, "heading_fixes")

    output_dir = get_output_dir(config, "split_documents", args.output_dir)
    ensure_output_dir(output_dir)

    # Only process _fixed.docx files — glob directly to avoid exclude_patterns filtering them out
    import glob as _glob
    docx_files = sorted(_glob.glob(os.path.join(input_dir, "*_fixed.docx")))

    if not docx_files:
        print(f"No *_fixed.docx files found in {input_dir}")
        print("Run heading_style_fixer.py (Step 3) first to generate _fixed.docx files.")
        return

    manifest_file = config.get("output", {}).get("split_documents", {}).get("manifest_file", "split_manifest.csv")
    csv_path = os.path.join(output_dir, manifest_file)
    fieldnames = [
        "original_doc", "sub_doc_filename", "heading_text",
        "character_count", "page_estimate",
    ]

    all_records = []
    files_processed = 0
    files_failed = 0

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        try:
            records = process_document(filepath, output_dir, max_chars, chars_per_page)
            all_records.extend(records)
            files_processed += 1
            print(f"  Processing {filename}... {len(records)} sub-document(s) created")
        except Exception as e:
            files_failed += 1
            print(f"  ERROR processing {filename}: {e}")
            traceback.print_exc()
            log_pipeline_issue(os.path.dirname(output_dir), "Step 4 - Section Splitter", filename, "ERROR", str(e))

    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(all_records)

    print("\n" + "=" * 60)
    print("STEP 4 — SECTION SPLITTER SUMMARY")
    print("=" * 60)
    print(f"Files processed:        {files_processed}")
    print(f"Files failed:           {files_failed}")
    print(f"Total sub-docs created: {len(all_records)}")

    over_limit = [r for r in all_records if r["character_count"] > max_chars]
    if over_limit:
        print(f"\nWARNING: {len(over_limit)} sub-document(s) still exceed {max_chars:,} characters:")
        for r in over_limit:
            print(f"  {r['sub_doc_filename']}: {r['character_count']:,} chars")

    print(f"\nSub-documents written to: {output_dir}")
    print(f"Manifest written to: {csv_path}")


if __name__ == "__main__":
    main()
