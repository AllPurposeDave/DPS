"""
Step 1: Compliance Document Control Extractor
=====================================================

Reads .docx compliance/security policy documents, extracts structured
control data (IDs, descriptions, supplemental guidance, and document
metadata), then writes everything into CSV and/or Excel output.

Features:
    - Multiple control ID patterns (any pattern triggers a match)
    - Whitelist/blacklist filtering by exact ID or prefix wildcard
    - 3-way text classification: description / supplemental guidance / miscellaneous
    - Configurable section heading detection (5 signal types)
    - CSV and/or Excel output (configurable)
    - Checkpointing for resumable batch runs

Usage (unified pipeline):
    python run_pipeline.py --step 1

Usage (standalone):
    python scripts/extract_controls.py
    python scripts/extract_controls.py --config dps_config.yaml
    python scripts/extract_controls.py ./input ./output

Output:
    controls_output.csv  — one row per extracted control (if CSV enabled)
    controls_output.xlsx — one row per extracted control (if Excel enabled)
    checkpoint.json      — tracks progress for resumable runs
    errors.log           — error and warning log

DEPENDENCIES:
    pip install python-docx pyyaml openpyxl
    Everything else is Python stdlib.
"""

import re
import csv
import os
import json
import logging
from pathlib import Path
from dataclasses import dataclass, asdict

from docx import Document

from shared_utils import (
    ensure_output_dir,
    get_input_dir,
    get_output_dir,
    iter_docx_files,
    load_config,
    log_pipeline_issue,
    resolve_path,
    setup_argparse,
)


# ============================================================
# SECTION: DATA MODEL
# ============================================================

@dataclass
class ControlRow:
    """Represents a single row in the output."""
    source_file: str = ""
    section_header: str = ""
    control_id: str = ""
    control_name: str = ""
    baseline: str = ""
    control_description: str = ""
    supplemental_guidance: str = ""
    miscellaneous: str = ""
    extraction_source: str = ""
    purpose: str = ""
    scope: str = ""
    applicability: str = ""
    compliance_date: str = ""
    published_url: str = ""


# ============================================================
# SECTION: REGEX PATTERNS (built from config)
# ============================================================

def build_patterns(config: dict) -> dict:
    """Build all regex patterns from config, with sensible defaults."""
    ctrl_cfg = config.get("control_extraction", {})

    id_patterns = ctrl_cfg.get("control_id_patterns", [r'\b[A-Z]{2,4}[-.]?\d{1,3}[-.]\d{2,4}\b'])

    # Build combined regex from all patterns
    combined_pattern = "|".join(f"({p})" for p in id_patterns)
    control_id_combined = re.compile(combined_pattern)

    impl_trigger = ctrl_cfg.get("implementation_trigger",
                                r'(?i)(implementation guidance|implementation:|guidelines:|how to implement|supplemental guidance)')

    # Build guidance regex from keywords list if available, else use impl_trigger
    guidance_keywords = ctrl_cfg.get("guidance_keywords", None)
    if guidance_keywords:
        escaped_keywords = [re.escape(kw) for kw in guidance_keywords]
        guidance_pattern = "|".join(escaped_keywords)
        guidance_regex = re.compile(f"(?i)({guidance_pattern})")
    else:
        guidance_regex = re.compile(impl_trigger)

    # Build metadata trigger regexes from config or defaults
    meta_cfg = ctrl_cfg.get("metadata_triggers", {})
    metadata_regexes = {}
    all_metadata_words = []
    default_meta = {
        "purpose": ["purpose", "objective", "intent"],
        "scope": ["scope", "coverage", "boundary"],
        "applicability": ["applicability", "applies to", "applies for"],
    }
    for field_name, default_keywords in default_meta.items():
        keywords = meta_cfg.get(field_name, default_keywords)
        escaped = [re.escape(kw) for kw in keywords]
        pattern = r'\b(' + "|".join(escaped) + r')\b'
        metadata_regexes[field_name] = re.compile(pattern, re.IGNORECASE)
        all_metadata_words.extend(keywords)

    # Build metadata label strip regex.
    # Handles compound headings like "Applicability and Scope" or
    # "Applicability & Scope" — after stripping the first keyword, also strip
    # "and <keyword>" or "& <keyword>" if it immediately follows.
    escaped_all = [re.escape(w) for w in all_metadata_words]
    kw_group = "|".join(escaped_all)
    strip_pattern = r'(?i)^\s*(' + kw_group + r')(\s+(?:and|&)\s+(' + kw_group + r'))?\s*[:\-]?\s*'

    # Build heading detection regexes from config
    hd_cfg = ctrl_cfg.get("heading_detection", {})
    section_kw = hd_cfg.get("section_keyword_pattern", r'^[Ss]ection\s+\d{1,2}')
    numbered_title = hd_cfg.get("numbered_title_pattern", r'^\d{1,2}\.?\d{0,2}\s+[A-Z][a-zA-Z\s]{3,50}$')

    # --- Bold-filtering settings ---
    # require_bold: when True, only control IDs appearing in bold runs are extracted
    #   from normal body paragraphs. This prevents false positives from inline
    #   references like "See also AC-1.002" in description text.
    require_bold = ctrl_cfg.get("require_bold_control_id", False)
    # bold_fallback: if bold-only extraction finds ZERO controls in a document,
    #   automatically retry without the bold requirement. Catches docs where
    #   control IDs are never bold-formatted.
    bold_fallback = ctrl_cfg.get("bold_fallback_if_zero", True)
    # tables_ignore_bold: table cells bypass the bold filter entirely. Table cells
    #   often contain control IDs in non-bold text within structured grids.
    tables_ignore_bold = ctrl_cfg.get("tables_ignore_bold", True)
    # headings_ignore_bold: Word heading-styled paragraphs (Heading 1/2/3/etc.)
    #   bypass the bold filter. Control IDs placed in headings (e.g.,
    #   "Control MON01.001 (L, M, H) - Authenticated Vulnerability Scanning")
    #   inherit the heading style but are NOT bold-formatted in the run data.
    #   Without this bypass, bold_fallback_if_zero won't help because other
    #   controls (from tables or bold body text) ARE found, so the fallback
    #   never triggers — silently dropping every heading-based control.
    headings_ignore_bold = ctrl_cfg.get("headings_ignore_bold", True)
    anchor_start = ctrl_cfg.get("control_id_anchor_start", False)
    anchor_chars = ctrl_cfg.get("control_id_start_chars", 25)
    min_block_lines = ctrl_cfg.get("min_control_block_lines", 0)

    return {
        "control_id": control_id_combined,
        "require_bold_control_id": require_bold,
        "bold_fallback_if_zero": bold_fallback,
        "tables_ignore_bold": tables_ignore_bold,
        "headings_ignore_bold": headings_ignore_bold,
        "control_id_anchor_start": anchor_start,
        "control_id_start_chars": anchor_chars,
        "min_control_block_lines": min_block_lines,
        "guidance_regex": guidance_regex,
        "section_keyword": re.compile(section_kw),
        "numbered_title": re.compile(numbered_title),
        "heading_detection": {
            "use_word_heading_style": hd_cfg.get("use_word_heading_style", True),
            "detect_allcaps": hd_cfg.get("detect_allcaps", True),
            "allcaps_max_length": hd_cfg.get("allcaps_max_length", 80),
            "allcaps_min_words": hd_cfg.get("allcaps_min_words", 3),
            "detect_bold_short": hd_cfg.get("detect_bold_short", True),
            "bold_max_length": hd_cfg.get("bold_max_length", 60),
        },
        "metadata_triggers": metadata_regexes,
        "metadata_label_strip": re.compile(strip_pattern),
    }


# ============================================================
# SECTION: LOGGING SETUP
# ============================================================

def setup_logging(output_dir: str):
    """Configure logging to write errors to the output directory."""
    os.makedirs(output_dir, exist_ok=True)

    logger = logging.getLogger("extract_controls")
    logger.setLevel(logging.DEBUG)

    # Remove existing handlers to avoid duplicates on re-run
    logger.handlers.clear()

    file_handler = logging.FileHandler(os.path.join(output_dir, "errors.log"), mode="a")
    file_handler.setLevel(logging.WARNING)
    file_handler.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
    logger.addHandler(file_handler)

    stream_handler = logging.StreamHandler()
    stream_handler.setLevel(logging.WARNING)
    stream_handler.setFormatter(logging.Formatter("%(levelname)s: %(message)s"))
    logger.addHandler(stream_handler)

    return logger


# ============================================================
# SECTION: PARAGRAPH EXTRACTION
# ============================================================

def _classify_table(table):
    """Classify a table by its header row to determine extraction strategy.

    Returns one of:
        "control"  — structured control-definition table (has Control ID + Title/Requirement/Description)
        "skip"     — noise table (crosswalk, revision history) that should be excluded from extraction
        "other"    — unrecognized structure; fall through to cell-by-cell extraction

    How it works:
        Reads the first row (header row) of the table and normalizes each cell to
        lowercase.  Then checks for column-name combinations:

        1. If the table has a column matching "control id" (or "org control id")
           AND a column matching "title", "requirement", or "description", it is a
           control-definition table → "control".

        2. If the table has a "control id"-like column but NO title/description
           column (e.g. only NIST / CIS / FedRAMP columns), it is a framework
           crosswalk → "skip".

        3. If the table has "version" + "date" + "changes" columns, it is a
           revision history table whose "Changes" cells often mention control IDs
           in prose ("Added MON04.001") → "skip".

        4. Everything else → "other" (cell-by-cell extraction, same as before).
    """
    if not table.rows:
        return "other"

    headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

    has_control_id_col = any("control" in h and "id" in h for h in headers)
    has_title_or_desc = any(
        kw in h for h in headers for kw in ("title", "requirement", "description")
    )
    has_revision_cols = (
        any("version" in h for h in headers)
        and any("date" in h for h in headers)
        and any("change" in h for h in headers)
    )

    if has_control_id_col and has_title_or_desc:
        return "control"
    if has_control_id_col and not has_title_or_desc:
        # Crosswalk / mapping table — has control IDs but only framework columns
        return "skip"
    if has_revision_cols:
        # Revision history — mentions control IDs in "Changes" prose
        return "skip"
    return "other"


def _extract_structured_table_rows(table, control_id_regex):
    """Extract structured control rows from a control-definition table.

    Reads the header row to find column indices for:
        - control_id:   first column whose header contains "control" and "id"
        - title:        column whose header contains "title"
        - description:  column whose header contains "requirement" or "description"

    For each data row, if the control ID cell matches the control_id_regex,
    a pre-formed control block dict is returned with properly separated fields.
    This avoids the cell-by-cell flattening that would otherwise merge Title,
    Description, NIST mapping, Status, and Owner into one blob of text.

    Returns a list of dicts with keys: control_id, control_name, baseline,
    control_description, supplemental_guidance, miscellaneous, section_header, source.
    """
    if len(table.rows) < 2:
        return []

    headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

    # Find column indices
    id_col = None
    title_col = None
    desc_col = None
    for i, h in enumerate(headers):
        if id_col is None and "control" in h and "id" in h:
            id_col = i
        elif title_col is None and "title" in h:
            title_col = i
        elif desc_col is None and ("requirement" in h or "description" in h):
            desc_col = i

    if id_col is None:
        return []

    rows = []
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        id_text = cells[id_col] if id_col < len(cells) else ""

        # The control ID cell may include baseline like "MON03.001 (L, M, H)"
        id_match = control_id_regex.search(id_text)
        if not id_match:
            continue

        # Flatten findall tuple if needed
        match_result = control_id_regex.findall(id_text)
        if match_result and isinstance(match_result[0], tuple):
            control_id = next(g for g in match_result[0] if g)
        else:
            control_id = match_result[0] if match_result else id_text

        # Parse baseline from the ID cell (e.g. "MON03.001 (L, M, H)" → "L,M,H")
        baseline_match = _BASELINE_RE.search(id_text)
        baseline = baseline_match.group(1).replace(" ", "") if baseline_match else ""

        title = cells[title_col].strip() if title_col is not None and title_col < len(cells) else ""
        description = cells[desc_col].strip() if desc_col is not None and desc_col < len(cells) else ""

        rows.append({
            "control_id": control_id,
            "control_name": clean_text(title),
            "baseline": baseline,
            "control_description": clean_text(description),
            "supplemental_guidance": "",
            "miscellaneous": "",
            "section_header": "",
            "source": "Table",
        })

    return rows


def extract_paragraphs_from_docx(filepath):
    """Open a .docx file and return paragraphs plus structured table controls.

    Returns (paragraphs, structured_table_controls) where:
        - paragraphs: list of paragraph dicts for find_control_blocks()
        - structured_table_controls: list of pre-formed control block dicts
          extracted from tables with recognized Control ID column structure

    Table handling strategy:
        - "control" tables  → structured extraction (column-aware), cells excluded
                               from paragraph list to prevent duplicates
        - "skip" tables     → excluded entirely (crosswalk, revision history noise)
        - "other" tables    → cell-by-cell extraction into paragraph list (original
                               behavior, catches controls in non-standard tables)
    """
    document = Document(str(filepath))
    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        first_run_bold = False
        bold_text = ""
        if paragraph.runs:
            first_run_bold = bool(paragraph.runs[0].bold)
            bold_text = "".join(r.text for r in paragraph.runs if r.bold)
        style_name = paragraph.style.name if paragraph.style else ""
        paragraphs.append({"text": text, "bold": first_run_bold, "bold_text": bold_text, "style": style_name, "source": "Text"})

    # --- Table extraction with classification ---
    structured_table_controls = []

    for table in document.tables:
        table_type = _classify_table(table)

        if table_type == "control":
            # Structured extraction: column-aware, returns pre-formed control dicts.
            # We need the control_id_regex but don't have patterns here — use a broad
            # pattern. The real regex filtering happens later in find_control_blocks
            # and apply_filters. For structured tables we just need to find the ID.
            # Import the compiled regex from the caller via a simple broad match.
            broad_id_re = re.compile(r'\b[A-Z]{2,4}[-.]?\d{1,3}[-.]?\d{1,4}\b')
            structured_table_controls.extend(
                _extract_structured_table_rows(table, broad_id_re)
            )
            # Do NOT add these cells to the paragraph list — they are fully handled.

        elif table_type == "skip":
            # Crosswalk or revision history — exclude entirely to prevent
            # false-positive control extraction from framework mappings or
            # changelog prose like "Added MON04.001".
            pass

        else:
            # Unrecognized table structure — fall through to cell-by-cell extraction
            # so controls in non-standard tables are still captured.
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        bold_text = "".join(
                            r.text for para in cell.paragraphs for r in para.runs if r.bold
                        )
                        paragraphs.append({"text": text, "bold": False, "bold_text": bold_text, "style": "", "source": "Table"})

    return paragraphs, structured_table_controls


# ============================================================
# SECTION: SECTION HEADER DETECTION
# ============================================================

def is_section_header(paragraph_dict, patterns):
    """Determine whether a paragraph is a section header.

    Checks multiple signals in priority order, each configurable via
    the heading_detection section in dps_config.yaml.
    """
    text = paragraph_dict["text"]
    style = paragraph_dict["style"]
    bold = paragraph_dict["bold"]
    hd = patterns["heading_detection"]

    # Skip empty paragraphs and lines containing a control ID
    if not text or patterns["control_id"].search(text):
        return (False, "")

    # Signal 1: Word heading style
    if hd["use_word_heading_style"] and "Heading" in style:
        return (True, text)

    # Signal 2: "Section N" keyword
    if patterns["section_keyword"].match(text):
        return (True, text)

    # Signal 3: Numbered title
    if patterns["numbered_title"].match(text):
        return (True, text)

    # Signal 4: ALL-CAPS short text
    if hd["detect_allcaps"]:
        word_count = len(text.split())
        if (text == text.upper()
                and len(text) < hd["allcaps_max_length"]
                and not text.endswith(".")
                and word_count >= hd["allcaps_min_words"]):
            return (True, text)

    # Signal 5: Bold short text
    # Exclude lines ending with ":" — those are in-control labels (e.g. "Implementation:",
    # "Description:") that would incorrectly close an active control block if treated as headers.
    if hd["detect_bold_short"]:
        if bold and len(text) < hd["bold_max_length"] and not text.endswith(".") and not text.endswith(":"):
            return (True, text)

    return (False, "")


# ============================================================
# SECTION: METADATA EXTRACTION
# ============================================================

def extract_metadata(paragraphs, patterns, config):
    """Scan the first N paragraphs for document-level metadata."""
    metadata = {field: "" for field in patterns["metadata_triggers"]}
    scan_limit_cfg = config.get("control_extraction", {}).get("metadata_scan_paragraphs", 40)
    scan_limit = min(scan_limit_cfg, len(paragraphs))

    # Collect ALL matching fields per paragraph so compound headings like
    # "Applicability and Scope" populate both the scope and applicability fields.
    trigger_positions = []  # list of (para_index, [field_name, ...])
    for index in range(scan_limit):
        text = paragraphs[index]["text"]
        if not text:
            continue
        matched = []
        for field_name, pattern in patterns["metadata_triggers"].items():
            if pattern.search(text):
                matched.append(field_name)
        if matched:
            trigger_positions.append((index, matched))

    for position_index, (para_index, field_names) in enumerate(trigger_positions):
        if all(metadata[f] for f in field_names):
            continue

        collected_lines = []
        # Find the next trigger at a DIFFERENT paragraph index for stop_index
        stop_index = scan_limit
        for next_idx in range(position_index + 1, len(trigger_positions)):
            if trigger_positions[next_idx][0] != para_index:
                stop_index = trigger_positions[next_idx][0]
                break

        for capture_index in range(para_index, stop_index):
            paragraph_text = paragraphs[capture_index]["text"]
            if not paragraph_text:
                continue
            if capture_index > para_index:
                is_header, _ = is_section_header(paragraphs[capture_index], patterns)
                if is_header:
                    break
            collected_lines.append(paragraph_text)

        raw_metadata = " ".join(collected_lines)
        cleaned = patterns["metadata_label_strip"].sub("", raw_metadata, count=1)
        cleaned_text = clean_text(cleaned)

        for field_name in field_names:
            if not metadata[field_name]:
                metadata[field_name] = cleaned_text

    return metadata


# ============================================================
# SECTION: COMPLIANCE DATE EXTRACTION
# ============================================================

# Matches "as of Month D, YYYY" or "as of Month DD, YYYY" (comma optional)
_COMPLIANCE_DATE_RE = re.compile(
    r'(?i)\bas\s+of\s+([A-Za-z]+\s+\d{1,2},?\s+\d{4})'
)

def extract_compliance_date(paragraphs):
    """Scan paragraphs for a 'Compliance Date' header, then extract the date.

    Looks for a heading whose text matches 'Compliance Date' (case-insensitive),
    then scans subsequent paragraphs for the pattern '...as of Month ##, ####.'
    Returns the extracted date string (e.g. 'December 31, 2025'), or '' if not found.
    """
    header_re = re.compile(r'(?i)^\s*compliance\s+date\s*$')

    def _parse_date(text):
        match = _COMPLIANCE_DATE_RE.search(text)
        if not match:
            return ""
        date_str = re.sub(r'\s+', ' ', match.group(1)).strip()
        # Ensure comma after day number: "December 31 2025" -> "December 31, 2025"
        date_str = re.sub(r'(\d{1,2})\s+(\d{4})$', r'\1, \2', date_str)
        return date_str

    # First pass: find the header, then search nearby paragraphs
    for i, para in enumerate(paragraphs):
        if header_re.match(para["text"]):
            for j in range(i + 1, min(i + 11, len(paragraphs))):
                date_str = _parse_date(paragraphs[j]["text"])
                if date_str:
                    return date_str
            break  # header found but no date nearby — stop

    # Fallback: scan the entire document for the "as of" pattern
    for para in paragraphs:
        date_str = _parse_date(para["text"])
        if date_str:
            return date_str

    return ""


# ============================================================
# SECTION: CONTROL BLOCK SEGMENTATION
# ============================================================

# Regex to capture baseline indicators like (L), (L, M), (L, M, H), (H), etc.
# Matches all permutations of L, M, H (comma-separated, inside parentheses).
_BASELINE_RE = re.compile(
    r'\s*\(\s*([LMH](?:\s*,\s*[LMH])*)\s*\)'
)

# After baseline, an optional " - Control Name" or " – Control Name"
_CONTROL_NAME_RE = re.compile(
    r'\s*[-\u2013\u2014]\s*(.+)'
)


def parse_baseline_and_name(text, control_id):
    """Extract baseline indicators and control name from the text surrounding a control ID.

    Handles multiple heading formats found in real compliance documents:

        Format 1 — dash-separated, baseline after ID:
            "AC-1.001 (L, M, H) - Access Control Policy"

        Format 2 — dash-separated, no baseline:
            "AC-1.001 - Access Control Policy"

        Format 3 — space-separated, baseline at end (no dash):
            "MON02.001 SIEM Log Ingestion and Correlation (L, M, H)"
            The name is everything between the ID and the trailing baseline.

        Format 4 — name BEFORE the ID:
            "Log Retention Requirements  Control MON02.002  (L, M, H)"
            Common when authors put the title first. The word "Control" or
            "Control ID" is stripped as a noise prefix.

        Format 5 — dash-separated, baseline at end of name:
            "Control ID MON01.002 - Vulnerability Remediation SLAs (L, M, H)"

    Returns (baseline, control_name) — both empty strings if not found.
    """
    # Find where the control ID ends in the text
    id_pos = text.find(control_id)
    if id_pos == -1:
        return ("", "")

    after_id = text[id_pos + len(control_id):]

    baseline = ""
    control_name = ""

    # Try to match baseline immediately after ID: "ID (L, M, H) ..."
    bl_match = _BASELINE_RE.match(after_id)
    if bl_match:
        baseline = bl_match.group(1).replace(" ", "")  # normalize to "L,M,H"
        after_id = after_id[bl_match.end():]

    # Try to match control name after a dash separator: "... - Control Name"
    name_match = _CONTROL_NAME_RE.match(after_id)
    if name_match:
        control_name = name_match.group(1).strip()
        # Handle trailing baseline inside the name: "Name (L, M, H)"
        if not baseline:
            trailing_bl = _BASELINE_RE.search(control_name)
            if trailing_bl:
                baseline = trailing_bl.group(1).replace(" ", "")
                control_name = control_name[:trailing_bl.start()].strip()

    # No dash separator — try space-separated name with trailing baseline:
    #   "MON02.001 SIEM Log Ingestion and Correlation (L, M, H)"
    # The name is all the text between the ID and the "(L, M, H)" at the end.
    if not control_name and not name_match:
        remaining = after_id.strip()
        if remaining:
            trailing_bl = _BASELINE_RE.search(remaining)
            if trailing_bl:
                # Text before the baseline is the name
                candidate_name = remaining[:trailing_bl.start()].strip()
                if candidate_name and not baseline:
                    baseline = trailing_bl.group(1).replace(" ", "")
                if candidate_name:
                    control_name = candidate_name
            elif remaining and not remaining[0] in '(.':
                # No baseline, no dash — treat remaining text as name if it looks
                # like a title (starts with a letter, not punctuation).
                # This catches "MON02.001 SIEM Log Ingestion and Correlation"
                control_name = remaining

    # Handle name BEFORE the ID: "Access Control Policy ACC10.111 - (L, M, H)"
    # Also: "Log Retention Requirements  Control MON02.002  (L,M, H)"
    if not control_name and id_pos > 0:
        before_id = re.sub(r'[\s\-\u2013\u2014]+$', '', text[:id_pos]).strip()
        if before_id:
            # Strip noise prefixes: "Control" or "Control ID" that authors
            # place before the ID (e.g. "Control MON02.002").
            before_id = re.sub(r'(?i)\bcontrol\s*(id)?\s*$', '', before_id).strip()
            if before_id:
                control_name = before_id

    return (baseline, control_name)


def _strip_control_prefix(text, control_id):
    """Remove the control ID, baseline, control name, and any before-ID prefix.

    Strips everything that parse_baseline_and_name would identify as metadata,
    leaving only the description body text.

    Examples:
        "AC-1.001 (L, M, H) - Access Control Policy The org shall..."
            → "The org shall..."
        "MON02.001 SIEM Log Ingestion and Correlation (L, M, H)"
            → "" (entire line is ID + name + baseline)
        "Log Retention Requirements  Control MON02.002  (L,M, H)"
            → "" (entire line is name + ID + baseline)
        "Control MON01.001 (L, M, H) - Authenticated Vulnerability Scanning"
            → "" (entire line is prefix + ID + baseline + name)
    """
    id_pos = text.find(control_id)
    if id_pos == -1:
        return text

    # Start after the control ID
    after = text[id_pos + len(control_id):]

    # Strip baseline like "(L, M, H)"
    bl_match = _BASELINE_RE.match(after)
    if bl_match:
        after = after[bl_match.end():]

    # Strip dash-separated control name: "- Access Control Policy"
    name_match = _CONTROL_NAME_RE.match(after)
    if name_match:
        after = after[name_match.end():]
    else:
        # Space-separated name with trailing baseline:
        #   "SIEM Log Ingestion and Correlation (L, M, H)"
        # Strip everything up to and including the trailing baseline.
        trailing_bl = _BASELINE_RE.search(after)
        if trailing_bl:
            after = after[trailing_bl.end():]

    # Also strip any before-ID prefix text (e.g. "Control" or "Log Retention...")
    # by starting from after the ID rather than from the original text start.
    return after.strip()


def find_control_blocks(paragraphs, patterns):
    """Identify and segment individual control blocks from paragraphs.

    Splits each block into description / supplemental_guidance / miscellaneous
    using guidance keyword boundaries.
    """
    control_id_regex = patterns["control_id"]
    require_bold = patterns.get("require_bold_control_id", False)
    tables_ignore_bold = patterns.get("tables_ignore_bold", True)
    headings_ignore_bold = patterns.get("headings_ignore_bold", True)
    anchor_start = patterns.get("control_id_anchor_start", False)
    anchor_chars = patterns.get("control_id_start_chars", 25)
    min_block_lines = patterns.get("min_control_block_lines", 0)
    blocks = []
    current_section_header = ""
    active_block = None

    for paragraph_dict in paragraphs:
        text = paragraph_dict["text"]
        source = paragraph_dict.get("source", "Text")
        is_header, header_text = is_section_header(paragraph_dict, patterns)
        if is_header:
            if active_block is not None:
                blocks.append(active_block)
                active_block = None
            current_section_header = header_text
            continue

        if not text:
            continue

        # --- Bold filtering decision ---
        # Three paragraph sources can contain control IDs:
        #   1. Normal body text  — bold filter APPLIES (prevents inline ref noise)
        #   2. Table cells       — bold filter SKIPPED (tables_ignore_bold)
        #   3. Heading-styled    — bold filter SKIPPED (headings_ignore_bold)
        #
        # Why headings need a bypass: Word heading styles (Heading 1/2/3) do not
        # store text as bold in the run-level XML, even if the heading renders
        # bold visually. So paragraph.runs[0].bold is False and bold_text is "".
        # Without this bypass, controls like "Control MON01.001 (L, M, H) -
        # Authenticated Vulnerability Scanning" in a Heading 3 paragraph are
        # silently dropped.
        #
        # Note: is_section_header() (called above) already returns (False, "")
        # for any paragraph containing a control ID, so heading-styled paragraphs
        # WITH control IDs flow through to here rather than being consumed as
        # section headers.
        is_table = (source == "Table")
        style = paragraph_dict.get("style", "")
        is_heading = ("Heading" in style)
        use_bold = (require_bold
                    and not (tables_ignore_bold and is_table)
                    and not (headings_ignore_bold and is_heading))

        # When bold filtering is active, only match control IDs that appear in
        # bold runs. This prevents false positives from control IDs referenced
        # by name inside description or guidance text (e.g., "See also AC-1.002").
        if use_bold:
            bold_text = paragraph_dict.get("bold_text", "")
            bold_matches = control_id_regex.findall(bold_text)
            if bold_matches and isinstance(bold_matches[0], tuple):
                bold_matches = [next(g for g in m if g) for m in bold_matches]
            # Only keep IDs that actually appear in bold runs
            valid_bold_ids = set(bold_matches)
            all_matches = control_id_regex.findall(text)
            if all_matches and isinstance(all_matches[0], tuple):
                all_matches = [next(g for g in m if g) for m in all_matches]
            control_id_matches = [m for m in all_matches if m in valid_bold_ids]
        else:
            control_id_matches = control_id_regex.findall(text)

        # control_id_anchor_start: only keep IDs that appear near the start of
        # the paragraph. Filters out inline references like "See also AC-1.002"
        # without requiring bold formatting.
        if anchor_start and control_id_matches:
            anchored = []
            for mid in (control_id_matches if not isinstance(control_id_matches[0], tuple) else control_id_matches):
                cid = mid if isinstance(mid, str) else next(g for g in mid if g)
                pos = text.find(cid)
                if pos != -1 and pos <= anchor_chars:
                    anchored.append(cid)
            control_id_matches = anchored

        if control_id_matches:
            # Flatten: findall with groups returns tuples — take first non-empty.
            # Always flatten unconditionally to handle all code paths.
            if isinstance(control_id_matches[0], tuple):
                control_id_matches = [
                    next(g for g in m if g) for m in control_id_matches
                ]

            if active_block is not None:
                blocks.append(active_block)

            for match_index, control_id in enumerate(control_id_matches):
                baseline, control_name = parse_baseline_and_name(text, control_id)
                first_line = _strip_control_prefix(text, control_id)
                if match_index == 0:
                    active_block = {
                        "control_id": control_id,
                        "control_name": control_name,
                        "baseline": baseline,
                        "raw_lines": [first_line] if first_line else [],
                        "section_header": current_section_header,
                        "source": source,
                    }
                else:
                    blocks.append({
                        "control_id": control_id,
                        "control_name": control_name,
                        "baseline": baseline,
                        "raw_lines": [first_line] if first_line else [],
                        "section_header": current_section_header,
                        "source": source,
                    })
        elif active_block is not None:
            active_block["raw_lines"].append(text)

    if active_block is not None:
        blocks.append(active_block)

    # min_control_block_lines: discard blocks with too few content lines.
    # Inline references (e.g. "See AC-1.002") produce blocks with 0-1 lines;
    # real control definitions typically have multiple lines of description.
    if min_block_lines > 0:
        blocks = [b for b in blocks if len(b["raw_lines"]) >= min_block_lines]

    # Split each block's raw text into description / guidance / miscellaneous
    guidance_regex = patterns["guidance_regex"]
    processed_blocks = []

    for block in blocks:
        joined_text = "\n".join(block["raw_lines"])
        split_result = guidance_regex.split(joined_text, maxsplit=1)

        if len(split_result) >= 3:
            description_text = clean_text(split_result[0])
            guidance_text = clean_text(split_result[2])

            # Check for a second guidance trigger — text after goes to miscellaneous
            second_split = guidance_regex.split(split_result[2], maxsplit=1)
            if len(second_split) >= 3:
                guidance_text = clean_text(second_split[0])
                miscellaneous_text = clean_text(second_split[2])
            else:
                miscellaneous_text = ""
        else:
            description_text = clean_text(joined_text)
            guidance_text = ""
            miscellaneous_text = ""

        processed_blocks.append({
            "control_id": block["control_id"],
            "control_name": block.get("control_name", ""),
            "baseline": block.get("baseline", ""),
            "control_description": description_text,
            "supplemental_guidance": guidance_text,
            "miscellaneous": miscellaneous_text,
            "section_header": block["section_header"],
            "source": block["source"],
        })

    return processed_blocks


# ============================================================
# SECTION: WHITELIST / BLACKLIST FILTERING
# ============================================================

def matches_filter(control_id, filter_list):
    """Check if a control ID matches any entry in a filter list.

    Supports exact matches and prefix wildcards (e.g., "AC-*").
    """
    for pattern in filter_list:
        if pattern.endswith("*"):
            if control_id.startswith(pattern[:-1]):
                return True
        else:
            if control_id == pattern:
                return True
    return False


def apply_filters(control_blocks, config):
    """Apply whitelist and blacklist filters to control blocks."""
    ctrl_cfg = config.get("control_extraction", {})
    whitelist = ctrl_cfg.get("whitelist", [])
    blacklist = ctrl_cfg.get("blacklist", [])

    filtered = control_blocks

    if whitelist:
        filtered = [b for b in filtered if matches_filter(b["control_id"], whitelist)]

    if blacklist:
        filtered = [b for b in filtered if not matches_filter(b["control_id"], blacklist)]

    return filtered


# ============================================================
# SECTION: TEXT CLEANING
# ============================================================

def clean_text(text):
    """Normalize whitespace and remove non-printable characters."""
    text = text.strip()
    text = re.sub(r'[^\x20-\x7E\n]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


# ============================================================
# SECTION: URL LOOKUP
# ============================================================

def load_url_mapping(config: dict) -> dict:
    """Load document-name-to-URL mapping from the Doc_URL Excel file.

    Reads the same file used by Step 5 (metadata injection), configured at
    metadata.url.lookup_file. Returns a dict of {lowercase_name: url}.
    Empty dict if no file is configured or found.
    """
    from openpyxl import load_workbook

    meta_cfg = config.get("metadata", {})
    url_cfg = meta_cfg.get("url", {})
    lookup_file = url_cfg.get("lookup_file", "")

    if not lookup_file:
        return {}

    lookup_path = resolve_path(config, lookup_file)

    if not os.path.isfile(lookup_path):
        print(f"  NOTE: URL lookup file not found: {lookup_path}")
        print("  Published URL column will be empty. Create input/Doc_URL.xlsx to populate.")
        return {}

    name_col = url_cfg.get("name_column", "Document_Name").lower()
    url_col = url_cfg.get("url_column", "URL").lower()
    sheet_ref = url_cfg.get("sheet", 0)

    try:
        wb = load_workbook(lookup_path, read_only=True, data_only=True)
        if isinstance(sheet_ref, int):
            ws = wb.worksheets[sheet_ref]
        else:
            ws = wb[sheet_ref]
    except Exception as e:
        print(f"  WARNING: Could not read URL lookup file: {e}")
        return {}

    # Find column indices from header row
    headers = {}
    for col_idx, cell in enumerate(next(ws.iter_rows(min_row=1, max_row=1, values_only=False))):
        if cell.value:
            headers[str(cell.value).strip().lower()] = col_idx

    name_idx = headers.get(name_col)
    url_idx = headers.get(url_col)

    if name_idx is None or url_idx is None:
        print(f"  WARNING: Could not find columns '{name_col}' and/or '{url_col}' in {lookup_path}")
        print(f"  Found columns: {list(headers.keys())}")
        return {}

    mapping = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        name_val = row[name_idx] if name_idx < len(row) else None
        url_val = row[url_idx] if url_idx < len(row) else None
        if name_val and url_val:
            # Store original name (will be normalized in resolve_url for comparison)
            mapping[str(name_val).strip()] = str(url_val).strip()

    wb.close()
    if mapping:
        print(f"  Loaded {len(mapping)} URL mappings from {os.path.basename(lookup_path)}")
    return mapping


def resolve_url(filename: str, url_mapping: dict) -> str:
    """Resolve the URL for a document using substring matching.

    Same matching logic as Step 5: the Excel name is checked as a substring
    of the document name (case-insensitive), and vice versa. Both are normalized
    (underscores converted to spaces) before comparison.
    """
    if not url_mapping:
        return ""

    # Derive a clean document name from the filename
    # Normalize: remove extension, convert underscores to spaces, lowercase
    doc_name = os.path.splitext(filename)[0].replace("_", " ").lower()

    for excel_name_orig, url in url_mapping.items():
        # Normalize Excel name the same way for comparison
        excel_name_normalized = excel_name_orig.replace("_", " ").lower()
        if excel_name_normalized in doc_name or doc_name in excel_name_normalized:
            return url

    return ""


# ============================================================
# SECTION: SINGLE DOCUMENT PROCESSING
# ============================================================

def process_single_document(filepath, patterns, config, url_mapping=None):
    """Run the full extraction pipeline on a single .docx file.

    Two extraction paths run and their results are merged:

        1. Structured table extraction — for tables with a recognized
           "Control ID" + "Title"/"Requirement"/"Description" column layout.
           Returns pre-formed control dicts with properly separated fields.
           Crosswalk and revision-history tables are skipped entirely.

        2. Paragraph-based extraction — scans body text (and cells from
           unrecognized tables) for control IDs using regex + bold filtering.
           Catches controls in headings, prose, and non-standard table layouts.

    Deduplication: if both paths find the same control_id, the structured
    table row wins because it has clean column-separated name/description.
    Paragraph-based rows for IDs already found in structured tables are dropped.
    """
    filepath = Path(filepath)
    source_file = filepath.name

    # extract_paragraphs_from_docx now returns two things:
    #   paragraphs — flat list for paragraph-based extraction
    #   structured_table_controls — pre-formed dicts from column-aware table parsing
    paragraphs, structured_table_controls = extract_paragraphs_from_docx(filepath)

    if not paragraphs and not structured_table_controls:
        print(f"  WARNING: {source_file} has 0 paragraphs, skipping.")
        return []

    metadata = extract_metadata(paragraphs, patterns, config)
    compliance_date = extract_compliance_date(paragraphs)

    purpose_found = "Yes" if metadata.get("purpose") else "No"
    scope_found = "Yes" if metadata.get("scope") else "No"
    applicability_found = "Yes" if metadata.get("applicability") else "No"
    print(f"  Metadata found -- Purpose: {purpose_found} | "
          f"Scope: {scope_found} | Applicability: {applicability_found}")

    if structured_table_controls:
        print(f"  Structured table controls: {len(structured_table_controls)}")

    # --- Paragraph-based extraction (body text + unrecognized table cells) ---
    control_blocks = find_control_blocks(paragraphs, patterns)

    # bold_fallback_if_zero: if bold-only extraction found nothing, retry
    # without the bold requirement so documents with non-bold IDs still extract.
    if (not control_blocks
            and patterns.get("require_bold_control_id", False)
            and patterns.get("bold_fallback_if_zero", True)):
        fallback_patterns = dict(patterns, require_bold_control_id=False)
        control_blocks = find_control_blocks(paragraphs, fallback_patterns)
        if control_blocks:
            print(f"  Bold fallback: re-extracted {len(control_blocks)} controls without bold requirement")

    # --- Merge: structured table rows take priority, then paragraph-based ---
    # Structured rows have clean column-separated data; paragraph rows for the
    # same ID would have polluted descriptions (NIST/Status/Owner concatenated).
    structured_ids = {b["control_id"] for b in structured_table_controls}
    paragraph_only_blocks = [b for b in control_blocks if b["control_id"] not in structured_ids]

    merged_blocks = structured_table_controls + paragraph_only_blocks
    print(f"  Controls extracted: {len(merged_blocks)} "
          f"({len(structured_table_controls)} from tables, {len(paragraph_only_blocks)} from text)")

    # Apply whitelist/blacklist filtering
    merged_blocks = apply_filters(merged_blocks, config)
    ctrl_cfg = config.get("control_extraction", {})
    if ctrl_cfg.get("whitelist") or ctrl_cfg.get("blacklist"):
        print(f"  Controls after filtering: {len(merged_blocks)}")

    published_url = resolve_url(source_file, url_mapping or {})

    rows = []
    for block in merged_blocks:
        row = ControlRow(
            source_file=source_file,
            section_header=block["section_header"],
            control_id=block["control_id"],
            control_name=block.get("control_name", ""),
            baseline=block.get("baseline", ""),
            control_description=block["control_description"],
            supplemental_guidance=block["supplemental_guidance"],
            miscellaneous=block["miscellaneous"],
            extraction_source=block.get("source", "Text"),
            purpose=metadata.get("purpose", ""),
            scope=metadata.get("scope", ""),
            applicability=metadata.get("applicability", ""),
            compliance_date=compliance_date,
            published_url=published_url,
        )
        rows.append(row)

    return rows


# ============================================================
# SECTION: CHECKPOINTING
# ============================================================

def load_checkpoint(checkpoint_path):
    if os.path.exists(checkpoint_path):
        with open(checkpoint_path, "r") as f:
            return json.load(f)
    return []


def save_checkpoint(checkpoint_path, completed_files):
    with open(checkpoint_path, "w") as f:
        json.dump(completed_files, f, indent=2)


# ============================================================
# SECTION: OUTPUT WRITERS
# ============================================================

CSV_COLUMNS = [
    "control_id", "control_name", "baseline",
    "control_description", "supplemental_guidance", "purpose", "scope", "applicability",
    "miscellaneous", "section_header", "source_file", "extraction_source",
    "compliance_date", "published_url",
]


def write_rows_to_csv(output_path, rows, append=True):
    """Write control rows to a CSV file."""
    file_exists = os.path.exists(output_path) and os.path.getsize(output_path) > 0
    write_mode = "a" if (file_exists and append) else "w"

    with open(output_path, write_mode, newline="", encoding="utf-8") as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=CSV_COLUMNS, quoting=csv.QUOTE_ALL)
        if not file_exists or not append:
            writer.writeheader()
        for row in rows:
            writer.writerow(asdict(row))


def write_rows_to_excel(output_path, rows):
    """Write control rows to an Excel file, creating it if needed."""
    from openpyxl import Workbook, load_workbook

    output_path = Path(output_path)
    if output_path.exists() and output_path.stat().st_size > 0:
        wb = load_workbook(output_path)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "Controls"
        ws.append(CSV_COLUMNS)

    for row in rows:
        row_dict = asdict(row)
        ws.append([row_dict[col] for col in CSV_COLUMNS])

    wb.save(output_path)


# ============================================================
# SECTION: MAIN ENTRY POINT
# ============================================================

def main():
    parser = setup_argparse("Step 1: Extract controls from compliance .docx documents")
    args = parser.parse_args()

    config = load_config(args.config)
    ctrl_cfg = config.get("control_extraction", {})
    input_dir = get_input_dir(config, args.input_dir)
    output_dir = get_output_dir(config, "controls", args.output_dir)
    ensure_output_dir(output_dir)

    logger = setup_logging(output_dir)
    patterns = build_patterns(config)

    # Determine output format
    output_format = ctrl_cfg.get("output_format", "both")
    output_csv_file = config.get("output", {}).get("controls", {}).get("output_file", "controls_output.csv")
    output_xlsx_file = config.get("output", {}).get("controls", {}).get("output_file_xlsx", "controls_output.xlsx")
    checkpoint_file = config.get("output", {}).get("controls", {}).get("checkpoint_file", "checkpoint.json")

    checkpoint_path = os.path.join(output_dir, checkpoint_file)
    output_csv_path = os.path.join(output_dir, output_csv_file)
    output_xlsx_path = os.path.join(output_dir, output_xlsx_file)

    url_mapping = load_url_mapping(config)

    docx_files = iter_docx_files(input_dir, config)
    if not docx_files:
        print(f"No .docx files found in {input_dir}")
        return

    enable_checkpoint = ctrl_cfg.get("enable_checkpoint", True)
    completed_files = load_checkpoint(checkpoint_path) if enable_checkpoint else []
    total_controls = 0
    files_processed = 0
    files_skipped = 0
    error_count = 0

    for filepath in docx_files:
        filename = os.path.basename(filepath)

        if filename in completed_files:
            print(f"  Skipping (already processed): {filename}")
            files_skipped += 1
            continue

        print(f"  Processing: {filename}")

        try:
            rows = process_single_document(filepath, patterns, config, url_mapping)

            if rows:
                if output_format in ("csv", "both"):
                    write_rows_to_csv(output_csv_path, rows)
                if output_format in ("xlsx", "both"):
                    write_rows_to_excel(output_xlsx_path, rows)

            total_controls += len(rows)

            if enable_checkpoint:
                completed_files.append(filename)
                save_checkpoint(checkpoint_path, completed_files)
            files_processed += 1

        except Exception as error:
            error_count += 1
            logger.error("Failed to process %s: %s", filename, error, exc_info=True)
            print(f"  ERROR processing {filename}: {error}")
            log_pipeline_issue(os.path.dirname(output_dir), "Step 1 - Controls", filename, "ERROR", str(error))

    print("\n" + "=" * 60)
    print("STEP 1 — CONTROL EXTRACTION SUMMARY")
    print("=" * 60)
    print(f"Files processed: {files_processed}")
    print(f"Controls found:  {total_controls}")
    print(f"Errors:          {error_count}")
    if files_skipped:
        print(f"Skipped (checkpoint): {files_skipped}")

    outputs = []
    if output_format in ("csv", "both"):
        outputs.append(output_csv_path)
    if output_format in ("xlsx", "both"):
        outputs.append(output_xlsx_path)
    print(f"\nOutput: {', '.join(outputs)}")


if __name__ == "__main__":
    main()
