#!/usr/bin/env python3
"""
Policy Document Profiler for Copilot KB Optimization
=====================================================
Scans all .docx files in a directory, extracts structural metadata,
classifies documents by type (A/B/C/D), scores optimization priority,
and outputs inventories for downstream processing.

Feeds into: Notebook 1 (profiling), Notebook 2 (transformation), 
            Notebook 3 (QA validation), deep thinking (calibration).

Usage:
    python run_pipeline.py --step 0                           # unified pipeline
    python scripts/policy_profiler.py --config dps_config.yaml  # standalone with unified config
    python scripts/policy_profiler.py --config profiler_config.yaml --input ./docs/
    python scripts/policy_profiler.py --help

Requirements:
    pip install python-docx pyyaml openpyxl
"""

import argparse
import csv
import glob
import json
import os
import re
import sys
import time
import traceback
from collections import OrderedDict
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# IMPORTANT: Do NOT import third-party packages here at the top level.
# They are imported inside check_environment() below so that missing packages
# produce a clear, actionable error message instead of a raw Python traceback.
# ---------------------------------------------------------------------------


def check_environment():
    """
    Run this before anything else.

    Checks that every required third-party package is installed.
    If anything is missing, prints a plain-English fix and exits cleanly
    — no Python traceback, no confusing error codes.

    WHAT THIS CHECKS:
      - python-docx  : reads .docx files
      - pyyaml       : reads the profiler_config.yaml settings file
      - openpyxl     : writes the Excel inventory spreadsheet

    IF YOU SEE "MISSING PACKAGE" ERRORS:
      1. Open a terminal / command prompt (Windows: search "cmd" or "PowerShell")
      2. Run exactly this command:
             pip install python-docx pyyaml openpyxl
      3. Wait for it to finish, then run this script again.
      If pip itself is not found, Python may not be on your PATH.
      Fix: re-run the Python installer and check "Add Python to PATH".
    """
    required = {
        "docx":     "python-docx",
        "yaml":     "pyyaml",
        "openpyxl": "openpyxl",
    }
    missing = []
    for import_name, package_name in required.items():
        try:
            __import__(import_name)
        except ImportError:
            missing.append(package_name)

    if missing:
        print("=" * 60)
        print("SETUP PROBLEM: Required packages are not installed.")
        print("=" * 60)
        for pkg in missing:
            print(f"  MISSING: {pkg}")
        print()
        print("TO FIX: Open a terminal/command prompt and run:")
        print()
        print("    pip install " + " ".join(missing))
        print()
        print("Then run this script again.")
        print("=" * 60)
        sys.exit(1)


# Run the check immediately — before any other imports that would crash.
check_environment()

# Now it is safe to import third-party packages.
import yaml
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class TableInfo:
    """Metadata for a single table in a document."""
    index: int
    rows: int
    cols: int
    header_row: list
    has_merged_cells: bool
    is_nested: bool
    classification: str  # control_matrix, applicability_table, etc.
    parent_section: str
    char_count: int

@dataclass
class SectionInfo:
    """Metadata for a heading-delimited section."""
    heading_text: str
    heading_level: int
    heading_style: str  # actual Word style name
    is_builtin_style: bool
    is_fake_heading: bool
    char_count: int
    table_count: int
    paragraph_count: int
    start_para_index: int
    standard_section: str  # purpose/scope/intent/controls/appendix/none
    cross_ref_count: int

@dataclass
class CrossRefInfo:
    """A single cross-reference instance."""
    text: str
    pattern_matched: str
    paragraph_index: int
    parent_section: str
    is_hyperlink: bool

@dataclass
class DocumentProfile:
    """Complete profile for one .docx file."""
    filename: str
    filepath: str
    file_size_bytes: int
    total_char_count: int
    total_word_count: int
    total_paragraphs: int
    approx_pages: float
    over_size_limit: bool

    # Heading analysis
    heading_count: int
    h1_count: int
    h2_count: int
    h3_count: int
    h4_count: int
    builtin_heading_count: int
    custom_heading_count: int
    fake_heading_count: int
    heading_styles_used: list

    # Section structure
    sections: list  # list of SectionInfo dicts
    has_purpose: bool
    has_scope: bool
    has_intent: bool
    has_controls: bool
    has_appendix: bool
    missing_sections: list
    section_count: int

    # Section size distribution
    largest_section_name: str
    largest_section_chars: int
    largest_section_pct: float
    controls_section_chars: int
    controls_section_pct: float
    appendix_section_chars: int
    appendix_section_pct: float

    # Tables
    table_count: int
    tables: list  # list of TableInfo dicts
    total_table_chars: int
    table_content_pct: float
    tables_with_merged_cells: int
    nested_table_count: int
    table_types: dict  # classification -> count

    # Cross-references
    cross_ref_count: int
    cross_refs: list  # list of CrossRefInfo dicts
    cross_ref_patterns: dict  # pattern -> count

    # Formatting anomalies
    has_text_boxes: bool
    has_tracked_changes: bool
    has_comments: bool
    has_images: bool
    has_embedded_objects: bool
    has_password_protection: bool

    # Classification
    doc_type: str  # A, B, C, D, E
    doc_type_reason: str

    # Profiling flags
    control_id_count: int
    control_density: float  # control IDs per approx page
    control_dense: bool
    level_skips: int  # heading level jumps > 1 (e.g., H1→H3)
    heading_variance: bool
    heading_variance_reason: str
    unique_sections: list  # H1 headings not matching standard section terms
    unique_section_count: int
    table_dense: bool

    # Priority scoring
    priority_score: float
    priority_rank: int  # filled after all docs scored

    # Key term search results
    search_term_hits: dict  # {term: count}

    # Errors
    errors: list

    def to_dict(self):
        """Convert to dict, handling nested dataclass lists."""
        d = {}
        for k, v in self.__dict__.items():
            if isinstance(v, list) and v and hasattr(v[0], '__dict__'):
                d[k] = [x.__dict__ if hasattr(x, '__dict__') else x for x in v]
            else:
                d[k] = v
        return d


# ============================================================================
# Config Loader
# ============================================================================

def load_config(config_path: str) -> dict:
    """
    Read and return the YAML configuration file.

    WHAT IT DOES:
      Opens profiler_config.yaml (or the path you passed with --config) and
      parses all the settings that drive the rest of the script: input/output
      paths, heading detection rules, cross-reference patterns, thresholds, etc.

    FAILURE MODES:

      "Config file not found" → You are running the script from the wrong folder,
        OR the --config path is wrong.
        FIX: Either (a) open your terminal, cd to the folder that contains both
        policy_profiler.py AND profiler_config.yaml, then run the script; or
        (b) pass the full path explicitly:
              python policy_profiler.py --config "C:/full/path/profiler_config.yaml"

      "Config file has a formatting error" → The YAML file was edited and a
        typo was introduced (wrong indentation, stray colon, etc.).
        FIX: Open profiler_config.yaml in a text editor. The error message
        includes the line number. YAML is sensitive to indentation — use
        spaces, never tabs. Each level of nesting = 2 more spaces.

      "Config is empty or invalid" → The file exists but is blank or contains
        only comments.
        FIX: Restore from a backup, or re-download the original config.
    """
    config_path = os.path.abspath(config_path)

    # --- File existence check ------------------------------------------------
    if not os.path.isfile(config_path):
        print("=" * 60)
        print("ERROR: Config file not found.")
        print(f"  Expected: {config_path}")
        print()
        print("Are you running the script from the right folder?")
        print("The config file must be in the same folder as the script,")
        print("OR you must pass the full path with --config.")
        print()
        print("EXAMPLE (full path):")
        print('  python policy_profiler.py --config "C:/MyDocs/profiler_config.yaml"')
        print("=" * 60)
        sys.exit(1)

    # --- Parse YAML ----------------------------------------------------------
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            cfg = yaml.safe_load(f)
    except yaml.YAMLError as e:
        print("=" * 60)
        print("ERROR: Config file has a formatting error.")
        print(f"  File: {config_path}")
        print()
        # yaml exceptions often include line/column info
        if hasattr(e, 'problem_mark') and e.problem_mark:
            print(f"  Problem near line {e.problem_mark.line + 1}, "
                  f"column {e.problem_mark.column + 1}")
        print(f"  Detail: {e}")
        print()
        print("YAML rules to check:")
        print("  - Use 2 spaces per indentation level (no tabs)")
        print("  - Colons must be followed by a space:  key: value")
        print("  - Lists start with '  - item' (2 spaces, dash, space)")
        print("=" * 60)
        sys.exit(1)
    except Exception as e:
        print(f"ERROR: Could not read config file: {e}")
        sys.exit(1)

    # --- Sanity check --------------------------------------------------------
    if not cfg or not isinstance(cfg, dict):
        print("=" * 60)
        print("ERROR: Config file appears to be empty or invalid.")
        print(f"  File: {config_path}")
        print("FIX: Restore profiler_config.yaml from a backup.")
        print("=" * 60)
        sys.exit(1)

    return cfg


# ============================================================================
# Core Profiling Functions
# ============================================================================

def detect_heading_info(paragraph, config: dict) -> tuple:
    """
    Decide whether a paragraph is a heading, and what kind.

    Returns a tuple: (level, style_name, is_builtin, is_custom, is_fake)
      - level=0 means "not a heading at all"
      - level=1,2,3,4 matches Word's Heading 1 / Heading 2 / etc.
      - is_builtin=True  → paragraph uses a standard Word heading style
      - is_custom=True   → paragraph uses an org-specific custom style
                           (add those names to profiler_config.yaml under
                            headings: custom_heading_styles)
      - is_fake=True     → paragraph is NOT a real heading style, but it
                           looks like one (bold + short + large font).
                           These break Copilot chunking and must be fixed.

    HOW HEADING DETECTION WORKS:
      1. Check the paragraph's Word style against the builtin list.
      2. Check against the custom list.
      3. If neither matches, check for "fake heading" signals:
         bold text + short length + font size >= fake_heading_min_font_size.

    FAILURE MODES:

      "0 fake headings" but the doc clearly uses bold headers →
        The font size may be at or below fake_heading_min_font_size (default 12).
        FIX: Lower fake_heading_min_font_size to 11 in profiler_config.yaml.
        Also check: if the bold text's font size is not explicitly set in Word
        (i.e., it inherits from the paragraph style), run.font.size returns None.
        The code handles this by still flagging it if it's bold + short (≤100 chars).

      Headings show as "fake" but they're real org-specific headings →
        Your org's Word template uses a custom style name not in the default list.
        FIX: In Word, click the heading paragraph, look at the Styles pane (Home tab).
        Note the exact style name. Add it to profiler_config.yaml under:
          headings:
            custom_heading_styles:
              - "Your Exact Style Name Here"

      All headings show at level 1 (H1) even if some are subheadings →
        The document uses custom styles without a number in the name (e.g. "Policy
        Header" instead of "Policy Heading 2"). The level is inferred from the digit
        in the style name. Without a digit, it defaults to level 1.
        FIX: This is informational — the section structure will still be captured,
        just without sub-level granularity.
    """
    style_name = paragraph.style.name if paragraph.style else ""
    builtin = config['headings']['builtin_styles']
    custom = config['headings']['custom_heading_styles']

    # --- Step 1: Check built-in Word heading styles (Heading 1, Heading 2, etc.) ---
    # These are the "real" headings that Word and Copilot both understand.
    for bs in builtin:
        if style_name.lower() == bs.lower():
            level = int(re.search(r'\d', style_name).group()) if re.search(r'\d', style_name) else 1
            return (level, style_name, True, False, False)

    # --- Step 2: Check org-specific custom heading styles ---------------------
    # If your org's Word template uses non-standard style names, add them to
    # profiler_config.yaml → headings → custom_heading_styles.
    for cs in custom:
        if style_name.lower() == cs.lower():
            level_match = re.search(r'\d', style_name)
            level = int(level_match.group()) if level_match else 1
            return (level, style_name, False, True, False)

    # --- Step 3: Detect fake headings ----------------------------------------
    # A "fake heading" is a normal paragraph that someone made look like a
    # heading by making it bold (and sometimes larger font), without applying
    # an actual Word heading style. These are invisible to Word's navigation
    # pane, invisible to Copilot's chunking, and must be fixed before upload.
    text = paragraph.text.strip()
    if not text:
        return (0, style_name, False, False, False)

    is_bold = False
    max_font_size = 0
    for run in paragraph.runs:
        if run.bold:
            is_bold = True
        if run.font.size:
            size_pt = run.font.size.pt
            if size_pt > max_font_size:
                max_font_size = size_pt

    min_size = config['headings']['fake_heading_min_font_size']
    max_chars = config['headings']['fake_heading_max_chars']

    if is_bold and len(text) <= max_chars and (max_font_size >= min_size or max_font_size == 0):
        # Bold, short, potentially large font = fake heading
        # If font size is 0 (not explicitly set), still flag if bold + short
        if max_font_size == 0 and len(text) <= 100:
            return (0, style_name, False, False, True)
        elif max_font_size >= min_size:
            return (0, style_name, False, False, True)

    return (0, style_name, False, False, False)


def match_standard_section(heading_text: str, config: dict) -> str:
    """Match heading text against standard section patterns. Returns section key or 'none'."""
    text_lower = heading_text.lower().strip()
    for section_key, patterns in config['sections'].items():
        for pattern in patterns:
            if pattern.lower() in text_lower or text_lower in pattern.lower():
                return section_key
    return "none"


def search_key_terms(full_text: str, table_text: str, config: dict) -> dict:
    """
    Search for configured key terms in the document's full text and table text.
    Returns a dict mapping each term to its occurrence count.
    """
    search_cfg = config.get("search_terms", {})
    if not search_cfg.get("enabled", False):
        return {}

    terms = search_cfg.get("terms", [])
    if not terms:
        return {}

    match_mode = search_cfg.get("match_mode", "word")
    combined_text = full_text + "\n" + table_text

    results = {}
    for term in terms:
        if match_mode == "word":
            pattern = r'\b' + re.escape(term) + r'\b'
            count = len(re.findall(pattern, combined_text, re.IGNORECASE))
        else:
            count = len(re.findall(re.escape(term), combined_text, re.IGNORECASE))
        results[term] = count

    return results


def classify_table(table, config: dict) -> tuple:
    """
    Classify a table by examining its first row (assumed to be the header).

    Returns: (classification_str, header_row_list, has_merged_cells, char_count)
      - classification_str: one of control_matrix / applicability_table /
        reference_table / crosswalk_table / role_responsibility / unclassified
      - header_row_list: list of cell text strings from the first row
      - has_merged_cells: True if any cells span multiple rows or columns
      - char_count: total characters across all cells (used for table_content_pct)

    HOW CLASSIFICATION WORKS:
      The first row is matched against keyword lists in profiler_config.yaml
      under tables: classification. The type with the most keyword hits wins.
      A minimum column count is also enforced per type.

    FAILURE MODES:

      Table classified as "unclassified" but you know what type it is →
        Either the header row keywords don't match, or the table has no header row
        (data starts in row 1 with no column labels).
        FIX (keyword mismatch): Add the actual header text as a keyword in
        profiler_config.yaml under tables: classification: [type]: keywords.
        FIX (no header row): The profiler can only read what's there. This is
        flagged in the tables_inventory.csv so you can handle it manually.

      Tables with merged cells show "has_merged_cells: True" unexpectedly →
        Merged cells are detected by checking if two cells in the same row share
        the same underlying XML element. This is reliable — if it fires, the
        merge is real. These tables need the explicit structural callout approach
        in Notebook 2 (per the workflow doc).

      Table row/column counts look wrong →
        python-docx counts all rows including header rows. A table that looks like
        "8 data rows" will show as 9 rows if it has a header. This is expected.
        Nested tables (a table inside a table cell) are not counted as sub-rows —
        they appear as a separate table entry in the list.
    """
    header_row = []
    has_merged = False
    char_count = 0

    # --- Count all characters across every cell (for table_content_pct) ------
    for row in table.rows:
        for cell in row.cells:
            char_count += len(cell.text)

    # --- Detect merged cells -------------------------------------------------
    # Merged cells in python-docx share the same underlying XML element (_tc).
    # If the same element appears twice in a row's cell list, a merge is present.
    try:
        for row in table.rows:
            seen_cells = set()
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    has_merged = True
                    break
                seen_cells.add(cell_id)
            if has_merged:
                break
    except Exception:
        pass  # If the table XML is malformed, skip merge detection silently

    # --- Extract first row (assumed to be column headers) --------------------
    if table.rows:
        first_row = table.rows[0]
        header_row = [cell.text.strip()[:100] for cell in first_row.cells]

    # --- Classify by matching header keywords --------------------------------
    # The classification with the most keyword hits in the header row wins.
    # Ties are broken by whichever type appears first in the config.
    header_text = " ".join(header_row).lower()
    classifications = config['tables']['classification']

    best_match = "unclassified"
    best_score = 0

    for cls_name, cls_config in classifications.items():
        score = 0
        for kw in cls_config['keywords']:
            if kw.lower() in header_text:
                score += 1
        min_cols = cls_config.get('min_columns', 2)
        if score > best_score and len(header_row) >= min_cols:
            best_score = score
            best_match = cls_name

    return (best_match, header_row, has_merged, char_count)


def pattern_to_readable_label(pattern: str) -> str:
    """Convert a regex pattern to a human-readable label for output."""
    # Map common patterns to readable names
    pattern_map = {
        r'see section\s+[\d\.]+[a-z]?': 'See section reference',
        r'refer to\s+(section|the)': 'Refer to clause/section',
        r'as described in\s+(the|section)': 'As described in reference',
        r'per section\s+[\d\.]+': 'Per section reference',
        r'in accordance with\s+(the|section)': 'In accordance with',
        r'as defined in\s+(the|section)': 'As defined in reference',
        r'as outlined in\s+(the|section)': 'As outlined in reference',
        r'per the organization\'?s\s+\w+': 'Per organization document',
        r'see the\s+\w+\s+policy': 'See policy reference',
        r'as specified in\s+(the|section)': 'As specified in reference',
    }
    return pattern_map.get(pattern, f'Pattern: {pattern}')


def detect_cross_references(text: str, config: dict) -> list:
    """Find all cross-reference patterns in a text string. Returns list of (match_text, pattern_name)."""
    results = []
    for pattern in config['cross_references']['profiler_patterns']:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            readable_label = pattern_to_readable_label(pattern)
            results.append((match.group(), readable_label))
    return results


def detect_formatting_anomalies(doc_path: str, doc) -> dict:
    """
    Scan the document for formatting elements that will cause problems in Copilot.

    Returns a dict of True/False flags:
      has_text_boxes        → content in floating text boxes is invisible to Copilot
      has_tracked_changes   → unaccepted changes create duplicate/conflicting text
      has_comments          → comments are stripped on upload; resolve before processing
      has_images            → images are ignored by Copilot; any image-only content is lost
      has_embedded_objects  → OLE objects (Excel sheets, Visio diagrams) are not readable
      has_password_protection → document is locked; must be unlocked before python-docx
                                can fully read it

    WHAT TO DO WITH EACH FLAG:
      has_tracked_changes = YES →
        Open in Word → Review tab → Accept All Changes. Save. Then re-run profiler.
        Unaccepted tracked changes can cause the same text to appear twice in the
        profiler output (once as deleted, once as inserted).

      has_comments = YES →
        Usually safe to ignore — comments are not part of the policy text.
        But if comments contain policy decisions not yet in the body text, resolve
        them in Word first.

      has_text_boxes = YES →
        Any text inside a text box will NOT appear in Copilot search results.
        In Word: Find the text box → cut its content → paste as normal paragraph text.
        The profiler cannot tell you which page the text box is on; you must visually
        scan the document.

      has_images = YES →
        Images themselves are fine (they're just ignored). Problem only if the image
        IS the policy content (e.g., a scanned diagram with requirement text).
        Flag those docs for manual review.

      has_embedded_objects = YES →
        Any OLE object content (linked Excel table, etc.) is invisible to Copilot.
        Convert those to static tables or prose in Word before processing.

      has_password_protection = YES →
        Word's "Restrict Editing" or document protection is on.
        FIX: In Word → Review tab → Restrict Editing → Stop Protection (enter password).
        Save as a new file. Use the unprotected version for processing.
        Note: IRM/Sensitivity Label protection is different and cannot be removed here.

    HOW DETECTION WORKS:
      The function inspects the raw XML inside the .docx file (which is a zip archive
      of XML files). Specific XML tags signal each anomaly. This is more reliable than
      using Word's UI because it works even on locked/restricted documents.
    """
    anomalies = {
        'has_text_boxes': False,
        'has_tracked_changes': False,
        'has_comments': False,
        'has_images': False,
        'has_embedded_objects': False,
        'has_password_protection': False,
    }

    # --- Images: check the document's relationship list ----------------------
    # Every image embedded in a .docx is listed as a "relationship" in the
    # package. If any relationship has "image" in its type URL, images exist.
    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                anomalies['has_images'] = True
                break
    except Exception:
        pass

    # --- Password/document protection: check settings.xml --------------------
    # Word stores document protection in a separate XML part called "settings".
    # The <w:documentProtection> element signals that editing is restricted.
    # This is a BUG FIX: the original code never actually set this flag to True.
    try:
        settings_part = doc.settings.element
        settings_xml = settings_part.xml if hasattr(settings_part, 'xml') else ""
        if '<w:documentProtection' in settings_xml:
            anomalies['has_password_protection'] = True
    except Exception:
        # If we can't read settings, the document may itself be encrypted.
        # A completely encrypted .docx would have already failed at Document()
        # with PackageNotFoundError, so reaching here means partially restricted.
        pass

    # --- Tracked changes, comments, text boxes, OLE objects: scan body XML ---
    # These are all detected by looking for specific XML element tags in the
    # document body. Two approaches: lxml (faster) or fallback paragraph scan.
    try:
        from lxml import etree
        body = doc.element.body
        xml_str = etree.tostring(body, encoding='unicode')

        # <w:ins> = inserted text, <w:del> = deleted text → tracked changes
        if '<w:ins ' in xml_str or '<w:del ' in xml_str:
            anomalies['has_tracked_changes'] = True
        # <w:commentRangeStart> = a comment anchor
        if '<w:commentRangeStart' in xml_str or '<w:comment ' in xml_str:
            anomalies['has_comments'] = True
        # <w:txbxContent> = inline text box, <wps:txbx> = drawing canvas text box
        if '<w:txbxContent' in xml_str or '<wps:txbx' in xml_str:
            anomalies['has_text_boxes'] = True
        # <o:OLEObject> = linked/embedded OLE, <w:object> = generic embedded object
        if '<o:OLEObject' in xml_str or '<w:object' in xml_str:
            anomalies['has_embedded_objects'] = True
    except ImportError:
        # lxml is not installed — fall back to per-paragraph XML scan.
        # This is slower and won't catch text boxes or OLE objects, but covers
        # the most important flags (tracked changes and comments).
        # To get lxml: pip install lxml  (it often comes with python-docx anyway)
        try:
            for para in doc.paragraphs:
                xml = para._element.xml
                if '<w:ins ' in xml or '<w:del ' in xml:
                    anomalies['has_tracked_changes'] = True
                if '<w:commentRangeStart' in xml:
                    anomalies['has_comments'] = True
                if anomalies['has_tracked_changes'] and anomalies['has_comments']:
                    break
        except Exception:
            pass

    return anomalies


def classify_document_type(profile: dict, config: dict) -> tuple:
    """
    Auto-classify a document into one of five types based on its content ratios.

    Returns: (type_letter, reason_string)
      - type_letter: "A", "B", "C", "D", or "E"
      - reason_string: human-readable explanation shown in the Excel inventory

    THE FIVE TYPES (from Solo Operation Workflow):
      A = Table-Heavy Control Docs     (>40% content in tables)
      B = Prose-Heavy Intent Docs      (<10% content in tables)
      C = Hybrid Docs                  (between A and B, often with procedures)
      D = Appendix-Dominant Docs       (>60% content in appendix sections)
      E = Unclassified / manual review needed

    HOW CLASSIFICATION WORKS:
      1. Check appendix % first (D beats everything else if dominant enough).
      2. Check table % for A (table-heavy).
      3. Check table % for B (prose-heavy), but scan for procedure keywords
         that would upgrade it to C.
      4. Anything between A and B thresholds becomes C.
      5. Anything that doesn't fit → E (requires your manual review).

    FAILURE MODES:

      Document classified as E (unclassified) but you know what type it is →
        The thresholds in profiler_config.yaml may not match this document's
        actual content distribution. Check the "Type Reason" column in the
        Excel inventory — it shows the actual percentages.
        FIX: You can override the type manually in the Excel spreadsheet.
        Do NOT force-fit; per the workflow, Type E gets its own calibration run.

      Document classified as B but has a large procedure section →
        The procedure keyword scan only fires if 2+ keywords match. If your
        org uses different terminology (e.g., "playbook" instead of "procedure"),
        add those terms to profiler_config.yaml under:
          classification: type_c: procedure_keywords

      Document classified as D but most of the "appendix" is actually body content →
        A section heading matched the appendix fuzzy patterns (e.g., "References"
        or "Definitions" mid-document). Check the section_inventory.csv to see
        which section triggered the appendix match.
        FIX: Add a more specific heading pattern to profiler_config.yaml under
        sections: appendix to avoid over-matching.
    """
    cls = config['classification']
    table_pct = profile.get('table_content_pct', 0)
    appendix_pct = profile.get('appendix_section_pct', 0)

    # --- Type D check first: appendix percentage dominates -------------------
    # D is checked before A because a doc can have both heavy tables AND a
    # dominant appendix — in that case, D is the more actionable classification.
    if appendix_pct >= cls['type_d']['min_appendix_content_pct']:
        return ("D", f"Appendix-dominant: {appendix_pct:.1f}% appendix content (threshold: {cls['type_d']['min_appendix_content_pct']}%)")

    # --- Type A: table-heavy -------------------------------------------------
    if table_pct >= cls['type_a']['min_table_content_pct']:
        return ("A", f"Table-heavy: {table_pct:.1f}% table content (threshold: {cls['type_a']['min_table_content_pct']}%)")

    # --- Type B: prose-heavy, unless procedure keywords push it to C ---------
    if table_pct <= cls['type_b']['max_table_content_pct']:
        full_text = profile.get('_full_text', '').lower()
        procedure_hits = 0
        for kw in cls['type_c']['procedure_keywords']:
            if kw.lower() in full_text:
                procedure_hits += 1
        if procedure_hits >= 2:
            return ("C", f"Hybrid: low table content ({table_pct:.1f}%) but {procedure_hits} procedure keywords detected")
        return ("B", f"Prose-heavy: {table_pct:.1f}% table content (threshold: {cls['type_b']['max_table_content_pct']}%)")

    # --- Between A and B thresholds: Type C or E ----------------------------
    # (table_pct is between 10% and 40% — neither prose-heavy nor table-heavy)
    full_text = profile.get('_full_text', '').lower()
    procedure_hits = 0
    for kw in cls['type_c']['procedure_keywords']:
        if kw.lower() in full_text:
            procedure_hits += 1
    if procedure_hits >= 2:
        return ("C", f"Hybrid: {table_pct:.1f}% table content, {procedure_hits} procedure keywords detected")

    if table_pct > cls['type_b']['max_table_content_pct'] and table_pct < cls['type_a']['min_table_content_pct']:
        return ("C", f"Hybrid: {table_pct:.1f}% table content (between A and B thresholds)")

    # --- Type E: doesn't fit any category ------------------------------------
    # Per the workflow: do NOT force-fit into A-D. Run a separate calibration
    # cycle for Type E documents before processing them.
    return ("E", f"Unclassified: table_pct={table_pct:.1f}%, appendix_pct={appendix_pct:.1f}%. Review manually.")


def compute_priority_score(profile: dict, config: dict) -> float:
    """
    Compute a combined priority score. Higher score = more optimization work needed
    = should be processed earlier in your project.

    HOW THE SCORE WORKS:
      Each signal adds weighted points. Defaults (from profiler_config.yaml):
        table_count        × 2.0   (more tables = more flattening work)
        cross_ref_count    × 1.5   (each cross-ref needs inline restatement)
        fake_heading_count × 1.0   (each fake heading needs style correction)
        approx_pages       × 0.5   (longer docs are harder to process)
        missing_sections   × 1.5   (missing standard sections = structural rework)
        merged_cells       × 1.0   (merged cell tables need special handling)
        over_size_limit    + 3.0   (flat bonus: doc MUST be split regardless)

      usage_frequency (optional): if you add scores in profiler_config.yaml under
        priority_scoring: usage_frequency: { "PolicyName.docx": 9 }
        that score is multiplied by 2.0 and added. This is the strongest signal.

    TO ADJUST PRIORITY WEIGHTING:
      Edit the weights in profiler_config.yaml → priority_scoring → weights.
      You do NOT need to touch this code.

    TO ADD USAGE FREQUENCY SCORES (recommended):
      In profiler_config.yaml, under priority_scoring: usage_frequency:
        "Access Control Policy.docx": 9
        "Acceptable Use Policy.docx": 7
      Scale: 0 (nobody asks about it) to 10 (most-queried policy).
      These scores push high-value documents to the top of the processing queue
      even if they're structurally simple.
    """
    weights = config['priority_scoring']['weights']
    score = 0.0

    score += profile.get('table_count', 0) * weights.get('table_count', 1)
    score += profile.get('cross_ref_count', 0) * weights.get('cross_ref_count', 1)
    score += profile.get('fake_heading_count', 0) * weights.get('fake_heading_count', 1)
    score += profile.get('approx_pages', 0) * weights.get('page_count', 0.5)
    score += len(profile.get('missing_sections', [])) * weights.get('missing_sections', 1)
    score += profile.get('tables_with_merged_cells', 0) * weights.get('merged_cells', 1)
    if profile.get('over_size_limit', False):
        score += weights.get('over_size_limit', 3)

    # Usage frequency is a strong signal — multiply by 2.0 to push frequently
    # queried docs to the top even if they're structurally simple.
    freq = config['priority_scoring'].get('usage_frequency', {})
    filename = profile.get('filename', '')
    if filename in freq:
        score += freq[filename] * 2.0

    return round(score, 2)


# ============================================================================
# Main Profiler
# ============================================================================

def profile_document(filepath: str, config: dict) -> DocumentProfile:
    """
    Profile a single .docx file end-to-end and return a DocumentProfile.

    This is the main per-document function. It runs all detection passes
    and assembles the result. If the document cannot be opened, it returns
    a profile with empty data and an error message — it does NOT crash the
    script. The batch loop continues with the next file.

    FAILURE MODES:

      "Cannot open file — possibly corrupted or password-protected" →
        python-docx raised PackageNotFoundError. This happens when:
        (a) The file is password-protected (IRM/encryption, not just restrict-editing)
        (b) The file is corrupted or not actually a .docx (e.g., renamed .doc)
        (c) Word has a temp lock file open (~$filename.docx) — excluded by config
            but worth checking if you see unexpected failures.
        FIX: Try opening the file directly in Word. If Word asks for a password,
        the doc needs decryption before processing. If Word can't open it either,
        the file is corrupted — get a fresh copy from your document management system.

      Unexpected error during profiling →
        A second exception handler catches any other Python error during document
        processing. The error text is logged in the "Errors" column of the inventory.
        These are usually caused by malformed XML inside a .docx — Microsoft Word
        sometimes creates files that are technically invalid but Word itself tolerates.
        FIX: Check the "Errors" column in document_inventory.xlsx. If the error says
        something about XML or element, try File → Save As → .docx from within Word
        to re-serialize the file cleanly, then re-run the profiler on the saved copy.

      Section character counts don't add up to the total →
        Section boundaries are determined by Heading 1 paragraphs. Characters in
        tables are NOT included in section character counts (python-docx separates
        tables from paragraphs in its object model). The table_content_pct is
        calculated separately. This is expected behavior, not a bug.

      Approx page count looks too high or too low →
        Page count is estimated as paragraph_count / paragraphs_per_page (default 30).
        Dense policy docs with long paragraphs will look like fewer pages; docs with
        lots of short bullet items will look like more.
        FIX: Adjust paragraphs_per_page in profiler_config.yaml.
        To calibrate: open one of your docs in Word, note the real page count,
        count the paragraphs (Ctrl+End and check status bar), then set
        paragraphs_per_page = paragraph_count / actual_page_count.
    """
    filename = os.path.basename(filepath)
    errors = []

    # --- Attempt to open the document ----------------------------------------
    # PackageNotFoundError = file is encrypted or not a valid .docx.
    # Any other Exception = something unexpected went wrong.
    try:
        doc = Document(filepath)
    except PackageNotFoundError:
        return DocumentProfile(
            filename=filename, filepath=filepath, file_size_bytes=os.path.getsize(filepath),
            total_char_count=0, total_word_count=0, total_paragraphs=0, approx_pages=0, over_size_limit=False,
            heading_count=0, h1_count=0, h2_count=0, h3_count=0, h4_count=0,
            builtin_heading_count=0, custom_heading_count=0, fake_heading_count=0,
            heading_styles_used=[], sections=[], has_purpose=False, has_scope=False,
            has_intent=False, has_controls=False, has_appendix=False, missing_sections=[],
            section_count=0, largest_section_name="", largest_section_chars=0,
            largest_section_pct=0, controls_section_chars=0, controls_section_pct=0,
            appendix_section_chars=0, appendix_section_pct=0, table_count=0, tables=[],
            total_table_chars=0, table_content_pct=0, tables_with_merged_cells=0,
            nested_table_count=0, table_types={}, cross_ref_count=0, cross_refs=[],
            cross_ref_patterns={}, has_text_boxes=False, has_tracked_changes=False,
            has_comments=False, has_images=False, has_embedded_objects=False,
            has_password_protection=False, doc_type="E", doc_type_reason="Failed to open",
            priority_score=0, priority_rank=0, search_term_hits={},
            errors=[f"Cannot open file: {filepath}. Possibly corrupted or password-protected."]
        )
    except Exception as e:
        return DocumentProfile(
            filename=filename, filepath=filepath, file_size_bytes=os.path.getsize(filepath),
            total_char_count=0, total_word_count=0, total_paragraphs=0, approx_pages=0, over_size_limit=False,
            heading_count=0, h1_count=0, h2_count=0, h3_count=0, h4_count=0,
            builtin_heading_count=0, custom_heading_count=0, fake_heading_count=0,
            heading_styles_used=[], sections=[], has_purpose=False, has_scope=False,
            has_intent=False, has_controls=False, has_appendix=False, missing_sections=[],
            section_count=0, largest_section_name="", largest_section_chars=0,
            largest_section_pct=0, controls_section_chars=0, controls_section_pct=0,
            appendix_section_chars=0, appendix_section_pct=0, table_count=0, tables=[],
            total_table_chars=0, table_content_pct=0, tables_with_merged_cells=0,
            nested_table_count=0, table_types={}, cross_ref_count=0, cross_refs=[],
            cross_ref_patterns={}, has_text_boxes=False, has_tracked_changes=False,
            has_comments=False, has_images=False, has_embedded_objects=False,
            has_password_protection=False, doc_type="E", doc_type_reason=f"Error: {str(e)}",
            priority_score=0, priority_rank=0, search_term_hits={},
            errors=[str(e)]
        )

    # ----- Basic stats -----
    file_size = os.path.getsize(filepath)
    paragraphs = doc.paragraphs
    total_chars = sum(len(p.text) for p in paragraphs)
    total_words = sum(len(p.text.split()) for p in paragraphs if p.text.strip())
    total_paras = len(paragraphs)
    ppp = config['thresholds']['paragraphs_per_page']
    approx_pages = round(total_paras / ppp, 1) if ppp > 0 else 0
    over_size = total_chars > config['thresholds']['max_characters']

    # ----- Full text for classification -----
    full_text = "\n".join(p.text for p in paragraphs)

    # ----- Heading analysis -----
    heading_data = []  # (para_index, level, style, is_builtin, is_custom, is_fake, text)
    styles_used = set()

    for i, para in enumerate(paragraphs):
        level, style, is_builtin, is_custom, is_fake = detect_heading_info(para, config)
        if level > 0 or is_fake:
            heading_data.append((i, level if level > 0 else 1, style, is_builtin, is_custom, is_fake, para.text.strip()))
            styles_used.add(style)

    h1s = [h for h in heading_data if h[1] == 1 and not h[5]]
    h2s = [h for h in heading_data if h[1] == 2 and not h[5]]
    h3s = [h for h in heading_data if h[1] == 3 and not h[5]]
    h4s = [h for h in heading_data if h[1] == 4 and not h[5]]
    builtins = [h for h in heading_data if h[3]]
    customs = [h for h in heading_data if h[4]]
    fakes = [h for h in heading_data if h[5]]

    # ----- Section structure (split at H1 boundaries) -----
    h1_indices = [(h[0], h[6]) for h in heading_data if h[1] == 1]
    sections_info = []

    if not h1_indices:
        # No H1 headings found: treat entire doc as one section
        sec_chars = total_chars
        sec_tables = len(doc.tables)
        sec_crossrefs = len(detect_cross_references(full_text, config))
        sections_info.append(SectionInfo(
            heading_text="(No H1 headings)",
            heading_level=0,
            heading_style="none",
            is_builtin_style=False,
            is_fake_heading=False,
            char_count=sec_chars,
            table_count=sec_tables,
            paragraph_count=total_paras,
            start_para_index=0,
            standard_section="none",
            cross_ref_count=sec_crossrefs,
        ))
    else:
        for idx, (para_idx, heading_text) in enumerate(h1_indices):
            # Determine end of this section
            if idx + 1 < len(h1_indices):
                end_idx = h1_indices[idx + 1][0]
            else:
                end_idx = total_paras

            sec_paras = paragraphs[para_idx:end_idx]
            sec_text = "\n".join(p.text for p in sec_paras)
            sec_chars = len(sec_text)
            sec_crossrefs = detect_cross_references(sec_text, config)

            # Count tables in this section range (approximate)
            sec_table_count = 0  # Hard to map tables to paragraph ranges in python-docx

            # Determine heading metadata
            hd = next((h for h in heading_data if h[0] == para_idx), None)
            std_section = match_standard_section(heading_text, config)

            sections_info.append(SectionInfo(
                heading_text=heading_text,
                heading_level=1,
                heading_style=hd[2] if hd else "unknown",
                is_builtin_style=hd[3] if hd else False,
                is_fake_heading=hd[5] if hd else False,
                char_count=sec_chars,
                table_count=sec_table_count,
                paragraph_count=len(sec_paras),
                start_para_index=para_idx,
                standard_section=std_section,
                cross_ref_count=len(sec_crossrefs),
            ))

    # Standard section detection
    found_sections = {si.standard_section for si in sections_info}
    has_purpose = "purpose" in found_sections
    has_scope = "scope" in found_sections
    has_intent = "intent" in found_sections
    has_controls = "controls" in found_sections
    has_appendix = "appendix" in found_sections
    missing = []
    for key in ["purpose", "scope", "intent", "controls", "appendix"]:
        if key not in found_sections:
            missing.append(key)

    # Section size distribution
    largest_sec = max(sections_info, key=lambda s: s.char_count) if sections_info else None
    controls_chars = sum(s.char_count for s in sections_info if s.standard_section == "controls")
    appendix_chars = sum(s.char_count for s in sections_info if s.standard_section == "appendix")

    # ----- Table analysis -----
    tables_info = []
    total_table_chars = 0
    total_table_words = 0
    merged_count = 0
    type_counts = {}

    for tidx, table in enumerate(doc.tables):
        cls_name, header_row, has_merged, tbl_chars = classify_table(table, config)
        total_table_chars += tbl_chars
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    total_table_words += len(cell_text.split())
        if has_merged:
            merged_count += 1
        type_counts[cls_name] = type_counts.get(cls_name, 0) + 1

        tables_info.append(TableInfo(
            index=tidx,
            rows=len(table.rows),
            cols=len(table.columns),
            header_row=header_row,
            has_merged_cells=has_merged,
            is_nested=False,  # python-docx doesn't easily expose nesting
            classification=cls_name,
            parent_section="",  # would require XML-level mapping
            char_count=tbl_chars,
        ))

    table_pct = round((total_table_chars / total_chars * 100), 1) if total_chars > 0 else 0

    # ----- Key term search -----
    table_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    search_term_hits = search_key_terms(full_text, table_text, config)

    # ----- Cross-references -----
    all_crossrefs = []
    crossref_pattern_counts = {}
    for i, para in enumerate(paragraphs):
        refs = detect_cross_references(para.text, config)
        for match_text, pattern in refs:
            # Find parent section
            parent = "(unknown)"
            for si in reversed(sections_info):
                if si.start_para_index <= i:
                    parent = si.heading_text
                    break
            all_crossrefs.append(CrossRefInfo(
                text=match_text,
                pattern_matched=pattern,
                paragraph_index=i,
                parent_section=parent,
                is_hyperlink=False,
            ))
            crossref_pattern_counts[pattern] = crossref_pattern_counts.get(pattern, 0) + 1

    # ----- Formatting anomalies -----
    anomalies = detect_formatting_anomalies(filepath, doc)

    # ----- Profiling flags -----
    flags_cfg = config.get('profiling_flags', {})

    # 1. ControlDense: count control IDs in full text using Step 1 patterns
    ctrl_patterns = config.get('control_extraction', {}).get('control_id_patterns', [])
    control_id_count = 0
    for pat in ctrl_patterns:
        try:
            control_id_count += len(re.findall(pat, full_text))
        except re.error:
            pass
    control_density = round(control_id_count / approx_pages, 1) if approx_pages > 0 else 0
    ctrl_cfg = flags_cfg.get('control_dense', {})
    control_dense = control_density >= ctrl_cfg.get('min_controls_per_page', 5.0)

    # 2. HeadingVariance: detect structural issues in heading hierarchy
    hv_cfg = flags_cfg.get('heading_variance', {})
    level_skips = 0
    sorted_headings = sorted(heading_data, key=lambda h: h[0])  # sort by paragraph index
    for i in range(1, len(sorted_headings)):
        prev_level = sorted_headings[i - 1][1]
        curr_level = sorted_headings[i][1]
        if curr_level > prev_level + 1:
            level_skips += 1
    total_heading_count = len(heading_data)
    fake_ratio = len(fakes) / total_heading_count if total_heading_count > 0 else 0
    h1_only_flat = len(h1s) > 0 and len(h2s) == 0

    hv_reasons = []
    if level_skips >= hv_cfg.get('max_level_skips', 2):
        hv_reasons.append(f"{level_skips} level skips")
    if fake_ratio >= hv_cfg.get('max_fake_ratio', 0.5) and total_heading_count > 0:
        hv_reasons.append(f"{fake_ratio:.0%} fake headings")
    if h1_only_flat:
        hv_reasons.append("flat H1-only structure")
    heading_variance = len(hv_reasons) > 0
    heading_variance_reason = "; ".join(hv_reasons) if hv_reasons else ""

    # 3. UniqueSections: H1 headings not matching any standard section
    unique_secs = [s.heading_text for s in sections_info if s.standard_section == "none"]
    unique_section_count = len(unique_secs)

    # 4. TableDense: lower-threshold table content warning
    td_cfg = flags_cfg.get('table_dense', {})
    table_dense = table_pct >= td_cfg.get('min_table_content_pct', 30)

    # ----- Build intermediate dict for classification -----
    profile_dict = {
        'table_content_pct': table_pct,
        'appendix_section_pct': round(appendix_chars / total_chars * 100, 1) if total_chars > 0 else 0,
        '_full_text': full_text,
        'table_count': len(doc.tables),
        'cross_ref_count': len(all_crossrefs),
        'fake_heading_count': len(fakes),
        'approx_pages': approx_pages,
        'missing_sections': missing,
        'tables_with_merged_cells': merged_count,
        'over_size_limit': over_size,
        'filename': filename,
    }

    doc_type, doc_type_reason = classify_document_type(profile_dict, config)
    priority = compute_priority_score(profile_dict, config)

    # ----- Assemble final profile -----
    return DocumentProfile(
        filename=filename,
        filepath=filepath,
        file_size_bytes=file_size,
        total_char_count=total_chars,
        total_word_count=total_words + total_table_words,
        total_paragraphs=total_paras,
        approx_pages=approx_pages,
        over_size_limit=over_size,
        heading_count=len(heading_data),
        h1_count=len(h1s),
        h2_count=len(h2s),
        h3_count=len(h3s),
        h4_count=len(h4s),
        builtin_heading_count=len(builtins),
        custom_heading_count=len(customs),
        fake_heading_count=len(fakes),
        heading_styles_used=list(styles_used),
        sections=[vars(s) for s in sections_info],
        has_purpose=has_purpose,
        has_scope=has_scope,
        has_intent=has_intent,
        has_controls=has_controls,
        has_appendix=has_appendix,
        missing_sections=missing,
        section_count=len(sections_info),
        largest_section_name=largest_sec.heading_text if largest_sec else "",
        largest_section_chars=largest_sec.char_count if largest_sec else 0,
        largest_section_pct=round(largest_sec.char_count / total_chars * 100, 1) if largest_sec and total_chars > 0 else 0,
        controls_section_chars=controls_chars,
        controls_section_pct=round(controls_chars / total_chars * 100, 1) if total_chars > 0 else 0,
        appendix_section_chars=appendix_chars,
        appendix_section_pct=round(appendix_chars / total_chars * 100, 1) if total_chars > 0 else 0,
        table_count=len(doc.tables),
        tables=[vars(t) for t in tables_info],
        total_table_chars=total_table_chars,
        table_content_pct=table_pct,
        tables_with_merged_cells=merged_count,
        nested_table_count=0,
        table_types=type_counts,
        cross_ref_count=len(all_crossrefs),
        cross_refs=[vars(c) for c in all_crossrefs],
        cross_ref_patterns=crossref_pattern_counts,
        has_text_boxes=anomalies['has_text_boxes'],
        has_tracked_changes=anomalies['has_tracked_changes'],
        has_comments=anomalies['has_comments'],
        has_images=anomalies['has_images'],
        has_embedded_objects=anomalies['has_embedded_objects'],
        has_password_protection=anomalies['has_password_protection'],
        doc_type=doc_type,
        doc_type_reason=doc_type_reason,
        control_id_count=control_id_count,
        control_density=control_density,
        control_dense=control_dense,
        level_skips=level_skips,
        heading_variance=heading_variance,
        heading_variance_reason=heading_variance_reason,
        unique_sections=unique_secs,
        unique_section_count=unique_section_count,
        table_dense=table_dense,
        priority_score=priority,
        priority_rank=0,
        search_term_hits=search_term_hits,
        errors=errors,
    )


# ============================================================================
# Output Writers
# ============================================================================

def _safe_write(write_func, output_path: str, *args, **kwargs) -> bool:
    """
    Call a write function and catch the most common failure modes.

    Returns True if the write succeeded, False if it failed.

    FAILURE MODES this handles:

      PermissionError ("file is open in another program") →
        The output file is already open in Excel, Word, or another app.
        FIX: Close the file in Excel/Word and run the script again.
        You do NOT need to re-process all documents — the script will
        overwrite the output files. All document processing results are
        still in memory and will be written correctly once the file is closed.

      OSError / disk full / path too long →
        The output directory may have a path that Windows considers too long
        (>260 characters), or the disk is out of space.
        FIX (path too long): Move the script and docs closer to the root,
        e.g., C:/DPS/ instead of deep nested folders.
        FIX (disk full): Free up disk space and run again.

      Any other error → prints the traceback so you can investigate.
    """
    try:
        write_func(output_path, *args, **kwargs)
        return True
    except PermissionError:
        print(f"\n  WARNING: Cannot write '{os.path.basename(output_path)}'")
        print(f"  The file is open in another program (probably Excel or Word).")
        print(f"  FIX: Close the file and run the script again.")
        print(f"  Full path: {output_path}")
        return False
    except OSError as e:
        print(f"\n  WARNING: Cannot write '{os.path.basename(output_path)}': {e}")
        print(f"  Possible causes: path too long, disk full, or no write permission.")
        return False
    except Exception as e:
        print(f"\n  ERROR writing '{os.path.basename(output_path)}': {e}")
        traceback.print_exc()
        return False


def write_inventory_xlsx(profiles: list, output_path: str, config: dict):
    """
    Write the master inventory spreadsheet (document_inventory.xlsx).

    This is the primary output of the profiler. It contains:
      - "Document Inventory" tab: one row per document, sorted by priority score,
        color-coded by document type. Autofilter enabled on all columns.
      - "Summary" tab: aggregate counts and averages across all documents.

    WHAT TO DO WITH THE SPREADSHEET:
      1. Open it in Excel (or upload to SharePoint and open in Excel Online).
      2. Start with the Summary tab for an overview.
      3. On the Document Inventory tab:
         - Column C (Type): your A/B/C/D/E classification
         - Column D (Priority Score): higher = more work needed = process first
         - Column G (Over 36k Limit): YES = must split before uploading to SharePoint
         - Column I (Merged Cell Tables): any count > 0 = complex table handling needed
         - Column T (Missing Sections): sections that don't exist in the doc
         - Column U (Tracked Changes): YES = accept/reject in Word before processing
      4. Add a "Usage Frequency" column manually (0-10 per doc, based on how often
         people ask about this policy). Use this to re-sort if needed.
      5. Add a "Processing Status" column as you work through each document.

    MOST COMMON FAILURE:
      If the script crashes while saving and you see "PermissionError" →
      document_inventory.xlsx is open in Excel. Close it and run again.
      (This is handled by _safe_write in the calling code, which gives a
      clear message instead of crashing.)
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Document Inventory"

    # Header style
    header_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style='thin', color='CCCCCC'),
        right=Side(style='thin', color='CCCCCC'),
        top=Side(style='thin', color='CCCCCC'),
        bottom=Side(style='thin', color='CCCCCC'),
    )

    # Columns
    columns = [
        ("Rank", 6),
        ("Filename", 35),
        ("Type", 6),
        ("Priority Score", 12),
        ("Pages (approx)", 10),
        ("Characters", 12),
        ("Words", 10),
        ("Over 36k Limit", 10),
        ("Tables", 7),
        ("Merged Cell Tables", 10),
        ("Table Content %", 10),
        ("Cross-Refs", 8),
        ("Fake Headings", 8),
        ("H1 Count", 7),
        ("H2 Count", 7),
        ("Has Purpose", 8),
        ("Has Scope", 8),
        ("Has Intent", 8),
        ("Has Controls", 8),
        ("Has Appendix", 8),
        ("Missing Sections", 18),
        ("Largest Section", 25),
        ("Largest Sec %", 10),
        ("Controls Sec %", 10),
        ("Appendix Sec %", 10),
        ("Tracked Changes", 8),
        ("Comments", 8),
        ("Images", 8),
        ("Text Boxes", 8),
        ("Control IDs", 8),
        ("Ctrls/Page", 8),
        ("Control Dense", 8),
        ("Heading Variance", 8),
        ("Level Skips", 7),
        ("Unique Sections", 30),
        ("Table Dense", 8),
        ("Type Reason", 45),
        ("Errors", 30),
    ]

    # Dynamic search term columns
    search_cfg = config.get("search_terms", {})
    search_terms = search_cfg.get("terms", []) if search_cfg.get("enabled", False) else []
    show_counts = search_cfg.get("show_counts", True)
    for term in search_terms:
        columns.append((f"Term: {term}", 10))

    # Write headers
    for col_idx, (col_name, width) in enumerate(columns, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    # Write data
    bool_to_str = lambda b: "YES" if b else ""
    fail_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    warn_fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
    type_fills = {
        "A": PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid"),
        "B": PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid"),
        "C": PatternFill(start_color="FCE4D6", end_color="FCE4D6", fill_type="solid"),
        "D": PatternFill(start_color="E4DFEC", end_color="E4DFEC", fill_type="solid"),
        "E": PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid"),
    }

    for row_idx, p in enumerate(profiles, 2):
        row_data = [
            p.priority_rank,
            p.filename,
            p.doc_type,
            p.priority_score,
            p.approx_pages,
            p.total_char_count,
            p.total_word_count,
            bool_to_str(p.over_size_limit),
            p.table_count,
            p.tables_with_merged_cells,
            p.table_content_pct,
            p.cross_ref_count,
            p.fake_heading_count,
            p.h1_count,
            p.h2_count,
            bool_to_str(p.has_purpose),
            bool_to_str(p.has_scope),
            bool_to_str(p.has_intent),
            bool_to_str(p.has_controls),
            bool_to_str(p.has_appendix),
            ", ".join(p.missing_sections) if p.missing_sections else "",
            p.largest_section_name[:40],
            p.largest_section_pct,
            p.controls_section_pct,
            p.appendix_section_pct,
            bool_to_str(p.has_tracked_changes),
            bool_to_str(p.has_comments),
            bool_to_str(p.has_images),
            bool_to_str(p.has_text_boxes),
            p.control_id_count,
            p.control_density,
            bool_to_str(p.control_dense),
            bool_to_str(p.heading_variance),
            p.level_skips,
            ", ".join(p.unique_sections)[:60] if p.unique_sections else "",
            bool_to_str(p.table_dense),
            p.doc_type_reason[:80],
            "; ".join(p.errors)[:60] if p.errors else "",
        ]
        # Append search term results
        for term in search_terms:
            count = p.search_term_hits.get(term, 0)
            if show_counts:
                row_data.append(count if count > 0 else "")
            else:
                row_data.append("YES" if count > 0 else "")

        for col_idx, val in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.font = Font(name="Arial", size=9)
            cell.border = thin_border
            cell.alignment = Alignment(vertical="center")

        # Conditional formatting
        type_cell = ws.cell(row=row_idx, column=3)
        type_cell.fill = type_fills.get(p.doc_type, type_fills["E"])

        if p.over_size_limit:
            ws.cell(row=row_idx, column=7).fill = fail_fill
        if p.fake_heading_count > 5:
            ws.cell(row=row_idx, column=12).fill = warn_fill
        if p.tables_with_merged_cells > 0:
            ws.cell(row=row_idx, column=9).fill = warn_fill
        if p.missing_sections:
            ws.cell(row=row_idx, column=20).fill = warn_fill

        # Profiling flag highlights (yellow for flagged)
        if p.control_dense:
            ws.cell(row=row_idx, column=33).fill = warn_fill
        if p.heading_variance:
            ws.cell(row=row_idx, column=34).fill = warn_fill
        if p.table_dense:
            ws.cell(row=row_idx, column=37).fill = warn_fill

        # Highlight search term hits with green fill
        hit_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
        static_col_count = 38  # number of static columns before search terms
        for t_idx, term in enumerate(search_terms):
            count = p.search_term_hits.get(term, 0)
            if count > 0:
                ws.cell(row=row_idx, column=static_col_count + t_idx + 1).fill = hit_fill

    # Summary sheet
    ws2 = wb.create_sheet("Summary")
    ws2.column_dimensions['A'].width = 30
    ws2.column_dimensions['B'].width = 15

    summary_data = [
        ("Total Documents", len(profiles)),
        ("", ""),
        ("Type A (Table-Heavy)", len([p for p in profiles if p.doc_type == "A"])),
        ("Type B (Prose-Heavy)", len([p for p in profiles if p.doc_type == "B"])),
        ("Type C (Hybrid)", len([p for p in profiles if p.doc_type == "C"])),
        ("Type D (Appendix-Dominant)", len([p for p in profiles if p.doc_type == "D"])),
        ("Type E (Unclassified)", len([p for p in profiles if p.doc_type == "E"])),
        ("", ""),
        ("Over 36k Character Limit", len([p for p in profiles if p.over_size_limit])),
        ("Has Fake Headings", len([p for p in profiles if p.fake_heading_count > 0])),
        ("Has Merged Cell Tables", len([p for p in profiles if p.tables_with_merged_cells > 0])),
        ("Has Cross-References", len([p for p in profiles if p.cross_ref_count > 0])),
        ("Has Tracked Changes", len([p for p in profiles if p.has_tracked_changes])),
        ("Has Comments", len([p for p in profiles if p.has_comments])),
        ("Missing 1+ Standard Sections", len([p for p in profiles if p.missing_sections])),
        ("", ""),
        ("Control Dense Docs", len([p for p in profiles if p.control_dense])),
        ("Heading Variance Docs", len([p for p in profiles if p.heading_variance])),
        ("Table Dense Docs", len([p for p in profiles if p.table_dense])),
        ("Docs with Unique Sections", len([p for p in profiles if p.unique_section_count > 0])),
        ("Avg Controls/Page", round(sum(p.control_density for p in profiles) / len(profiles), 1) if profiles else 0),
        ("", ""),
        ("Avg Characters", round(sum(p.total_char_count for p in profiles) / len(profiles)) if profiles else 0),
        ("Avg Words", round(sum(p.total_word_count for p in profiles) / len(profiles)) if profiles else 0),
        ("Total Words Across All Docs", sum(p.total_word_count for p in profiles)),
        ("Avg Pages (approx)", round(sum(p.approx_pages for p in profiles) / len(profiles), 1) if profiles else 0),
        ("Avg Tables", round(sum(p.table_count for p in profiles) / len(profiles), 1) if profiles else 0),
        ("Avg Cross-Refs", round(sum(p.cross_ref_count for p in profiles) / len(profiles), 1) if profiles else 0),
        ("Total Tables Across All Docs", sum(p.table_count for p in profiles)),
        ("Total Cross-Refs Across All Docs", sum(p.cross_ref_count for p in profiles)),
    ]

    for row_idx, (label, value) in enumerate(summary_data, 1):
        ws2.cell(row=row_idx, column=1, value=label).font = Font(name="Arial", bold=True, size=10)
        ws2.cell(row=row_idx, column=2, value=value).font = Font(name="Arial", size=10)

    # Freeze pane (row 1 = headers stay visible when scrolling)
    ws.freeze_panes = "C2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(columns))}{len(profiles)+1}"

    # Save is called here — this is the step most likely to fail with PermissionError
    # if the file is open in Excel. The _safe_write wrapper in main() handles that.
    wb.save(output_path)


def write_sections_csv(profiles: list, output_path: str):
    """
    Write section_inventory.csv: one row per section per document.

    WHAT IT CONTAINS:
      Each Heading 1 section in each document gets its own row, with:
        - heading text, style, fake-heading flag
        - character count and paragraph count for that section
        - cross-reference count within that section
        - which standard section it matches (purpose/scope/intent/controls/appendix)

    USE THIS FILE TO:
      - Identify which sections are the largest (sort by char_count)
      - Find documents where the Controls section dominates (controls char_count
        > 50% of document total)
      - See which documents have sections the profiler didn't recognize as
        standard (standard_section = "none")
    """
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow([
            "filename", "section_heading", "heading_level", "heading_style",
            "is_builtin_style", "is_fake_heading", "standard_section",
            "char_count", "paragraph_count", "cross_ref_count"
        ])
        for p in profiles:
            for s in p.sections:
                writer.writerow([
                    p.filename, s['heading_text'], s['heading_level'], s['heading_style'],
                    s['is_builtin_style'], s['is_fake_heading'], s['standard_section'],
                    s['char_count'], s['paragraph_count'], s['cross_ref_count']
                ])


def write_tables_csv(profiles: list, output_path: str):
    """
    Write table_inventory.csv: one row per table per document.

    WHAT IT CONTAINS:
      Each table in each document gets its own row, with:
        - table index (0-based position in the document)
        - row and column count
        - classification (control_matrix, reference_table, etc.)
        - whether it has merged cells
        - character count
        - first row cell text (the header row as seen by the profiler)

    USE THIS FILE TO:
      - Find all control_matrix tables across your document set
      - Identify which tables have merged cells (sort by has_merged_cells)
      - Verify the profiler's classification of each table against what you
        see in the actual document — if it's wrong, add keywords to the config
    """
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow([
            "filename", "table_index", "rows", "cols", "classification",
            "has_merged_cells", "char_count", "header_row"
        ])
        for p in profiles:
            for t in p.tables:
                writer.writerow([
                    p.filename, t['index'], t['rows'], t['cols'],
                    t['classification'], t['has_merged_cells'], t['char_count'],
                    " | ".join(t['header_row'])
                ])


def write_crossrefs_csv(profiles: list, output_path: str):
    """
    Write crossref_inventory.csv: one row per cross-reference instance found.

    WHAT IT CONTAINS:
      Each detected "See Section X", "refer to the...", etc. gets its own row:
        - the exact text that was matched
        - which regex pattern caught it
        - paragraph index and parent section heading

    USE THIS FILE TO:
      - Build your cross-reference resolution queue for Notebook 2
      - Identify which documents have the most cross-references (sort by filename
        and count rows)
      - Find all references to a specific policy (filter by cross_ref_text)

    NOTE: This only catches text-pattern cross-references. Hyperlinked
    cross-references that use section titles as visible link text are noted
    in the config (detect_hyperlink_crossrefs) but require XML-level parsing
    that is not implemented in this version. Visually scan each doc for
    blue underlined links that point to other sections.
    """
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow([
            "filename", "cross_ref_text", "pattern_matched",
            "paragraph_index", "parent_section"
        ])
        for p in profiles:
            for cr in p.cross_refs:
                writer.writerow([
                    p.filename, cr['text'], cr['pattern_matched'],
                    cr['paragraph_index'], cr['parent_section']
                ])


def write_json(profiles: list, output_path: str):
    """Write full profiles as JSON for downstream automation."""
    data = []
    for p in profiles:
        d = p.to_dict()
        # Remove large internal fields
        d.pop('_full_text', None)
        data.append(d)
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, default=str)


# ============================================================================
# CLI
# ============================================================================

def main():
    # -------------------------------------------------------------------------
    # COMMAND-LINE ARGUMENTS
    # You can run this script three ways:
    #
    #   Simplest (uses config file defaults):
    #     python policy_profiler.py
    #
    #   Override the input folder:
    #     python policy_profiler.py --input "C:/MyDocs/PolicyDocs"
    #
    #   Override both input and output folders:
    #     python policy_profiler.py --input "C:/Docs" --output "C:/Output"
    #
    #   Use a different config file:
    #     python policy_profiler.py --config "C:/path/my_config.yaml"
    #
    #   Get help:
    #     python policy_profiler.py --help
    # -------------------------------------------------------------------------
    parser = argparse.ArgumentParser(
        description=(
            "Policy Document Profiler — scans .docx files and produces an\n"
            "Excel inventory ranked by Copilot optimization priority.\n\n"
            "Run with no arguments to use profiler_config.yaml defaults.\n"
            "Run with --help to see all options."
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    # Find the best default config: dps_config.yaml (unified) or profiler_config.yaml (legacy)
    default_config = 'profiler_config.yaml'
    for candidate in ['dps_config.yaml', '../dps_config.yaml']:
        if os.path.isfile(candidate):
            default_config = candidate
            break

    parser.add_argument(
        '--config', '-c',
        default=default_config,
        help='Path to YAML config file (default: dps_config.yaml or profiler_config.yaml)'
    )
    parser.add_argument(
        '--input', '-i',
        default=None,
        help='Input folder containing .docx files (overrides config input.directory)'
    )
    parser.add_argument(
        '--output', '-o',
        default=None,
        help='Output folder for results (overrides config output.directory)'
    )
    args = parser.parse_args()

    # -------------------------------------------------------------------------
    # STARTUP BANNER
    # -------------------------------------------------------------------------
    print()
    print("=" * 60)
    print("  Policy Document Profiler")
    print("  Copilot KB Optimization — Document Inventory Generator")
    print("=" * 60)

    # -------------------------------------------------------------------------
    # LOAD CONFIG
    # load_config() handles all error cases with clear messages.
    # -------------------------------------------------------------------------
    config = load_config(args.config)
    print(f"  Config loaded: {os.path.abspath(args.config)}")

    # -------------------------------------------------------------------------
    # RESOLVE PATHS
    # Convert relative paths (like ./policy_docs) to absolute paths so you
    # can see exactly where the script is looking, regardless of which folder
    # your terminal is in when you run the script.
    # -------------------------------------------------------------------------
    # Resolve paths — handle both unified dps_config.yaml and legacy profiler_config.yaml
    input_dir = os.path.abspath(args.input or config['input']['directory'])

    if args.output:
        output_dir = os.path.abspath(args.output)
    elif 'profiler' in config.get('output', {}):
        # Unified config: output.directory + output.profiler.directory
        output_root = config['output'].get('directory', './output')
        profiler_subdir = config['output']['profiler'].get('directory', '0 - profiler')
        output_dir = os.path.abspath(os.path.join(output_root, profiler_subdir))
    else:
        # Legacy config: output.directory directly
        output_dir = os.path.abspath(config['output']['directory'])

    print(f"  Input folder:  {input_dir}")
    print(f"  Output folder: {output_dir}")
    print()

    # -------------------------------------------------------------------------
    # VALIDATE INPUT DIRECTORY
    # -------------------------------------------------------------------------
    if not os.path.isdir(input_dir):
        print("=" * 60)
        print("ERROR: Input folder does not exist.")
        print(f"  Looking for: {input_dir}")
        print()
        print("This usually means one of three things:")
        print("  (a) You are running the script from the wrong folder.")
        print("      FIX: cd to the folder that contains your .docx files,")
        print("           or pass the full path with --input:")
        print(f'           python policy_profiler.py --input "C:/full/path/to/docs"')
        print()
        print("  (b) The path in profiler_config.yaml is wrong.")
        print("      FIX: Open profiler_config.yaml and update input: directory")
        print("           to the full path where your .docx files live.")
        print()
        print("  (c) The folder hasn't been created yet.")
        print("      FIX: Create the folder and put your .docx files in it.")
        print("=" * 60)
        sys.exit(1)

    # Create output directory if it doesn't exist
    try:
        os.makedirs(output_dir, exist_ok=True)
    except OSError as e:
        print(f"ERROR: Cannot create output folder: {output_dir}")
        print(f"  Detail: {e}")
        print("  FIX: Check that you have write permission to that location,")
        print("       or change output: directory in profiler_config.yaml.")
        sys.exit(1)

    # -------------------------------------------------------------------------
    # FIND .DOCX FILES
    # -------------------------------------------------------------------------
    pattern   = config['input']['pattern']
    recursive = config['input']['recursive']
    exclude   = [e.lower() for e in config['input'].get('exclude_patterns', [])]

    if recursive:
        glob_pattern = os.path.join(input_dir, '**', pattern)
        files = glob.glob(glob_pattern, recursive=True)
    else:
        glob_pattern = os.path.join(input_dir, pattern)
        files = glob.glob(glob_pattern)

    # Remove excluded files (Word temp files, already-processed files, etc.)
    excluded_files = [f for f in files if any(ex in os.path.basename(f).lower() for ex in exclude)]
    files = [f for f in files if f not in excluded_files]

    if not files:
        print("=" * 60)
        print("ERROR: No .docx files found.")
        print(f"  Searched in: {input_dir}")
        print(f"  Pattern:     {pattern}")
        if excluded_files:
            print(f"  Note: {len(excluded_files)} file(s) were found but excluded by")
            print(f"        the exclude_patterns setting in config.")
        print()
        print("POSSIBLE FIXES:")
        print("  - Make sure your .docx files are in the input folder shown above")
        print("  - Check that files have the .docx extension (not .doc or .pdf)")
        print("  - If recursive: true in config, subfolders are also scanned")
        print("=" * 60)
        sys.exit(1)

    if excluded_files:
        print(f"  Skipping {len(excluded_files)} excluded file(s) (temp files, already-processed, etc.)")
    print(f"  Found {len(files)} .docx file(s) to profile.")
    print("-" * 60)

    # -------------------------------------------------------------------------
    # PROFILE EACH DOCUMENT
    # Each file is processed independently. If one fails, the script keeps
    # going and logs the error — it does NOT stop on a single bad file.
    # -------------------------------------------------------------------------
    profiles = []
    total_start = time.time()

    for i, filepath in enumerate(sorted(files)):
        filename = os.path.basename(filepath)
        print(f"  [{i+1:3d}/{len(files)}] {filename}", end="  ")
        sys.stdout.flush()  # Force print to appear before the processing starts

        doc_start = time.time()
        try:
            profile = profile_document(filepath, config)
            profiles.append(profile)
            elapsed = time.time() - doc_start

            if profile.errors:
                # Document was opened but had issues during analysis
                print(f"WARN  ({elapsed:.1f}s) — {len(profile.errors)} issue(s): {profile.errors[0][:60]}")
            else:
                print(
                    f"OK  [Type {profile.doc_type}]  "
                    f"{profile.total_char_count:,} chars  "
                    f"{profile.total_word_count:,} words  "
                    f"{profile.table_count} tables  "
                    f"({elapsed:.1f}s)"
                )
        except Exception as e:
            # Unexpected error — this should be rare since profile_document()
            # has its own error handling. Log it and continue.
            elapsed = time.time() - doc_start
            print(f"FAIL  ({elapsed:.1f}s) — {e}")
            traceback.print_exc()

    total_elapsed = time.time() - total_start
    print("-" * 60)
    print(f"  Profiling complete: {len(profiles)} documents in {total_elapsed:.1f}s")

    if not profiles:
        print("ERROR: No documents were successfully profiled. Nothing to write.")
        sys.exit(1)

    # -------------------------------------------------------------------------
    # ASSIGN PRIORITY RANKS
    # Sort by priority score (highest first) and assign rank numbers 1..N.
    # These ranks appear in the Excel inventory's "Rank" column.
    # -------------------------------------------------------------------------
    profiles.sort(key=lambda p: p.priority_score, reverse=True)
    for rank, p in enumerate(profiles, 1):
        p.priority_rank = rank

    # -------------------------------------------------------------------------
    # WRITE OUTPUT FILES
    # Each write is wrapped in _safe_write() which catches PermissionError
    # (file open in Excel) and gives a clear message instead of crashing.
    # All five output files are attempted even if one fails.
    # -------------------------------------------------------------------------
    print("\n  Writing output files...")

    # Handle both unified config (output.profiler.*) and legacy config (output.*)
    profiler_out = config['output'].get('profiler', config['output'])
    inv_path  = os.path.join(output_dir, profiler_out.get('inventory_file', 'document_inventory.xlsx'))
    json_path = os.path.join(output_dir, profiler_out.get('json_file', 'document_profiles.json'))
    sec_path  = os.path.join(output_dir, profiler_out.get('sections_file', 'section_inventory.csv'))
    tbl_path  = os.path.join(output_dir, profiler_out.get('tables_file', 'table_inventory.csv'))
    xref_path = os.path.join(output_dir, profiler_out.get('crossrefs_file', 'crossref_inventory.csv'))

    # Master Excel inventory (most important output — open this first)
    ok = _safe_write(lambda p: write_inventory_xlsx(profiles, p, config), inv_path)
    status = "OK" if ok else "FAILED"
    print(f"  [{status}] Inventory spreadsheet : {inv_path}")

    # JSON (used by downstream automation scripts; ignore if you don't need it)
    ok = _safe_write(lambda p: write_json(profiles, p), json_path)
    status = "OK" if ok else "FAILED"
    print(f"  [{status}] JSON profiles          : {json_path}")

    # Section detail CSV (one row per section per doc)
    ok = _safe_write(lambda p: write_sections_csv(profiles, p), sec_path)
    status = "OK" if ok else "FAILED"
    print(f"  [{status}] Section inventory      : {sec_path}")

    # Table detail CSV (one row per table per doc)
    ok = _safe_write(lambda p: write_tables_csv(profiles, p), tbl_path)
    status = "OK" if ok else "FAILED"
    print(f"  [{status}] Table inventory        : {tbl_path}")

    # Cross-reference detail CSV
    ok = _safe_write(lambda p: write_crossrefs_csv(profiles, p), xref_path)
    status = "OK" if ok else "FAILED"
    print(f"  [{status}] Cross-ref inventory    : {xref_path}")

    # -------------------------------------------------------------------------
    # SUMMARY
    # -------------------------------------------------------------------------
    print("\n" + "=" * 60)
    print("  RESULTS SUMMARY")
    print("=" * 60)
    type_counts = {t: len([p for p in profiles if p.doc_type == t]) for t in "ABCDE"}
    print(f"  Total documents profiled : {len(profiles)}")
    print(f"  Type A  Table-Heavy      : {type_counts['A']}")
    print(f"  Type B  Prose-Heavy      : {type_counts['B']}")
    print(f"  Type C  Hybrid           : {type_counts['C']}")
    print(f"  Type D  Appendix-Dom     : {type_counts['D']}")
    print(f"  Type E  Unclassified     : {type_counts['E']}"
          + ("  ← needs manual review" if type_counts['E'] else ""))
    print()
    over_limit = len([p for p in profiles if p.over_size_limit])
    tracked    = len([p for p in profiles if p.has_tracked_changes])
    protected  = len([p for p in profiles if p.has_password_protection])
    print(f"  Over 36k char limit      : {over_limit}"
          + ("  ← MUST SPLIT before SharePoint upload" if over_limit else ""))
    print(f"  Has tracked changes      : {tracked}"
          + ("  ← accept/reject in Word before processing" if tracked else ""))
    print(f"  Password protected       : {protected}"
          + ("  ← must unlock before Notebook 2 can process" if protected else ""))
    print(f"  Fake headings (total)    : {sum(p.fake_heading_count for p in profiles)}")
    print(f"  Total tables             : {sum(p.table_count for p in profiles)}")
    print(f"  Tables w/ merged cells   : {sum(p.tables_with_merged_cells for p in profiles)}")
    print(f"  Total cross-references   : {sum(p.cross_ref_count for p in profiles)}")

    # Top 10 priority
    print()
    print("  TOP 10 PRIORITY DOCUMENTS (process these first):")
    for p in profiles[:10]:
        print(f"    #{p.priority_rank:2d}  [{p.doc_type}]  Score {p.priority_score:6.1f}  {p.filename}")

    # -------------------------------------------------------------------------
    # NEXT STEPS
    # -------------------------------------------------------------------------
    print()
    print("=" * 60)
    print("  NEXT STEPS")
    print("=" * 60)
    print(f"  1. Open: {inv_path}")
    print("     → Start with the 'Summary' tab for an overview")
    print("     → 'Document Inventory' tab is pre-sorted by priority score")
    print()
    if type_counts['E']:
        print(f"  2. REQUIRED: {type_counts['E']} Type E document(s) need manual classification.")
        print("     → See 'Type Reason' column for why they didn't auto-classify")
        print("     → Do NOT force-fit into A-D; run a calibration cycle for Type E docs")
        print()
    if over_limit:
        print(f"  3. REQUIRED: {over_limit} document(s) exceed the 36k character limit.")
        print("     → These must be split at Heading 1 boundaries before uploading")
        print("     → Use Notebook 1 to generate a splitting script")
        print()
    if tracked:
        print(f"  4. ACTION NEEDED: {tracked} document(s) have unaccepted tracked changes.")
        print("     → Open each in Word → Review tab → Accept All Changes → Save")
        print()
    if protected:
        print(f"  5. ACTION NEEDED: {protected} document(s) are password protected.")
        print("     → Open in Word → Review → Restrict Editing → Stop Protection")
        print()
    print("  Add a 'Usage Frequency' column (0-10) for policies people query most.")
    print("  This helps re-rank docs by business value, not just structural complexity.")
    print()
    print(f"  All output files are in: {output_dir}")
    print("=" * 60)
    print()


if __name__ == "__main__":
    main()
