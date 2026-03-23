"""
Step 5 of 5: Word Counter
============================

Scans a folder for .docx files, counts the words in each document,
and writes a CSV with per-document and total word counts.

Use this as a QA step — run it on the final split sub-documents
to verify sizes are reasonable.

Usage (unified pipeline):
    python run_pipeline.py --step 5

Usage (standalone):
    python scripts/word_counter.py
    python scripts/word_counter.py --config dps_config.yaml
    python scripts/word_counter.py ./output/3\ -\ split_documents/

Output:
    word_counts.csv — one row per document + TOTAL row

REQUIREMENTS:
    pip install python-docx
    Python 3.10 or later
"""

import csv
import os
import sys

from docx import Document

from shared_utils import (
    ensure_output_dir,
    get_output_dir,
    iter_docx_files,
    load_config,
    setup_argparse,
)


def count_words(doc_path: str) -> int:
    """Count total words in a .docx file (paragraphs + table cells)."""
    doc = Document(doc_path)
    total = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            total += len(text.split())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    total += len(text.split())

    return total


def main():
    parser = setup_argparse("Step 5: Count words in .docx files")
    args = parser.parse_args()

    config = load_config(args.config)

    # Input: by default, count the split sub-documents from Step 3
    if args.input_dir:
        input_dir = args.input_dir
    else:
        input_dir = get_output_dir(config, "split_documents")

    output_dir = get_output_dir(config, "reports", args.output_dir)
    ensure_output_dir(output_dir)

    docx_files = iter_docx_files(input_dir, config)
    if not docx_files:
        print(f"No .docx files found in '{input_dir}'.")
        return

    word_counts_file = config.get("output", {}).get("reports", {}).get("word_counts_file", "word_counts.csv")
    output_csv = os.path.join(output_dir, word_counts_file)

    print(f"Found {len(docx_files)} .docx file(s) in '{input_dir}'.\n")

    results: list[tuple[str, int]] = []
    grand_total = 0

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        try:
            wc = count_words(filepath)
            results.append((filename, wc))
            grand_total += wc
            print(f"  {filename:.<60s} {wc:>8,} words")
        except Exception as e:
            print(f"  ERROR processing {filename}: {e}")
            results.append((filename, 0))

    print(f"\n  {'TOTAL':.<60s} {grand_total:>8,} words")

    with open(output_csv, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["filename", "word_count"])
        for filename, wc in results:
            writer.writerow([filename, wc])
        writer.writerow(["TOTAL", grand_total])

    print(f"\n" + "=" * 60)
    print("STEP 5 — WORD COUNT SUMMARY")
    print("=" * 60)
    print(f"Documents counted: {len(results)}")
    print(f"Total words:       {grand_total:,}")
    print(f"\nCSV written to: {output_csv}")


if __name__ == "__main__":
    main()
