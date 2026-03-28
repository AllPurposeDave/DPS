"""
Step 9: Control Validation
============================

Validates that all controls extracted by Step 2 (extract_controls.py) are
present in the split documents produced by Step 5 (section_splitter.py).

Also generates a human-review validation workbook with confidence scoring
to help reviewers efficiently triage the ~10-20% of controls that may have
extraction errors.

Reads:
    - output/2 - controls/controls_output.csv  (source controls)
    - output/5 - split_documents/*.docx         (split documents)
    - output/5 - split_documents/split_manifest.csv (sub-doc → parent mapping)
    - input/*.docx                              (source documents for context)

Outputs:
    - output/9 - validation/control_validation.csv
    - output/9 - validation/validation_review.xlsx

Usage (unified pipeline):
    python run_pipeline.py --step 9

Usage (standalone):
    python scripts/validate_controls.py
    python scripts/validate_controls.py --config ../dps_config.xlsx

PREREQUISITES: Steps 2, 4, and 5 must have run first.
"""

from __future__ import annotations

import csv
import os
import re
import traceback
from collections import Counter
from typing import Optional

from docx import Document

from shared_utils import (
    ensure_output_dir,
    get_input_dir,
    load_config,
    setup_argparse,
)
from extract_controls import extract_paragraphs_from_docx


CSV_COLUMNS = [
    "control_id",
    "source_file",
    "source_section",
    "found_in_split_docs",
    "split_doc_filenames",
    "parent_doc_match",
    "status",
]


def build_control_id_regex(config: dict) -> re.Pattern:
    """Build a combined regex from the control_id_patterns in config."""
    ctrl_cfg = config.get("control_extraction", {})
    id_patterns = ctrl_cfg.get("control_id_patterns", None)
    if id_patterns is None:
        legacy = ctrl_cfg.get("control_id_pattern", r'\b[A-Z]{2,4}[-.]?\d{1,3}[-.]\d{2,4}\b')
        id_patterns = [legacy]
    combined = "|".join(f"({p})" for p in id_patterns)
    return re.compile(combined)


def load_source_controls(controls_csv_path: str) -> list[dict]:
    """
    Read controls_output.csv and extract unique control records.
    Returns list of dicts with control_id, source_file, source_section.
    """
    controls = []
    seen = set()
    with open(controls_csv_path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            control_id = row.get("control_id", "").strip()
            source_file = row.get("source_file", "").strip()
            section = row.get("section_header", "").strip()
            if not control_id:
                continue
            key = (control_id, source_file)
            if key not in seen:
                seen.add(key)
                controls.append({
                    "control_id": control_id,
                    "source_file": source_file,
                    "source_section": section,
                })
    return controls


def load_split_manifest(manifest_path: str) -> dict:
    """
    Read split_manifest.csv to map sub-doc filenames to parent documents.
    Returns dict: {sub_doc_filename: original_doc}
    """
    mapping = {}
    if not os.path.exists(manifest_path):
        return mapping
    with open(manifest_path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            sub_doc = row.get("sub_doc_filename", "").strip()
            original = row.get("original_doc", "").strip()
            if sub_doc:
                mapping[sub_doc] = original
    return mapping


def scan_split_documents(split_doc_dir: str, control_regex: re.Pattern) -> dict:
    """
    Open each .docx in the split documents directory, extract all text,
    and find all control IDs present.
    Returns dict: {control_id: [list of split doc filenames]}
    """
    control_to_docs = {}
    docx_files = sorted([
        f for f in os.listdir(split_doc_dir)
        if f.lower().endswith(".docx") and not f.startswith("~$")
    ])

    for filename in docx_files:
        filepath = os.path.join(split_doc_dir, filename)
        try:
            doc = Document(filepath)
        except Exception as e:
            print(f"  WARNING: Could not open {filename}: {e}")
            continue

        # Extract all text from paragraphs and tables
        text_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        full_text = "\n".join(text_parts)

        # Find all control IDs in this document
        matches = control_regex.findall(full_text)
        # findall with groups returns tuples; flatten to get actual matched strings
        found_ids = set()
        for match in matches:
            if isinstance(match, tuple):
                # Take the first non-empty group
                for group in match:
                    if group:
                        found_ids.add(group)
                        break
            else:
                found_ids.add(match)

        for cid in found_ids:
            control_to_docs.setdefault(cid, []).append(filename)

    return control_to_docs


def normalize_parent_name(original_doc: str, source_file: str) -> bool:
    """
    Check if a split doc's parent matches the source file.
    Handles _fixed suffix and common naming variations.
    """
    # Strip _fixed suffix from original_doc
    orig = original_doc.replace("_fixed", "").strip()
    src = source_file.strip()
    # Compare base names without extension
    orig_base = os.path.splitext(orig)[0].lower()
    src_base = os.path.splitext(src)[0].lower()
    return orig_base == src_base


def validate_controls(
    source_controls: list[dict],
    control_to_docs: dict,
    manifest_mapping: dict,
) -> list[dict]:
    """
    Cross-reference source controls against split document findings.
    Returns validation records.
    """
    results = []

    for ctrl in source_controls:
        cid = ctrl["control_id"]
        source_file = ctrl["source_file"]
        source_section = ctrl["source_section"]

        found_docs = control_to_docs.get(cid, [])
        found = len(found_docs) > 0

        # Check parent document match
        parent_match = False
        if found and manifest_mapping:
            for doc_name in found_docs:
                parent = manifest_mapping.get(doc_name, "")
                if parent and normalize_parent_name(parent, source_file):
                    parent_match = True
                    break

        # Determine status
        if not found:
            status = "MISSING"
            parent_match_str = ""
        elif not manifest_mapping:
            status = "PASS"
            parent_match_str = "N/A"
        elif parent_match:
            status = "PASS"
            parent_match_str = "YES"
        else:
            status = "RELOCATED"
            parent_match_str = "NO"

        results.append({
            "control_id": cid,
            "source_file": source_file,
            "source_section": source_section,
            "found_in_split_docs": "YES" if found else "NO",
            "split_doc_filenames": ", ".join(found_docs) if found_docs else "",
            "parent_doc_match": parent_match_str,
            "status": status,
        })

    return results


# ============================================================
# SECTION: VALIDATION REVIEW WORKBOOK
# ============================================================

# Sections that almost always contain reference/crosswalk tables, not real controls
SUSPECT_SECTIONS = [
    "revision history", "framework crosswalk", "references",
    "glossary", "definitions", "document control",
]

# Table header keywords that indicate leaked table structure in descriptions
TABLE_HEADER_KEYWORDS = [
    "control id", "nist mapping", "implementation status",
    "responsible party", "framework ref", "control family",
]


def load_full_controls(csv_path: str) -> list[dict]:
    """Read controls_output.csv keeping all columns."""
    controls = []
    with open(csv_path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for i, row in enumerate(reader):
            row["_row_index"] = i
            controls.append(row)
    return controls


def build_source_index(docx_path: str, control_regex: re.Pattern) -> dict:
    """
    Open a source .docx and build a lookup for control ID locations.
    Returns {"paragraphs": [...], "control_locations": {id: [indices]}}
    """
    paragraphs = extract_paragraphs_from_docx(docx_path)
    control_locations = {}
    for idx, para in enumerate(paragraphs):
        matches = control_regex.findall(para["text"])
        for match in matches:
            if isinstance(match, tuple):
                cid = next((g for g in match if g), None)
            else:
                cid = match
            if cid:
                control_locations.setdefault(cid, []).append(idx)
    return {"paragraphs": paragraphs, "control_locations": control_locations}


def find_source_context(source_index: dict, control_id: str, context_n: int = 3) -> str:
    """
    Return raw text of paragraphs surrounding the first occurrence of
    a control ID in the source document. The control paragraph is
    marked with '>>> ' prefix for quick visual identification.
    """
    paragraphs = source_index["paragraphs"]
    locations = source_index["control_locations"].get(control_id, [])
    if not locations:
        return "(Control ID not found in source document)"

    center = locations[0]
    start = max(0, center - context_n)
    end = min(len(paragraphs), center + context_n + 1)

    lines = []
    for i in range(start, end):
        text = paragraphs[i]["text"]
        if not text:
            continue
        prefix = ">>> " if i == center else "    "
        lines.append(f"{prefix}{text}")
    return "\n".join(lines)


def compute_confidence(
    control: dict,
    source_index: dict | None,
    all_controls: list[dict],
    config: dict,
) -> tuple[int, list[str]]:
    """
    Score a control's extraction confidence (0-100). Returns (score, [flags]).
    """
    score = 100
    flags = []

    description = control.get("control_description", "").strip()
    section = control.get("section_header", "").strip()
    baseline = control.get("baseline", "").strip()
    name = control.get("control_name", "").strip()
    source = control.get("extraction_source", "").strip()
    cid = control.get("control_id", "").strip()
    source_file = control.get("source_file", "").strip()

    # Empty or short description
    if not description:
        score -= 40
        flags.append("EMPTY_DESCRIPTION")
    elif len(description) < 20:
        score -= 30
        flags.append("SHORT_DESCRIPTION")
    elif len(description) > 2000:
        score -= 20
        flags.append("LONG_DESCRIPTION")

    # Suspect section headers
    section_lower = section.lower()
    if any(s in section_lower for s in SUSPECT_SECTIONS):
        score -= 35
        flags.append("SUSPECT_SECTION")
    elif "appendix" in section_lower:
        score -= 15
        flags.append("APPENDIX_SECTION")

    # Duplicate control IDs — same ID within same source file is very suspicious
    # (e.g. body text + crosswalk table). Same ID across different files is less so.
    same_file_dupes = sum(
        1 for c in all_controls
        if c.get("control_id", "").strip() == cid
        and c.get("source_file", "").strip() == source_file
    )
    cross_file_dupes = sum(
        1 for c in all_controls
        if c.get("control_id", "").strip() == cid
        and c.get("source_file", "").strip() != source_file
    )
    if same_file_dupes > 1:
        score -= 25
        flags.append("DUPLICATE_ID_SAME_DOC")
    elif cross_file_dupes > 0:
        score -= 10
        flags.append("DUPLICATE_ID_CROSS_DOC")

    # Table header keywords in description
    desc_lower = description.lower()
    if any(kw in desc_lower for kw in TABLE_HEADER_KEYWORDS):
        score -= 20
        flags.append("TABLE_HEADERS_IN_DESC")

    # Guidance keywords in description (boundary likely wrong)
    ctrl_cfg = config.get("control_extraction", {})
    guidance_keywords = ctrl_cfg.get("guidance_keywords", [
        "implementation guidance", "implementation:", "guidelines:",
        "how to implement", "supplemental guidance",
    ])
    if any(kw.lower() in desc_lower for kw in guidance_keywords):
        score -= 15
        flags.append("GUIDANCE_IN_DESC")

    # Table extraction source
    if source.lower() == "table":
        score -= 10
        flags.append("TABLE_SOURCE")

    # Empty baseline when siblings (same source_file) have baselines
    if not baseline:
        siblings = [c for c in all_controls if c.get("source_file", "").strip() == source_file]
        has_baseline_siblings = any(c.get("baseline", "").strip() for c in siblings)
        if has_baseline_siblings:
            score -= 10
            flags.append("EMPTY_BASELINE")

    # Empty control name when siblings have names
    if not name:
        siblings = [c for c in all_controls if c.get("source_file", "").strip() == source_file]
        has_name_siblings = any(c.get("control_name", "").strip() for c in siblings)
        if has_name_siblings:
            score -= 5
            flags.append("EMPTY_NAME")

    # Multiple controls in same paragraph (source_index check)
    if source_index:
        locations = source_index["control_locations"]
        for para_indices in locations.values():
            for pidx in para_indices:
                # Check if another control ID shares this paragraph
                ids_at_para = [
                    k for k, indices in locations.items()
                    if pidx in indices and k != cid
                ]
                if cid in [k for k, indices in locations.items() if pidx in indices]:
                    if ids_at_para:
                        score -= 10
                        flags.append("MULTI_CONTROL_PARAGRAPH")
                        break
            else:
                continue
            break

    return (max(0, score), flags)


def generate_validation_xlsx(
    scored_controls: list[dict],
    output_path: str,
) -> None:
    """
    Generate a validation review workbook with two sheets:
    - 'Validation Review': one row per control, sorted by confidence
    - 'Summary': per-document and per-flag statistics
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.datavalidation import DataValidation

    wb = Workbook()

    # ---- Sheet 1: Validation Review ----
    ws = wb.active
    ws.title = "Validation Review"

    headers = [
        "Row #", "Confidence", "Flags", "Control ID", "Source File",
        "Section Header", "Extraction Source", "Source Context",
        "Extracted Description", "Extracted Guidance", "Baseline",
        "Control Name", "Validation Status", "Reviewer Notes",
    ]
    ws.append(headers)

    # Color fills for confidence bands
    red_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    yellow_fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
    green_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")

    # Sort: confidence ascending (worst first), then source_file
    scored_controls.sort(key=lambda c: (c["_confidence"], c.get("source_file", "")))

    for i, ctrl in enumerate(scored_controls, start=1):
        confidence = ctrl["_confidence"]
        row_data = [
            i,
            confidence,
            ", ".join(ctrl["_flags"]),
            ctrl.get("control_id", ""),
            ctrl.get("source_file", ""),
            ctrl.get("section_header", ""),
            ctrl.get("extraction_source", ""),
            ctrl.get("_source_context", ""),
            ctrl.get("control_description", ""),
            ctrl.get("supplemental_guidance", ""),
            ctrl.get("baseline", ""),
            ctrl.get("control_name", ""),
            "",  # Validation Status (human fills in)
            "",  # Reviewer Notes (human fills in)
        ]
        ws.append(row_data)

        # Color-code the row
        row_num = i + 1  # +1 for header
        if confidence <= 30:
            fill = red_fill
        elif confidence <= 60:
            fill = yellow_fill
        else:
            fill = green_fill

        for col in range(1, len(headers) + 1):
            ws.cell(row=row_num, column=col).fill = fill

    # Header styling
    header_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )
    for cell in ws[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    # Freeze pane and autofilter
    ws.freeze_panes = "A2"
    last_col = get_column_letter(len(headers))
    ws.auto_filter.ref = f"A1:{last_col}{ws.max_row}"

    # Column widths
    col_widths = {
        1: 7, 2: 12, 3: 35, 4: 16, 5: 40, 6: 30, 7: 12,
        8: 70, 9: 60, 10: 40, 11: 10, 12: 30, 13: 20, 14: 30,
    }
    wrap_align = Alignment(vertical="top", wrap_text=True)
    for col_idx, width in col_widths.items():
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    # Wrap text on long columns (Source Context, Description, Guidance)
    for data_row in ws.iter_rows(min_row=2):
        for cell in data_row:
            if cell.column in (8, 9, 10, 14):  # context, desc, guidance, notes
                cell.alignment = wrap_align

    # Data validation dropdown for Validation Status (column M = 13)
    status_options = (
        "Correct,Wrong-FalsePositive,Wrong-Description,Wrong-Guidance,"
        "Wrong-Baseline,Wrong-Section,Missing-Content,Needs-Review"
    )
    dv = DataValidation(type="list", formula1=f'"{status_options}"', allow_blank=True)
    dv.prompt = "Select validation status"
    dv.promptTitle = "Validation Status"
    ws.add_data_validation(dv)
    for row_num in range(2, ws.max_row + 1):
        dv.add(ws.cell(row=row_num, column=13))

    # ---- Sheet 2: Summary ----
    ws2 = wb.create_sheet(title="Summary")

    # Confidence band counts
    total = len(scored_controls)
    red_count = sum(1 for c in scored_controls if c["_confidence"] <= 30)
    yellow_count = sum(1 for c in scored_controls if 31 <= c["_confidence"] <= 60)
    green_count = sum(1 for c in scored_controls if c["_confidence"] > 60)

    ws2.append(["Validation Review Summary"])
    ws2.append([])
    ws2.append(["Total Controls", total])
    ws2.append(["Red (0-30) - Review First", red_count])
    ws2.append(["Yellow (31-60) - Suspicious", yellow_count])
    ws2.append(["Green (61-100) - Likely Correct", green_count])
    ws2.append([])

    # Per-document breakdown
    ws2.append(["Per-Document Breakdown"])
    ws2.append(["Source File", "Count", "Avg Confidence", "Red", "Yellow", "Green"])
    docs = {}
    for c in scored_controls:
        sf = c.get("source_file", "")
        docs.setdefault(sf, []).append(c["_confidence"])
    for sf, confs in sorted(docs.items()):
        avg = round(sum(confs) / len(confs), 1)
        r = sum(1 for x in confs if x <= 30)
        y = sum(1 for x in confs if 31 <= x <= 60)
        g = sum(1 for x in confs if x > 60)
        ws2.append([sf, len(confs), avg, r, y, g])

    ws2.append([])

    # Per-flag breakdown
    ws2.append(["Flag Frequency"])
    ws2.append(["Flag", "Count"])
    flag_counter = Counter()
    for c in scored_controls:
        for f in c["_flags"]:
            flag_counter[f] += 1
    for flag, count in flag_counter.most_common():
        ws2.append([flag, count])

    ws2.append([])
    ws2.append(["Workflow Instructions"])
    ws2.append(["1. Red rows appear first in the Validation Review sheet — review these first."])
    ws2.append(["2. Compare Source Context (col H) with Extracted Description (col I)."])
    ws2.append(["3. Set Validation Status (col M) using the dropdown."])
    ws2.append(["4. Add notes in col N as needed."])
    ws2.append(["5. Spot-check 10-20% of green rows for calibration."])
    ws2.append(["6. Use the 'Add Missing Controls' sheet to enter controls the extractor missed."])

    # Style summary header
    ws2["A1"].font = Font(name="Arial", bold=True, size=14)
    for cell in ws2[9]:  # doc breakdown header row
        if cell.value:
            cell.font = Font(name="Arial", bold=True, size=10)

    # Auto-width for summary
    for col_idx in range(1, 7):
        ws2.column_dimensions[get_column_letter(col_idx)].width = 35

    # ---- Sheet 3: Add Missing Controls ----
    ws3 = wb.create_sheet(title="Add Missing Controls")

    # Instruction row
    ws3.append([
        "Enter controls that exist in source documents but were NOT extracted by the pipeline. "
        "Required fields: Control ID, Source File, Section Header, Control Description."
    ])
    ws3.merge_cells(start_row=1, start_column=1, end_row=1, end_column=8)
    ws3["A1"].font = Font(name="Arial", italic=True, size=10, color="555555")
    ws3["A1"].alignment = Alignment(wrap_text=True)
    ws3.row_dimensions[1].height = 35

    # Headers
    missing_headers = [
        "Control ID", "Source File", "Section Header", "Control Description",
        "Supplemental Guidance", "Baseline", "Control Name", "Reviewer Notes",
    ]
    ws3.append(missing_headers)
    for cell in ws3[2]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    # Column widths
    missing_col_widths = {1: 16, 2: 40, 3: 30, 4: 60, 5: 40, 6: 10, 7: 30, 8: 30}
    for col_idx, width in missing_col_widths.items():
        ws3.column_dimensions[get_column_letter(col_idx)].width = width

    # Example placeholder row (light gray italic)
    example_row = [
        "AC-2.3", "Access_Control_Policy_POL-AC-2026-001.docx",
        "Access Enforcement", "The organization enforces approved authorizations...",
        "", "L, M, H", "Access Enforcement", "Missed because control ID was in a table cell",
    ]
    ws3.append(example_row)
    placeholder_font = Font(name="Arial", italic=True, size=10, color="999999")
    for cell in ws3[3]:
        cell.font = placeholder_font

    # Data validation: Source File dropdown from known source files
    source_files = sorted(set(
        c.get("source_file", "").strip() for c in scored_controls
        if c.get("source_file", "").strip()
    ))
    if source_files:
        sf_list = ",".join(source_files)
        # Excel data validation has a 255-char formula limit; use a range if too long
        if len(sf_list) <= 250:
            sf_dv = DataValidation(type="list", formula1=f'"{sf_list}"', allow_blank=True)
        else:
            # Too many files for inline list — skip dropdown, reviewers type manually
            sf_dv = None
        if sf_dv:
            sf_dv.prompt = "Select the source document"
            sf_dv.promptTitle = "Source File"
            ws3.add_data_validation(sf_dv)
            # Apply to rows 3-102 (100 empty rows for reviewer input)
            for row_num in range(3, 103):
                sf_dv.add(ws3.cell(row=row_num, column=2))

    # Freeze header
    ws3.freeze_panes = "A3"

    wb.save(output_path)


def write_review_readme(validation_dir: str) -> None:
    """Write a README.md into the validation output directory to guide reviewers."""
    readme_path = os.path.join(validation_dir, "README.md")
    content = """\
# Validation Review Workflow

## What This Folder Contains

| File | Purpose |
|------|---------|
| **control_validation.csv** | Machine-readable validation status for each control (PASS / MISSING / RELOCATED) |
| **validation_review.xlsx** | Human review workbook with confidence scoring, source context, and input sheets |
| **README.md** | This file — instructions for the review workflow |

After review, the feedback script produces:
| File | Purpose |
|------|---------|
| **confirmed_controls.csv** | The authoritative control set (reviewed + manually added, minus false positives) |
| **feedback_report.txt** | Summary of review findings, error patterns, and config suggestions |
| **suggested_config_changes.yaml** | Machine-readable config patch (only if clear patterns found) |

---

## How to Review validation_review.xlsx

> **Important:** Open in Microsoft Excel (not Google Sheets — data validation dropdowns may not work).

### Step 1: Check the Summary Sheet
Open the **Summary** sheet to understand the scope:
- Total controls, confidence band counts (Red / Yellow / Green)
- Per-document breakdown and flag frequency

### Step 2: Understand the Color Coding
The **Validation Review** sheet is sorted worst-first:
- **Red rows (0-30):** Review these first — likely extraction errors
- **Yellow rows (31-60):** Suspicious — worth spot-checking
- **Green rows (61-100):** Likely correct — sample 10-20% for calibration

### Step 3: Review Each Control (Red First)
For each row, compare:
- **Source Context (col H):** The surrounding paragraphs from the original document, with `>>>` marking the control line
- **Extracted Description (col I):** What the pipeline extracted

Also check:
- **Section Header (col F):** Is this the real section, or a suspect one (crosswalk, glossary)?
- **Extraction Source (col G):** "Table" extractions are more error-prone than "Text"
- **Baseline (col K) / Control Name (col L):** Verify if present

### Step 4: Set Validation Status (Column M — Dropdown)

| Status | When to Use |
|--------|-------------|
| **Correct** | Extraction is accurate — control ID, description, and metadata are right |
| **Wrong-FalsePositive** | This is NOT a real control (e.g., crosswalk table reference, inline mention) |
| **Wrong-Description** | Control is real but the extracted description text is wrong or truncated |
| **Wrong-Guidance** | Guidance text is wrong, or guidance was mixed into the description |
| **Wrong-Baseline** | Baseline (L/M/H) is incorrect |
| **Wrong-Section** | Control was placed under the wrong section header |
| **Missing-Content** | Control is real but key content was not captured |
| **Needs-Review** | Unsure — flag for a second reviewer |

### Step 5: Add Reviewer Notes (Column N)
Be specific about what is wrong and what the correct value should be:
- Good: "Description should stop at 'Implementation Guidance:' — everything after is guidance"
- Good: "This is a crosswalk table entry, not a real control"
- Bad: "wrong"

### Step 6: Add Missing Controls (Third Sheet)
If you spot a control in a source document that the extractor **missed entirely**, go to
the **Add Missing Controls** sheet and enter it manually. Required fields:
- **Control ID** — the ID as it appears in the document
- **Source File** — which .docx it came from (use the dropdown)
- **Section Header** — which section it appears under
- **Control Description** — the full control text

Optional: Supplemental Guidance, Baseline, Control Name, Reviewer Notes.

### Step 7: Save and Hand Off
- Save the file — keep the **.xlsx format** and **do not rename it**
- The feedback script expects `validation_review.xlsx` in this folder

---

## Flag Reference

| Flag | Confidence Penalty | What It Means |
|------|--------------------|---------------|
| EMPTY_DESCRIPTION | -40 | No control description was extracted |
| SHORT_DESCRIPTION | -30 | Description is under 20 characters |
| LONG_DESCRIPTION | -20 | Description exceeds 2000 characters |
| SUSPECT_SECTION | -35 | Found in a section like Revision History, Glossary, or Framework Crosswalk |
| APPENDIX_SECTION | -15 | Found in an Appendix section |
| DUPLICATE_ID_SAME_DOC | -25 | Same control ID extracted multiple times from the same document |
| DUPLICATE_ID_CROSS_DOC | -10 | Same control ID appears in different source documents |
| TABLE_HEADERS_IN_DESC | -20 | Description contains table header text (extraction boundary error) |
| GUIDANCE_IN_DESC | -15 | Description contains guidance keywords (boundary set too far) |
| TABLE_SOURCE | -10 | Control was extracted from a Word table (less context available) |
| EMPTY_BASELINE | -10 | No baseline when sibling controls have baselines |
| EMPTY_NAME | -5 | No control name when sibling controls have names |
| MULTI_CONTROL_PARAGRAPH | -10 | Multiple control IDs share the same paragraph |

---

## How This Feeds Back Into the Pipeline

After completing your review:

```
python scripts/ingest_review_feedback.py
python scripts/ingest_review_feedback.py --config dps_config.xlsx
```

This reads your reviewed `validation_review.xlsx` and produces:

1. **confirmed_controls.csv** — The authoritative control set:
   - Controls marked "Correct" + unreviewed controls (assumed correct)
   - Manually added controls from the "Add Missing Controls" sheet
   - Excludes controls marked "Wrong-FalsePositive"

2. **feedback_report.txt** — Analysis of review findings:
   - Review coverage and status breakdown
   - Missing controls added by reviewer
   - Error patterns and per-document accuracy
   - Config improvement suggestions

3. **suggested_config_changes.yaml** — Actionable config changes (if patterns detected)

---

## Tips for Efficient Review

- **Sort by Confidence** (default) — worst rows first, so you catch real problems early
- **Filter by Source File** — focus on one document at a time for consistency
- **Use the Summary sheet's flag frequency** — if one flag dominates, investigate that pattern
- **If you see the same error 5+ times**, note it in Reviewer Notes — it likely indicates a config fix
- **Don't review every green row** — sample 10-20% to calibrate, then trust the rest
- **Use the "Add Missing Controls" sheet** whenever you notice the extractor missed something
"""
    with open(readme_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"  Review README written to: {readme_path}")


def main():
    parser = setup_argparse("Step 6: Validate controls in split documents")
    args = parser.parse_args()

    config = load_config(args.config)
    output_cfg = config.get("output", {})
    config_dir = config.get("_config_dir", os.getcwd())

    # Resolve output root
    output_root = output_cfg.get("directory", "./output")
    if not os.path.isabs(output_root):
        output_root = os.path.normpath(os.path.join(config_dir, output_root))

    # Resolve Step 2 controls CSV path
    controls_cfg = output_cfg.get("controls", {})
    controls_dir = os.path.join(output_root, controls_cfg.get("directory", "2 - controls"))
    controls_file = controls_cfg.get("output_file", "controls_output.csv")
    controls_csv_path = os.path.join(controls_dir, controls_file)

    if not os.path.exists(controls_csv_path):
        # Try xlsx → csv fallback
        alt_path = controls_csv_path.replace(".xlsx", ".csv")
        if os.path.exists(alt_path):
            controls_csv_path = alt_path
        else:
            print(f"ERROR: Controls output not found at {controls_csv_path}")
            print("Step 2 (extract_controls.py) must run first.")
            return

    # Resolve Step 5 split documents directory
    split_cfg = output_cfg.get("split_documents", {})
    split_dir = os.path.join(output_root, split_cfg.get("directory", "5 - split_documents"))
    manifest_file = split_cfg.get("manifest_file", "split_manifest.csv")
    manifest_path = os.path.join(split_dir, manifest_file)

    if not os.path.isdir(split_dir):
        print(f"ERROR: Split documents directory not found at {split_dir}")
        print("Steps 4-5 (heading_style_fixer + section_splitter) must run first.")
        return

    # Resolve Step 9 output directory
    validation_cfg = output_cfg.get("validation", {})
    validation_dir = os.path.join(output_root, validation_cfg.get("directory", "9 - validation"))
    validation_file = validation_cfg.get("output_file", "control_validation.csv")
    ensure_output_dir(validation_dir)
    output_path = os.path.join(validation_dir, validation_file)

    # Build control ID regex
    control_regex = build_control_id_regex(config)

    # Step A: Load source controls
    print("  Loading source controls from Step 1...")
    source_controls = load_source_controls(controls_csv_path)
    if not source_controls:
        print("  WARNING: No controls found in controls output. Nothing to validate.")
        return
    print(f"  Found {len(source_controls)} unique control(s) from {len(set(c['source_file'] for c in source_controls))} document(s)")

    # Step B: Load split manifest
    manifest_mapping = load_split_manifest(manifest_path)
    if not manifest_mapping:
        print("  WARNING: Split manifest not found or empty. Parent-doc matching will be skipped.")

    # Step C: Scan split documents
    print("  Scanning split documents for control IDs...")
    control_to_docs = scan_split_documents(split_dir, control_regex)
    unique_ids_found = len(control_to_docs)
    print(f"  Found {unique_ids_found} unique control ID(s) across split documents")

    # Step D: Validate
    print("  Cross-referencing...")
    results = validate_controls(source_controls, control_to_docs, manifest_mapping)

    # Write output CSV
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=CSV_COLUMNS)
        writer.writeheader()
        writer.writerows(results)

    # Summary
    total = len(results)
    passed = sum(1 for r in results if r["status"] == "PASS")
    missing = sum(1 for r in results if r["status"] == "MISSING")
    relocated = sum(1 for r in results if r["status"] == "RELOCATED")

    print()
    print("=" * 60)
    print("STEP 6 — CONTROL VALIDATION SUMMARY")
    print("=" * 60)
    print(f"Total source controls: {total}")
    print(f"Found (PASS):          {passed}  ({round(passed/total*100, 1) if total else 0}%)")
    print(f"Missing:               {missing}  ({round(missing/total*100, 1) if total else 0}%)")
    print(f"Relocated:             {relocated}  ({round(relocated/total*100, 1) if total else 0}%)")

    if missing > 0:
        print(f"\nMISSING CONTROLS:")
        for r in results:
            if r["status"] == "MISSING":
                print(f"  {r['control_id']}  from  {r['source_file']}")

    if relocated > 0:
        print(f"\nRELOCATED CONTROLS (found in different parent doc):")
        for r in results:
            if r["status"] == "RELOCATED":
                print(f"  {r['control_id']}  from  {r['source_file']}  →  {r['split_doc_filenames'][:60]}")

    print(f"\nValidation output written to: {output_path}")

    # ---- Generate Validation Review Workbook ----
    print("\n  Generating validation review workbook...")

    # Load full controls (all columns)
    full_controls = load_full_controls(controls_csv_path)
    if not full_controls:
        print("  WARNING: No controls to review.")
        return

    # Resolve input directory for source .docx files
    input_dir = get_input_dir(config)

    # Build source indices (one per unique source file)
    source_indices = {}
    for ctrl in full_controls:
        sf = ctrl.get("source_file", "").strip()
        if sf and sf not in source_indices:
            docx_path = os.path.join(input_dir, sf)
            if os.path.exists(docx_path):
                try:
                    source_indices[sf] = build_source_index(docx_path, control_regex)
                except Exception as e:
                    print(f"  WARNING: Could not index {sf}: {e}")
                    source_indices[sf] = None
            else:
                source_indices[sf] = None

    # Compute confidence and source context for each control
    for ctrl in full_controls:
        sf = ctrl.get("source_file", "").strip()
        src_idx = source_indices.get(sf)
        confidence, flags = compute_confidence(ctrl, src_idx, full_controls, config)
        ctrl["_confidence"] = confidence
        ctrl["_flags"] = flags
        if src_idx:
            ctrl["_source_context"] = find_source_context(src_idx, ctrl.get("control_id", "").strip())
        else:
            ctrl["_source_context"] = "(Source document not available)"

    # Write the validation review workbook
    review_cfg = output_cfg.get("validation", {})
    review_file = review_cfg.get("review_file", "validation_review.xlsx")
    review_path = os.path.join(validation_dir, review_file)

    generate_validation_xlsx(full_controls, review_path)

    red_count = sum(1 for c in full_controls if c["_confidence"] <= 30)
    yellow_count = sum(1 for c in full_controls if 31 <= c["_confidence"] <= 60)
    green_count = sum(1 for c in full_controls if c["_confidence"] > 60)

    print()
    print("=" * 60)
    print("VALIDATION REVIEW WORKBOOK")
    print("=" * 60)
    print(f"Total controls scored: {len(full_controls)}")
    print(f"Red (review first):    {red_count}")
    print(f"Yellow (suspicious):   {yellow_count}")
    print(f"Green (likely correct): {green_count}")
    print(f"\nReview workbook written to: {review_path}")

    # Write the reviewer README
    write_review_readme(validation_dir)


if __name__ == "__main__":
    main()
