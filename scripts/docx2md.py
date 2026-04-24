"""
docx2md — DOCX to Markdown converter optimised for RAG pipelines (DPS Step 7).

Converts .docx policy documents into clean, well-structured markdown files with
YAML frontmatter. Part of the DPS pipeline; can also run standalone.

Supports Pure Conversion (read from input/) or Optimized (read from a pipeline
step's output, e.g. heading_fixes).

REQUIREMENTS:
  pip install python-docx pyyaml openpyxl
  Python 3.8 or later

USAGE:
  Via pipeline:  python run_pipeline.py --step 7
  Standalone:    python scripts/docx2md.py --config dps_config.xlsx input/ output/7\ -\ markdown/
"""

from __future__ import annotations

import os
import re
import sys
import time
import traceback
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Import shared_utils from the DPS scripts/ directory
# ---------------------------------------------------------------------------
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.normpath(os.path.join(_SCRIPT_DIR, ".."))
_SCRIPTS_DIR = os.path.join(_PROJECT_ROOT, "scripts")
if _SCRIPTS_DIR not in sys.path:
    sys.path.insert(0, _SCRIPTS_DIR)

from shared_utils import (
    build_custom_style_map,
    ensure_output_dir,
    get_heading_level,
    is_paragraph_bold,
    iter_docx_files,
    load_config,
    match_doc_name,
    resolve_path,
    sanitize_filename,
    setup_argparse,
)
from add_metadata import load_url_mapping, resolve_url




# ============================================================================
# Unicode / cleanup helpers
# ============================================================================

_SMART_QUOTE_MAP = {
    "\u201c": '"', "\u201d": '"',   # left/right double quotes
    "\u2018": "'", "\u2019": "'",   # left/right single quotes
    "\u2013": "--",                  # en-dash
    "\u2014": "---",                # em-dash
    "\u2026": "...",                # ellipsis
}

_ZERO_WIDTH_RE = re.compile("[\u200b\u200c\u200d\ufeff]")


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")


def _split_body_at_sections(body_lines: list, char_limit: int,
                             base_chars: int,
                             control_limit: int = 0,
                             control_regex: "re.Pattern | None" = None) -> list[dict]:
    """Pack body lines into chunks at heading boundaries.

    Each chunk is a dict {"lines": [...], "last_header": str|None}.
    A section (heading + its content, up to the next heading of equal or
    shallower level) is the smallest splittable unit — a single oversized
    section is kept whole (overflow is preferred to splitting mid-section).

    Splits when EITHER limit would be exceeded:
      - char_limit > 0   and accumulated chars + next section > char_limit
      - control_limit > 0 and accumulated controls + next section > control_limit
    """
    # Parse into sections. A section starts at a heading line; content before
    # the first heading forms a headerless preamble section.
    sections: list[dict] = []
    current = {"header": None, "level": 0, "lines": []}
    for line in body_lines:
        m = _HEADING_RE.match(line)
        if m:
            if current["lines"] or current["header"]:
                sections.append(current)
            current = {
                "header": m.group(2).strip(),
                "level": len(m.group(1)),
                "lines": [line],
            }
        else:
            current["lines"].append(line)
    if current["lines"] or current["header"]:
        sections.append(current)

    # Greedy pack: add sections until the next one would exceed a limit,
    # then start a new chunk. An oversized section goes in alone.
    chunks: list[dict] = []
    cur_lines: list[str] = []
    cur_chars = base_chars
    cur_controls = 0
    cur_last_header: "str | None" = None

    for section in sections:
        sec_chars = sum(len(s) + 1 for s in section["lines"])
        sec_controls = 0
        if control_limit > 0 and control_regex is not None:
            sec_controls = len(control_regex.findall("\n".join(section["lines"])))

        char_overflow = char_limit > 0 and (cur_chars + sec_chars > char_limit)
        ctrl_overflow = control_limit > 0 and (cur_controls + sec_controls > control_limit)
        if cur_lines and (char_overflow or ctrl_overflow):
            chunks.append({"lines": cur_lines, "last_header": cur_last_header})
            cur_lines = []
            cur_chars = base_chars
            cur_controls = 0
            cur_last_header = None

        cur_lines.extend(section["lines"])
        cur_chars += sec_chars
        cur_controls += sec_controls
        if section["header"]:
            cur_last_header = section["header"]

    if cur_lines:
        chunks.append({"lines": cur_lines, "last_header": cur_last_header})

    return chunks


def _clean_text(text: str, config_d2m: dict) -> str:
    """Apply smart-quote and zero-width cleanup to a text string."""
    if config_d2m.get("clean_smart_quotes", True):
        for old, new in _SMART_QUOTE_MAP.items():
            text = text.replace(old, new)
    if config_d2m.get("strip_zero_width_chars", True):
        text = _ZERO_WIDTH_RE.sub("", text)
    return text


# ============================================================================
# Metadata frontmatter
# ============================================================================

def _build_frontmatter(doc, filepath: str, config: dict,
                       url_mapping: dict | None = None) -> str:
    """Build YAML frontmatter string from configurable metadata_fields."""
    d2m = config.get("docx2md", {})
    if not d2m.get("include_metadata_frontmatter", True):
        return ""

    fields = list(d2m.get("metadata_fields", _default_metadata_fields()))
    static_tags = config.get("metadata", {}).get("tags", {}).get("static_tags", []) or []

    lines = ["---"]

    for field in fields:
        name = field.get("name", "")
        source = field.get("source", "")
        default = field.get("default", "")
        value = _resolve_metadata_value(doc, filepath, source, default,
                                        url_mapping=url_mapping, config=config)

        # Merge org-wide static_tags into the Tags field so metadata.tags.static_tags
        # is the single source of truth for both Word-doc tags (Step 6) and
        # .md frontmatter tags (Step 7). Case-insensitive dedup; existing values kept.
        if static_tags and name.lower() == "tags":
            if isinstance(value, list):
                merged = list(value)
            elif value and value != default:
                merged = [value]
            else:
                merged = []
            seen = {t.lower() for t in merged}
            for t in static_tags:
                if t.lower() not in seen:
                    merged.append(t)
                    seen.add(t.lower())
            value = merged

        if isinstance(value, list):
            items = ", ".join(f'"{v}"' for v in value)
            value = f"[{items}]"
        else:
            value = f'"{value}"'
        lines.append(f"{name}: {value}")

    lines.append("---")
    lines.append("")
    return "\n".join(lines)


def _resolve_metadata_value(doc, filepath: str, source: str, default: str,
                            url_mapping: dict | None = None,
                            config: dict | None = None) -> "str | list":
    """Resolve a single metadata field value from its source descriptor.

    Returns a string for scalar sources, or a list of strings for
    ``excel_lookup_list`` sources.
    """
    try:
        if source == "filename":
            return os.path.basename(filepath)

        if source == "converted_date":
            return datetime.now().strftime("%Y-%m-%dT%H:%M:%S")

        if source == "doc_url" and url_mapping is not None and config is not None:
            stem = os.path.splitext(os.path.basename(filepath))[0]
            url, _src = resolve_url(stem, os.path.basename(filepath),
                                    url_mapping, config)
            return url if url else default

        if source.startswith("core:"):
            prop_name = source[5:]
            cp = doc.core_properties
            val = getattr(cp, prop_name, None)
            if val is None:
                return default
            if hasattr(val, "strftime"):
                return val.strftime("%Y-%m-%d")
            return str(val).strip() or default

        if source.startswith("filename_regex:"):
            pattern = source[15:]
            basename = os.path.splitext(os.path.basename(filepath))[0]
            m = re.search(pattern, basename)
            return m.group(1) if m else default

        if source.startswith("static:"):
            return source[7:]

        if source.startswith("excel_lookup_list:") and config is not None:
            # Format: excel_lookup_list:<excel_path>:<sheet>:<key_col_header>:<value_col_header>
            # Reads all rows in <sheet> where <key_col_header> matches the current
            # doc filename and returns a list of <value_col_header> values.
            params = source[len("excel_lookup_list:"):].split(":", 3)
            if len(params) != 4:
                return default
            xl_path_raw, sheet_name, key_col_name, value_col_name = params
            xl_path = resolve_path(config, xl_path_raw)
            if not os.path.isfile(xl_path):
                return default
            import openpyxl
            wb = openpyxl.load_workbook(xl_path, read_only=True, data_only=True)
            if sheet_name not in wb.sheetnames:
                return default
            ws = wb[sheet_name]
            rows = ws.iter_rows(values_only=True)
            headers = list(next(rows, []))
            try:
                key_idx = headers.index(key_col_name)
                val_idx = headers.index(value_col_name)
            except ValueError:
                return default
            fname = os.path.basename(filepath)
            results = []
            for row in rows:
                cell_key = str(row[key_idx]) if row[key_idx] is not None else ""
                if match_doc_name(fname, cell_key):
                    val = row[val_idx]
                    if val is not None and str(val) != "_ERROR":
                        results.append(str(val))
            return results if results else default

        if source.startswith("excel_lookup_dict:") and config is not None:
            # Format: excel_lookup_dict:<excel_path>:<sheet>:<key_col>:<acro_col>:<def_col>
            # Reads all rows in <sheet> where <key_col> matches the current
            # doc filename. Returns a list of "ACRONYM = Definition" strings.
            # Only rows with a non-empty definition column are included.
            params = source[len("excel_lookup_dict:"):].split(":", 4)
            if len(params) != 5:
                return default
            xl_path_raw, sheet_name, key_col_name, acro_col_name, def_col_name = params
            xl_path = resolve_path(config, xl_path_raw)
            if not os.path.isfile(xl_path):
                return default
            import openpyxl
            wb = openpyxl.load_workbook(xl_path, read_only=True, data_only=True)
            if sheet_name not in wb.sheetnames:
                wb.close()
                return default
            ws = wb[sheet_name]
            rows = ws.iter_rows(values_only=True)
            headers = list(next(rows, []))
            try:
                key_idx = headers.index(key_col_name)
                acro_idx = headers.index(acro_col_name)
                def_idx = headers.index(def_col_name)
            except ValueError:
                wb.close()
                return default
            fname = os.path.basename(filepath)
            results = []
            for row in rows:
                cell_key = str(row[key_idx]) if row[key_idx] is not None else ""
                if match_doc_name(fname, cell_key):
                    acro = row[acro_idx]
                    defn = row[def_idx]
                    if acro is not None and defn is not None:
                        acro_str = str(acro).strip()
                        defn_str = str(defn).strip()
                        if acro_str and defn_str:
                            results.append(f"{acro_str} = {defn_str}")
            wb.close()
            return results if results else default

    except Exception:
        pass
    return default


def _default_metadata_fields() -> list[dict]:
    """Fallback metadata fields when none configured."""
    return [
        {"name": "title", "source": "core:title", "default": ""},
        {"name": "source_file", "source": "filename"},
        {"name": "author", "source": "core:author", "default": "Unknown"},
        {"name": "created", "source": "core:created", "default": ""},
        {"name": "modified", "source": "core:modified", "default": ""},
        {"name": "doc_id", "source": r"filename_regex:([A-Z]+-[A-Z]+-\d{4}-\d+)", "default": ""},
        {"name": "converted", "source": "converted_date"},
    ]


# ============================================================================
# DocxToMarkdown converter class
# ============================================================================

class DocxToMarkdown:
    """Converts a single .docx file to markdown."""

    def __init__(self, filepath: str, config: dict, output_dir: str,
                 url_mapping: dict | None = None):
        self.filepath = filepath
        self.filename = os.path.basename(filepath)
        self.config = config
        self.d2m = config.get("docx2md", {})
        self.output_dir = output_dir
        self.url_mapping = url_mapping or {}

        # Heading detection config
        self.custom_style_map = build_custom_style_map(config)

        # Control ID patterns — loaded from control_extraction config and reused
        # for both heading promotion and the max_controls_per_file split trigger.
        ctrl_cfg = config.get("control_extraction", {})
        ctrl_patterns_raw = ctrl_cfg.get("control_id_patterns", [])
        self.control_id_patterns: list[re.Pattern] = [re.compile(p) for p in ctrl_patterns_raw]
        self.promote_control_ids = self.d2m.get("promote_control_ids_to_heading", False)
        self.require_bold_control_id = ctrl_cfg.get("require_bold_control_id", False)
        # Combined regex for counting controls when splitting (mirrors extract_controls.py)
        if ctrl_patterns_raw:
            combined = "|".join(f"({p})" for p in ctrl_patterns_raw)
            self._combined_control_re: "re.Pattern | None" = re.compile(combined)
        else:
            self._combined_control_re = None

        # Scope-clone + heading-delete filters (applied pre-split)
        self.scope_statements_keep = bool(self.d2m.get("scope_statements_keep", False))
        self.scope_statement_headings = {
            str(h).strip().lower()
            for h in (self.d2m.get("scope_statement_headings") or [])
            if str(h).strip()
        }
        self.headings_to_delete = {
            str(h).strip().lower()
            for h in (self.d2m.get("headings_to_delete") or [])
            if str(h).strip()
        }
        self._scope_clone_lines: list[str] = []

        # State
        self.lines: list[str] = []
        self.warnings: list[dict] = []
        self.inventory: list[dict] = []
        self.image_counter = 0
        self.heading_levels_used: list[int] = []
        self.stats = {
            "paragraphs": 0,
            "tables": 0,
            "images": 0,
            "links": 0,
            "headings": 0,
            "control_headings": 0,
            "has_merged_cells": False,
        }
        self._in_list = False
        self._prev_was_blank = False

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def convert(self) -> dict:
        """Convert the document and write output files. Returns a result dict."""
        start = time.time()

        try:
            doc = Document(self.filepath)
        except Exception as exc:
            msg = f"Cannot open document: {exc}"
            self._warn("file", msg, "")
            return self._result("error", time.time() - start, msg)

        # Build frontmatter
        frontmatter = _build_frontmatter(doc, self.filepath, self.config,
                                         url_mapping=self.url_mapping)
        if frontmatter:
            self.lines.append(frontmatter)

        # Build element maps for interleaved iteration
        para_map = {id(p._element): p for p in doc.paragraphs}
        table_map = {id(t._element): t for t in doc.tables}

        # Walk body children in document order
        for child in doc.element.body:
            child_id = id(child)
            try:
                if child_id in para_map:
                    self._process_paragraph(para_map[child_id], doc)
                elif child_id in table_map:
                    self._end_list()
                    self._process_table(table_map[child_id])
            except Exception as exc:
                has_img = False
                try:
                    xml_str = child.xml
                    has_img = "a:blip" in xml_str or "v:imagedata" in xml_str
                except Exception:
                    pass
                if has_img:
                    self.image_counter += 1
                    self.stats["images"] += 1
                    self._warn("image", f"Image paragraph failed: {exc}", "")
                    self.lines.append(f"--image {self.image_counter} placeholder--")
                else:
                    self._warn("element", f"Processing error: {exc}", "")
                    self.lines.append(f"[Conversion error: {exc}]")
                self.lines.append("")

        # Extract text boxes
        if self.d2m.get("extract_text_boxes", True):
            self._extract_text_boxes(doc)

        # Post-process
        self._post_process()

        # Append bottom metadata block if placement is "top_and_bottom"
        if self.d2m.get("metadata_placement", "top") == "top_and_bottom":
            self.lines.extend(self._build_bottom_metadata(doc))

        # Write .md file(s) — optionally split at section boundaries
        stem = os.path.splitext(self.filename)[0]
        safe_name = sanitize_filename(stem, max_len=100)
        ensure_output_dir(self.output_dir)

        # Apply heading-delete removals and collect scope-clone sections
        self._apply_heading_filters()

        output_files = self._write_markdown(safe_name)

        # Extract images
        if self.d2m.get("image_handling", "extract") == "extract":
            self._extract_images(doc, safe_name)

        elapsed = time.time() - start
        status = "warning" if self.warnings else "success"
        return self._result(status, elapsed, output_file=output_files)

    # ------------------------------------------------------------------
    # Bottom metadata block
    # ------------------------------------------------------------------

    def _build_bottom_metadata(self, doc) -> list[str]:
        """Build a readable metadata block for the bottom of the document."""
        d2m = self.d2m
        fields = list(d2m.get("metadata_fields", _default_metadata_fields()))

        lines = ["", "---", ""]
        for field in fields:
            name = field.get("name", "")
            source = field.get("source", "")
            default = field.get("default", "")
            value = _resolve_metadata_value(doc, self.filepath, source, default,
                                            url_mapping=self.url_mapping,
                                            config=self.config)
            if isinstance(value, list):
                value = ", ".join(value)
            lines.append(f"**{name}:** {value}")
        lines.append("")
        return lines

    # ------------------------------------------------------------------
    # Paragraph processing
    # ------------------------------------------------------------------

    def _process_paragraph(self, para, doc):
        """Process a single paragraph element."""
        self.stats["paragraphs"] += 1
        text = para.text.strip()

        # Check for page break
        if self._has_page_break(para):
            pass  # skip page break markers

        # Empty paragraph
        if not text and not self._has_image(para):
            self._end_list()
            if not self._prev_was_blank:
                self.lines.append("")
                self._prev_was_blank = True
            return

        self._prev_was_blank = False

        # Check for image in paragraph
        if self._has_image(para):
            self._end_list()
            self._handle_image_paragraph(para, doc)

        # Heading detection
        heading_level = self._detect_heading_level(para)
        if heading_level:
            self._end_list()
            prefix = "#" * heading_level + " "
            heading_text = self._render_inline(para, doc)
            # Strip redundant bold/italic wrapping from headings
            heading_text = re.sub(r"^\*{1,3}(.+?)\*{1,3}$", r"\1", heading_text.strip())
            self.lines.append("")
            self.lines.append(prefix + heading_text)
            self.lines.append("")
            self.stats["headings"] += 1
            self.heading_levels_used.append(heading_level)
            self._add_inventory("heading", f"H{heading_level}: {text[:80]}")
            return

        # List detection
        list_info = self._detect_list(para)
        if list_info:
            list_type, indent_level = list_info
            indent = "  " * indent_level if list_type == "bullet" else "   " * indent_level
            marker = "- " if list_type == "bullet" else "1. "
            self.lines.append(indent + marker + self._render_inline(para, doc))
            self._in_list = True
            return

        # Regular paragraph
        self._end_list()
        rendered = self._render_inline(para, doc)
        if rendered:
            self.lines.append(rendered)
            self.lines.append("")

    def _detect_heading_level(self, para) -> Optional[int]:
        """Return heading level (1-6) or None."""
        style = para.style
        style_name = style.name if style else ""

        # Built-in heading style
        level = get_heading_level(style)
        if level:
            return level

        # Custom style map
        if style_name:
            mapped = self.custom_style_map.get(style_name.strip().lower())
            if mapped:
                m = re.match(r"Heading (\d)", mapped)
                if m:
                    return int(m.group(1))

        # Control ID heading promotion
        if self.promote_control_ids and self.control_id_patterns:
            text = para.text.strip()
            if self.require_bold_control_id and not is_paragraph_bold(para):
                pass  # skip non-bold paragraphs when bold is required
            else:
                for pat in self.control_id_patterns:
                    if pat.search(text):
                        self.stats["control_headings"] += 1
                        return 2  # Always H2

        return None

    def _detect_list(self, para) -> Optional[tuple]:
        """Return (list_type, indent_level) or None."""
        style_name = (para.style.name or "").lower() if para.style else ""

        # Check XML for numPr (numbering properties)
        num_pr = para._element.find(qn("w:pPr"))
        if num_pr is not None:
            num_pr = num_pr.find(qn("w:numPr"))

        if num_pr is not None:
            # Get indent level
            ilvl = num_pr.find(qn("w:ilvl"))
            indent = int(ilvl.get(qn("w:val"))) if ilvl is not None else 0

            # Determine bullet vs numbered from style name
            if "bullet" in style_name:
                return ("bullet", indent)
            elif "number" in style_name or "list number" in style_name:
                return ("numbered", indent)
            # Default to bullet if we can't determine type
            return ("bullet", indent)

        # Fallback: check style name patterns
        if "list bullet" in style_name:
            # List Bullet 2 -> indent 1, List Bullet 3 -> indent 2
            m = re.search(r"list bullet\s*(\d)?", style_name)
            indent = int(m.group(1)) - 1 if m and m.group(1) else 0
            return ("bullet", max(0, indent))
        if "list number" in style_name:
            m = re.search(r"list number\s*(\d)?", style_name)
            indent = int(m.group(1)) - 1 if m and m.group(1) else 0
            return ("numbered", max(0, indent))

        return None

    def _end_list(self):
        """Insert blank line after list block ends."""
        if self._in_list:
            self.lines.append("")
            self._in_list = False

    # ------------------------------------------------------------------
    # Inline rendering (runs + hyperlinks interleaved)
    # ------------------------------------------------------------------

    def _render_inline(self, para, doc) -> str:
        """Walk paragraph XML children to render runs and hyperlinks in order."""
        # Italics are intentionally dropped: DOCX often italicises incidental
        # text, which adds noise to RAG. Bold and strike are preserved.
        # Segments are merged across runs with matching (bold, strike) to avoid
        # mid-word `**` artifacts from Word splitting formatted words.
        segments: list = []  # each item: (text, bold, strike) or ("hyperlink", str)

        for child in para._element:
            tag = child.tag

            if tag == qn("w:r"):
                text_parts = []
                for t in child.findall(qn("w:t")):
                    if t.text:
                        text_parts.append(t.text)
                text = "".join(text_parts)
                if not text:
                    continue
                rpr = child.find(qn("w:rPr"))
                bold = strike = False
                if rpr is not None:
                    bold = rpr.find(qn("w:b")) is not None
                    strike = rpr.find(qn("w:strike")) is not None
                segments.append((text, bold, strike))

            elif tag == qn("w:hyperlink"):
                segments.append(("hyperlink", self._render_hyperlink(child, doc)))

        # Merge consecutive runs with identical formatting
        parts: list[str] = []
        i = 0
        while i < len(segments):
            seg = segments[i]
            if seg[0] == "hyperlink":
                parts.append(seg[1])
                i += 1
                continue
            text, bold, strike = seg
            j = i + 1
            while j < len(segments) and segments[j][0] != "hyperlink":
                _, b2, st2 = segments[j]
                if b2 == bold and st2 == strike:
                    text += segments[j][0]
                    j += 1
                else:
                    break
            i = j
            # Lift leading/trailing whitespace outside emphasis so we don't
            # emit invalid markdown like `**Description: **`.
            lead_len = len(text) - len(text.lstrip())
            trail_len = len(text) - len(text.rstrip())
            lead = text[:lead_len]
            trail = text[len(text) - trail_len:] if trail_len else ""
            core = text.strip()
            if core:
                if strike:
                    core = f"~~{core}~~"
                if bold:
                    core = f"**{core}**"
            parts.append(lead + core + trail)

        text = "".join(parts)
        return _clean_text(text, self.d2m)

    def _render_run(self, run_elem) -> str:
        """Render a single w:r element with inline formatting (used externally)."""
        text_parts = []
        for t in run_elem.findall(qn("w:t")):
            if t.text:
                text_parts.append(t.text)

        text = "".join(text_parts)
        if not text:
            return ""

        rpr = run_elem.find(qn("w:rPr"))
        bold = strike = False
        if rpr is not None:
            bold = rpr.find(qn("w:b")) is not None
            strike = rpr.find(qn("w:strike")) is not None

        lead_len = len(text) - len(text.lstrip())
        trail_len = len(text) - len(text.rstrip())
        lead = text[:lead_len]
        trail = text[len(text) - trail_len:] if trail_len else ""
        core = text.strip()
        if core:
            if strike:
                core = f"~~{core}~~"
            if bold:
                core = f"**{core}**"
        return lead + core + trail

    def _render_hyperlink(self, hyperlink_elem, doc) -> str:
        """Render a w:hyperlink element as a markdown link."""
        # Extract display text
        display_parts = []
        for run_elem in hyperlink_elem.findall(qn("w:r")):
            for t in run_elem.findall(qn("w:t")):
                if t.text:
                    display_parts.append(t.text)
        display_text = "".join(display_parts).strip()
        if not display_text:
            return ""

        self.stats["links"] += 1

        # External URL via r:id
        r_id = hyperlink_elem.get(qn("r:id"))
        if r_id:
            try:
                target = doc.part.rels[r_id].target_ref
                self._add_inventory("link", f"[{display_text[:50]}]({target[:80]})")
                return f"[{display_text}]({target})"
            except (KeyError, AttributeError):
                pass

        # Internal bookmark via w:anchor
        anchor = hyperlink_elem.get(qn("w:anchor"))
        if anchor:
            self._add_inventory("link", f"[{display_text[:50]}](#{anchor})")
            return f"[{display_text}](#{anchor})"

        # Fallback: just return the text
        return display_text

    # ------------------------------------------------------------------
    # Table processing
    # ------------------------------------------------------------------

    def _process_table(self, table):
        """Convert a table to markdown (pipe table or HTML)."""
        self.stats["tables"] += 1
        strategy = self.d2m.get("table_strategy", "auto")

        has_merged = self._detect_merged_cells(table)
        if has_merged:
            self.stats["has_merged_cells"] = True

        # Determine strategy
        use_html = (strategy == "html") or (strategy == "auto" and has_merged)

        if use_html:
            self._table_to_html(table)
        else:
            self._table_to_pipe(table)

        header_text = ""
        if table.rows:
            header_text = " | ".join(c.text.strip()[:40] for c in table.rows[0].cells)
        self._add_inventory("table", f"{'[merged] ' if has_merged else ''}{header_text[:100]}")

    def _detect_merged_cells(self, table) -> bool:
        """Check if any row has merged cells."""
        try:
            for row in table.rows:
                seen = set()
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid in seen:
                        return True
                    seen.add(cid)
        except Exception:
            pass
        return False

    def _table_to_pipe(self, table):
        """Render table as markdown pipe table."""
        if not table.rows:
            return

        rows_data = []
        for row in table.rows:
            # Deduplicate merged cells in a row
            cells = []
            seen = set()
            for cell in row.cells:
                cid = id(cell._tc)
                if cid not in seen:
                    seen.add(cid)
                    cell_text = cell.text.strip().replace("\n", "<br>").replace("|", "\\|")
                    cells.append(cell_text)
            rows_data.append(cells)

        if not rows_data:
            return

        # Normalize column count
        max_cols = max(len(r) for r in rows_data)
        for r in rows_data:
            while len(r) < max_cols:
                r.append("")

        self.lines.append("")

        # Header row
        self.lines.append("| " + " | ".join(rows_data[0]) + " |")
        self.lines.append("| " + " | ".join("---" for _ in range(max_cols)) + " |")

        # Data rows
        for row_cells in rows_data[1:]:
            self.lines.append("| " + " | ".join(row_cells) + " |")

        self.lines.append("")

    def _table_to_html(self, table):
        """Render table as HTML for complex/merged-cell tables."""
        if not table.rows:
            return

        self.lines.append("")
        self.lines.append("<table>")

        for row_idx, row in enumerate(table.rows):
            self.lines.append("  <tr>")
            seen = set()
            for cell in row.cells:
                cid = id(cell._tc)
                if cid in seen:
                    continue
                seen.add(cid)

                # Detect colspan
                tc = cell._tc
                grid_span = tc.find(qn("w:tcPr"))
                colspan = 1
                if grid_span is not None:
                    gs = grid_span.find(qn("w:gridSpan"))
                    if gs is not None:
                        colspan = int(gs.get(qn("w:val"), 1))

                cell_text = cell.text.strip().replace("\n", "<br>")
                tag = "th" if row_idx == 0 else "td"
                span_attr = f' colspan="{colspan}"' if colspan > 1 else ""
                self.lines.append(f"    <{tag}{span_attr}>{cell_text}</{tag}>")

            self.lines.append("  </tr>")

        self.lines.append("</table>")
        self.lines.append("")

    # ------------------------------------------------------------------
    # Image handling
    # ------------------------------------------------------------------

    def _has_image(self, para) -> bool:
        """Check if paragraph contains an image element."""
        xml_str = para._element.xml
        return "a:blip" in xml_str or "v:imagedata" in xml_str

    def _has_page_break(self, para) -> bool:
        """Check for page break in paragraph."""
        for br in para._element.findall(".//" + qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
        return False

    def _handle_image_paragraph(self, para, doc):
        """Insert image reference or placeholder for images in a paragraph."""
        handling = self.d2m.get("image_handling", "extract")

        if handling == "skip":
            return

        self.image_counter += 1
        self.stats["images"] += 1

        # Try to get alt text from drawing properties
        alt_text = ""
        for doc_pr in para._element.findall(".//" + qn("wp:docPr")):
            descr = doc_pr.get("descr", "")
            name = doc_pr.get("name", "")
            alt_text = descr or name
            break

        if not alt_text:
            alt_text = f"Image {self.image_counter}"

        # Check for SmartArt
        if "dgm:relIds" in para._element.xml:
            self.lines.append(f"[Diagram: SmartArt - not fully extractable as image]")
            self.lines.append("")
            self._add_inventory("image", f"SmartArt diagram #{self.image_counter}")
            return

        if handling == "extract":
            # Find the relationship ID for the image
            r_id = None
            for blip in para._element.findall(".//" + qn("a:blip")):
                r_id = blip.get(qn("r:embed"))
                break
            if r_id is None:
                for img_data in para._element.findall(".//" + qn("v:imagedata")):
                    r_id = img_data.get(qn("r:id"))
                    break

            if r_id:
                try:
                    rel = doc.part.rels[r_id]
                    blob = rel.target_part.blob
                    content_type = rel.target_part.content_type
                    ext = _content_type_to_ext(content_type)
                    stem = os.path.splitext(self.filename)[0]
                    safe_stem = sanitize_filename(stem, max_len=100)
                    img_dir_name = f"{safe_stem}_images"
                    img_dir = os.path.join(self.output_dir, img_dir_name)
                    ensure_output_dir(img_dir)
                    img_filename = f"image_{self.image_counter}{ext}"
                    img_path = os.path.join(img_dir, img_filename)
                    with open(img_path, "wb") as f:
                        f.write(blob)
                    rel_path = f"{img_dir_name}/{img_filename}"
                    self.lines.append(f"![{alt_text}]({rel_path})")
                    self.lines.append("")
                    self._add_inventory("image", f"Extracted: {rel_path}")
                    return
                except Exception as exc:
                    self._warn("image", f"Failed to extract image #{self.image_counter}: {exc}",
                               f"paragraph")

        # Placeholder fallback
        self.lines.append(f"--image {self.image_counter} placeholder--")
        self.lines.append("")
        self._add_inventory("image", f"Placeholder: {alt_text}")

    def _extract_images(self, doc, safe_name: str):
        """Extract all images from document relationships (batch extraction fallback)."""
        # Images are already extracted inline during paragraph processing.
        # This method is kept for potential future batch operations.
        pass

    # ------------------------------------------------------------------
    # Text box extraction
    # ------------------------------------------------------------------

    def _extract_text_boxes(self, doc):
        """Extract text from text boxes and append to output."""
        try:
            ns_map = {
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
            }
            # Find text box content elements
            txbx_elements = doc.element.body.findall(".//" + qn("w:txbxContent"))
            wps_elements = doc.element.body.findall(
                ".//{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx"
            )

            all_txbx = txbx_elements + wps_elements
            if not all_txbx:
                return

            self._add_inventory("text_box", f"Found {len(all_txbx)} text box(es)")

            for txbx in all_txbx:
                # Get paragraphs within the text box
                paras = txbx.findall(qn("w:p"))
                if not paras:
                    continue

                self.lines.append("")
                self.lines.append("<!-- Text Box Content -->")
                for p_elem in paras:
                    text_parts = []
                    for t in p_elem.findall(".//" + qn("w:t")):
                        if t.text:
                            text_parts.append(t.text)
                    text = "".join(text_parts).strip()
                    if text:
                        text = _clean_text(text, self.d2m)
                        self.lines.append(text)
                self.lines.append("")

        except Exception as exc:
            self._warn("text_box", f"Error extracting text boxes: {exc}", "")

    # ------------------------------------------------------------------
    # Output writing & splitting
    # ------------------------------------------------------------------

    def _apply_heading_filters(self) -> None:
        """Remove sections whose heading matches headings_to_delete, and collect
        scope-statement sections for cloning into split chunks.

        A section is a heading line plus every following line up to (but not
        including) the next heading of any level. Matching is case-insensitive
        and exact on the trimmed heading text.
        """
        want_delete = bool(self.headings_to_delete)
        want_clone = self.scope_statements_keep and bool(self.scope_statement_headings)
        if not (want_delete or want_clone):
            return

        has_fm = bool(self.lines) and self.lines[0].lstrip().startswith("---")
        fm_block = self.lines[0] if has_fm else None
        body = self.lines[1:] if has_fm else list(self.lines)

        out: list[str] = []
        scope_sections: list[list[str]] = []
        i = 0
        n = len(body)
        while i < n:
            line = body[i]
            m = _HEADING_RE.match(line)
            if not m:
                out.append(line)
                i += 1
                continue
            heading_text = m.group(2).strip().lower()
            j = i + 1
            while j < n and not _HEADING_RE.match(body[j]):
                j += 1
            section = body[i:j]
            if want_delete and heading_text in self.headings_to_delete:
                i = j
                continue
            if want_clone and heading_text in self.scope_statement_headings:
                scope_sections.append(section)
            out.extend(section)
            i = j

        self.lines = ([fm_block] + out) if has_fm else out

        if scope_sections:
            clone: list[str] = []
            for sec in scope_sections:
                clone.extend(sec)
                if clone and clone[-1] != "":
                    clone.append("")
            self._scope_clone_lines = clone

    def _write_markdown(self, safe_name: str) -> list[str]:
        """Write self.lines to one or more .md files. Returns list of filenames.

        If docx2md.character_limit is set and total content exceeds it, splits
        at heading boundaries (never mid-section). Each chunk keeps the
        frontmatter, with the last section header appended to title and filename.
        """
        try:
            char_limit = int(self.d2m.get("character_limit", 0) or 0)
        except (TypeError, ValueError):
            char_limit = 0
        try:
            control_limit = int(self.d2m.get("max_controls_per_file", 0) or 0)
        except (TypeError, ValueError):
            control_limit = 0

        control_regex = self._combined_control_re if control_limit > 0 else None

        total_chars = sum(len(s) + 1 for s in self.lines)
        total_controls = (
            len(control_regex.findall("\n".join(self.lines))) if control_regex else 0
        )
        char_split_needed = char_limit > 0 and total_chars > char_limit
        ctrl_split_needed = control_limit > 0 and total_controls > control_limit

        if not (char_split_needed or ctrl_split_needed):
            md_path = os.path.join(self.output_dir, f"{safe_name}.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("\n".join(self.lines))
            return [f"{safe_name}.md"]

        # Separate frontmatter (single multi-line element at index 0) from body
        frontmatter = ""
        body = list(self.lines)
        if body and body[0].lstrip().startswith("---"):
            frontmatter = body[0]
            body = body[1:]

        base_chars = len(frontmatter) + 1 if frontmatter else 0
        chunks = _split_body_at_sections(
            body, char_limit, base_chars,
            control_limit=control_limit, control_regex=control_regex,
        )

        # Fall back to single file if splitting didn't find any boundary
        if len(chunks) <= 1:
            md_path = os.path.join(self.output_dir, f"{safe_name}.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("\n".join(self.lines))
            return [f"{safe_name}.md"]

        written: list[str] = []
        for idx, chunk in enumerate(chunks, 1):
            suffix = chunk["last_header"] or f"Part {idx}"
            safe_suffix = sanitize_filename(suffix, max_len=80) or f"Part {idx}"
            chunk_name = f"{safe_name} - {safe_suffix}"
            if len(chunk_name) > 180:
                chunk_name = chunk_name[:180].rstrip()

            # Avoid overwrite if two chunks sanitize to the same header
            unique_name = chunk_name
            dup = 2
            while f"{unique_name}.md" in written or os.path.exists(
                    os.path.join(self.output_dir, f"{unique_name}.md")):
                unique_name = f"{chunk_name} ({dup})"
                dup += 1

            patched_fm = self._patch_frontmatter_title(frontmatter, safe_suffix) \
                if frontmatter else ""

            parts: list[str] = []
            if patched_fm:
                parts.append(patched_fm)
            if idx > 1 and self._scope_clone_lines:
                parts.extend(self._scope_clone_lines)
            parts.extend(chunk["lines"])

            md_path = os.path.join(self.output_dir, f"{unique_name}.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("\n".join(parts))
            written.append(f"{unique_name}.md")

        return written

    @staticmethod
    def _patch_frontmatter_title(frontmatter: str, suffix: str) -> str:
        """Append ' - <suffix>' to the title: line in a YAML frontmatter block."""
        if not frontmatter or not suffix:
            return frontmatter
        out = []
        patched = False
        for line in frontmatter.splitlines():
            if not patched:
                m = re.match(r'^(title:\s*)"(.*)"(\s*)$', line)
                if m:
                    line = f'{m.group(1)}"{m.group(2)} - {suffix}"{m.group(3)}'
                    patched = True
            out.append(line)
        return "\n".join(out)

    # ------------------------------------------------------------------
    # Post-processing
    # ------------------------------------------------------------------

    def _post_process(self):
        """Clean up the accumulated lines."""
        max_blank = self.d2m.get("max_consecutive_blank_lines", 2)

        # Heading level normalization (shift so minimum = H1)
        if self.d2m.get("heading_normalization", True) and self.heading_levels_used:
            min_level = min(self.heading_levels_used)
            if min_level > 1:
                shift = min_level - 1
                new_lines = []
                for line in self.lines:
                    m = re.match(r"^(#{1,6})\s", line)
                    if m:
                        current_level = len(m.group(1))
                        new_level = max(1, current_level - shift)
                        line = "#" * new_level + line[current_level:]
                    new_lines.append(line)
                self.lines = new_lines

        # Heading cap — collapse anything deeper than max_heading_level
        max_level = self.d2m.get("max_heading_level", 6)
        if max_level < 6:
            new_lines = []
            for line in self.lines:
                m = re.match(r"^(#{1,6})\s", line)
                if m:
                    current_level = len(m.group(1))
                    if current_level > max_level:
                        line = "#" * max_level + line[current_level:]
                new_lines.append(line)
            self.lines = new_lines

        # Collapse blank lines and strip trailing whitespace
        cleaned = []
        blank_count = 0
        for line in self.lines:
            stripped = line.rstrip()
            if not stripped:
                blank_count += 1
                if blank_count <= max_blank:
                    cleaned.append("")
            else:
                blank_count = 0
                cleaned.append(stripped)

        # Remove trailing blank lines
        while cleaned and not cleaned[-1]:
            cleaned.pop()
        cleaned.append("")  # single trailing newline

        self.lines = cleaned

    # ------------------------------------------------------------------
    # Logging helpers
    # ------------------------------------------------------------------

    def _warn(self, element_type: str, message: str, location: str):
        """Record a warning."""
        self.warnings.append({
            "filename": self.filename,
            "severity": "warning",
            "element_type": element_type,
            "message": message,
            "location": location,
        })

    def _add_inventory(self, element_type: str, details: str):
        """Record an element in the inventory."""
        idx = len(self.inventory) + 1
        self.inventory.append({
            "filename": self.filename,
            "element_type": element_type,
            "element_index": idx,
            "details": details[:200],
        })

    def _result(self, status: str, elapsed: float, error_msg: str = "",
                output_file=""):
        """Build the result dict for logging.

        output_file may be a string (single file) or a list (split output).
        """
        if isinstance(output_file, list):
            files = [f for f in output_file if f]
            display = "; ".join(files)
        else:
            files = [output_file] if output_file else []
            display = output_file

        size_bytes = 0
        for f in files:
            p = os.path.join(self.output_dir, f)
            if os.path.exists(p):
                size_bytes += os.path.getsize(p)
        file_size = round(size_bytes / 1024, 1) if size_bytes else 0

        return {
            "filename": self.filename,
            "status": status,
            "output_file": display,
            "error_msg": error_msg,
            "paragraphs": self.stats["paragraphs"],
            "tables": self.stats["tables"],
            "images": self.stats["images"],
            "links": self.stats["links"],
            "headings": self.stats["headings"],
            "control_headings": self.stats["control_headings"],
            "has_merged_cells": self.stats["has_merged_cells"],
            "file_size_kb": file_size,
            "conversion_time_s": round(elapsed, 2),
            "warnings": self.warnings,
            "inventory": self.inventory,
        }


# ============================================================================
# Image utility
# ============================================================================

def _content_type_to_ext(content_type: str) -> str:
    """Map MIME content type to file extension."""
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/tiff": ".tiff",
        "image/x-emf": ".emf",
        "image/x-wmf": ".wmf",
        "image/svg+xml": ".svg",
    }
    return mapping.get(content_type, ".png")


# ============================================================================
# Excel log writer
# ============================================================================

def write_excel_log(results: list[dict], output_dir: str, prefix: str):
    """Write the three-sheet Excel log file."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        print("  WARNING: openpyxl not installed — skipping Excel log.")
        return

    wb = Workbook()
    wb.remove(wb.active)

    # Styling
    header_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )

    def _style_sheet(ws):
        """Apply standard styling to a worksheet."""
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border
        ws.freeze_panes = "A2"
        if ws.max_column and ws.max_row:
            last_col = get_column_letter(ws.max_column)
            ws.auto_filter.ref = f"A1:{last_col}{ws.max_row}"
        for col_idx in range(1, ws.max_column + 1):
            max_width = 0
            for row in ws.iter_rows(min_col=col_idx, max_col=col_idx, values_only=True):
                val = row[0]
                if val is not None:
                    max_width = max(max_width, len(str(val)))
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max_width + 2, 60)

    # Sheet 1: Conversion Summary
    ws1 = wb.create_sheet("Conversion Summary")
    ws1.append([
        "Filename", "Status", "Output File", "Paragraphs", "Tables",
        "Images", "Links", "Headings", "Control Headings", "Merged Cells",
        "File Size (KB)", "Time (s)", "Error"
    ])
    for r in results:
        ws1.append([
            r["filename"], r["status"], r["output_file"],
            r["paragraphs"], r["tables"], r["images"], r["links"],
            r["headings"], r["control_headings"],
            "Yes" if r["has_merged_cells"] else "No",
            r["file_size_kb"], r["conversion_time_s"],
            r.get("error_msg", ""),
        ])
    _style_sheet(ws1)

    # Sheet 2: Errors/Warnings
    ws2 = wb.create_sheet("Errors & Warnings")
    ws2.append(["Filename", "Severity", "Element Type", "Message", "Location"])
    for r in results:
        for w in r.get("warnings", []):
            ws2.append([
                w["filename"], w["severity"], w["element_type"],
                w["message"], w.get("location", ""),
            ])
    _style_sheet(ws2)

    # Sheet 3: Element Inventory
    ws3 = wb.create_sheet("Element Inventory")
    ws3.append(["Filename", "Element Type", "Index", "Details"])
    for r in results:
        for inv in r.get("inventory", []):
            ws3.append([
                inv["filename"], inv["element_type"],
                inv["element_index"], inv["details"],
            ])
    _style_sheet(ws3)

    # Save
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    filename = f"{prefix}_{timestamp}.xlsx"
    save_path = os.path.join(output_dir, filename)
    try:
        wb.save(save_path)
        print(f"  Log saved: {save_path}")
    except PermissionError:
        print(f"  WARNING: Could not save {filename} — file may be open in another app.")


# ============================================================================
# Main
# ============================================================================

def main():
    parser = setup_argparse("DPS_docx2md — Convert .docx files to RAG-optimised markdown")
    args = parser.parse_args()

    config = load_config(args.config)
    d2m = config.get("docx2md", {})

    # Load Doc URL mapping if enabled via legacy flag OR any metadata field uses source "doc_url"
    url_mapping = {}
    _fields_use_doc_url = any(
        f.get("source") == "doc_url"
        for f in d2m.get("metadata_fields", [])
    )
    if d2m.get("include_doc_url", False) or _fields_use_doc_url:
        url_mapping = load_url_mapping(config)

    # Resolve input directory — supports Pure Conversion vs Optimized toggle
    input_dir = args.input_dir
    if not input_dir:
        if d2m.get("pure_conversion", False):
            # Pure Conversion: read from original input/
            input_dir = config.get("input", {}).get("directory", "./input")
        else:
            # Optimized: read from a pipeline step's output
            source_step = d2m.get("optimized_source_step", "heading_fixes")
            step_dir = config.get("output", {}).get(source_step, {}).get("directory", "")
            if step_dir:
                output_root = config.get("output", {}).get("directory", "./output")
                input_dir = os.path.join(output_root, step_dir)
            else:
                input_dir = config.get("input", {}).get("directory", "./input")
        input_dir = resolve_path(config, input_dir)

    # Resolve output directory
    output_dir = args.output_dir
    if not output_dir:
        output_dir = d2m.get("output_directory", "./output/7 - markdown")
        output_dir = resolve_path(config, output_dir)

    ensure_output_dir(output_dir)

    print("=" * 70)
    print("  DPS_docx2md — DOCX to Markdown Converter")
    print("=" * 70)
    print(f"  Input:  {input_dir}")
    print(f"  Output: {output_dir}")
    print()

    # When reading from step output, only exclude temp files (not _fixed, _optimized etc.)
    exclude_ovr = ["~$"] if not d2m.get("pure_conversion", False) else None
    files = iter_docx_files(input_dir, config, exclude_override=exclude_ovr)
    if not files:
        print("  No .docx files found in input directory.")
        print("  Check your input path and exclude patterns in dps_config.xlsx.")
        return

    print(f"  Found {len(files)} document(s) to convert.")
    print("-" * 70)

    results = []
    for i, filepath in enumerate(files, 1):
        fname = os.path.basename(filepath)
        print(f"  [{i}/{len(files)}] {fname} ... ", end="", flush=True)

        converter = DocxToMarkdown(filepath, config, output_dir,
                                   url_mapping=url_mapping)
        result = converter.convert()
        results.append(result)

        status_icon = "OK" if result["status"] == "success" else result["status"].upper()
        print(f"{status_icon} ({result['conversion_time_s']}s)")

        if result["warnings"]:
            for w in result["warnings"]:
                print(f"         WARNING: {w['message']}")

    print("-" * 70)

    # Summary
    success = sum(1 for r in results if r["status"] == "success")
    warnings = sum(1 for r in results if r["status"] == "warning")
    errors = sum(1 for r in results if r["status"] == "error")
    print(f"  Done: {success} success, {warnings} with warnings, {errors} errors")

    # Write Excel log
    log_prefix = d2m.get("log_file_prefix", "docx2md_log")
    write_excel_log(results, output_dir, log_prefix)

    print("=" * 70)


if __name__ == "__main__":
    main()
