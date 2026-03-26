#!/usr/bin/env python3
"""
DPS — Step 5: Metadata Injector
=================================

Stamps each sub-document with a structured metadata block (name, URL,
scope, intent, tags) so Copilot/RAG always knows what document a chunk
came from, where to find the original, and what it covers.

WHY:
    After splitting (Step 4), sub-documents lose identity context.
    This step restores it in a RAG-friendly format — each metadata field
    is a Heading 1 followed by a Normal paragraph.  Chunkers parse this
    reliably (unlike tables, which get stripped or mangled).

INPUTS:
    Reads .docx files from output/4 - split_documents/ (or any directory
    you point it at via the CLI).

OUTPUTS:
    Writes copies of each .docx with the metadata block prepended (and/or
    appended) to output/5 - metadata/.  Also writes metadata_manifest.csv
    logging what metadata was applied to each file.

USAGE:
    # Run as part of the pipeline:
    python run_pipeline.py --step 5

    # Run standalone with config:
    python add_metadata.py --config ../dps_config.yaml

    # Run standalone with explicit directories:
    python add_metadata.py ../output/4\ -\ split_documents ../output/5\ -\ metadata

    # Run on any folder of .docx files:
    python add_metadata.py /path/to/my/docs /path/to/output

DATA SOURCES (all optional — the script degrades gracefully):
    - document_profiles.json (Step 0) — scope, intent, doc type, sections
    - split_manifest.csv (Step 4) — maps sub-docs to parent documents
    - URL lookup Excel file — maps document names to SharePoint URLs
    - Acronym audit Excel file — per-document acronym lists for tags

FAILURE POINTS:
    - If Step 0 hasn't run: scope/intent/tags use direct scan or fallback.
    - If Step 4 hasn't run: parent-doc mapping uses filename heuristic.
    - If URL Excel doesn't exist: URLs show "(URL not configured)".
    - If a .docx is corrupt: that file is skipped, others continue.

REQUIREMENTS:
    pip install python-docx pyyaml openpyxl
    Python 3.10 or later
"""

from __future__ import annotations

import csv
import json
import os
import re
import shutil
import sys
from copy import deepcopy
from typing import Optional

# ---------------------------------------------------------------------------
# Environment check — run this BEFORE importing docx/yaml/openpyxl so users
# get a clear error message with install instructions, not a traceback.
# ---------------------------------------------------------------------------


def check_environment():
    """
    Verify that required packages are installed.

    WHAT IT DOES:
        Tries to import python-docx, pyyaml, and openpyxl.  If any are
        missing, prints beginner-friendly install instructions and exits.

    FAILURE POINT:
        If you're in a virtual environment that doesn't have these packages,
        activate the right venv first, then run: pip install -r requirements.txt
    """
    missing = []
    try:
        import docx  # noqa: F401
    except ImportError:
        missing.append("python-docx")
    try:
        import yaml  # noqa: F401
    except ImportError:
        missing.append("pyyaml")
    try:
        import openpyxl  # noqa: F401
    except ImportError:
        missing.append("openpyxl")

    if missing:
        print("ERROR: Missing required packages:", ", ".join(missing))
        print()
        print("FIX: Run this command to install them:")
        print(f"  pip install {' '.join(missing)}")
        print()
        print("Or install everything at once:")
        print("  pip install -r requirements.txt")
        sys.exit(1)


check_environment()

import docx  # noqa: E402
from docx.shared import Pt, RGBColor  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

# Import shared utilities (works whether run from scripts/ or project root)
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from shared_utils import (  # noqa: E402
    load_config,
    setup_argparse,
    get_output_dir,
    iter_docx_files,
    ensure_output_dir,
    resolve_path,
)


# ===========================================================================
# Data loading — these run once at startup and cache results in memory
# ===========================================================================


def load_url_mapping(config: dict) -> dict[str, str]:
    """
    Load the Excel file that maps document names to URLs.

    WHAT IT DOES:
        Reads the Excel lookup file specified in config (metadata.url.lookup_file),
        finds the name and URL columns by header text, and returns a dict mapping
        lowercase document names to URLs.

    RETURNS:
        Dict of {lowercase_name: url}.  Empty dict if no lookup file configured.

    FAILURE POINT:
        If the column names in your Excel don't match the config settings
        (metadata.url.name_column / url_column), you'll get an empty dict
        and a warning message.  Check your column headers.
    """
    import openpyxl

    meta_cfg = config.get("metadata", {})
    url_cfg = meta_cfg.get("url", {})
    lookup_file = url_cfg.get("lookup_file", "")

    if not lookup_file:
        return {}

    lookup_path = resolve_path(config, lookup_file)
    if not os.path.isfile(lookup_path):
        print(f"  WARNING: URL lookup file not found: {lookup_path}")
        print("  URLs will use the fallback template or show '(URL not configured)'.")
        return {}

    name_col = url_cfg.get("name_column", "Document Name").lower()
    url_col = url_cfg.get("url_column", "SharePoint URL").lower()
    sheet_ref = url_cfg.get("sheet", 0)

    try:
        wb = openpyxl.load_workbook(lookup_path, read_only=True, data_only=True)
        if isinstance(sheet_ref, int):
            ws = wb.worksheets[sheet_ref]
        else:
            ws = wb[sheet_ref]
    except Exception as e:
        print(f"  WARNING: Could not read URL lookup file: {e}")
        return {}

    # Find column indices from header row
    headers = {}
    for col_idx, cell in enumerate(next(ws.iter_rows(min_row=1, max_row=1, values_only=False))):
        if cell.value:
            headers[str(cell.value).strip().lower()] = col_idx

    name_idx = headers.get(name_col)
    url_idx = headers.get(url_col)

    if name_idx is None or url_idx is None:
        print(f"  WARNING: Could not find columns '{name_col}' and/or '{url_col}' in {lookup_path}")
        print(f"  Found columns: {list(headers.keys())}")
        return {}

    mapping = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        name_val = row[name_idx] if name_idx < len(row) else None
        url_val = row[url_idx] if url_idx < len(row) else None
        if name_val and url_val:
            # Store original name (will be normalized in resolve_url for comparison)
            mapping[str(name_val).strip()] = str(url_val).strip()

    wb.close()
    print(f"  Loaded {len(mapping)} URL mappings from {os.path.basename(lookup_path)}")
    return mapping


def load_profiler_data(config: dict) -> dict:
    """
    Load document_profiles.json from Step 0 output.

    WHAT IT DOES:
        Reads the profiler JSON and returns a dict keyed by lowercase filename.
        Each value contains the full profile for that document (sections found,
        doc type, scope/intent text, etc.).

    RETURNS:
        Dict of {lowercase_filename: profile_dict}.  Empty dict if file not found.

    FAILURE POINT:
        If Step 0 hasn't been run yet, this returns an empty dict and prints
        a note.  The script will still work — it just falls back to direct
        document scanning for scope/intent.
    """
    output_root = config.get("output", {}).get("directory", "./output")
    profiler_dir = config.get("output", {}).get("profiler", {}).get("directory", "0 - profiler")
    json_file = config.get("output", {}).get("profiler", {}).get("json_file", "document_profiles.json")
    json_path = resolve_path(config, os.path.join(output_root, profiler_dir, json_file))

    if not os.path.isfile(json_path):
        print("  NOTE: Profiler output not found (Step 0 not run yet?).")
        print(f"        Expected: {json_path}")
        print("        Scope/intent/tags will use direct scan or fallback values.")
        return {}

    try:
        with open(json_path, "r", encoding="utf-8") as f:
            profiles_list = json.load(f)
    except Exception as e:
        print(f"  WARNING: Could not read profiler JSON: {e}")
        return {}

    # Index by lowercase filename for easy lookup
    profiles = {}
    if isinstance(profiles_list, list):
        for p in profiles_list:
            fname = p.get("filename", "")
            if fname:
                profiles[fname.lower()] = p
    elif isinstance(profiles_list, dict):
        # Maybe it's already a dict keyed by filename
        for k, v in profiles_list.items():
            profiles[k.lower()] = v

    print(f"  Loaded {len(profiles)} document profiles from Step 0 output.")
    return profiles


def load_split_manifest(config: dict) -> dict[str, str]:
    """
    Load split_manifest.csv from Step 4 output.

    WHAT IT DOES:
        Maps sub-document filenames to their original parent document filename.
        This is the most reliable way to know which parent a sub-doc came from.

    RETURNS:
        Dict of {lowercase_subdoc_filename: original_doc_filename}.
        Empty dict if manifest not found.
    """
    output_root = config.get("output", {}).get("directory", "./output")
    split_dir = config.get("output", {}).get("split_documents", {}).get("directory", "3 - split_documents")
    manifest_file = config.get("output", {}).get("split_documents", {}).get("manifest_file", "split_manifest.csv")
    manifest_path = resolve_path(config, os.path.join(output_root, split_dir, manifest_file))

    if not os.path.isfile(manifest_path):
        print("  NOTE: Split manifest not found (Step 4 not run yet?).")
        print("        Parent document mapping will use filename heuristic.")
        return {}

    mapping = {}
    try:
        with open(manifest_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                subdoc = row.get("sub_doc_filename", "").strip()
                original = row.get("original_file", row.get("original_doc", "")).strip()
                if subdoc and original:
                    mapping[subdoc.lower()] = original
    except Exception as e:
        print(f"  WARNING: Could not read split manifest: {e}")
        return {}

    print(f"  Loaded {len(mapping)} sub-doc → parent mappings from split manifest.")
    return mapping


def load_acronym_data(config: dict) -> dict[str, list[str]]:
    """
    Load per-document acronyms from the Acronym Finder output Excel.

    WHAT IT DOES:
        Reads the "Per Document" sheet from the acronym audit Excel file.
        Returns a dict mapping lowercase filenames to lists of acronyms
        sorted by frequency (most common first).

    RETURNS:
        Dict of {lowercase_filename: [acronym1, acronym2, ...]}.
        Empty dict if no audit file configured or not found.
    """
    import openpyxl

    meta_cfg = config.get("metadata", {})
    tags_cfg = meta_cfg.get("tags", {})
    audit_file = tags_cfg.get("acronym_audit_file", "")

    if not audit_file:
        return {}

    audit_path = resolve_path(config, audit_file)
    if not os.path.isfile(audit_path):
        print(f"  NOTE: Acronym audit file not found: {audit_path}")
        print("        Acronym tags will be skipped.")
        return {}

    try:
        wb = openpyxl.load_workbook(audit_path, read_only=True, data_only=True)
    except Exception as e:
        print(f"  WARNING: Could not read acronym audit file: {e}")
        return {}

    # Try to find the "Per Document" sheet
    sheet_name = None
    for name in wb.sheetnames:
        if "per document" in name.lower():
            sheet_name = name
            break

    if not sheet_name:
        # Fall back to "Global Summary" or first sheet
        for name in wb.sheetnames:
            if "global" in name.lower() or "summary" in name.lower():
                sheet_name = name
                break
        if not sheet_name:
            sheet_name = wb.sheetnames[0]

    ws = wb[sheet_name]

    # Find relevant columns
    headers = {}
    for col_idx, cell in enumerate(next(ws.iter_rows(min_row=1, max_row=1, values_only=False))):
        if cell.value:
            headers[str(cell.value).strip().lower()] = col_idx

    # Look for filename and acronym columns
    file_idx = None
    acro_idx = None
    count_idx = None
    for key, idx in headers.items():
        if "file" in key or "document" in key or "doc" in key:
            file_idx = idx
        if "acronym" in key or "abbrev" in key:
            acro_idx = idx
        if "count" in key or "occur" in key or "freq" in key:
            count_idx = idx

    if file_idx is None or acro_idx is None:
        print(f"  WARNING: Could not find filename/acronym columns in '{sheet_name}' sheet.")
        wb.close()
        return {}

    # Collect acronyms per document with counts
    doc_acronyms: dict[str, list[tuple[str, int]]] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        fname = row[file_idx] if file_idx < len(row) else None
        acro = row[acro_idx] if acro_idx < len(row) else None
        count = row[count_idx] if count_idx is not None and count_idx < len(row) else 1
        if fname and acro:
            key = str(fname).strip().lower()
            if key not in doc_acronyms:
                doc_acronyms[key] = []
            try:
                count = int(count)
            except (TypeError, ValueError):
                count = 1
            doc_acronyms[key].append((str(acro).strip(), count))

    wb.close()

    # Sort by count descending and extract just the acronym strings
    result = {}
    for key, acro_list in doc_acronyms.items():
        acro_list.sort(key=lambda x: x[1], reverse=True)
        result[key] = [a[0] for a in acro_list]

    total_acronyms = sum(len(v) for v in result.values())
    print(f"  Loaded {total_acronyms} acronyms across {len(result)} documents from audit file.")
    return result


# ===========================================================================
# Field resolvers — each "auto" field key has a resolver function
# ===========================================================================


def resolve_document_name(
    filename: str,
    split_manifest: dict[str, str],
) -> str:
    """
    Resolve the human-readable document name for a file.

    WHAT IT DOES:
        Uses the split manifest to find the parent document name.
        Falls back to parsing the filename if no manifest entry exists.

    TWEAK:
        If your filenames use a different separator than " - ", adjust
        the split logic in the fallback branch.
    """
    # Try split manifest first (most reliable)
    parent = split_manifest.get(filename.lower(), "")
    if parent:
        # Strip extension and common suffixes from parent name
        name = os.path.splitext(parent)[0]
        for suffix in ("_fixed", "_optimized", "_backup"):
            if name.lower().endswith(suffix):
                name = name[: -len(suffix)]
        return name.replace("_", " ")

    # Fallback: parse the filename directly
    name = os.path.splitext(filename)[0]
    # Remove sub-doc suffix added by splitter (e.g., " - Scope")
    if " - " in name:
        name = name.split(" - ")[0]
    for suffix in ("_fixed", "_optimized", "_backup"):
        if name.lower().endswith(suffix):
            name = name[: -len(suffix)]
    return name.replace("_", " ")


def resolve_url(
    doc_name: str,
    filename: str,
    url_mapping: dict[str, str],
    config: dict,
) -> tuple[str, str]:
    """
    Resolve the URL for a document.

    RETURNS:
        Tuple of (url, source) where source is "excel", "fallback", or "none".

    WHAT IT DOES:
        1. Checks if any key in the URL Excel mapping is a substring of the
           document name (case-insensitive). Both are normalized (underscores
           converted to spaces) before comparison.
        2. Falls back to the URL template with {filename} substitution.
        3. Returns "(URL not configured)" if neither works.
    """
    # Normalize doc_name: convert underscores to spaces, lowercase
    doc_name_normalized = doc_name.replace("_", " ").lower()

    # 1. Excel lookup — substring match
    for excel_name, url in url_mapping.items():
        # Normalize excel_name the same way for comparison
        excel_name_normalized = excel_name.replace("_", " ").lower()
        if excel_name_normalized in doc_name_normalized or doc_name_normalized in excel_name_normalized:
            return url, "excel"

    # 2. Fallback template
    meta_cfg = config.get("metadata", {})
    fallback = meta_cfg.get("url", {}).get("fallback_template", "")
    if fallback:
        # Use the base document name (no extension) for substitution
        base_name = doc_name.replace(" ", "_")
        url = fallback.replace("{filename}", base_name)
        return url, "fallback"

    # 3. Nothing configured
    return "(URL not configured)", "none"


def extract_section_text(
    doc_name: str,
    section_type: str,
    profiler_data: dict,
    split_manifest: dict[str, str],
    filename: str,
    doc_path: str,
    config: dict,
    max_chars: int = 300,
) -> tuple[str, str]:
    """
    Extract scope or intent text for a document.

    WHAT IT DOES:
        1. Looks up the parent document in profiler data.
        2. Finds the section classified as scope/intent.
        3. Opens the source document and extracts text under that heading.
        4. Falls back to scanning the current sub-document directly.

    RETURNS:
        Tuple of (text, source) where source is "profiler", "direct_scan",
        or "not_detected".

    TWEAK:
        Adjust max_chars in config (metadata.max_scope_chars / max_intent_chars)
        if the extracted text is too short or too long.
    """
    # Determine the parent document filename for profiler lookup
    parent = split_manifest.get(filename.lower(), filename)
    parent_lower = parent.lower()

    # --- Strategy 1: Profiler data ---
    profile = profiler_data.get(parent_lower)
    if profile:
        sections = profile.get("sections", [])
        for section in sections:
            std_section = section.get("standard_section", "")
            if std_section == section_type:
                # Found the section — get its text from the profile
                # Some profilers store section text directly
                section_text = section.get("text", "")
                if section_text:
                    return _truncate(section_text, max_chars), "profiler"
                # Otherwise, we know the heading — try to find it in the doc
                heading_text = section.get("heading_text", "")
                if heading_text:
                    text = _extract_text_under_heading(doc_path, heading_text, max_chars)
                    if text:
                        return text, "profiler"

    # --- Strategy 2: Direct scan of the current document ---
    section_keywords = config.get("sections", {}).get(section_type, [])
    if section_keywords:
        text = _scan_for_section(doc_path, section_keywords, max_chars)
        if text:
            return text, "direct_scan"

    return "(Not detected)", "not_detected"


def _extract_text_under_heading(doc_path: str, heading_text: str, max_chars: int) -> str:
    """
    Open a .docx and extract body paragraphs under a specific heading.

    Walks paragraphs until it finds one matching heading_text, then collects
    subsequent Normal paragraphs until the next heading or max_chars is reached.
    """
    try:
        doc = docx.Document(doc_path)
    except Exception:
        return ""

    found = False
    collected = []
    heading_lower = heading_text.strip().lower()

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""

        if not found:
            # Check if this paragraph matches the heading
            if para.text.strip().lower() == heading_lower:
                found = True
                continue
        else:
            # Stop at the next heading
            if style_name.startswith("Heading"):
                break
            text = para.text.strip()
            if text:
                collected.append(text)
                total = sum(len(t) for t in collected)
                if total >= max_chars:
                    break

    return _truncate(" ".join(collected), max_chars)


def _scan_for_section(doc_path: str, keywords: list[str], max_chars: int) -> str:
    """
    Scan a .docx for a heading matching any of the given keywords,
    then extract the body text below it.

    WHAT IT DOES:
        Walks the first 60 paragraphs looking for a heading (bold text or
        Heading style) that contains one of the keywords.  Extracts body
        text under that heading.
    """
    try:
        doc = docx.Document(doc_path)
    except Exception:
        return ""

    paragraphs = doc.paragraphs[:60]  # Don't scan the whole doc

    for i, para in enumerate(paragraphs):
        text = para.text.strip().lower()
        if not text:
            continue

        # Check if this looks like a heading with a matching keyword
        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading")

        if not is_heading:
            continue

        for kw in keywords:
            if kw.lower() in text:
                # Found it — collect body text below
                collected = []
                for j in range(i + 1, min(i + 15, len(paragraphs))):
                    next_style = paragraphs[j].style.name if paragraphs[j].style else ""
                    if next_style.startswith("Heading"):
                        break
                    body = paragraphs[j].text.strip()
                    if body:
                        collected.append(body)
                        total = sum(len(t) for t in collected)
                        if total >= max_chars:
                            break
                if collected:
                    return _truncate(" ".join(collected), max_chars)

    return ""


def _truncate(text: str, max_chars: int) -> str:
    """Truncate text at the last sentence boundary before max_chars."""
    if len(text) <= max_chars:
        return text

    # Try to cut at a sentence boundary
    truncated = text[:max_chars]
    last_period = truncated.rfind(".")
    last_question = truncated.rfind("?")
    last_exclaim = truncated.rfind("!")
    boundary = max(last_period, last_question, last_exclaim)

    if boundary > max_chars * 0.5:  # Only use boundary if it's in the latter half
        return truncated[: boundary + 1] + "..."
    return truncated.rstrip() + "..."


def generate_tags(
    doc_name: str,
    filename: str,
    profiler_data: dict,
    split_manifest: dict[str, str],
    acronym_data: dict[str, list[str]],
    config: dict,
) -> str:
    """
    Generate a comma-separated tag string from multiple sources.

    WHAT IT DOES:
        Combines tags from:
        1. Document type (A/B/C/D/E) from profiler
        2. Standard sections found (has-scope, has-controls, etc.)
        3. Per-document acronyms from the Acronym Finder output
        4. Static tags from config

    RETURNS:
        Comma-separated string of unique tags.

    TWEAK:
        - Disable individual tag sources in config (metadata.tags.include_doc_type, etc.)
        - Limit acronym tags with metadata.tags.max_acronym_tags
        - Add org-level labels via metadata.tags.static_tags
    """
    meta_cfg = config.get("metadata", {})
    tags_cfg = meta_cfg.get("tags", {})
    tags = []

    # Determine parent document for profiler lookup
    parent = split_manifest.get(filename.lower(), filename)
    parent_lower = parent.lower()
    profile = profiler_data.get(parent_lower, {})

    # 1. Document type tag
    if tags_cfg.get("include_doc_type", True) and profile:
        doc_type = profile.get("doc_type", profile.get("document_type", ""))
        if doc_type:
            tags.append(f"Type-{doc_type}")

    # 2. Standard sections found
    if tags_cfg.get("include_sections_found", True) and profile:
        sections = profile.get("sections", [])
        for section in sections:
            std = section.get("standard_section", "")
            if std and std != "none":
                tags.append(f"has-{std}")

    # 3. Acronym tags
    if acronym_data:
        max_tags = tags_cfg.get("max_acronym_tags", 15)
        # Try matching by parent filename
        for key, acronyms in acronym_data.items():
            if key in parent_lower or parent_lower in key:
                limited = acronyms[:max_tags] if max_tags > 0 else acronyms
                tags.extend(limited)
                break

    # 4. Static tags
    static = tags_cfg.get("static_tags", [])
    if static:
        tags.extend(static)

    # Deduplicate while preserving order
    seen = set()
    unique = []
    for tag in tags:
        tag_lower = tag.lower()
        if tag_lower not in seen:
            seen.add(tag_lower)
            unique.append(tag)

    return ", ".join(unique)


def resolve_excel_field(
    doc_name: str,
    excel_column: str,
    url_mapping_wb_path: str,
    config: dict,
) -> str:
    """
    Resolve a custom field value from the URL lookup Excel file.

    WHAT IT DOES:
        Reads a specific column from the Excel file, matching by document name.
        Used for custom fields with source: "excel".
    """
    import openpyxl

    meta_cfg = config.get("metadata", {})
    url_cfg = meta_cfg.get("url", {})
    lookup_file = url_cfg.get("lookup_file", "")

    if not lookup_file:
        return "(Excel lookup not configured)"

    lookup_path = resolve_path(config, lookup_file)
    if not os.path.isfile(lookup_path):
        return "(Excel file not found)"

    name_col = url_cfg.get("name_column", "Document Name").lower()
    target_col = excel_column.lower()
    sheet_ref = url_cfg.get("sheet", 0)

    try:
        wb = openpyxl.load_workbook(lookup_path, read_only=True, data_only=True)
        if isinstance(sheet_ref, int):
            ws = wb.worksheets[sheet_ref]
        else:
            ws = wb[sheet_ref]
    except Exception:
        return "(Could not read Excel)"

    # Find column indices
    headers = {}
    for col_idx, cell in enumerate(next(ws.iter_rows(min_row=1, max_row=1, values_only=False))):
        if cell.value:
            headers[str(cell.value).strip().lower()] = col_idx

    name_idx = headers.get(name_col)
    target_idx = headers.get(target_col)

    if name_idx is None or target_idx is None:
        wb.close()
        return f"(Column '{excel_column}' not found)"

    doc_name_lower = doc_name.lower()
    for row in ws.iter_rows(min_row=2, values_only=True):
        row_name = row[name_idx] if name_idx < len(row) else None
        if row_name and str(row_name).strip().lower() in doc_name_lower:
            val = row[target_idx] if target_idx < len(row) else None
            wb.close()
            return str(val).strip() if val else "(empty)"

    wb.close()
    return "(No match found)"


# ===========================================================================
# Metadata block construction and insertion
# ===========================================================================


def resolve_field_value(
    field: dict,
    filename: str,
    doc_path: str,
    doc_name: str,
    url_mapping: dict[str, str],
    profiler_data: dict,
    split_manifest: dict[str, str],
    acronym_data: dict[str, list[str]],
    config: dict,
) -> tuple[str, str]:
    """
    Resolve the value for a single metadata field based on its source type.

    RETURNS:
        Tuple of (value, source_hint) for manifest logging.
    """
    key = field.get("key", "")
    source = field.get("source", "auto")

    if source == "static":
        return field.get("value", ""), "static"

    if source == "excel":
        excel_col = field.get("excel_column", "")
        if not excel_col:
            return "(No excel_column specified)", "error"
        val = resolve_excel_field(doc_name, excel_col, "", config)
        return val, "excel"

    # source == "auto" — use built-in resolvers based on key
    if key == "name":
        return doc_name, "auto"

    if key == "url":
        url, url_source = resolve_url(doc_name, filename, url_mapping, config)
        return url, url_source

    if key == "scope":
        meta_cfg = config.get("metadata", {})
        max_chars = meta_cfg.get("max_scope_chars", 300)
        text, text_source = extract_section_text(
            doc_name, "scope", profiler_data, split_manifest,
            filename, doc_path, config, max_chars,
        )
        return text, text_source

    if key == "intent":
        meta_cfg = config.get("metadata", {})
        max_chars = meta_cfg.get("max_intent_chars", 300)
        # Try "intent" first, then "purpose" (profiler may classify either way)
        text, text_source = extract_section_text(
            doc_name, "intent", profiler_data, split_manifest,
            filename, doc_path, config, max_chars,
        )
        if text == "(Not detected)":
            text, text_source = extract_section_text(
                doc_name, "purpose", profiler_data, split_manifest,
                filename, doc_path, config, max_chars,
            )
        return text, text_source

    if key == "tags":
        tags = generate_tags(
            doc_name, filename, profiler_data, split_manifest,
            acronym_data, config,
        )
        return tags if tags else "(No tags generated)", "auto"

    # Unknown key — return empty
    return f"(Unknown field key: {key})", "error"


def build_metadata_elements(
    doc: docx.Document,
    fields: list[dict],
    resolved_values: list[tuple[str, str]],
    config: dict,
) -> list:
    """
    Build the metadata block as a list of XML elements (heading + paragraph pairs).

    WHAT IT DOES:
        Creates Heading 1 + Normal paragraph pairs for each enabled field.
        Returns a list of XML elements ready for insertion into the document body.

    FORMAT:
        [Heading 1] Document
        [Normal]    Access Control Policy

        [Heading 1] URL
        [Normal]    https://...

        ... etc.
    """
    meta_cfg = config.get("metadata", {})
    font_size = meta_cfg.get("font_size", 8)
    label_color_hex = meta_cfg.get("label_color", "2F5496")

    # Parse the hex color
    try:
        r = int(label_color_hex[0:2], 16)
        g = int(label_color_hex[2:4], 16)
        b = int(label_color_hex[4:6], 16)
        label_color = RGBColor(r, g, b)
    except (ValueError, IndexError):
        label_color = RGBColor(0x2F, 0x54, 0x96)

    elements = []

    for field, (value, _source) in zip(fields, resolved_values):
        label = field.get("label", field.get("key", "Field"))

        # Create the heading paragraph.
        # Some .docx files don't have built-in "Heading 1" style defined,
        # so we fall back to a bold paragraph styled to look like a heading.
        try:
            heading_para = doc.add_heading(label, level=1)
        except KeyError:
            # "Heading 1" style doesn't exist in this document —
            # create a bold paragraph instead (works with any .docx)
            heading_para = doc.add_paragraph()
            run = heading_para.add_run(label)
            run.bold = True

        # Style the heading: set color and size
        for run in heading_para.runs:
            run.font.color.rgb = label_color
            run.font.size = Pt(font_size + 4)  # Headings slightly larger

        # Remove heading from body (we'll re-insert at the right position later)
        heading_elem = heading_para._element
        heading_elem.getparent().remove(heading_elem)
        elements.append(heading_elem)

        # Create the value paragraph
        value_para = doc.add_paragraph(value)
        for run in value_para.runs:
            run.font.size = Pt(font_size)

        value_elem = value_para._element
        value_elem.getparent().remove(value_elem)
        elements.append(value_elem)

    # Optional separator
    if meta_cfg.get("add_separator", True):
        sep_para = doc.add_paragraph()
        # Add a bottom border to simulate a horizontal rule
        pPr = sep_para._element.get_or_add_pPr()
        pBdr = docx.oxml.OxmlElement("w:pBdr")
        bottom = docx.oxml.OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        pBdr.append(bottom)
        pPr.append(pBdr)

        sep_elem = sep_para._element
        sep_elem.getparent().remove(sep_elem)
        elements.append(sep_elem)

    return elements


def insert_metadata_top(doc: docx.Document, elements: list):
    """
    Insert metadata elements at the very top of the document body.

    WHAT IT DOES:
        Inserts each element at the beginning of the body, in order,
        so the metadata block appears before any existing content.
    """
    body = doc.element.body
    for i, elem in enumerate(elements):
        body.insert(i, elem)


def insert_metadata_bottom(doc: docx.Document, elements: list):
    """
    Append metadata elements at the bottom of the document body.

    WHAT IT DOES:
        Appends each element to the end of the body.
    """
    body = doc.element.body
    for elem in elements:
        body.append(elem)


def insert_metadata_header(doc: docx.Document, fields: list[dict], resolved_values: list[tuple[str, str]], config: dict):
    """
    Insert metadata into Word headers so it appears on every page.

    WHAT IT DOES:
        Builds a compact single-line summary and places it in the header
        of every section in the document.  Uses a condensed format since
        headers have limited space.

    FORMAT:
        Document: Access Control Policy | URL: https://... | Tags: Type-B, MFA
    """
    meta_cfg = config.get("metadata", {})
    font_size = meta_cfg.get("font_size", 8)

    # Build compact one-liner
    parts = []
    for field, (value, _) in zip(fields, resolved_values):
        label = field.get("label", field.get("key", ""))
        # Truncate long values for header display
        if len(value) > 80:
            value = value[:77] + "..."
        parts.append(f"{label}: {value}")
    header_text = " | ".join(parts)

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # Check if header already has content — append below it
        if header.paragraphs and header.paragraphs[0].text.strip():
            # Add a new paragraph below existing header content
            para = header.add_paragraph(header_text)
        else:
            # Header is empty — use the default paragraph
            if header.paragraphs:
                para = header.paragraphs[0]
                para.text = header_text
            else:
                para = header.add_paragraph(header_text)

        for run in para.runs:
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ===========================================================================
# Main processing
# ===========================================================================


def process_document(
    doc_path: str,
    output_dir: str,
    url_mapping: dict[str, str],
    profiler_data: dict,
    split_manifest: dict[str, str],
    acronym_data: dict[str, list[str]],
    config: dict,
) -> Optional[dict]:
    """
    Process a single .docx file: resolve metadata, insert block, save copy.

    RETURNS:
        A dict with manifest data for this file, or None if processing failed.

    WHAT IT DOES:
        1. Opens the .docx
        2. Resolves each enabled metadata field
        3. Builds the metadata block (H1 headings + Normal paragraphs)
        4. Inserts at top, bottom, or header based on config placement setting
        5. Saves to the output directory
    """
    filename = os.path.basename(doc_path)
    meta_cfg = config.get("metadata", {})
    placement = meta_cfg.get("placement", "top")
    fields_cfg = meta_cfg.get("fields", [])

    # Filter to enabled fields only
    enabled_fields = [f for f in fields_cfg if f.get("enabled", True)]
    if not enabled_fields:
        print(f"    WARNING: No metadata fields enabled — skipping {filename}")
        return None

    # Open the document
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"    ERROR: Could not open {filename}: {e}")
        return None

    # Resolve the document name first (needed by other resolvers)
    doc_name = resolve_document_name(filename, split_manifest)

    # Resolve all field values
    resolved_values = []
    for field in enabled_fields:
        value, source = resolve_field_value(
            field, filename, doc_path, doc_name,
            url_mapping, profiler_data, split_manifest, acronym_data, config,
        )
        resolved_values.append((value, source))

    # Print progress
    print(f"\n    Processing {filename}...")
    for field, (value, source) in zip(enabled_fields, resolved_values):
        label = field.get("label", field.get("key", ""))
        preview = value[:60] + "..." if len(value) > 60 else value
        print(f"      {label + ':':<12s} {preview} (from {source})")
    print(f"      Placed:      {placement}")

    # Build and insert metadata block
    if placement == "each_page":
        # Header mode — don't insert body elements
        insert_metadata_header(doc, enabled_fields, resolved_values, config)
    else:
        # Build elements for body insertion
        elements_top = build_metadata_elements(doc, enabled_fields, resolved_values, config)

        if placement == "top":
            insert_metadata_top(doc, elements_top)
        elif placement == "top_and_bottom":
            insert_metadata_top(doc, elements_top)
            # Build a second copy for the bottom
            elements_bottom = build_metadata_elements(doc, enabled_fields, resolved_values, config)
            insert_metadata_bottom(doc, elements_bottom)

    # Save to output directory
    output_path = os.path.join(output_dir, filename)
    try:
        doc.save(output_path)
    except Exception as e:
        print(f"    ERROR: Could not save {filename}: {e}")
        return None

    # Build manifest row
    manifest = {
        "source_file": filename,
        "output_file": filename,
        "name": doc_name,
        "placement": placement,
    }

    # Add per-field data to manifest
    for field, (value, source) in zip(enabled_fields, resolved_values):
        key = field.get("key", "")
        if key == "url":
            manifest["url"] = value
            manifest["url_source"] = source
        elif key == "scope":
            manifest["scope_source"] = source
            manifest["scope_preview"] = value[:80]
        elif key == "intent":
            manifest["intent_source"] = source
            manifest["intent_preview"] = value[:80]
        elif key == "tags":
            manifest["tags"] = value
            manifest["tag_count"] = len(value.split(", ")) if value and value != "(No tags generated)" else 0

    return manifest


def main():
    """
    Main entry point — orchestrates the full metadata injection pipeline.

    WHAT IT DOES:
        1. Loads config and all data sources (URL Excel, profiler, manifest, acronyms)
        2. Iterates over .docx files in the input directory
        3. Processes each file (resolve metadata, insert block, save copy)
        4. Writes metadata_manifest.csv
        5. Prints summary banner
    """
    # --- Parse CLI arguments ---
    parser = setup_argparse("Step 5 — Add metadata blocks to sub-documents")
    args = parser.parse_args()

    config = load_config(args.config)
    meta_cfg = config.get("metadata", {})

    # Resolve directories
    input_dir = args.input_dir
    output_dir = args.output_dir

    if not input_dir:
        # Default: read from Step 4 split_documents output
        input_dir = get_output_dir(config, "split_documents")
    if not output_dir:
        output_dir = get_output_dir(config, "metadata")

    ensure_output_dir(output_dir)

    # --- Banner ---
    print()
    print("=" * 70)
    print("  STEP 5 — METADATA INJECTOR")
    print("=" * 70)
    print(f"  Input:     {os.path.abspath(input_dir)}")
    print(f"  Output:    {os.path.abspath(output_dir)}")
    print(f"  Placement: {meta_cfg.get('placement', 'top')}")

    enabled_fields = [f for f in meta_cfg.get("fields", []) if f.get("enabled", True)]
    field_labels = [f.get("label", f.get("key", "?")) for f in enabled_fields]
    print(f"  Fields:    {', '.join(field_labels)}")
    print("=" * 70)

    # --- Load data sources ---
    print("\n  Loading data sources...")
    url_mapping = load_url_mapping(config)
    profiler_data = load_profiler_data(config)
    split_manifest = load_split_manifest(config)
    acronym_data = load_acronym_data(config)

    # --- Find input files ---
    files = iter_docx_files(input_dir, config)
    if not files:
        print(f"\n  No .docx files found in: {input_dir}")
        print("  TIP: Run Steps 3-4 first to generate split sub-documents,")
        print("       or point --input_dir at a folder with .docx files.")
        return

    print(f"\n  Found {len(files)} .docx files to process.")

    # --- Process each file ---
    manifests = []
    files_ok = 0
    files_failed = 0
    url_excel = 0
    url_fallback = 0
    scope_detected = 0
    intent_detected = 0
    total_tags = 0

    for doc_path in files:
        result = process_document(
            doc_path, output_dir,
            url_mapping, profiler_data, split_manifest, acronym_data, config,
        )
        if result:
            manifests.append(result)
            files_ok += 1
            # Track stats
            if result.get("url_source") == "excel":
                url_excel += 1
            elif result.get("url_source") == "fallback":
                url_fallback += 1
            if result.get("scope_source", "not_detected") != "not_detected":
                scope_detected += 1
            if result.get("intent_source", "not_detected") != "not_detected":
                intent_detected += 1
            total_tags += result.get("tag_count", 0)
        else:
            files_failed += 1

    # --- Write manifest CSV ---
    manifest_file = meta_cfg.get("manifest_file",
                                  config.get("output", {}).get("metadata", {}).get("manifest_file", "metadata_manifest.csv"))
    manifest_path = os.path.join(output_dir, os.path.basename(manifest_file))

    if manifests:
        # Collect all unique keys across manifests
        all_keys = []
        seen_keys = set()
        for m in manifests:
            for k in m.keys():
                if k not in seen_keys:
                    seen_keys.add(k)
                    all_keys.append(k)

        with open(manifest_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=all_keys)
            writer.writeheader()
            writer.writerows(manifests)

    # --- Summary banner ---
    print()
    print("=" * 70)
    print("  STEP 5 — METADATA INJECTION SUMMARY")
    print("=" * 70)
    print(f"  Files processed:          {files_ok}")
    print(f"  Files failed:             {files_failed}")
    if "url" in [f.get("key") for f in enabled_fields]:
        print(f"  URL resolved (Excel):     {url_excel}")
        print(f"  URL resolved (fallback):  {url_fallback}")
        print(f"  URL not configured:       {files_ok - url_excel - url_fallback}")
    if "scope" in [f.get("key") for f in enabled_fields]:
        print(f"  Scope detected:           {scope_detected}")
    if "intent" in [f.get("key") for f in enabled_fields]:
        print(f"  Intent detected:          {intent_detected}")
    if "tags" in [f.get("key") for f in enabled_fields]:
        avg_tags = total_tags / files_ok if files_ok else 0
        print(f"  Avg tags per document:    {avg_tags:.1f}")
    print()
    print(f"  Output written to: {os.path.abspath(output_dir)}")
    print(f"  Manifest written to: {manifest_path}")
    print("=" * 70)
    print()


if __name__ == "__main__":
    main()
