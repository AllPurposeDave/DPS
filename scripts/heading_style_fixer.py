"""
Step 3: Heading Style Fixer
=============================

RUN THIS AFTER cross_reference_extractor.py, before section_splitter.py.

Converts fake headings (bold text with no Word Heading style) to real
Word Heading 1/2/3 styles. The splitter (Step 4) depends on correct
Heading styles to know where to cut.

Usage (unified pipeline):
    python run_pipeline.py --step 3

Usage (standalone):
    python scripts/heading_style_fixer.py
    python scripts/heading_style_fixer.py --config dps_config.yaml
    python scripts/heading_style_fixer.py ./input ./output

Output:
    *_fixed.docx files — one per input doc, with corrected heading styles
    heading_changes.csv — change log of every style modification

FAILURE POINT: After this script runs, ALWAYS review heading_changes.csv.
False positives (bold labels, bold table headers, bold "NOTE:" text) will
be incorrectly promoted to heading styles.
"""
# pip install python-docx

from __future__ import annotations

import csv
import os
import re
import traceback

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from shared_utils import (
    ensure_output_dir,
    get_input_dir,
    get_output_dir,
    is_heading_style,
    is_paragraph_bold,
    iter_docx_files,
    load_config,
    log_pipeline_issue,
    setup_argparse,
)


def build_heading_patterns(config: dict) -> tuple:
    """Build heading level regex patterns from config."""
    headings_cfg = config.get("headings", {})
    h1 = re.compile(headings_cfg.get("heading1_pattern", r"^(?:\d+\.0\s+|[IVXLC]+\.\s+)"))
    h2 = re.compile(headings_cfg.get("heading2_pattern", r"^(?:\d+\.\d+\s+|[A-Z]\.\s+)"))
    h3 = re.compile(headings_cfg.get("heading3_pattern", r"^\d+\.\d+\.\d+\s+"))
    return h1, h2, h3


def build_custom_style_map(config: dict) -> dict:
    """Build custom style map from config, falling back to defaults."""
    headings_cfg = config.get("headings", {})
    config_map = headings_cfg.get("custom_style_map", {})
    if config_map:
        # Normalize keys to lowercase
        return {k.strip().lower(): v for k, v in config_map.items()}

    # Fallback defaults
    return {
        "policy heading 1": "Heading 1",
        "policy heading 2": "Heading 2",
        "policy heading 3": "Heading 3",
        "policyheading1": "Heading 1",
        "policyheading2": "Heading 2",
        "policyheading3": "Heading 3",
        "doc heading 1": "Heading 1",
        "doc heading 2": "Heading 2",
        "doc heading 3": "Heading 3",
        "heading1": "Heading 1",
        "heading2": "Heading 2",
        "heading3": "Heading 3",
        "title heading": "Heading 1",
        "section heading": "Heading 1",
        "subsection heading": "Heading 2",
    }


def determine_heading_level(text: str, re_h1, re_h2, re_h3, default_level: int = 2) -> str:
    """
    Determine the appropriate heading style based on numbering patterns.
    Returns 'Heading 1', 'Heading 2', 'Heading 3', or default.

    CHECK ORDER IS IMPORTANT:
    RE_HEADING3 must be checked before RE_HEADING2 because "1.1.1 " also matches
    the RE_HEADING2 pattern.
    """
    stripped = text.strip()
    if re_h3.match(stripped):
        return "Heading 3"
    if re_h1.match(stripped):
        return "Heading 1"
    if re_h2.match(stripped):
        return "Heading 2"
    return f"Heading {default_level}"


def is_fake_heading(paragraph, max_chars: int = 120) -> bool:
    """
    Detect fake headings: paragraphs that look like headings but don't
    have a Word Heading style applied.

    A paragraph is a fake heading if ALL of these are true:
    - Bold formatting (paragraph-level or all runs bold)
    - Short (under max_chars characters)
    - Does NOT end with a period
    - Does NOT already have a Heading style
    """
    text = paragraph.text.strip()
    if not text:
        return False
    if is_heading_style(paragraph.style):
        return False
    if len(text) >= max_chars:
        return False
    if text.endswith("."):
        return False
    if not is_paragraph_bold(paragraph):
        return False
    return True


def get_custom_style_mapping(style_name: str, style_map: dict):
    """Check if a style name matches a known custom heading style."""
    if not style_name:
        return None
    normalized = style_name.strip().lower()
    return style_map.get(normalized)


def apply_text_deletions(doc, config: dict) -> list[dict]:
    """
    Delete configured phrases from all paragraph and table cell text.
    Modifies run.text to preserve formatting (bold, font, etc.).
    Returns a list of deletion records for the change log.
    """
    deletion_cfg = config.get("text_deletions", {})
    if not deletion_cfg.get("enabled", False):
        return []

    phrases = deletion_cfg.get("phrases", [])
    if not phrases:
        return []

    case_sensitive = deletion_cfg.get("case_sensitive", True)
    deletions = []

    def delete_from_run(run, phrase, location):
        """Delete a phrase from a run's text. Returns True if changed."""
        original = run.text
        if not original:
            return False
        if case_sensitive:
            if phrase not in original:
                return False
            new_text = original.replace(phrase, "")
        else:
            pattern = re.compile(re.escape(phrase), re.IGNORECASE)
            if not pattern.search(original):
                return False
            new_text = pattern.sub("", original)
        # Collapse multiple spaces
        new_text = re.sub(r'  +', ' ', new_text).strip()
        deletions.append({
            "location": location,
            "phrase_deleted": phrase,
            "original_preview": original[:80],
        })
        run.text = new_text
        return True

    # Process paragraphs
    for idx, para in enumerate(doc.paragraphs):
        for run in para.runs:
            for phrase in phrases:
                delete_from_run(run, phrase, f"Paragraph {idx}")

    # Process table cells
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        for phrase in phrases:
                            delete_from_run(run, phrase, f"Table {t_idx}, Row {r_idx}, Cell {c_idx}")

    return deletions


def _get_heading_level(para) -> int | None:
    """Return heading level (1-9) if the paragraph has a Heading style, else None."""
    if para.style and is_heading_style(para.style):
        # Style name is "Heading N"
        try:
            return int(para.style.name.split()[-1])
        except (ValueError, IndexError):
            pass
    return None


def apply_section_deletions(doc, config: dict, filename: str = "") -> list[dict]:
    """
    Delete entire sections (heading + body) based on configured section headings.

    Two sources are combined (union):
      - text_deletions.section_deletions           — global list, case-insensitive
                                                      SUBSTRING match against heading text.
      - text_deletions.per_doc_section_deletions   — per-document list, keyed by
                                                      exact basename. Uses EXACT match
                                                      (stripped, case-insensitive) so
                                                      headings copied from the profiler
                                                      remove only the intended section.

    Removes the matching heading paragraph and all subsequent body elements
    (paragraphs and tables) until the next heading of the same or higher level.
    Works on the document's XML body directly so both paragraphs and tables
    within the section are removed.

    Returns a list of deletion records for the change log.
    """
    deletion_cfg = config.get("text_deletions", {})
    if not deletion_cfg.get("enabled", False):
        return []

    section_deletions = deletion_cfg.get("section_deletions", [])
    per_doc_deletions = deletion_cfg.get("per_doc_section_deletions", [])

    # Global targets — substring match
    global_targets = []
    for entry in section_deletions:
        if entry.get("delete", True):
            global_targets.append(entry["heading"].strip().lower())

    # Per-doc targets — exact match, scoped to the current filename
    per_doc_targets: list[str] = []
    if filename:
        for entry in per_doc_deletions:
            if str(entry.get("doc_name", "")).strip() == filename:
                for h in entry.get("headings") or []:
                    h_clean = str(h).strip().lower()
                    if h_clean:
                        per_doc_targets.append(h_clean)

    if not global_targets and not per_doc_targets:
        return []

    # Build the scan list: (target_value, match_type). Order doesn't matter for
    # correctness — we re-scan the document per target because removals shift
    # sibling indices.
    scan_list = [(t, "global_substring") for t in global_targets] + \
                [(t, "per_doc_exact") for t in per_doc_targets]

    from lxml import etree

    body = doc.element.body
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    records = []

    # We iterate repeatedly because removing elements shifts indices.
    # Process one target section at a time.
    for target, match_type in scan_list:
        # Scan paragraphs for a heading match
        while True:
            found = False
            for para in doc.paragraphs:
                level = _get_heading_level(para)
                if level is None:
                    continue
                para_text_cmp = para.text.strip().lower()
                if match_type == "per_doc_exact":
                    if para_text_cmp != target:
                        continue
                else:
                    if target not in para_text_cmp:
                        continue

                # Found a matching heading — collect elements to remove
                found = True
                heading_text = para.text.strip()
                elements_to_remove = [para._element]
                count_paras = 0
                count_tables = 0

                # Walk subsequent siblings in the XML body
                sibling = para._element.getnext()
                while sibling is not None:
                    tag = etree.QName(sibling.tag).localname if isinstance(sibling.tag, str) else ""
                    if tag == "p":
                        # Check if this paragraph is a heading of same or higher level
                        style_elem = sibling.find(".//w:pStyle", nsmap)
                        if style_elem is not None:
                            style_val = style_elem.get(f"{{{nsmap['w']}}}val") or ""
                            if style_val.startswith("Heading"):
                                try:
                                    # XML stores "Heading1" (no space) or "Heading 1"
                                    digits = style_val.replace("Heading", "").strip()
                                    sib_level = int(digits)
                                    if sib_level <= level:
                                        break  # Stop — next section at same/higher level
                                except (ValueError, IndexError):
                                    pass
                        elements_to_remove.append(sibling)
                        count_paras += 1
                    elif tag == "tbl":
                        elements_to_remove.append(sibling)
                        count_tables += 1
                    # Skip non-paragraph/table elements (e.g. section breaks)
                    sibling = sibling.getnext()

                # Remove all collected elements
                for elem in elements_to_remove:
                    parent = elem.getparent()
                    if parent is not None:
                        parent.remove(elem)

                records.append({
                    "heading_deleted": heading_text,
                    "heading_level": level,
                    "paragraphs_removed": count_paras,
                    "tables_removed": count_tables,
                    "match_type": match_type,
                })
                break  # Restart scan (indices shifted)

            if not found:
                break  # No more matches for this target

    return records


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_NSMAP = {"w": _W_NS}

DEFAULT_PAGE_NUMBER_PATTERNS = [
    r"^Page\s+\d+(\s+of\s+\d+)?$",
    r"^\d+$",
    r"^-\s*\d+\s*-$",
]


def strip_tracked_changes(doc, config: dict) -> list[dict]:
    """
    Accept all tracked changes in the document body:
      - <w:ins> wrappers are unwrapped (their <w:r> children are kept)
      - <w:del> elements (and their <w:delText> content) are removed
      - <w:moveTo> is unwrapped; <w:moveFrom> is removed
      - Formatting-change markers (<w:pPrChange>, <w:rPrChange>, etc.) are removed

    Must run BEFORE any paragraph/text operation because python-docx does not
    iterate runs nested inside <w:ins>; unwrapping makes that content visible
    to the rest of the pipeline.
    """
    deletion_cfg = config.get("text_deletions", {})
    if not deletion_cfg.get("strip_tracked_changes", False):
        return []

    body_root = doc.element.body
    counts = {"ins_unwrapped": 0, "del_removed": 0, "moves_handled": 0, "change_markers_removed": 0}

    # Accept inserts: unwrap <w:ins> (move children up, remove wrapper)
    for ins in body_root.xpath(".//w:ins"):
        parent = ins.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ins)
        for child in list(ins):
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins)
        counts["ins_unwrapped"] += 1

    # Reject deletes: remove <w:del> entirely
    for d in body_root.xpath(".//w:del"):
        parent = d.getparent()
        if parent is not None:
            parent.remove(d)
            counts["del_removed"] += 1

    # Moves: treat moveTo like ins, moveFrom like del
    for mt in body_root.xpath(".//w:moveTo"):
        parent = mt.getparent()
        if parent is None:
            continue
        idx = list(parent).index(mt)
        for child in list(mt):
            parent.insert(idx, child)
            idx += 1
        parent.remove(mt)
        counts["moves_handled"] += 1
    for mf in body_root.xpath(".//w:moveFrom"):
        parent = mf.getparent()
        if parent is not None:
            parent.remove(mf)
            counts["moves_handled"] += 1

    # Formatting-change markers
    change_tags = ["pPrChange", "rPrChange", "sectPrChange", "tblPrChange",
                   "tblGridChange", "trPrChange", "tcPrChange", "numberingChange"]
    for tag in change_tags:
        for elem in body_root.xpath(f".//w:{tag}"):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                counts["change_markers_removed"] += 1

    if not any(counts.values()):
        return []
    return [{
        "removal_type": "tracked_changes_accepted",
        **counts,
    }]


def remove_comments(doc, config: dict) -> list[dict]:
    """
    Strip inline comment markers from the body AND drop the comments part
    from the docx package so no comment annotation text can leak into
    downstream extraction. Removes:
      - <w:commentRangeStart>, <w:commentRangeEnd>, <w:commentReference> inline markers
      - relationships to comments.xml / commentsExtended.xml / commentsIds.xml /
        commentsExtensible.xml (drops the parts from the package on save)
    """
    deletion_cfg = config.get("text_deletions", {})
    if not deletion_cfg.get("remove_comments", False):
        return []

    body_root = doc.element.body
    inline_markers_removed = 0
    for tag in ["commentRangeStart", "commentRangeEnd", "commentReference"]:
        for elem in body_root.xpath(f".//w:{tag}"):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                inline_markers_removed += 1

    parts_dropped = []
    for rel_id, rel in list(doc.part.rels.items()):
        if "comments" in rel.reltype.lower():
            parts_dropped.append(rel.reltype.rsplit("/", 1)[-1])
            doc.part.drop_rel(rel_id)

    if not inline_markers_removed and not parts_dropped:
        return []
    return [{
        "removal_type": "comments_removed",
        "inline_markers_removed": inline_markers_removed,
        "parts_dropped": ", ".join(parts_dropped),
    }]


def remove_inline_page_numbers(doc, config: dict) -> list[dict]:
    """
    Delete paragraphs whose ENTIRE stripped text matches any configured
    page-number regex. Useful for cleaning exports where page numbers end
    up as standalone body paragraphs (e.g. 'Page 5', 'Page 5 of 20', bare
    digit '5', '- 5 -'). Matching is case-insensitive.
    """
    deletion_cfg = config.get("text_deletions", {})
    if not deletion_cfg.get("remove_inline_page_numbers", False):
        return []

    raw_patterns = deletion_cfg.get("inline_page_number_patterns") or DEFAULT_PAGE_NUMBER_PATTERNS
    compiled = []
    for p in raw_patterns:
        try:
            compiled.append(re.compile(p, re.IGNORECASE))
        except re.error:
            continue  # Skip invalid regex silently — operator can see it's absent in the change log
    if not compiled:
        return []

    records = []
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for pat in compiled:
            if pat.fullmatch(text):
                records.append({
                    "removal_type": "inline_page_number",
                    "text_preview": text[:80],
                    "matched_pattern": pat.pattern,
                })
                parent = para._element.getparent()
                if parent is not None:
                    parent.remove(para._element)
                break
    return records


def remove_table_of_content(doc, config: dict) -> list[dict]:
    """
    Remove paragraphs with Word TOC styles (TOC 1, TOC 2, TOC 3, etc.).
    These are generated by Word's Insert Table of Contents feature and use
    special paragraph styles that section deletion cannot catch.
    Returns a list of deletion records for the change log.
    """
    deletion_cfg = config.get("text_deletions", {})
    if not deletion_cfg.get("remove_table_of_content", False):
        return []

    records = []
    for para in list(doc.paragraphs):
        if para.style and para.style.name.upper().startswith("TOC"):
            text_preview = para.text.strip()[:80]
            records.append({
                "style_removed": para.style.name,
                "text_preview": text_preview,
                "removal_type": "toc_paragraph",
            })
            para._element.getparent().remove(para._element)

    return records


def remove_headers_footers(doc, config: dict) -> list[dict]:
    """
    Clear all page header and footer content (branding, page numbers,
    classification banners). Clears run text rather than removing XML
    parts to avoid corrupting the document structure.
    Returns a list of deletion records for the change log.
    """
    deletion_cfg = config.get("text_deletions", {})
    if not deletion_cfg.get("remove_headers_footers", False):
        return []

    records = []
    for sec_idx, section in enumerate(doc.sections):
        hf_attrs = [
            "header", "first_page_header", "even_page_header",
            "footer", "first_page_footer", "even_page_footer",
        ]
        for attr in hf_attrs:
            hf = getattr(section, attr, None)
            if hf is None:
                continue
            cleared_text = []
            for para in hf.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        cleared_text.append(run.text.strip())
                        run.text = ""
            if cleared_text:
                records.append({
                    "section_index": sec_idx,
                    "header_footer": attr,
                    "text_cleared": " | ".join(cleared_text)[:80],
                    "removal_type": "header_footer",
                })

    return records


def remove_revision_tables(doc, config: dict) -> list[dict]:
    """
    Remove revision/change history tables that don't live under a heading
    (missed by section deletion). Detects tables by column headers matching
    a revision-table signature: 'version' AND 'date' AND one of
    'changes', 'description', 'author'.
    Returns a list of deletion records for the change log.
    """
    deletion_cfg = config.get("text_deletions", {})
    if not deletion_cfg.get("remove_revision_tables", False):
        return []

    records = []
    for table in list(doc.tables):
        try:
            if not table.rows:
                continue
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
        except Exception:
            continue

        has_version = any("version" in h for h in header_cells)
        has_date = any("date" in h for h in header_cells)
        has_changes = any(
            kw in h for h in header_cells
            for kw in ("changes", "description", "author", "modified by")
        )

        if has_version and has_date and has_changes:
            row_count = len(table.rows)
            records.append({
                "columns": " | ".join(header_cells)[:80],
                "row_count": row_count,
                "removal_type": "revision_table",
            })
            table._element.getparent().remove(table._element)

    return records


def flatten_definition_tables(doc, config: dict) -> list[dict]:
    """
    Find tables under headings matching 'Terms', 'Definitions', or 'Glossary'
    and convert each row to a prose paragraph: bold Term followed by ': Definition'.
    The section heading is preserved. The table is replaced with paragraphs.
    Returns a list of records for the change log.
    """
    deletion_cfg = config.get("text_deletions", {})
    if not deletion_cfg.get("flatten_definition_tables", False):
        return []

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from lxml import etree

    definition_keywords = ("terms", "definitions", "glossary")
    body = doc.element.body
    records = []

    # Find heading paragraphs that match definition keywords
    definition_headings = []
    for para in doc.paragraphs:
        if not is_heading_style(para.style):
            continue
        text_lower = para.text.strip().lower()
        if any(kw in text_lower for kw in definition_keywords):
            definition_headings.append(para)

    for heading_para in definition_headings:
        # Find the next table after this heading in the XML body
        target_table = None
        sibling = heading_para._element.getnext()
        while sibling is not None:
            tag = etree.QName(sibling.tag).localname if isinstance(sibling.tag, str) else ""
            if tag == "tbl":
                target_table = sibling
                break
            elif tag == "p":
                # Check if this is another heading — stop searching
                nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                style_elem = sibling.find(".//w:pStyle", nsmap)
                if style_elem is not None:
                    style_val = style_elem.get(f"{{{nsmap['w']}}}val") or ""
                    if style_val.startswith("Heading"):
                        break
            sibling = sibling.getnext()

        if target_table is None:
            continue

        # Read table rows and build prose paragraphs
        from docx.table import Table
        tbl_obj = Table(target_table, doc)
        rows_flattened = 0

        # Insert prose paragraphs before the table
        insert_point = target_table
        for row_idx, row in enumerate(tbl_obj.rows):
            cells = [cell.text.strip() for cell in row.cells]
            # Skip header row (first row) and empty rows
            if row_idx == 0:
                continue
            if len(cells) < 2 or not cells[0]:
                continue

            term = cells[0]
            definition = cells[1] if len(cells) > 1 else ""

            # Create a new paragraph element
            new_p = OxmlElement("w:p")
            # Bold run for term
            r_bold = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            rPr.append(b)
            r_bold.append(rPr)
            t_bold = OxmlElement("w:t")
            t_bold.text = term
            t_bold.set(qn("xml:space"), "preserve")
            r_bold.append(t_bold)
            new_p.append(r_bold)

            # Normal run for ': definition'
            if definition:
                r_normal = OxmlElement("w:r")
                t_normal = OxmlElement("w:t")
                t_normal.text = f": {definition}"
                t_normal.set(qn("xml:space"), "preserve")
                r_normal.append(t_normal)
                new_p.append(r_normal)

            body.insert(list(body).index(insert_point), new_p)
            rows_flattened += 1

        # Remove the table
        if rows_flattened > 0:
            body.remove(target_table)
            records.append({
                "heading": heading_para.text.strip()[:80],
                "rows_flattened": rows_flattened,
                "removal_type": "definition_table_flattened",
            })

    return records


def remove_italics(doc, config: dict) -> list[dict]:
    """
    Explicitly disable italic formatting on every run in the document body,
    table cells, and headers/footers. Sets run.italic = False so italics
    inherited from paragraph styles are also suppressed.
    """
    deletion_cfg = config.get("text_deletions", {})
    if not deletion_cfg.get("remove_italics", False):
        return []

    runs_cleared = 0

    def clear(para):
        nonlocal runs_cleared
        for run in para.runs:
            if run.italic is not False:
                run.italic = False
                runs_cleared += 1

    for para in doc.paragraphs:
        clear(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    clear(para)
    for section in doc.sections:
        for attr in ("header", "first_page_header", "even_page_header",
                     "footer", "first_page_footer", "even_page_footer"):
            hf = getattr(section, attr, None)
            if hf is None:
                continue
            for para in hf.paragraphs:
                clear(para)

    if runs_cleared == 0:
        return []
    return [{
        "removal_type": "italics_removed",
        "runs_cleared": runs_cleared,
    }]


def process_document(filepath: str, output_dir: str, config: dict) -> dict:
    """
    Process a single .docx file. Pipeline order:
      0. Accept tracked changes (must run first — affects XML structure)
      1. Strip comment markers + drop comments part
      2. Remove TOC-styled paragraphs
      3. Clear headers/footers
      4. Remove inline page-number paragraphs
      5. Apply phrase-level text deletions
      6. Fix heading styles
      7. Remove orphan revision tables
      8. Flatten definition tables to prose
      9. Apply section deletions (needs correct heading styles)
     10. Strip italic formatting (final formatting pass)
    Saves as _fixed.docx.

    Returns a dict with keys for every step's records.
    """
    headings_cfg = config.get("headings", {})
    max_chars = headings_cfg.get("fake_heading_max_chars_fixer", 120)
    default_level = headings_cfg.get("default_heading_level", 2)
    style_map = build_custom_style_map(config)
    re_h1, re_h2, re_h3 = build_heading_patterns(config)

    doc = Document(filepath)
    filename = os.path.basename(filepath)
    base_name = os.path.splitext(filename)[0]

    # 0. Accept all tracked changes (must run before anything else touches text)
    tracked_change_records = strip_tracked_changes(doc, config)

    # 1. Strip comments (inline markers + comments.xml part)
    comment_records = remove_comments(doc, config)

    # 2. Remove TOC-styled paragraphs (no heading dependency)
    toc_records = remove_table_of_content(doc, config)

    # 3. Clear headers and footers (no heading dependency)
    hf_records = remove_headers_footers(doc, config)

    # 4. Remove paragraphs that are only page-number text
    page_number_records = remove_inline_page_numbers(doc, config)

    # 5. Apply text deletions (before heading analysis)
    deletion_records = apply_text_deletions(doc, config)

    # 6. Fix heading styles
    changes = []
    for idx, para in enumerate(doc.paragraphs):
        original_style = para.style.name if para.style else "None"
        new_style = None

        # Check for custom named styles that should be standard headings
        if not is_heading_style(para.style) and para.style:
            mapped = get_custom_style_mapping(para.style.name, style_map)
            if mapped:
                new_style = mapped

        # Check for fake headings (bold + short + no period + no heading style)
        if new_style is None and is_fake_heading(para, max_chars):
            new_style = determine_heading_level(para.text, re_h1, re_h2, re_h3, default_level)

        # Apply the style change if needed
        if new_style:
            if new_style not in doc.styles:
                doc.styles.add_style(new_style, WD_STYLE_TYPE.PARAGRAPH)
            para.style = doc.styles[new_style]
            text_preview = para.text.strip()[:80]
            changes.append({
                "doc_name": filename,
                "paragraph_index": idx,
                "original_style": original_style,
                "new_style": new_style,
                "paragraph_text_preview": text_preview,
            })

    # 7. Remove orphan revision tables (after heading fixes)
    revision_table_records = remove_revision_tables(doc, config)

    # 8. Flatten definition tables to prose (after heading fixes)
    definition_records = flatten_definition_tables(doc, config)

    # 9. Apply section deletions (needs correct heading styles)
    section_deletion_records = apply_section_deletions(doc, config, filename)

    # 10. Strip italics from the surviving content (final formatting pass)
    italic_records = remove_italics(doc, config)

    # Save the fixed document
    output_path = os.path.join(output_dir, f"{base_name}_fixed.docx")
    doc.save(output_path)

    return {
        "heading_changes": changes,
        "deletion_records": deletion_records,
        "section_deletion_records": section_deletion_records,
        "toc_records": toc_records,
        "hf_records": hf_records,
        "revision_table_records": revision_table_records,
        "definition_records": definition_records,
        "tracked_change_records": tracked_change_records,
        "comment_records": comment_records,
        "page_number_records": page_number_records,
        "italic_records": italic_records,
    }


def main():
    parser = setup_argparse("Step 3: Fix heading styles in .docx policy documents")
    args = parser.parse_args()

    config = load_config(args.config)
    input_dir = get_input_dir(config, args.input_dir)
    output_dir = get_output_dir(config, "heading_fixes", args.output_dir)
    ensure_output_dir(output_dir)

    docx_files = iter_docx_files(input_dir, config)
    if not docx_files:
        print(f"No .docx files found in {input_dir}")
        return

    changes_file = config.get("output", {}).get("heading_fixes", {}).get("changes_file", "heading_changes.csv")
    csv_path = os.path.join(output_dir, changes_file)
    fieldnames = [
        "doc_name", "paragraph_index", "original_style",
        "new_style", "paragraph_text_preview",
        "change_type", "phrase_deleted",
    ]

    all_changes = []
    total_deletions = 0
    total_section_deletions = 0
    total_toc_removals = 0
    total_hf_clearings = 0
    total_revision_tables = 0
    total_definitions_flattened = 0
    total_tracked_change_docs = 0
    total_comment_strips = 0
    total_page_number_removals = 0
    total_italic_runs_cleared = 0
    files_processed = 0
    files_failed = 0

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        try:
            result = process_document(filepath, output_dir, config)
            changes = result["heading_changes"]
            deletion_records = result["deletion_records"]
            section_deletion_records = result["section_deletion_records"]
            toc_records = result["toc_records"]
            hf_records = result["hf_records"]
            revision_table_records = result["revision_table_records"]
            definition_records = result["definition_records"]
            tracked_change_records = result["tracked_change_records"]
            comment_records = result["comment_records"]
            page_number_records = result["page_number_records"]
            italic_records = result["italic_records"]

            # Tag heading changes
            for c in changes:
                c["change_type"] = "heading_fix"
                c["phrase_deleted"] = ""
            # Convert text deletion records to CSV rows
            for d in deletion_records:
                all_changes.append({
                    "doc_name": filename,
                    "paragraph_index": d["location"],
                    "original_style": "",
                    "new_style": "[DELETED]",
                    "paragraph_text_preview": d["original_preview"],
                    "change_type": "text_deletion",
                    "phrase_deleted": d["phrase_deleted"],
                })
            # Convert section deletion records to CSV rows
            for s in section_deletion_records:
                match_type = s.get("match_type", "global_substring")
                all_changes.append({
                    "doc_name": filename,
                    "paragraph_index": "",
                    "original_style": f"Heading {s['heading_level']}",
                    "new_style": "[SECTION DELETED]",
                    "paragraph_text_preview": s["heading_deleted"][:80],
                    "change_type": f"section_deletion ({match_type})",
                    "phrase_deleted": f"{s['paragraphs_removed']} para(s) + {s['tables_removed']} table(s) removed",
                })
            # Convert TOC removal records to CSV rows
            for t in toc_records:
                all_changes.append({
                    "doc_name": filename,
                    "paragraph_index": "",
                    "original_style": t["style_removed"],
                    "new_style": "[TOC REMOVED]",
                    "paragraph_text_preview": t["text_preview"],
                    "change_type": "toc_removal",
                    "phrase_deleted": "",
                })
            # Convert header/footer clearing records to CSV rows
            for h in hf_records:
                all_changes.append({
                    "doc_name": filename,
                    "paragraph_index": "",
                    "original_style": h["header_footer"],
                    "new_style": "[HEADER/FOOTER CLEARED]",
                    "paragraph_text_preview": h["text_cleared"],
                    "change_type": "header_footer_cleared",
                    "phrase_deleted": "",
                })
            # Convert revision table records to CSV rows
            for r in revision_table_records:
                all_changes.append({
                    "doc_name": filename,
                    "paragraph_index": "",
                    "original_style": "",
                    "new_style": "[REVISION TABLE REMOVED]",
                    "paragraph_text_preview": r["columns"],
                    "change_type": "revision_table_removed",
                    "phrase_deleted": f"{r['row_count']} row(s)",
                })
            # Convert definition flattening records to CSV rows
            for df in definition_records:
                all_changes.append({
                    "doc_name": filename,
                    "paragraph_index": "",
                    "original_style": "",
                    "new_style": "[TABLE FLATTENED TO PROSE]",
                    "paragraph_text_preview": df["heading"],
                    "change_type": "definition_table_flattened",
                    "phrase_deleted": f"{df['rows_flattened']} row(s) converted",
                })
            # Convert tracked-change records to CSV rows
            for tc in tracked_change_records:
                preview = (
                    f"ins:{tc.get('ins_unwrapped', 0)} del:{tc.get('del_removed', 0)} "
                    f"moves:{tc.get('moves_handled', 0)} markers:{tc.get('change_markers_removed', 0)}"
                )
                all_changes.append({
                    "doc_name": filename,
                    "paragraph_index": "",
                    "original_style": "",
                    "new_style": "[TRACKED CHANGES ACCEPTED]",
                    "paragraph_text_preview": preview,
                    "change_type": "tracked_changes_accepted",
                    "phrase_deleted": "",
                })
            # Convert comment-strip records to CSV rows
            for cm in comment_records:
                preview = (
                    f"markers:{cm.get('inline_markers_removed', 0)} "
                    f"parts:{cm.get('parts_dropped', '')}"
                )
                all_changes.append({
                    "doc_name": filename,
                    "paragraph_index": "",
                    "original_style": "",
                    "new_style": "[COMMENTS REMOVED]",
                    "paragraph_text_preview": preview,
                    "change_type": "comments_removed",
                    "phrase_deleted": "",
                })
            # Convert inline-page-number records to CSV rows
            for pn in page_number_records:
                all_changes.append({
                    "doc_name": filename,
                    "paragraph_index": "",
                    "original_style": "",
                    "new_style": "[PAGE NUMBER REMOVED]",
                    "paragraph_text_preview": pn["text_preview"],
                    "change_type": "inline_page_number_removed",
                    "phrase_deleted": pn["matched_pattern"],
                })
            # Convert italic-removal records to CSV rows
            for it in italic_records:
                all_changes.append({
                    "doc_name": filename,
                    "paragraph_index": "",
                    "original_style": "",
                    "new_style": "[ITALICS REMOVED]",
                    "paragraph_text_preview": f"{it['runs_cleared']} run(s) de-italicized",
                    "change_type": "italics_removed",
                    "phrase_deleted": "",
                })

            all_changes.extend(changes)
            total_deletions += len(deletion_records)
            total_section_deletions += len(section_deletion_records)
            total_toc_removals += len(toc_records)
            total_hf_clearings += len(hf_records)
            total_revision_tables += len(revision_table_records)
            total_definitions_flattened += len(definition_records)
            if tracked_change_records:
                total_tracked_change_docs += 1
            if comment_records:
                total_comment_strips += 1
            total_page_number_removals += len(page_number_records)
            for it in italic_records:
                total_italic_runs_cleared += it.get("runs_cleared", 0)
            files_processed += 1

            parts = [f"{len(changes)} heading(s) fixed"]
            if tracked_change_records:
                parts.append("tracked changes accepted")
            if comment_records:
                parts.append("comments stripped")
            if page_number_records:
                parts.append(f"{len(page_number_records)} page-number para(s) removed")
            if deletion_records:
                parts.append(f"{len(deletion_records)} deletion(s)")
            if section_deletion_records:
                parts.append(f"{len(section_deletion_records)} section(s) deleted")
            if toc_records:
                parts.append(f"{len(toc_records)} TOC paragraph(s) removed")
            if hf_records:
                parts.append(f"{len(hf_records)} header/footer(s) cleared")
            if revision_table_records:
                parts.append(f"{len(revision_table_records)} revision table(s) removed")
            if definition_records:
                parts.append(f"{len(definition_records)} definition table(s) flattened")
            if italic_records:
                parts.append(f"{italic_records[0]['runs_cleared']} italic run(s) cleared")
            print(f"  Processing {filename}... {', '.join(parts)}")
        except Exception as e:
            files_failed += 1
            print(f"  ERROR processing {filename}: {e}")
            traceback.print_exc()
            log_pipeline_issue(os.path.dirname(output_dir), "Step 3 - Heading Fixer", filename, "ERROR", str(e))

    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(all_changes)

    heading_count = sum(1 for c in all_changes if c.get("change_type") == "heading_fix")
    print("\n" + "=" * 60)
    print("STEP 3 — HEADING STYLE FIXER SUMMARY")
    print("=" * 60)
    print(f"Files processed:            {files_processed}")
    print(f"Files failed:               {files_failed}")
    print(f"Total headings fixed:       {heading_count}")
    print(f"Total text deletions:       {total_deletions}")
    print(f"Total sections deleted:     {total_section_deletions}")
    print(f"Total TOC paragraphs removed: {total_toc_removals}")
    print(f"Total headers/footers cleared: {total_hf_clearings}")
    print(f"Total revision tables removed: {total_revision_tables}")
    print(f"Total definition tables flattened: {total_definitions_flattened}")
    print(f"Docs w/ tracked changes accepted: {total_tracked_change_docs}")
    print(f"Docs w/ comments stripped:  {total_comment_strips}")
    print(f"Total page-number paragraphs removed: {total_page_number_removals}")
    print(f"Total italic runs cleared:  {total_italic_runs_cleared}")
    print(f"\nFixed documents written to: {output_dir}")
    print(f"Change log written to: {csv_path}")


if __name__ == "__main__":
    main()
