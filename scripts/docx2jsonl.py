"""
docx2jsonl — DOCX to JSONL converter optimised for RAG pipelines (DPS Step 8).

Converts .docx policy documents into chunked JSONL files (saved as *.jsonl.txt)
with per-chunk metadata (tags, PublishedURL, heading context).
Part of the DPS pipeline; can also run standalone.

Supports Pure Conversion (read from input/) or Optimized (read from a pipeline
step's output, e.g. heading_fixes).

REQUIREMENTS:
  pip install python-docx pyyaml openpyxl
  Python 3.8 or later

USAGE:
  Via pipeline:  python run_pipeline.py --step 8
  Standalone:    python scripts/docx2jsonl.py --config dps_config.xlsx input/ output/8\ -\ jsonl/
"""

from __future__ import annotations

import json
import os
import re
import sys
import time
import traceback
from typing import Optional

# ---------------------------------------------------------------------------
# Import shared_utils from the DPS scripts/ directory
# ---------------------------------------------------------------------------
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_PROJECT_ROOT = os.path.normpath(os.path.join(_SCRIPT_DIR, ".."))
_SCRIPTS_DIR = os.path.join(_PROJECT_ROOT, "scripts")
if _SCRIPTS_DIR not in sys.path:
    sys.path.insert(0, _SCRIPTS_DIR)

from docx import Document
from docx.oxml.ns import qn

from shared_utils import (
    ensure_output_dir,
    get_heading_level,
    is_heading_style,
    is_paragraph_bold,
    iter_docx_files,
    load_config,
    match_doc_name,
    normalize_doc_name,
    resolve_path,
    sanitize_filename,
    setup_argparse,
)
from add_metadata import load_url_mapping, resolve_url


# ============================================================================
# Constants
# ============================================================================

_SMART_QUOTE_MAP = {
    "\u201c": '"', "\u201d": '"',
    "\u2018": "'", "\u2019": "'",
    "\u2013": "--",
    "\u2014": "---",
    "\u2026": "...",
}

_ZERO_WIDTH_RE = re.compile("[\u200b\u200c\u200d\ufeff]")


# ============================================================================
# Text cleanup
# ============================================================================

def _clean_text(text: str) -> str:
    """Apply smart-quote and zero-width cleanup."""
    for old, new in _SMART_QUOTE_MAP.items():
        text = text.replace(old, new)
    text = _ZERO_WIDTH_RE.sub("", text)
    return text


# ============================================================================
# Heading helpers (ported from docx2md.py / heading_style_fixer.py)
# ============================================================================

def _build_heading_patterns(config: dict):
    """Build heading level regex patterns from config."""
    hcfg = config.get("headings", {})
    h1 = re.compile(hcfg.get("heading1_pattern", r"^(?:\d+\.0\s+|[IVXLC]+\.\s+)"))
    h2 = re.compile(hcfg.get("heading2_pattern", r"^(?:\d+\.\d+\s+|[A-Z]\.\s+)"))
    h3 = re.compile(hcfg.get("heading3_pattern", r"^\d+\.\d+\.\d+\s+"))
    return h1, h2, h3


def _build_custom_style_map(config: dict) -> dict:
    """Build lowercase-key custom style map from config."""
    hcfg = config.get("headings", {})
    cmap = hcfg.get("custom_style_map", {})
    if cmap:
        return {k.strip().lower(): v for k, v in cmap.items()}
    return {}


def _determine_heading_level(text: str, re_h1, re_h2, re_h3, default_level: int = 2) -> int:
    """Return the numeric heading level (1-3) for a fake heading."""
    stripped = text.strip()
    if re_h3.match(stripped):
        return 3
    if re_h1.match(stripped):
        return 1
    if re_h2.match(stripped):
        return 2
    return default_level


def _is_fake_heading(paragraph, max_chars: int = 120) -> bool:
    """Detect bold-text paragraphs that look like headings but lack a heading style."""
    text = paragraph.text.strip()
    if not text:
        return False
    if is_heading_style(paragraph.style):
        return False
    if len(text) >= max_chars:
        return False
    if text.endswith("."):
        return False
    if not is_paragraph_bold(paragraph):
        return False
    return True


# ============================================================================
# Excel mapping loaders
# ============================================================================

def _resolve_url_for_jsonl(url_mapping: dict, doc_name: str, config: dict) -> str:
    """Resolve URL for a document name using shared add_metadata.resolve_url().

    Thin wrapper that adapts the shared (url, source) return to a plain string.
    """
    stem = os.path.splitext(doc_name)[0]
    url, _source = resolve_url(stem, doc_name, url_mapping, config)
    return url if url != "(URL not configured)" else ""


def load_tag_mapping(tag_file: str) -> dict[str, list[str]]:
    """Load per-document tags from a Tag Config Excel file.

    Expected columns: Document_Name | Tags (comma-separated).
    Returns {lowercase_doc_name: [tag1, tag2, ...]}.
    """
    import openpyxl

    if not tag_file or not os.path.isfile(tag_file):
        if tag_file:
            print(f"  NOTE: Tag file not found: {tag_file}")
        return {}

    try:
        wb = openpyxl.load_workbook(tag_file, read_only=True, data_only=True)
        ws = wb.worksheets[0]
    except Exception as e:
        print(f"  WARNING: Could not read tag file: {e}")
        return {}

    # Find column indices from header row
    headers = {}
    for col_idx, cell in enumerate(next(ws.iter_rows(min_row=1, max_row=1, values_only=False))):
        if cell.value:
            headers[str(cell.value).strip().lower()] = col_idx

    name_idx = headers.get("document_name")
    tags_idx = headers.get("tags")

    if name_idx is None or tags_idx is None:
        print(f"  WARNING: Tag file must have 'Document_Name' and 'Tags' columns.")
        print(f"  Found columns: {list(headers.keys())}")
        return {}

    mapping = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        name_val = row[name_idx] if name_idx < len(row) else None
        tags_val = row[tags_idx] if tags_idx < len(row) else None
        if name_val:
            name_key = str(name_val).strip().lower()
            if tags_val:
                tags = [t.strip() for t in str(tags_val).split(",") if t.strip()]
            else:
                tags = []
            mapping[name_key] = tags

    wb.close()
    print(f"  Loaded tags for {len(mapping)} documents from {os.path.basename(tag_file)}")
    return mapping


def _resolve_tags(tag_mapping: dict, doc_name: str) -> list[str]:
    """Resolve tags for a document name."""
    stem = os.path.splitext(doc_name)[0].lower()

    # Exact match
    if stem in tag_mapping:
        return tag_mapping[stem]

    # Try with spaces instead of underscores
    stem_spaced = stem.replace("_", " ")
    if stem_spaced in tag_mapping:
        return tag_mapping[stem_spaced]

    # Partial match
    for key, tags in tag_mapping.items():
        if key in stem or stem in key:
            return tags

    return []


def load_acronym_mapping(acronym_file: str) -> dict[str, list[str]]:
    """Load per-document acronyms from confirmed Acronym Definitions Excel.

    Reads the 'Per Document' sheet with columns: Document | Acronym | ...
    Returns {lowercase_doc_name: [acronym1, acronym2, ...]}.
    Only acronyms for that specific document are included.
    """
    import openpyxl

    if not acronym_file or not os.path.isfile(acronym_file):
        return {}

    try:
        wb = openpyxl.load_workbook(acronym_file, read_only=True, data_only=True)
    except Exception as e:
        print(f"  WARNING: Could not read acronym definitions file: {e}")
        return {}

    # Look for "Per Document" sheet
    target_sheet = None
    for name in wb.sheetnames:
        if "per document" in name.lower():
            target_sheet = name
            break
    if not target_sheet:
        wb.close()
        return {}

    ws = wb[target_sheet]

    # Find column indices from header row
    headers = {}
    for col_idx, cell in enumerate(next(ws.iter_rows(min_row=1, max_row=1, values_only=False))):
        if cell.value:
            headers[str(cell.value).strip().lower()] = col_idx

    doc_idx = headers.get("document")
    acr_idx = headers.get("acronym")

    if doc_idx is None or acr_idx is None:
        wb.close()
        return {}

    mapping: dict[str, list[str]] = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        doc_val = row[doc_idx] if doc_idx < len(row) else None
        acr_val = row[acr_idx] if acr_idx < len(row) else None
        if doc_val and acr_val:
            doc_key = normalize_doc_name(str(doc_val).strip())
            acronym = str(acr_val).strip()
            mapping.setdefault(doc_key, []).append(acronym)

    wb.close()
    if mapping:
        print(f"  Loaded acronym definitions for {len(mapping)} documents from {os.path.basename(acronym_file)}")
    return mapping


def _resolve_acronyms(acronym_mapping: dict, doc_name: str) -> list[str]:
    """Resolve acronyms for a document name using shared matching logic."""
    if not acronym_mapping:
        return []

    for key, acronyms in acronym_mapping.items():
        if match_doc_name(doc_name, key):
            return acronyms

    return []


# ============================================================================
# Tag template generator
# ============================================================================

def generate_tag_template(config: dict, output_path: str):
    """Generate a Doc_Tags.xlsx template pre-populated with names from Doc_URL.xlsx."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill

    url_mapping = load_url_mapping(config)
    if not url_mapping:
        print("  No URL mappings found. Creating empty template.")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Tags"

    # Header styling
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=11)

    headers = ["Document_Name", "Tags"]
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font

    # Populate document names from URL mapping
    for row_idx, name in enumerate(sorted(url_mapping.keys()), 2):
        ws.cell(row=row_idx, column=1, value=name)
        ws.cell(row=row_idx, column=2, value="")

    # Auto-width columns
    ws.column_dimensions["A"].width = 60
    ws.column_dimensions["B"].width = 80

    # Freeze header row
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    ensure_output_dir(os.path.dirname(output_path) or ".")
    wb.save(output_path)
    print(f"\n  Tag template created: {output_path}")
    print(f"  {len(url_mapping)} document names pre-populated from Doc_URL.xlsx")
    print("  Fill in the 'Tags' column with comma-separated tags per document.")


# ============================================================================
# DocxToJsonl converter class
# ============================================================================

class DocxToJsonl:
    """Converts a single .docx file to chunked JSONL records."""

    def __init__(
        self,
        filepath: str,
        config: dict,
        output_dir: str,
        url_mapping: dict,
        tag_mapping: dict,
        acronym_mapping: Optional[dict] = None,
        max_chunk_chars: int = 1500,
        overlap_words: int = 30,
        min_chunk_chars: int = 100,
    ):
        self.filepath = filepath
        self.filename = os.path.basename(filepath)
        self.config = config
        self.output_dir = output_dir
        self.url_mapping = url_mapping
        self.tag_mapping = tag_mapping
        self.acronym_mapping = acronym_mapping or {}
        self.max_chunk_chars = max_chunk_chars
        self.overlap_words = overlap_words
        self.min_chunk_chars = min_chunk_chars

        # Heading detection config
        self.re_h1, self.re_h2, self.re_h3 = _build_heading_patterns(config)
        self.custom_style_map = _build_custom_style_map(config)
        hcfg = config.get("headings", {})
        self.default_heading_level = hcfg.get("default_heading_level", 2)
        self.fake_heading_max_chars = hcfg.get("fake_heading_max_chars_fixer", 120)

        # State
        self.chunks: list[dict] = []
        self.current_parts: list[str] = []
        self.current_heading: str = "(Document Start)"
        self.current_heading_level: int = 0
        self.image_counter: int = 0
        self.warnings: list[str] = []
        self.stats = {
            "paragraphs": 0,
            "tables": 0,
            "images": 0,
            "headings": 0,
            "chunks": 0,
        }

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def convert(self) -> dict:
        """Convert the document and write the .jsonl.txt file. Returns a result dict."""
        start = time.time()

        try:
            doc = Document(self.filepath)
        except Exception as exc:
            msg = f"Cannot open document: {exc}"
            self.warnings.append(msg)
            return self._result("error", time.time() - start, msg)

        # Build element maps for interleaved iteration
        para_map = {id(p._element): p for p in doc.paragraphs}
        table_map = {id(t._element): t for t in doc.tables}

        # Walk body children in document order
        for child in doc.element.body:
            child_id = id(child)
            try:
                if child_id in para_map:
                    self._process_paragraph(para_map[child_id], doc)
                elif child_id in table_map:
                    self._process_table(table_map[child_id])
            except Exception as exc:
                self.warnings.append(f"Element processing error: {exc}")
                self.current_parts.append(f"[Processing error: {exc}]")

        # Flush remaining content
        self._flush_chunk()

        # Merge tiny trailing chunk if below min threshold
        if len(self.chunks) >= 2:
            last = self.chunks[-1]
            if len(last["text"]) < self.min_chunk_chars:
                prev = self.chunks[-2]
                prev["text"] = prev["text"].rstrip() + "\n\n" + last["text"]
                prev["char_count"] = len(prev["text"])
                self.chunks.pop()

        # Resolve metadata and backfill total_chunks
        stem = os.path.splitext(self.filename)[0]
        doc_name = stem.replace("_", " ")
        url = _resolve_url_for_jsonl(self.url_mapping, self.filename, self.config)
        tags = _resolve_tags(self.tag_mapping, self.filename)
        total = len(self.chunks)

        for idx, chunk in enumerate(self.chunks):
            chunk["chunk_id"] = f"{sanitize_filename(stem, max_len=80)}_{idx + 1:03d}"
            chunk["doc_name"] = doc_name
            chunk["source_file"] = self.filename
            chunk["PublishedURL"] = url
            chunk["tags"] = tags
            chunk["acronyms"] = _resolve_acronyms(self.acronym_mapping, self.filename)
            chunk["chunk_index"] = idx
            chunk["total_chunks"] = total

        self.stats["chunks"] = total

        # Write output
        safe_name = sanitize_filename(stem, max_len=100)
        out_path = os.path.join(self.output_dir, f"{safe_name}.jsonl.txt")
        ensure_output_dir(self.output_dir)

        with open(out_path, "w", encoding="utf-8") as f:
            for chunk in self.chunks:
                f.write(json.dumps(chunk, ensure_ascii=False) + "\n")

        elapsed = time.time() - start
        status = "warning" if self.warnings else "success"
        return self._result(status, elapsed, output_file=f"{safe_name}.jsonl.txt")

    # ------------------------------------------------------------------
    # Paragraph processing
    # ------------------------------------------------------------------

    def _process_paragraph(self, para, doc):
        """Process a single paragraph."""
        self.stats["paragraphs"] += 1
        text = para.text.strip()

        # Image detection
        if self._has_image(para):
            img_text = self._handle_image(para)
            if img_text:
                self.current_parts.append(img_text)

        # Empty paragraph (skip)
        if not text and not self._has_image(para):
            return

        # Heading detection
        heading_level = self._detect_heading_level(para)
        if heading_level:
            self.stats["headings"] += 1
            # Flush current chunk before starting new section
            self._flush_chunk()
            self.current_heading = _clean_text(text)
            self.current_heading_level = heading_level
            return

        # Regular paragraph / list item -- extract plain text
        rendered = self._render_plain_text(para, doc)
        if rendered:
            self.current_parts.append(rendered)
            self._split_chunk_if_needed()

    def _detect_heading_level(self, para) -> Optional[int]:
        """Return heading level (1-6) or None."""
        style = para.style
        style_name = style.name if style else ""

        # Built-in heading style
        level = get_heading_level(style)
        if level:
            return level

        # Custom style map
        if style_name:
            mapped = self.custom_style_map.get(style_name.strip().lower())
            if mapped:
                m = re.match(r"Heading (\d)", mapped)
                if m:
                    return int(m.group(1))

        # Fake heading detection
        if _is_fake_heading(para, self.fake_heading_max_chars):
            return _determine_heading_level(
                para.text, self.re_h1, self.re_h2, self.re_h3,
                self.default_heading_level
            )

        return None

    # ------------------------------------------------------------------
    # Plain text rendering (no markdown formatting -- RAG-optimised)
    # ------------------------------------------------------------------

    def _render_plain_text(self, para, doc) -> str:
        """Extract plain text from paragraph, stripping formatting markers."""
        parts: list[str] = []

        for child in para._element:
            tag = child.tag

            if tag == qn("w:r"):
                # Extract text from run
                for t in child.findall(qn("w:t")):
                    if t.text:
                        parts.append(t.text)

            elif tag == qn("w:hyperlink"):
                # Extract display text from hyperlink runs
                for run_elem in child.findall(qn("w:r")):
                    for t in run_elem.findall(qn("w:t")):
                        if t.text:
                            parts.append(t.text)

        text = "".join(parts)
        return _clean_text(text).strip()

    # ------------------------------------------------------------------
    # Table processing
    # ------------------------------------------------------------------

    def _process_table(self, table):
        """Convert table to pipe-format markdown and add to current chunk."""
        self.stats["tables"] += 1
        table_text = self._table_to_text(table)
        if not table_text:
            return

        # If adding this table would exceed limit, flush first
        current_len = sum(len(p) for p in self.current_parts)
        if current_len > 0 and current_len + len(table_text) > self.max_chunk_chars:
            self._flush_chunk()

        self.current_parts.append(table_text)
        self._split_chunk_if_needed()

    def _table_to_text(self, table) -> str:
        """Render table as markdown pipe format."""
        if not table.rows:
            return ""

        rows_data = []
        for row in table.rows:
            cells = []
            seen = set()
            for cell in row.cells:
                cid = id(cell._tc)
                if cid not in seen:
                    seen.add(cid)
                    cell_text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
                    cells.append(cell_text)
            rows_data.append(cells)

        if not rows_data:
            return ""

        # Normalize column count
        max_cols = max(len(r) for r in rows_data)
        for r in rows_data:
            while len(r) < max_cols:
                r.append("")

        lines = []
        # Header row
        lines.append("| " + " | ".join(rows_data[0]) + " |")
        lines.append("| " + " | ".join("---" for _ in range(max_cols)) + " |")
        # Data rows
        for row_cells in rows_data[1:]:
            lines.append("| " + " | ".join(row_cells) + " |")

        return "\n".join(lines)

    # ------------------------------------------------------------------
    # Image handling
    # ------------------------------------------------------------------

    def _has_image(self, para) -> bool:
        """Check if paragraph contains an image element."""
        xml_str = para._element.xml
        return "a:blip" in xml_str or "v:imagedata" in xml_str

    def _handle_image(self, para) -> str:
        """Return a placeholder string for images in a paragraph."""
        self.image_counter += 1
        self.stats["images"] += 1

        # Try to get alt text
        alt_text = ""
        for doc_pr in para._element.findall(".//" + qn("wp:docPr")):
            descr = doc_pr.get("descr", "")
            name = doc_pr.get("name", "")
            alt_text = descr or name
            break

        if not alt_text:
            alt_text = f"Image {self.image_counter}"

        # Check for SmartArt
        if "dgm:relIds" in para._element.xml:
            return f"[DIAGRAM: SmartArt - content not extractable]"

        return f"[IMAGE: {alt_text}]"

    # ------------------------------------------------------------------
    # Chunking logic
    # ------------------------------------------------------------------

    def _current_text(self) -> str:
        """Join current parts into chunk text."""
        return "\n".join(self.current_parts).strip()

    def _flush_chunk(self):
        """Finalize the current chunk and start a new one."""
        text = self._current_text()
        if not text:
            self.current_parts = []
            return

        # Build overlap prefix from previous chunk
        overlap_prefix = ""
        if self.overlap_words > 0 and self.chunks:
            prev_text = self.chunks[-1]["text"]
            words = prev_text.split()
            if len(words) > self.overlap_words:
                overlap_prefix = " ".join(words[-self.overlap_words:]) + "\n\n"

        chunk = {
            "heading": self.current_heading,
            "heading_level": self.current_heading_level,
            "text": overlap_prefix + text,
            "char_count": len(overlap_prefix + text),
        }
        self.chunks.append(chunk)
        self.current_parts = []

    def _split_chunk_if_needed(self):
        """If current accumulated text exceeds max, split at paragraph boundary."""
        text = self._current_text()
        if len(text) <= self.max_chunk_chars:
            return

        # Find the last paragraph boundary that fits
        accumulated = ""
        split_at = 0
        for i, part in enumerate(self.current_parts):
            candidate = "\n".join(self.current_parts[: i + 1]).strip()
            if len(candidate) <= self.max_chunk_chars:
                split_at = i + 1
                accumulated = candidate
            else:
                break

        if split_at == 0:
            # Single part exceeds limit -- flush it as-is
            split_at = 1
            accumulated = self.current_parts[0]

        # Flush the portion that fits
        remaining = self.current_parts[split_at:]
        self.current_parts = self.current_parts[:split_at]
        self._flush_chunk()

        # Continue with the remaining parts (inherit same heading)
        self.current_parts = remaining

        # Check if remaining still exceeds (recursive split)
        if remaining:
            self._split_chunk_if_needed()

    # ------------------------------------------------------------------
    # Result helpers
    # ------------------------------------------------------------------

    def _result(self, status: str, elapsed: float, error: str = "", output_file: str = "") -> dict:
        return {
            "file": self.filename,
            "status": status,
            "elapsed": round(elapsed, 2),
            "error": error,
            "output_file": output_file,
            "stats": dict(self.stats),
            "warnings": list(self.warnings),
        }


# ============================================================================
# Main
# ============================================================================

def main():
    parser = setup_argparse(
        "Convert DOCX files to chunked JSONL (*.jsonl.txt) for RAG ingestion."
    )
    parser.add_argument(
        "--tags", default=None,
        help="Path to tag config Excel file (default: input/Doc_Tags.xlsx)"
    )
    parser.add_argument(
        "--max-chunk-chars", type=int, default=1500,
        help="Maximum characters per chunk (default: 1500)"
    )
    parser.add_argument(
        "--overlap-words", type=int, default=30,
        help="Number of overlap words between consecutive chunks (default: 30)"
    )
    parser.add_argument(
        "--min-chunk-chars", type=int, default=100,
        help="Merge chunks smaller than this into the previous chunk (default: 100)"
    )
    parser.add_argument(
        "--generate-tags", action="store_true",
        help="Generate a Doc_Tags.xlsx template from Doc_URL.xlsx and exit"
    )
    args = parser.parse_args()

    # Load config
    config = load_config(args.config) if args.config else {}
    d2j = config.get("docx2jsonl", {})

    # Apply config defaults to CLI args (CLI overrides config)
    if args.max_chunk_chars == 1500 and d2j.get("max_chunk_chars"):
        args.max_chunk_chars = d2j["max_chunk_chars"]
    if args.overlap_words == 30 and d2j.get("overlap_words"):
        args.overlap_words = d2j["overlap_words"]
    if args.min_chunk_chars == 100 and d2j.get("min_chunk_chars"):
        args.min_chunk_chars = d2j["min_chunk_chars"]

    # Handle --generate-tags
    if args.generate_tags:
        tag_output = args.tags or os.path.join(
            _PROJECT_ROOT, "input", "Doc_Tags.xlsx"
        )
        generate_tag_template(config, tag_output)
        return

    # Resolve input directory — supports Pure Conversion vs Optimized toggle
    input_dir = args.input_dir
    if not input_dir:
        if d2j.get("pure_conversion", False):
            # Pure Conversion: read from original input/
            input_dir = config.get("input", {}).get("directory", "./input")
        else:
            # Optimized: read from a pipeline step's output
            source_step = d2j.get("optimized_source_step", "heading_fixes")
            step_dir = config.get("output", {}).get(source_step, {}).get("directory", "")
            if step_dir:
                output_root = config.get("output", {}).get("directory", "./output")
                input_dir = os.path.join(output_root, step_dir)
            else:
                input_dir = config.get("input", {}).get("directory", "./input")
        input_dir = resolve_path(config, input_dir)

    # Resolve output directory
    output_dir = args.output_dir
    if not output_dir:
        output_dir = d2j.get("output_directory", "./output/8 - jsonl")
        output_dir = resolve_path(config, output_dir)

    if not os.path.isdir(input_dir):
        print(f"ERROR: Input directory not found: {input_dir}")
        sys.exit(1)

    ensure_output_dir(output_dir)

    print("=" * 60)
    print("docx2jsonl -- DOCX to JSONL Converter")
    print("=" * 60)
    print(f"  Input:           {input_dir}")
    print(f"  Output:          {output_dir}")
    print(f"  Max chunk chars: {args.max_chunk_chars}")
    print(f"  Overlap words:   {args.overlap_words}")
    print(f"  Min chunk chars: {args.min_chunk_chars}")

    # Load mappings
    print("\nLoading mappings...")
    url_mapping = load_url_mapping(config)

    tag_file = args.tags
    if not tag_file:
        tag_file = d2j.get("tag_file", "")
        if tag_file:
            tag_file = resolve_path(config, tag_file)
    if not tag_file:
        default_tag_path = os.path.join(_PROJECT_ROOT, "input", "Doc_Tags.xlsx")
        if os.path.isfile(default_tag_path):
            tag_file = default_tag_path
    tag_mapping = load_tag_mapping(tag_file) if tag_file else {}

    # Load acronym definitions (per-document matching)
    acronym_file = d2j.get("acronym_definitions_file", "")
    if acronym_file:
        acronym_file = resolve_path(config, acronym_file)
    acronym_mapping = load_acronym_mapping(acronym_file) if acronym_file else {}

    # Find input files
    # When reading from step output, only exclude temp files (not _fixed, _optimized etc.)
    exclude_ovr = ["~$"] if not d2j.get("pure_conversion", False) else None
    docx_files = iter_docx_files(input_dir, config, exclude_override=exclude_ovr)
    if not docx_files:
        print("\nNo .docx files found in input directory.")
        return

    print(f"\nProcessing {len(docx_files)} document(s)...\n")

    # Process each file
    results = []
    success_count = 0
    error_count = 0
    total_chunks = 0

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        print(f"  {filename} ... ", end="", flush=True)

        try:
            converter = DocxToJsonl(
                filepath=filepath,
                config=config,
                output_dir=output_dir,
                url_mapping=url_mapping,
                tag_mapping=tag_mapping,
                acronym_mapping=acronym_mapping,
                max_chunk_chars=args.max_chunk_chars,
                overlap_words=args.overlap_words,
                min_chunk_chars=args.min_chunk_chars,
            )
            result = converter.convert()
            results.append(result)

            if result["status"] == "error":
                error_count += 1
                print(f"ERROR: {result['error']}")
            else:
                success_count += 1
                chunks = result["stats"]["chunks"]
                total_chunks += chunks
                print(f"{chunks} chunks ({result['elapsed']}s)")
                if result["warnings"]:
                    for w in result["warnings"]:
                        print(f"    WARNING: {w}")
        except Exception as exc:
            error_count += 1
            print(f"ERROR: {exc}")
            traceback.print_exc()

    # Summary
    print("\n" + "=" * 60)
    print(f"Complete: {success_count} succeeded, {error_count} failed")
    print(f"Total chunks generated: {total_chunks}")
    print(f"Output directory: {output_dir}")
    print("=" * 60)


if __name__ == "__main__":
    main()
