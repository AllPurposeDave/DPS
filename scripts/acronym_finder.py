#!/usr/bin/env python3
"""
ACRONYM FINDER — DPS Pipeline Step 1
======================================
Scans .docx files for acronym candidates. Outputs Excel report.
Part of the DPS pipeline; can also run standalone.

pip install python-docx openpyxl pyyaml
"""

import argparse
import os
import re
import sys
import glob
from collections import defaultdict
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Add scripts/ to path for shared_utils import
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from shared_utils import load_config as _load_dps_config


def load_config_from_dps(config, input_dir, output_dir):
    """Build internal config dict from the DPS master config."""
    af_cfg = config.get("acronym_finder", {})

    output_file = af_cfg.get("output_file", "acronym_audit.xlsx")
    cfg = {
        "input_folder": input_dir,
        "output_file": os.path.join(output_dir, output_file),
        "_exclude_patterns": config.get("input", {}).get("exclude_patterns", ["~$"]),
    }

    # Search settings with defaults
    search = af_cfg.get("search", {})
    cfg["search"] = {
        "min_length": search.get("min_length", 2),
        "max_length": search.get("max_length", 10),
        "scan_tables": search.get("scan_tables", True),
        "scan_headers_footers": search.get("scan_headers_footers", True),
        "scan_textboxes": search.get("scan_textboxes", True),
        "min_global_occurrences": search.get("min_global_occurrences", 1),
        "min_doc_occurrences": search.get("min_doc_occurrences", 1),
    }

    # Pattern settings with defaults
    patterns = af_cfg.get("patterns", {})
    cfg["patterns"] = {
        "pure_caps": patterns.get("pure_caps", True),
        "caps_with_numbers": patterns.get("caps_with_numbers", True),
        "caps_with_hyphens": patterns.get("caps_with_hyphens", True),
        "caps_with_slashes": patterns.get("caps_with_slashes", True),
        "parenthetical_defs": patterns.get("parenthetical_defs", True),
    }

    cfg["ignore_list"] = af_cfg.get("ignore_list", [])
    return cfg


def build_regex(cfg):
    parts = []
    p = cfg['patterns']
    mn, mx = cfg['search']['min_length'], cfg['search']['max_length']
    if p['pure_caps']:
        parts.append(rf'[A-Z]{{{mn},{mx}}}')
    if p['caps_with_numbers']:
        parts.append(rf'[A-Z][A-Z]{{{mn - 1},{mx - 1}}}')
    if p['caps_with_hyphens']:
        parts.append(rf'[A-Z][A-Z\-]{{{mn - 1},{mx - 1}}}')
    if p['caps_with_slashes']:
        parts.append(rf'[A-Z]{{2,}}/[A-Z]{{2,}}')
    combined = '|'.join(parts)
    acronym_re = re.compile(rf'\b({combined})\b')
    paren_re = None
    if p['parenthetical_defs']:
        paren_re = re.compile(r'([A-Za-z\s\-]+?)\s*\(([A-Z][A-Z\-/]{1,9})\)')
    return acronym_re, paren_re


def extract_text_from_docx(doc_path, cfg):
    doc = Document(doc_path)
    sources = []

    # Body paragraphs
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            sources.append(('body', txt))

    # Tables
    if cfg['search']['scan_tables']:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        sources.append(('table', txt))

    # Headers/footers
    if cfg['search']['scan_headers_footers']:
        for section in doc.sections:
            for hdr in [section.header, section.first_page_header, section.even_page_header]:
                if hdr and hdr.is_linked_to_previous is False:
                    for para in hdr.paragraphs:
                        txt = para.text.strip()
                        if txt:
                            sources.append(('header', txt))
            for ftr in [section.footer, section.first_page_footer, section.even_page_footer]:
                if ftr and ftr.is_linked_to_previous is False:
                    for para in ftr.paragraphs:
                        txt = para.text.strip()
                        if txt:
                            sources.append(('footer', txt))

    # Text boxes (from XML fallback)
    if cfg['search']['scan_textboxes']:
        try:
            body_xml = doc.element.body
            for txbx in body_xml.iter(qn('w:txbxContent')):
                for p in txbx.iter(qn('w:p')):
                    runs = [r.text for r in p.iter(qn('w:r')) if r.text]
                    txt = ''.join(runs).strip()
                    if txt:
                        sources.append(('textbox', txt))
        except Exception:
            pass

    return sources


def find_acronyms(text_sources, acronym_re, paren_re, ignore_set, cfg):
    acronyms = defaultdict(lambda: {'count': 0, 'locations': set(), 'definitions': set()})
    min_len = cfg['search']['min_length']

    for source_type, text in text_sources:
        # Standard acronym matches
        for m in acronym_re.finditer(text):
            candidate = m.group(0).strip('-').strip('/')
            if len(candidate) < min_len:
                continue
            if any(c.isdigit() for c in candidate):
                continue
            if candidate.upper() in ignore_set:
                continue
            acronyms[candidate]['count'] += 1
            acronyms[candidate]['locations'].add(source_type)

        # Parenthetical definitions
        if paren_re:
            for m in paren_re.finditer(text):
                full_form = m.group(1).strip()
                short_form = m.group(2).strip('-').strip('/')
                if len(short_form) < min_len:
                    continue
                if any(c.isdigit() for c in short_form):
                    continue
                if short_form.upper() in ignore_set:
                    continue
                acronyms[short_form]['count'] += 1
                acronyms[short_form]['locations'].add(source_type)
                if full_form and len(full_form) > len(short_form):
                    acronyms[short_form]['definitions'].add(full_form)

    return acronyms


def process_all_docs(cfg):
    folder = cfg['input_folder']
    pattern = os.path.join(folder, '**', '*.docx')
    exclude_patterns = cfg.get('_exclude_patterns', ['~$'])

    def is_excluded(filepath):
        name = os.path.basename(filepath).lower()
        return any(pat.lower() in name for pat in exclude_patterns)

    files = [f for f in glob.glob(pattern, recursive=True) if not is_excluded(f)]

    if not files:
        print(f"ERROR: No .docx files found in '{folder}'")
        sys.exit(1)

    dps_note = f" (input from {cfg['_dps_config']})" if cfg.get('_dps_config') else ''
    print(f"Found {len(files)} .docx files in '{folder}'{dps_note}")

    ignore_set = {item.upper() for item in cfg.get('ignore_list', [])}
    acronym_re, paren_re = build_regex(cfg)

    # Per-doc results
    all_doc_results = {}
    # Global aggregation
    global_acronyms = defaultdict(lambda: {
        'total_count': 0, 'doc_count': 0, 'docs': [],
        'locations': set(), 'definitions': set()
    })

    for filepath in sorted(files):
        fname = os.path.basename(filepath)
        print(f"  Scanning: {fname}")
        try:
            text_sources = extract_text_from_docx(filepath, cfg)
            doc_acronyms = find_acronyms(text_sources, acronym_re, paren_re, ignore_set, cfg)

            min_doc_occ = cfg['search']['min_doc_occurrences']
            filtered = {k: v for k, v in doc_acronyms.items() if v['count'] >= min_doc_occ}
            all_doc_results[fname] = filtered

            for acr, data in filtered.items():
                global_acronyms[acr]['total_count'] += data['count']
                global_acronyms[acr]['doc_count'] += 1
                global_acronyms[acr]['docs'].append(fname)
                global_acronyms[acr]['locations'].update(data['locations'])
                global_acronyms[acr]['definitions'].update(data['definitions'])

        except Exception as e:
            print(f"  ERROR on {fname}: {e}")
            all_doc_results[fname] = {'_ERROR': {'count': 0, 'locations': set(), 'definitions': {str(e)}}}

    # Apply global minimum
    min_global = cfg['search']['min_global_occurrences']
    global_acronyms = {k: v for k, v in global_acronyms.items() if v['total_count'] >= min_global}

    return all_doc_results, global_acronyms, files


def write_excel(all_doc_results, global_acronyms, files, cfg):
    wb = Workbook()

    # --- Styles ---
    header_font = Font(name='Arial', bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill('solid', fgColor='2F5496')
    data_font = Font(name='Arial', size=10)
    warn_fill = PatternFill('solid', fgColor='FFF2CC')
    count_high_fill = PatternFill('solid', fgColor='FCE4EC')
    border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    center = Alignment(horizontal='center', vertical='center')
    wrap = Alignment(vertical='top', wrap_text=True)

    def style_header(ws, row_num, col_count):
        for col in range(1, col_count + 1):
            cell = ws.cell(row=row_num, column=col)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center
            cell.border = border

    def style_data(ws, row_num, col_count):
        for col in range(1, col_count + 1):
            cell = ws.cell(row=row_num, column=col)
            cell.font = data_font
            cell.border = border

    # =============================================
    # SHEET 1: ACRONYM DEFINITIONS (primary — consumers read this)
    # One row per document-acronym pair. Users curate here.
    # Status column lets reviewers mark rows without deleting them.
    # Only rows with Status=Confirmed (or blank) are consumed downstream.
    # =============================================
    ws1 = wb.active
    ws1.title = "Acronym Definitions"
    headers1 = ["Document", "Acronym", "Definition", "Status", "Notes"]
    for col, h in enumerate(headers1, 1):
        ws1.cell(row=1, column=col, value=h)
    style_header(ws1, 1, len(headers1))

    # Instruction row (light gray background, italic) — explains how to use this sheet
    instruct_fill = PatternFill('solid', fgColor='F2F2F2')
    instruct_font = Font(name='Arial', size=9, italic=True, color='666666')
    instructions = [
        "(source file)",
        "(acronym found)",
        "Yellow = no definition detected. Fill in manually.",
        "Set to Confirmed, False Positive, or Needs Review",
        "(optional reviewer notes)",
    ]
    for col, txt in enumerate(instructions, 1):
        cell = ws1.cell(row=2, column=col, value=txt)
        cell.font = instruct_font
        cell.fill = instruct_fill
        cell.alignment = wrap
        cell.border = border

    # Data validation dropdown for Status column (column D)
    from openpyxl.worksheet.datavalidation import DataValidation
    status_dv = DataValidation(
        type="list",
        formula1='"Confirmed,False Positive,Needs Review"',
        allow_blank=True,
        showErrorMessage=True,
        errorTitle="Invalid Status",
        error="Use: Confirmed, False Positive, or Needs Review",
    )
    status_dv.prompt = "Mark this acronym as Confirmed, False Positive, or Needs Review"
    status_dv.promptTitle = "Review Status"

    row = 3  # Data starts after header + instruction row
    for fname in sorted(all_doc_results.keys()):
        doc_acrs = all_doc_results[fname]
        for acr, data in sorted(doc_acrs.items(), key=lambda x: x[1]['count'], reverse=True):
            if acr == '_ERROR':
                continue
            ws1.cell(row=row, column=1, value=fname)
            ws1.cell(row=row, column=2, value=acr)
            defs = '; '.join(sorted(data['definitions'])) if data['definitions'] else ''
            ws1.cell(row=row, column=3, value=defs)
            ws1.cell(row=row, column=4, value='')  # Status — blank = unreviewed
            ws1.cell(row=row, column=5, value='')
            style_data(ws1, row, len(headers1))
            ws1.cell(row=row, column=3).alignment = wrap
            if not data['definitions']:
                ws1.cell(row=row, column=3).fill = warn_fill
            row += 1

    # Apply status dropdown to all data rows
    if row > 3:
        status_dv.sqref = f"D3:D{row - 1}"
        ws1.add_data_validation(status_dv)

    ws1.column_dimensions['A'].width = 40
    ws1.column_dimensions['B'].width = 16
    ws1.column_dimensions['C'].width = 50
    ws1.column_dimensions['D'].width = 18
    ws1.column_dimensions['E'].width = 30
    ws1.auto_filter.ref = f"A1:E{row - 1}"
    ws1.freeze_panes = 'A2'

    # =============================================
    # SHEET 2: PER-DOC BREAKDOWN (occurrence details for reference)
    # =============================================
    ws2 = wb.create_sheet("Per Document")
    headers2 = ["Document", "Acronym", "Occurrences", "Found In", "Definition(s) Detected"]
    for col, h in enumerate(headers2, 1):
        ws2.cell(row=1, column=col, value=h)
    style_header(ws2, 1, len(headers2))

    row = 2
    for fname in sorted(all_doc_results.keys()):
        doc_acrs = all_doc_results[fname]
        for acr, data in sorted(doc_acrs.items(), key=lambda x: x[1]['count'], reverse=True):
            if acr == '_ERROR':
                ws2.cell(row=row, column=1, value=fname)
                ws2.cell(row=row, column=2, value='ERROR')
                ws2.cell(row=row, column=5, value='; '.join(data['definitions']))
                style_data(ws2, row, len(headers2))
                row += 1
                continue
            ws2.cell(row=row, column=1, value=fname)
            ws2.cell(row=row, column=2, value=acr)
            ws2.cell(row=row, column=3, value=data['count'])
            ws2.cell(row=row, column=4, value=', '.join(sorted(data['locations'])))
            defs = '; '.join(sorted(data['definitions'])) if data['definitions'] else ''
            ws2.cell(row=row, column=5, value=defs)
            style_data(ws2, row, len(headers2))
            ws2.cell(row=row, column=5).alignment = wrap
            if not data['definitions']:
                ws2.cell(row=row, column=5).fill = warn_fill
            row += 1

    ws2.column_dimensions['A'].width = 40
    ws2.column_dimensions['B'].width = 16
    ws2.column_dimensions['C'].width = 14
    ws2.column_dimensions['D'].width = 22
    ws2.column_dimensions['E'].width = 50
    ws2.auto_filter.ref = f"A1:E{row - 1}"
    ws2.freeze_panes = 'A2'

    # =============================================
    # SHEET 3: GLOBAL SUMMARY (all acronyms across all docs)
    # =============================================
    ws3g = wb.create_sheet("Global Summary")
    headers3g = ["Acronym", "Total Occurrences", "Found In # Docs", "Documents", "Found In", "Definition(s) Detected"]
    for col, h in enumerate(headers3g, 1):
        ws3g.cell(row=1, column=col, value=h)
    style_header(ws3g, 1, len(headers3g))

    sorted_global = sorted(global_acronyms.items(), key=lambda x: x[1]['total_count'], reverse=True)
    for i, (acr, data) in enumerate(sorted_global, 2):
        ws3g.cell(row=i, column=1, value=acr)
        ws3g.cell(row=i, column=2, value=data['total_count'])
        ws3g.cell(row=i, column=3, value=data['doc_count'])
        ws3g.cell(row=i, column=4, value=', '.join(sorted(data['docs'])))
        ws3g.cell(row=i, column=5, value=', '.join(sorted(data['locations'])))
        defs = '; '.join(sorted(data['definitions'])) if data['definitions'] else ''
        ws3g.cell(row=i, column=6, value=defs)
        style_data(ws3g, i, len(headers3g))
        ws3g.cell(row=i, column=4).alignment = wrap
        ws3g.cell(row=i, column=6).alignment = wrap
        if not data['definitions']:
            ws3g.cell(row=i, column=6).fill = warn_fill
        if data['total_count'] >= 20:
            ws3g.cell(row=i, column=2).fill = count_high_fill

    ws3g.column_dimensions['A'].width = 16
    ws3g.column_dimensions['B'].width = 18
    ws3g.column_dimensions['C'].width = 16
    ws3g.column_dimensions['D'].width = 50
    ws3g.column_dimensions['E'].width = 22
    ws3g.column_dimensions['F'].width = 50
    ws3g.auto_filter.ref = f"A1:F{len(sorted_global) + 1}"
    ws3g.freeze_panes = 'A2'

    # =============================================
    # SHEET 4: UNDEFINED ACRONYMS (no parenthetical definition found)
    # =============================================
    ws3 = wb.create_sheet("Undefined Acronyms")
    headers3 = ["Acronym", "Total Occurrences", "# Docs", "Documents"]
    for col, h in enumerate(headers3, 1):
        ws3.cell(row=1, column=col, value=h)
    style_header(ws3, 1, len(headers3))

    undefined = [(acr, data) for acr, data in sorted_global if not data['definitions']]
    for i, (acr, data) in enumerate(undefined, 2):
        ws3.cell(row=i, column=1, value=acr)
        ws3.cell(row=i, column=2, value=data['total_count'])
        ws3.cell(row=i, column=3, value=data['doc_count'])
        ws3.cell(row=i, column=4, value=', '.join(sorted(data['docs'])))
        style_data(ws3, i, len(headers3))
        ws3.cell(row=i, column=4).alignment = wrap

    ws3.column_dimensions['A'].width = 16
    ws3.column_dimensions['B'].width = 18
    ws3.column_dimensions['C'].width = 10
    ws3.column_dimensions['D'].width = 60
    ws3.auto_filter.ref = f"A1:D{len(undefined) + 1}"
    ws3.freeze_panes = 'A2'

    # =============================================
    # SHEET 5: ACRONYM CROSS-REFERENCE MATRIX
    # =============================================
    ws4 = wb.create_sheet("Cross-Reference Matrix")
    all_acrs = sorted(global_acronyms.keys())
    doc_names = sorted(all_doc_results.keys())

    ws4.cell(row=1, column=1, value="Acronym \\ Document")
    for j, dname in enumerate(doc_names, 2):
        short = Path(dname).stem[:30]
        ws4.cell(row=1, column=j, value=short)
        ws4.cell(row=1, column=j).alignment = Alignment(text_rotation=90, horizontal='center')

    for i, acr in enumerate(all_acrs, 2):
        ws4.cell(row=i, column=1, value=acr)
        ws4.cell(row=i, column=1).font = data_font
        for j, dname in enumerate(doc_names, 2):
            doc_data = all_doc_results.get(dname, {})
            if acr in doc_data and acr != '_ERROR':
                count = doc_data[acr]['count']
                ws4.cell(row=i, column=j, value=count)
                ws4.cell(row=i, column=j).alignment = center
                ws4.cell(row=i, column=j).font = data_font

    style_header(ws4, 1, len(doc_names) + 1)
    ws4.column_dimensions['A'].width = 16
    for j in range(2, len(doc_names) + 2):
        ws4.column_dimensions[get_column_letter(j)].width = 6
    ws4.freeze_panes = 'B2'

    # =============================================
    # SHEET 6: CUSTOM TAGS (empty template for user to fill in)
    # Instruction row explains format; document names pre-populated.
    # =============================================
    ws_ct = wb.create_sheet("Custom Tags")
    ct_headers = ["Document_Name", "Tags"]
    for col, h in enumerate(ct_headers, 1):
        ws_ct.cell(row=1, column=col, value=h)
    style_header(ws_ct, 1, len(ct_headers))

    # Instruction row — tells the user exactly what format to use
    instruct_fill = PatternFill('solid', fgColor='F2F2F2')
    instruct_font = Font(name='Arial', size=9, italic=True, color='666666')
    ct_instructions = [
        "(do not edit — document names from scan)",
        "Enter comma-separated tags, e.g.: access control, authentication, CUI, FedRAMP-High",
    ]
    for col, txt in enumerate(ct_instructions, 1):
        cell = ws_ct.cell(row=2, column=col, value=txt)
        cell.font = instruct_font
        cell.fill = instruct_fill
        cell.alignment = wrap
        cell.border = border

    # Pre-populate document names from scanned files (data starts at row 3)
    for i, fname in enumerate(sorted(all_doc_results.keys()), 3):
        ws_ct.cell(row=i, column=1, value=fname).font = data_font
        ws_ct.cell(row=i, column=2, value="").font = data_font
        style_data(ws_ct, i, len(ct_headers))
    ct_last = max(len(all_doc_results) + 2, 3)
    ws_ct.auto_filter.ref = f"A1:B{ct_last}"
    ws_ct.freeze_panes = 'A2'
    ws_ct.column_dimensions['A'].width = 40
    ws_ct.column_dimensions['B'].width = 60

    # =============================================
    # SHEET 7: CONFIG SNAPSHOT
    # =============================================
    ws5 = wb.create_sheet("Config Used")
    ws5.cell(row=1, column=1, value="Setting")
    ws5.cell(row=1, column=2, value="Value")
    style_header(ws5, 1, 2)
    config_rows = [
        ("Input Folder", cfg['input_folder']),
        ("Input Source", cfg.get('_dps_config', 'acronym_config.yaml (fallback)')),
        ("Output File", cfg['output_file']),
        ("Min Acronym Length", cfg['search']['min_length']),
        ("Max Acronym Length", cfg['search']['max_length']),
        ("Scan Tables", cfg['search']['scan_tables']),
        ("Scan Headers/Footers", cfg['search']['scan_headers_footers']),
        ("Scan Textboxes", cfg['search']['scan_textboxes']),
        ("Min Global Occurrences", cfg['search']['min_global_occurrences']),
        ("Min Per-Doc Occurrences", cfg['search']['min_doc_occurrences']),
        ("Docs Scanned", len(files)),
        ("Unique Acronyms Found", len(global_acronyms)),
        ("Ignore List Size", len(cfg.get('ignore_list', []))),
    ]
    for i, (setting, val) in enumerate(config_rows, 2):
        ws5.cell(row=i, column=1, value=setting).font = data_font
        ws5.cell(row=i, column=2, value=str(val)).font = data_font
        style_data(ws5, i, 2)
    ws5.column_dimensions['A'].width = 28
    ws5.column_dimensions['B'].width = 50

    # Save
    out = cfg['output_file']
    wb.save(out)
    print(f"\nDone. Report saved to: {out}")
    print(f"  {len(global_acronyms)} unique acronyms across {len(files)} documents")
    undefined_count = sum(1 for _, d in global_acronyms.items() if not d['definitions'])
    print(f"  {undefined_count} acronyms with no detected definition (yellow-flagged)")


def main():
    parser = argparse.ArgumentParser(
        description="DPS Step 1 — Acronym Finder: scan .docx files for acronym candidates",
    )
    parser.add_argument("--config", "-c", default=None,
                        help="Path to dps_config.xlsx or .yaml")
    parser.add_argument("--input", "-i", default=None,
                        help="Input directory containing .docx files")
    parser.add_argument("--output", "-o", default=None,
                        help="Output directory for acronym_audit.xlsx")
    args = parser.parse_args()

    # Load DPS config
    config = _load_dps_config(args.config)
    if not config:
        print("ERROR: Could not load DPS config. Use --config to specify path.")
        sys.exit(1)

    # Resolve input/output directories
    config_dir = config.get("_config_dir", os.getcwd())
    input_dir = args.input or os.path.join(
        config_dir, config.get("input", {}).get("directory", "./input")
    )
    output_dir = args.output or os.path.join(
        config_dir, config.get("output", {}).get("directory", "./output"),
        config.get("output", {}).get("acronyms", {}).get("directory", "1 - acronyms"),
    )

    # Ensure output directory exists
    os.makedirs(output_dir, exist_ok=True)

    cfg = load_config_from_dps(config, input_dir, output_dir)
    all_doc_results, global_acronyms, files = process_all_docs(cfg)
    write_excel(all_doc_results, global_acronyms, files, cfg)


if __name__ == '__main__':
    main()
