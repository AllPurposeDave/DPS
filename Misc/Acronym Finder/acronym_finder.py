#!/usr/bin/env python3
"""
ACRONYM FINDER PROFILING SCRIPT
================================
Scans .docx files for acronym candidates. Outputs Excel report.
Uses YAML config for ignore list and search settings.

pip install python-docx openpyxl pyyaml
"""

import os
import re
import sys
import yaml
import glob
from collections import defaultdict
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


def load_config(config_path):
    with open(config_path, 'r') as f:
        cfg = yaml.safe_load(f)
    cfg.setdefault('input_folder', './policy_docs')
    cfg.setdefault('output_file', './acronym_audit.xlsx')
    cfg.setdefault('search', {})
    s = cfg['search']
    s.setdefault('min_length', 2)
    s.setdefault('max_length', 10)
    s.setdefault('scan_tables', True)
    s.setdefault('scan_headers_footers', True)
    s.setdefault('scan_textboxes', True)
    s.setdefault('min_global_occurrences', 1)
    s.setdefault('min_doc_occurrences', 1)
    cfg.setdefault('patterns', {})
    p = cfg['patterns']
    p.setdefault('pure_caps', True)
    p.setdefault('caps_with_numbers', True)
    p.setdefault('caps_with_hyphens', True)
    p.setdefault('caps_with_slashes', True)
    p.setdefault('parenthetical_defs', True)
    cfg.setdefault('ignore_list', [])
    return cfg


def build_regex(cfg):
    parts = []
    p = cfg['patterns']
    mn, mx = cfg['search']['min_length'], cfg['search']['max_length']
    if p['pure_caps']:
        parts.append(rf'[A-Z]{{{mn},{mx}}}')
    if p['caps_with_numbers']:
        parts.append(rf'[A-Z][A-Z0-9]{{{mn - 1},{mx - 1}}}')
    if p['caps_with_hyphens']:
        parts.append(rf'[A-Z][A-Z0-9\-]{{{mn - 1},{mx - 1}}}')
    if p['caps_with_slashes']:
        parts.append(rf'[A-Z]{{2,}}/[A-Z]{{2,}}')
    combined = '|'.join(parts)
    acronym_re = re.compile(rf'\b({combined})\b')
    paren_re = None
    if p['parenthetical_defs']:
        paren_re = re.compile(r'([A-Za-z\s\-]+?)\s*\(([A-Z][A-Z0-9\-/]{1,9})\)')
    return acronym_re, paren_re


def extract_text_from_docx(doc_path, cfg):
    doc = Document(doc_path)
    sources = []

    # Body paragraphs
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            sources.append(('body', txt))

    # Tables
    if cfg['search']['scan_tables']:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        sources.append(('table', txt))

    # Headers/footers
    if cfg['search']['scan_headers_footers']:
        for section in doc.sections:
            for hdr in [section.header, section.first_page_header, section.even_page_header]:
                if hdr and hdr.is_linked_to_previous is False:
                    for para in hdr.paragraphs:
                        txt = para.text.strip()
                        if txt:
                            sources.append(('header', txt))
            for ftr in [section.footer, section.first_page_footer, section.even_page_footer]:
                if ftr and ftr.is_linked_to_previous is False:
                    for para in ftr.paragraphs:
                        txt = para.text.strip()
                        if txt:
                            sources.append(('footer', txt))

    # Text boxes (from XML fallback)
    if cfg['search']['scan_textboxes']:
        try:
            body_xml = doc.element.body
            for txbx in body_xml.iter(qn('w:txbxContent')):
                for p in txbx.iter(qn('w:p')):
                    runs = [r.text for r in p.iter(qn('w:r')) if r.text]
                    txt = ''.join(runs).strip()
                    if txt:
                        sources.append(('textbox', txt))
        except Exception:
            pass

    return sources


def find_acronyms(text_sources, acronym_re, paren_re, ignore_set, cfg):
    acronyms = defaultdict(lambda: {'count': 0, 'locations': set(), 'definitions': set()})
    min_len = cfg['search']['min_length']

    for source_type, text in text_sources:
        # Standard acronym matches
        for m in acronym_re.finditer(text):
            candidate = m.group(0).strip('-').strip('/')
            if len(candidate) < min_len:
                continue
            if candidate.upper() in ignore_set:
                continue
            acronyms[candidate]['count'] += 1
            acronyms[candidate]['locations'].add(source_type)

        # Parenthetical definitions
        if paren_re:
            for m in paren_re.finditer(text):
                full_form = m.group(1).strip()
                short_form = m.group(2).strip('-').strip('/')
                if len(short_form) < min_len:
                    continue
                if short_form.upper() in ignore_set:
                    continue
                acronyms[short_form]['count'] += 1
                acronyms[short_form]['locations'].add(source_type)
                if full_form and len(full_form) > len(short_form):
                    acronyms[short_form]['definitions'].add(full_form)

    return acronyms


def process_all_docs(cfg):
    folder = cfg['input_folder']
    pattern = os.path.join(folder, '**', '*.docx')
    files = [f for f in glob.glob(pattern, recursive=True) if not os.path.basename(f).startswith('~$')]

    if not files:
        print(f"ERROR: No .docx files found in '{folder}'")
        sys.exit(1)

    print(f"Found {len(files)} .docx files in '{folder}'")

    ignore_set = {item.upper() for item in cfg.get('ignore_list', [])}
    acronym_re, paren_re = build_regex(cfg)

    # Per-doc results
    all_doc_results = {}
    # Global aggregation
    global_acronyms = defaultdict(lambda: {
        'total_count': 0, 'doc_count': 0, 'docs': [],
        'locations': set(), 'definitions': set()
    })

    for filepath in sorted(files):
        fname = os.path.basename(filepath)
        print(f"  Scanning: {fname}")
        try:
            text_sources = extract_text_from_docx(filepath, cfg)
            doc_acronyms = find_acronyms(text_sources, acronym_re, paren_re, ignore_set, cfg)

            min_doc_occ = cfg['search']['min_doc_occurrences']
            filtered = {k: v for k, v in doc_acronyms.items() if v['count'] >= min_doc_occ}
            all_doc_results[fname] = filtered

            for acr, data in filtered.items():
                global_acronyms[acr]['total_count'] += data['count']
                global_acronyms[acr]['doc_count'] += 1
                global_acronyms[acr]['docs'].append(fname)
                global_acronyms[acr]['locations'].update(data['locations'])
                global_acronyms[acr]['definitions'].update(data['definitions'])

        except Exception as e:
            print(f"  ERROR on {fname}: {e}")
            all_doc_results[fname] = {'_ERROR': {'count': 0, 'locations': set(), 'definitions': {str(e)}}}

    # Apply global minimum
    min_global = cfg['search']['min_global_occurrences']
    global_acronyms = {k: v for k, v in global_acronyms.items() if v['total_count'] >= min_global}

    return all_doc_results, global_acronyms, files


def write_excel(all_doc_results, global_acronyms, files, cfg):
    wb = Workbook()

    # --- Styles ---
    header_font = Font(name='Arial', bold=True, size=11, color='FFFFFF')
    header_fill = PatternFill('solid', fgColor='2F5496')
    data_font = Font(name='Arial', size=10)
    warn_fill = PatternFill('solid', fgColor='FFF2CC')
    count_high_fill = PatternFill('solid', fgColor='FCE4EC')
    border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    center = Alignment(horizontal='center', vertical='center')
    wrap = Alignment(vertical='top', wrap_text=True)

    def style_header(ws, row_num, col_count):
        for col in range(1, col_count + 1):
            cell = ws.cell(row=row_num, column=col)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center
            cell.border = border

    def style_data(ws, row_num, col_count):
        for col in range(1, col_count + 1):
            cell = ws.cell(row=row_num, column=col)
            cell.font = data_font
            cell.border = border

    # =============================================
    # SHEET 1: GLOBAL SUMMARY (all acronyms across all docs)
    # =============================================
    ws1 = wb.active
    ws1.title = "Global Summary"
    headers1 = ["Acronym", "Total Occurrences", "Found In # Docs", "Documents", "Found In", "Definition(s) Detected"]
    for col, h in enumerate(headers1, 1):
        ws1.cell(row=1, column=col, value=h)
    style_header(ws1, 1, len(headers1))

    sorted_global = sorted(global_acronyms.items(), key=lambda x: x[1]['total_count'], reverse=True)
    for i, (acr, data) in enumerate(sorted_global, 2):
        ws1.cell(row=i, column=1, value=acr)
        ws1.cell(row=i, column=2, value=data['total_count'])
        ws1.cell(row=i, column=3, value=data['doc_count'])
        ws1.cell(row=i, column=4, value=', '.join(sorted(data['docs'])))
        ws1.cell(row=i, column=5, value=', '.join(sorted(data['locations'])))
        defs = '; '.join(sorted(data['definitions'])) if data['definitions'] else ''
        ws1.cell(row=i, column=6, value=defs)
        style_data(ws1, i, len(headers1))
        ws1.cell(row=i, column=4).alignment = wrap
        ws1.cell(row=i, column=6).alignment = wrap
        if not data['definitions']:
            ws1.cell(row=i, column=6).fill = warn_fill
        if data['total_count'] >= 20:
            ws1.cell(row=i, column=2).fill = count_high_fill

    ws1.column_dimensions['A'].width = 16
    ws1.column_dimensions['B'].width = 18
    ws1.column_dimensions['C'].width = 16
    ws1.column_dimensions['D'].width = 50
    ws1.column_dimensions['E'].width = 22
    ws1.column_dimensions['F'].width = 50
    ws1.auto_filter.ref = f"A1:F{len(sorted_global) + 1}"
    ws1.freeze_panes = 'A2'

    # =============================================
    # SHEET 2: PER-DOC BREAKDOWN
    # =============================================
    ws2 = wb.create_sheet("Per Document")
    headers2 = ["Document", "Acronym", "Occurrences", "Found In", "Definition(s) Detected"]
    for col, h in enumerate(headers2, 1):
        ws2.cell(row=1, column=col, value=h)
    style_header(ws2, 1, len(headers2))

    row = 2
    for fname in sorted(all_doc_results.keys()):
        doc_acrs = all_doc_results[fname]
        for acr, data in sorted(doc_acrs.items(), key=lambda x: x[1]['count'], reverse=True):
            if acr == '_ERROR':
                ws2.cell(row=row, column=1, value=fname)
                ws2.cell(row=row, column=2, value='ERROR')
                ws2.cell(row=row, column=5, value='; '.join(data['definitions']))
                style_data(ws2, row, len(headers2))
                row += 1
                continue
            ws2.cell(row=row, column=1, value=fname)
            ws2.cell(row=row, column=2, value=acr)
            ws2.cell(row=row, column=3, value=data['count'])
            ws2.cell(row=row, column=4, value=', '.join(sorted(data['locations'])))
            defs = '; '.join(sorted(data['definitions'])) if data['definitions'] else ''
            ws2.cell(row=row, column=5, value=defs)
            style_data(ws2, row, len(headers2))
            ws2.cell(row=row, column=5).alignment = wrap
            if not data['definitions']:
                ws2.cell(row=row, column=5).fill = warn_fill
            row += 1

    ws2.column_dimensions['A'].width = 40
    ws2.column_dimensions['B'].width = 16
    ws2.column_dimensions['C'].width = 14
    ws2.column_dimensions['D'].width = 22
    ws2.column_dimensions['E'].width = 50
    ws2.auto_filter.ref = f"A1:E{row - 1}"
    ws2.freeze_panes = 'A2'

    # =============================================
    # SHEET 3: UNDEFINED ACRONYMS (no parenthetical definition found)
    # =============================================
    ws3 = wb.create_sheet("Undefined Acronyms")
    headers3 = ["Acronym", "Total Occurrences", "# Docs", "Documents"]
    for col, h in enumerate(headers3, 1):
        ws3.cell(row=1, column=col, value=h)
    style_header(ws3, 1, len(headers3))

    undefined = [(acr, data) for acr, data in sorted_global if not data['definitions']]
    for i, (acr, data) in enumerate(undefined, 2):
        ws3.cell(row=i, column=1, value=acr)
        ws3.cell(row=i, column=2, value=data['total_count'])
        ws3.cell(row=i, column=3, value=data['doc_count'])
        ws3.cell(row=i, column=4, value=', '.join(sorted(data['docs'])))
        style_data(ws3, i, len(headers3))
        ws3.cell(row=i, column=4).alignment = wrap

    ws3.column_dimensions['A'].width = 16
    ws3.column_dimensions['B'].width = 18
    ws3.column_dimensions['C'].width = 10
    ws3.column_dimensions['D'].width = 60
    ws3.auto_filter.ref = f"A1:D{len(undefined) + 1}"
    ws3.freeze_panes = 'A2'

    # =============================================
    # SHEET 4: ACRONYM CROSS-REFERENCE MATRIX
    # =============================================
    ws4 = wb.create_sheet("Cross-Reference Matrix")
    all_acrs = sorted(global_acronyms.keys())
    doc_names = sorted(all_doc_results.keys())

    ws4.cell(row=1, column=1, value="Acronym \\ Document")
    for j, dname in enumerate(doc_names, 2):
        short = Path(dname).stem[:30]
        ws4.cell(row=1, column=j, value=short)
        ws4.cell(row=1, column=j).alignment = Alignment(text_rotation=90, horizontal='center')

    for i, acr in enumerate(all_acrs, 2):
        ws4.cell(row=i, column=1, value=acr)
        ws4.cell(row=i, column=1).font = data_font
        for j, dname in enumerate(doc_names, 2):
            doc_data = all_doc_results.get(dname, {})
            if acr in doc_data and acr != '_ERROR':
                count = doc_data[acr]['count']
                ws4.cell(row=i, column=j, value=count)
                ws4.cell(row=i, column=j).alignment = center
                ws4.cell(row=i, column=j).font = data_font

    style_header(ws4, 1, len(doc_names) + 1)
    ws4.column_dimensions['A'].width = 16
    for j in range(2, len(doc_names) + 2):
        ws4.column_dimensions[get_column_letter(j)].width = 6
    ws4.freeze_panes = 'B2'

    # =============================================
    # SHEET 5: CONFIG SNAPSHOT
    # =============================================
    ws5 = wb.create_sheet("Config Used")
    ws5.cell(row=1, column=1, value="Setting")
    ws5.cell(row=1, column=2, value="Value")
    style_header(ws5, 1, 2)
    config_rows = [
        ("Input Folder", cfg['input_folder']),
        ("Output File", cfg['output_file']),
        ("Min Acronym Length", cfg['search']['min_length']),
        ("Max Acronym Length", cfg['search']['max_length']),
        ("Scan Tables", cfg['search']['scan_tables']),
        ("Scan Headers/Footers", cfg['search']['scan_headers_footers']),
        ("Scan Textboxes", cfg['search']['scan_textboxes']),
        ("Min Global Occurrences", cfg['search']['min_global_occurrences']),
        ("Min Per-Doc Occurrences", cfg['search']['min_doc_occurrences']),
        ("Docs Scanned", len(files)),
        ("Unique Acronyms Found", len(global_acronyms)),
        ("Ignore List Size", len(cfg.get('ignore_list', []))),
    ]
    for i, (setting, val) in enumerate(config_rows, 2):
        ws5.cell(row=i, column=1, value=setting).font = data_font
        ws5.cell(row=i, column=2, value=str(val)).font = data_font
        style_data(ws5, i, 2)
    ws5.column_dimensions['A'].width = 28
    ws5.column_dimensions['B'].width = 50

    # Save
    out = cfg['output_file']
    wb.save(out)
    print(f"\nDone. Report saved to: {out}")
    print(f"  {len(global_acronyms)} unique acronyms across {len(files)} documents")
    undefined_count = sum(1 for _, d in global_acronyms.items() if not d['definitions'])
    print(f"  {undefined_count} acronyms with no detected definition (yellow-flagged)")


def main():
    config_path = sys.argv[1] if len(sys.argv) > 1 else 'acronym_config.yaml'
    if not os.path.exists(config_path):
        print(f"ERROR: Config file not found: {config_path}")
        print("Expected: acronym_config.yaml in the same folder, or pass path as argument.")
        sys.exit(1)
    cfg = load_config(config_path)
    all_doc_results, global_acronyms, files = process_all_docs(cfg)
    write_excel(all_doc_results, global_acronyms, files, cfg)


if __name__ == '__main__':
    main()
