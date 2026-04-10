#!/usr/bin/env python3
"""
xlsx2docx.py — Convert Excel files to formatted Word documents.

Primary use case: controls matrices where each row is a control.
The control ID column becomes a Word heading; remaining columns
(description, guidance, notes, status) render as key-value content below.

All formatting is driven by xlsx2docx_config.yaml — fonts, colors, spacing,
heading styles, headers/footers, cover page, and sheet-to-section mapping.

Usage:
  python xlsx2docx.py
  python xlsx2docx.py --config my_config.yaml
  python xlsx2docx.py --input controls.xlsx
  python xlsx2docx.py --input ../../input/ --output ../../output/xlsx2docx/
"""

from __future__ import annotations

import argparse
import copy
import fnmatch
import logging
import os
import re
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


# ── Default configuration ────────────────────────────────────────────────────

DEFAULT_CONFIG: dict = {
    "input": {
        "files": [],
        "directory": "../../input/",
        "exclude_patterns": ["~$*"],
    },
    "output": {
        "directory": "../../output/xlsx2docx/",
        "filename_template": "{input_stem}.docx",
    },
    "identity": {
        "document_title": "",
        "organization_name": "",
        "classification": "Unclassified",
        "author": "",
        "version": "1.0",
        "date": "",
    },
    "cover_page": {
        "enabled": False,
        "title": "{document_title}",
        "subtitle": "{organization_name}",
        "classification": "{classification}",
        "classification_color": "FF0000",
        "date": "{date}",
        "prepared_by": "",
        "approved_by": "",
        "logo_placeholder": False,
    },
    "toc": {
        "enabled": False,
        "title": "Table of Contents",
        "depth": 2,
        "note": "Press Ctrl+A then F9 in Word to update page numbers after opening.",
    },
    "page": {
        "size": "Letter",
        "orientation": "Portrait",
        "margin_top": 1.0,
        "margin_bottom": 1.0,
        "margin_left": 1.25,
        "margin_right": 1.25,
    },
    "body": {
        "font_name": "Calibri",
        "font_size": 11,
        "line_spacing": 1.15,
        "space_after": 8,
        "alignment": "left",
    },
    "headings": {
        "h1": {
            "font_name": "Calibri", "font_size": 16, "bold": True, "italic": False,
            "color": "2F5496", "space_before": 12, "space_after": 6, "keep_with_next": True,
        },
        "h2": {
            "font_name": "Calibri", "font_size": 14, "bold": True, "italic": False,
            "color": "2F5496", "space_before": 10, "space_after": 4, "keep_with_next": True,
        },
        "h3": {
            "font_name": "Calibri", "font_size": 12, "bold": True, "italic": False,
            "color": "2F5496", "space_before": 8, "space_after": 3, "keep_with_next": False,
        },
        "h4": {
            "font_name": "Calibri", "font_size": 11, "bold": True, "italic": True,
            "color": "2F5496", "space_before": 6, "space_after": 2, "keep_with_next": False,
        },
    },
    "header": {
        "left": "{organization_name}",
        "center": "",
        "right": "{document_title}",
        "font_name": "Arial",
        "font_size": 8,
        "color": "666666",
    },
    "footer": {
        "left": "{classification}",
        "center": "Page {page} of {pages}",
        "right": "{date}",
        "font_name": "Arial",
        "font_size": 8,
        "color": "666666",
    },
    "different_first_page": True,
    "table": {
        "header_font_name": "Arial",
        "header_font_size": 10,
        "header_bold": True,
        "header_font_color": "FFFFFF",
        "header_fill_color": "2F5496",
        "body_font_name": "Calibri",
        "body_font_size": 10,
        "row_banding": True,
        "banding_color": "D6E4F0",
        "border_style": "single",
        "border_color": "CCCCCC",
        "fit_to_page": True,
        "max_col_width": 3.0,
    },
    "include_unmatched": True,
    "sheets": [],
}

# Defaults applied to each sheet rule that doesn't specify a given key.
DEFAULT_SHEET_RULE: dict = {
    "render_as": "table",
    "heading_column": "",
    "heading_level": 1,
    "body_render": "key_value",
    "include_columns": [],
    "rename_columns": {},
    "exclude_columns": [],
    "skip_empty_heading": True,
    "header_row": 1,
    "data_start_row": 2,
    "repeat_header": True,
    "add_section_heading": None,   # None = auto (False for row_as_heading, True otherwise)
    "section_heading_level": 1,
}


# ── Config loading ───────────────────────────────────────────────────────────

def _deep_merge(base: dict, override: dict) -> dict:
    """Recursively merge override into base. Returns a new dict."""
    result = copy.deepcopy(base)
    for key, val in override.items():
        if key in result and isinstance(result[key], dict) and isinstance(val, dict):
            result[key] = _deep_merge(result[key], val)
        else:
            result[key] = copy.deepcopy(val)
    return result


def load_config(config_path: str) -> dict:
    """Load xlsx2docx_config.yaml and deep-merge with DEFAULT_CONFIG."""
    config_path = os.path.abspath(config_path)
    if not os.path.isfile(config_path):
        logging.warning("Config not found at '%s' — using defaults.", config_path)
        return copy.deepcopy(DEFAULT_CONFIG)

    with open(config_path, encoding="utf-8") as fh:
        loaded = yaml.safe_load(fh) or {}

    config = _deep_merge(DEFAULT_CONFIG, loaded)
    logging.info("Loaded config: %s", config_path)
    return config


# ── Token resolution ─────────────────────────────────────────────────────────

def _build_token_map(identity: dict) -> dict[str, str]:
    """Build {token} → value map from identity config."""
    today = date.today().isoformat()
    date_val = str(identity.get("date", "")).strip() or today
    return {
        "{document_title}":    str(identity.get("document_title", "")),
        "{organization_name}": str(identity.get("organization_name", "")),
        "{classification}":    str(identity.get("classification", "Unclassified")),
        "{author}":            str(identity.get("author", "")),
        "{version}":           str(identity.get("version", "1.0")),
        "{date}":              date_val,
        # Aliases used by cover page functions
        "{effective_date}":    date_val,
        "{doc_id}":            str(identity.get("document_id", "")),
    }


def _resolve_tokens(text: str, token_map: dict[str, str]) -> str:
    """Replace {token} placeholders with values from token_map."""
    if not text:
        return text
    for token, value in token_map.items():
        text = text.replace(token, value)
    return text


# ── Color / alignment helpers ─────────────────────────────────────────────────

def _parse_hex_color(hex_str: Any) -> RGBColor:
    """Parse a 6-digit hex string (e.g. '2F5496') to RGBColor. Falls back to DPS blue."""
    try:
        s = str(hex_str).strip().lstrip("#")
        return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
    except (ValueError, IndexError, TypeError):
        logging.warning("Could not parse color '%s' — using default blue.", hex_str)
        return RGBColor(0x2F, 0x54, 0x96)


def _get_alignment(value: Any) -> WD_ALIGN_PARAGRAPH | None:
    """Convert alignment string to WD_ALIGN_PARAGRAPH enum."""
    s = str(value).strip().lower() if value else ""
    return {
        "left":      WD_ALIGN_PARAGRAPH.LEFT,
        "center":    WD_ALIGN_PARAGRAPH.CENTER,
        "right":     WD_ALIGN_PARAGRAPH.RIGHT,
        "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "justify":   WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(s)


# ── Style application ─────────────────────────────────────────────────────────

def _apply_heading_style(doc: Document, style_name: str, hcfg: dict) -> None:
    """Modify a built-in heading style using values from hcfg."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        logging.warning("Style '%s' not found — skipping.", style_name)
        return

    font = style.font
    if "font_name"  in hcfg: font.name       = str(hcfg["font_name"])
    if "font_size"  in hcfg: font.size        = Pt(hcfg["font_size"])
    if "bold"       in hcfg: font.bold        = bool(hcfg["bold"])
    if "italic"     in hcfg: font.italic      = bool(hcfg["italic"])
    if "underline"  in hcfg: font.underline   = bool(hcfg["underline"])
    if "all_caps"   in hcfg: font.all_caps    = bool(hcfg["all_caps"])
    if "color"      in hcfg: font.color.rgb   = _parse_hex_color(hcfg["color"])

    pf = style.paragraph_format
    if "space_before"   in hcfg: pf.space_before   = Pt(hcfg["space_before"])
    if "space_after"    in hcfg: pf.space_after     = Pt(hcfg["space_after"])
    if "keep_with_next" in hcfg: pf.keep_with_next  = bool(hcfg["keep_with_next"])
    if "alignment"      in hcfg:
        align = _get_alignment(hcfg["alignment"])
        if align is not None:
            pf.alignment = align


def _apply_body_defaults(doc: Document, body_cfg: dict) -> None:
    """Apply body font and paragraph settings to the Normal style."""
    try:
        normal = doc.styles["Normal"]
    except KeyError:
        logging.warning("'Normal' style not found — body defaults not applied.")
        return

    if "font_name" in body_cfg: normal.font.name = str(body_cfg["font_name"])
    if "font_size" in body_cfg: normal.font.size = Pt(body_cfg["font_size"])

    pf = normal.paragraph_format
    if "line_spacing" in body_cfg: pf.line_spacing = float(body_cfg["line_spacing"])
    if "space_after"  in body_cfg: pf.space_after  = Pt(body_cfg["space_after"])
    if "alignment"    in body_cfg:
        align = _get_alignment(body_cfg["alignment"])
        if align is not None:
            pf.alignment = align


def _apply_page_layout(doc: Document, page_cfg: dict) -> None:
    """Apply page size, orientation, and margins."""
    section = doc.sections[0]

    # Margins
    if "margin_top"    in page_cfg: section.top_margin    = Inches(page_cfg["margin_top"])
    if "margin_bottom" in page_cfg: section.bottom_margin = Inches(page_cfg["margin_bottom"])
    if "margin_left"   in page_cfg: section.left_margin   = Inches(page_cfg["margin_left"])
    if "margin_right"  in page_cfg: section.right_margin  = Inches(page_cfg["margin_right"])

    # Page size
    size_str = str(page_cfg.get("size", "Letter")).strip().upper()
    if size_str == "A4":
        width, height = Inches(8.27), Inches(11.69)
    else:  # Letter
        width, height = Inches(8.5), Inches(11.0)

    orientation = str(page_cfg.get("orientation", "Portrait")).strip().lower()
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width  = height
        section.page_height = width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width  = width
        section.page_height = height


def _apply_core_properties(doc: Document, identity: dict, token_map: dict) -> None:
    """Write identity values into Word core_properties."""
    props = doc.core_properties
    title = _resolve_tokens(str(identity.get("document_title", "")), token_map)
    if title: props.title = title
    if identity.get("author"):          props.author   = str(identity["author"])
    if identity.get("classification"):  props.subject  = _resolve_tokens(str(identity["classification"]), token_map)
    if identity.get("organization_name"): props.category = str(identity["organization_name"])
    parts = []
    if identity.get("version"):  parts.append(f"Version: {identity['version']}")
    if parts: props.comments = " | ".join(parts)


# ── Header / footer builders ─────────────────────────────────────────────────

def _add_page_number_field(paragraph) -> None:
    """Insert a Word PAGE field code into paragraph."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    run._element.append(begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run3._element.append(end)


def _add_num_pages_field(paragraph) -> None:
    """Insert a Word NUMPAGES field code into paragraph."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    run._element.append(begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " NUMPAGES "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run3._element.append(end)


def _build_tab_stop_hf(hf, left_text: str, center_text: str, right_text: str,
                        font_name: str, font_size: int, color: str, token_map: dict,
                        page_width_inches: float = 6.0) -> None:
    """Build a header/footer paragraph with left / center / right text via tab stops."""
    from docx.enum.text import WD_TAB_ALIGNMENT

    paragraph = hf.paragraphs[0] if hf.paragraphs else hf.add_paragraph()
    for run in paragraph.runs:
        run.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pf = paragraph.paragraph_format
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Inches(page_width_inches / 2), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(page_width_inches),     WD_TAB_ALIGNMENT.RIGHT)

    def _emit(text_val: str):
        resolved = _resolve_tokens(text_val, token_map)
        parts = re.split(r"(\{page\}|\{pages\})", resolved)
        for part in parts:
            if part == "{page}":
                _add_page_number_field(paragraph)
            elif part == "{pages}":
                _add_num_pages_field(paragraph)
            elif part:
                run = paragraph.add_run(part)
                run.font.name  = font_name
                run.font.size  = Pt(font_size)
                run.font.color.rgb = _parse_hex_color(color)

    if left_text:   _emit(left_text)
    paragraph.add_run("\t").font.size = Pt(font_size)
    if center_text: _emit(center_text)
    paragraph.add_run("\t").font.size = Pt(font_size)
    if right_text:  _emit(right_text)


def _apply_header_footer(doc: Document, config: dict, token_map: dict) -> None:
    """Apply header and footer from config to all document sections."""
    header_cfg = config.get("header", {})
    footer_cfg = config.get("footer", {})
    different_first = bool(config.get("different_first_page", True))

    for section in doc.sections:
        section.different_first_page_header_footer = different_first

        usable_inches = (section.page_width - section.left_margin - section.right_margin) / 914400

        hdr = section.header
        hdr.is_linked_to_previous = False
        _build_tab_stop_hf(
            hdr,
            left_text    = str(header_cfg.get("left",      "")),
            center_text  = str(header_cfg.get("center",    "")),
            right_text   = str(header_cfg.get("right",     "")),
            font_name    = str(header_cfg.get("font_name", "Arial")),
            font_size    = int(header_cfg.get("font_size", 8)),
            color        = str(header_cfg.get("color",     "666666")),
            token_map    = token_map,
            page_width_inches = usable_inches,
        )

        ftr = section.footer
        ftr.is_linked_to_previous = False
        _build_tab_stop_hf(
            ftr,
            left_text    = str(footer_cfg.get("left",      "")),
            center_text  = str(footer_cfg.get("center",    "")),
            right_text   = str(footer_cfg.get("right",     "")),
            font_name    = str(footer_cfg.get("font_name", "Arial")),
            font_size    = int(footer_cfg.get("font_size", 8)),
            color        = str(footer_cfg.get("color",     "666666")),
            token_map    = token_map,
            page_width_inches = usable_inches,
        )


# ── Cover page ────────────────────────────────────────────────────────────────

def _build_cover_page(doc: Document, cover_cfg: dict, token_map: dict) -> None:
    """Build an optional title/cover page followed by a section break."""
    if not bool(cover_cfg.get("enabled", False)):
        return

    def _resolve(key: str, default: str = "") -> str:
        return _resolve_tokens(str(cover_cfg.get(key, default)), token_map)

    # Classification banner (top)
    classification = _resolve("classification")
    if classification:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(36)
        run = p.add_run(classification)
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = _parse_hex_color(cover_cfg.get("classification_color", "FF0000"))

    # Logo placeholder
    if bool(cover_cfg.get("logo_placeholder", False)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(48)
        run = p.add_run("[Organization Logo]")
        run.font.size  = Pt(12)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Vertical spacer
    doc.add_paragraph().paragraph_format.space_before = Pt(72)

    # Title
    title_text = _resolve("title")
    if title_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(title_text)
        run.font.size  = Pt(28)
        run.font.bold  = True
        run.font.color.rgb = _parse_hex_color("2F5496")

    # Subtitle
    subtitle = _resolve("subtitle")
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(48)
        run = p.add_run(subtitle)
        run.font.size  = Pt(16)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Prepared By / Approved By
    for field_key, label in [("prepared_by", "Prepared By"), ("approved_by", "Approved By")]:
        val = _resolve(field_key)
        if val:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"{label}: {val}").font.size = Pt(11)

    # Date
    date_text = _resolve("date")
    if date_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        run = p.add_run(date_text)
        run.font.size  = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Classification banner (bottom)
    if classification:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(72)
        run = p.add_run(classification)
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = _parse_hex_color(cover_cfg.get("classification_color", "FF0000"))

    # Section break — new page after cover
    new_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    first   = doc.sections[0]
    new_sec.top_margin    = first.top_margin
    new_sec.bottom_margin = first.bottom_margin
    new_sec.left_margin   = first.left_margin
    new_sec.right_margin  = first.right_margin
    new_sec.page_width    = first.page_width
    new_sec.page_height   = first.page_height
    new_sec.orientation   = first.orientation


# ── Table of Contents ─────────────────────────────────────────────────────────

def _insert_toc(doc: Document, toc_cfg: dict) -> None:
    """Insert a Word TOC field. Requires Ctrl+A → F9 in Word to populate."""
    title = str(toc_cfg.get("title", "Table of Contents"))
    depth = int(toc_cfg.get("depth", 2))
    note  = str(toc_cfg.get("note",  ""))

    # TOC title — styled as a bold paragraph (not a real heading, to stay out of the TOC)
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = _parse_hex_color("2F5496")
        p.paragraph_format.space_after = Pt(6)

    # TOC field code
    para = doc.add_paragraph()
    run  = para.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"),       "true")
    run._element.append(begin)

    run2  = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "1-{depth}" \\h \\z \\u '
    run2._element.append(instr)

    run3 = para.add_run()
    end  = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run3._element.append(end)

    # Update note
    if note:
        p = doc.add_paragraph(note)
        if p.runs:
            p.runs[0].font.italic    = True
            p.runs[0].font.size      = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.paragraph_format.space_before = Pt(4)

    # Page break after TOC
    doc.add_section(WD_SECTION_START.NEW_PAGE)


# ── Table styling ─────────────────────────────────────────────────────────────

def _style_table(table, table_cfg: dict) -> None:
    """Apply header row styling, row banding, and borders to a Word table."""
    header_font        = str(table_cfg.get("header_font_name",  "Arial"))
    header_size        = int(table_cfg.get("header_font_size",  10))
    header_bold        = bool(table_cfg.get("header_bold",      True))
    header_font_color  = _parse_hex_color(table_cfg.get("header_font_color",  "FFFFFF"))
    header_fill        = str(table_cfg.get("header_fill_color", "2F5496")).lstrip("#")

    body_font          = str(table_cfg.get("body_font_name",   "Calibri"))
    body_size          = int(table_cfg.get("body_font_size",   10))
    banding            = bool(table_cfg.get("row_banding",     True))
    banding_color      = str(table_cfg.get("banding_color",    "D6E4F0")).lstrip("#")
    border_style       = str(table_cfg.get("border_style",     "single"))
    border_color       = str(table_cfg.get("border_color",     "CCCCCC")).lstrip("#")

    def _shade_cell(cell, fill_hex: str):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:val"),  "clear")
        cell._element.get_or_add_tcPr().append(shd)

    # Header row
    if table.rows:
        for cell in table.rows[0].cells:
            _shade_cell(cell, header_fill)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name      = header_font
                    run.font.size      = Pt(header_size)
                    run.font.bold      = header_bold
                    run.font.color.rgb = header_font_color

    # Body rows
    for i, row in enumerate(table.rows[1:], start=1):
        for cell in row.cells:
            if banding and i % 2 == 0:
                _shade_cell(cell, banding_color)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = body_font
                    run.font.size = Pt(body_size)

    # Borders
    if border_style != "none":
        tbl    = table._element
        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), border_color)
            borders.append(el)
        tbl_pr.append(borders)

    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def _set_header_row_repeat(table_row) -> None:
    """Mark a table row to repeat as a header on each page."""
    tr    = table_row._tr
    trPr  = tr.get_or_add_trPr()
    hdr   = OxmlElement("w:tblHeader")
    hdr.set(qn("w:val"), "true")
    trPr.append(hdr)


def _fit_columns_to_page(table, doc: Document, max_col_width: float = 3.0) -> None:
    """Proportionally scale table column widths to fit the usable page width."""
    try:
        section = doc.sections[-1]
        usable  = (section.page_width - section.left_margin - section.right_margin) / 914400  # inches

        n_cols = len(table.columns)
        if n_cols == 0:
            return

        # Use max content length per column as weight hints
        weights = [0] * n_cols
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < n_cols:
                    weights[i] = max(weights[i], len(cell.text))
        weights = [max(w, 3) for w in weights]   # floor at 3 chars
        total   = sum(weights)

        # Proportional widths
        raw = [w / total * usable for w in weights]
        # Cap and rescale
        capped = [min(w, max_col_width) for w in raw]
        capped_sum = sum(capped)
        if capped_sum > 0:
            scale  = usable / capped_sum
            capped = [w * scale for w in capped]

        for i, col in enumerate(table.columns):
            if i < len(capped):
                for cell in col.cells:
                    cell.width = Inches(capped[i])
    except Exception as exc:
        logging.warning("Column fit failed: %s", exc)


# ── Cell value helpers ────────────────────────────────────────────────────────

def _cell_to_str(val: Any) -> str:
    """Convert an openpyxl cell value to a clean display string."""
    if val is None:
        return ""
    if isinstance(val, datetime):
        if val.hour == 0 and val.minute == 0 and val.second == 0:
            return val.strftime("%Y-%m-%d")
        return val.strftime("%Y-%m-%d %H:%M")
    if isinstance(val, date):
        return val.strftime("%Y-%m-%d")
    return str(val).strip()


def _add_soft_return(paragraph) -> None:
    """Insert a soft line break (<w:br/>) into a paragraph."""
    run = paragraph.add_run()
    br  = OxmlElement("w:br")
    run._element.append(br)


# ── Sheet rule matching ───────────────────────────────────────────────────────

def match_sheet_rule(sheet_name: str, rules: list[dict]) -> dict | None:
    """Return the first rule whose 'match' glob pattern matches sheet_name (case-insensitive)."""
    name_lower = sheet_name.lower()
    for rule in rules:
        pattern = str(rule.get("match", "*")).lower()
        if fnmatch.fnmatch(name_lower, pattern):
            return _deep_merge(DEFAULT_SHEET_RULE, rule)
    return None


# ── Column resolution ─────────────────────────────────────────────────────────

def _resolve_columns(all_headers: list[str], rule: dict, heading_col_idx: int | None
                     ) -> list[tuple[int, str]]:
    """
    Return an ordered list of (col_index, display_name) for body columns.
    Respects include_columns, exclude_columns, and rename_columns.
    The heading column is always excluded from body columns.
    """
    include = [str(c).strip() for c in (rule.get("include_columns") or [])]
    exclude = {str(c).strip().lower() for c in (rule.get("exclude_columns") or [])}
    rename  = {str(k).strip(): str(v).strip() for k, v in (rule.get("rename_columns") or {}).items()}

    def _display(name: str) -> str:
        return rename.get(name, name)

    # Case-insensitive header lookup
    header_lower = [h.lower() for h in all_headers]

    if include:
        result = []
        for name in include:
            name_l = name.lower()
            for i, h_l in enumerate(header_lower):
                if h_l == name_l and i != heading_col_idx:
                    result.append((i, _display(all_headers[i])))
                    break
        return result
    else:
        result = []
        for i, header in enumerate(all_headers):
            if i == heading_col_idx:
                continue
            if header.lower() in exclude or header in exclude:
                continue
            result.append((i, _display(header)))
        return result


def _find_column_index(headers: list[str], col_name: str) -> int | None:
    """Find a column index by case-insensitive exact match, then partial match."""
    if not col_name:
        return None
    name_l = col_name.strip().lower()
    for i, h in enumerate(headers):
        if h.lower() == name_l:
            return i
    # Fallback: partial match
    for i, h in enumerate(headers):
        if name_l in h.lower():
            logging.info("heading_column '%s' matched '%s' by partial match.", col_name, h)
            return i
    return None


# ── Render helpers ────────────────────────────────────────────────────────────

def _add_key_value_para(doc: Document, label: str, value: str, body_cfg: dict) -> None:
    """
    Add a paragraph with a bold label run followed by the value.
    Multi-line values use soft line breaks after the first line.

    Example:  Description: Access Control Policy
    """
    font_name = str(body_cfg.get("font_name", "Calibri"))
    font_size = Pt(body_cfg.get("font_size", 11))
    space_aft = Pt(body_cfg.get("space_after", 4))

    para = doc.add_paragraph()
    para.paragraph_format.space_after = space_aft

    # Bold label
    lbl_run = para.add_run(f"{label}: ")
    lbl_run.bold        = True
    lbl_run.font.name   = font_name
    lbl_run.font.size   = font_size

    # Value — handle multi-line cells
    lines = value.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    for idx, line in enumerate(lines):
        if idx > 0:
            _add_soft_return(para)
        if line.strip():
            val_run = para.add_run(line)
            val_run.font.name = font_name
            val_run.font.size = font_size


def _read_worksheet_data(ws, rule: dict
                          ) -> tuple[list[str], list[list[Any]]]:
    """
    Read a worksheet and return (headers, data_rows).
    Both are raw values; callers convert with _cell_to_str().
    """
    header_row_idx = int(rule.get("header_row", 1))
    data_start_idx = int(rule.get("data_start_row", 2))

    headers: list[str] = []
    for row in ws.iter_rows(min_row=header_row_idx, max_row=header_row_idx, values_only=True):
        headers = [str(c).strip() if c is not None else f"Col{i+1}" for i, c in enumerate(row)]
        break

    data: list[list[Any]] = []
    for row in ws.iter_rows(min_row=data_start_idx, values_only=True):
        data.append(list(row))

    return headers, data


# ── Render modes ──────────────────────────────────────────────────────────────

def render_row_as_heading(doc: Document, ws, rule: dict, config: dict) -> None:
    """
    PRIMARY MODE: Each data row becomes a Word heading + body content.

    heading_column  → which column's value is the heading text
    body_render     → how remaining columns appear: key_value | table_row | prose
    """
    heading_col_name = str(rule.get("heading_column", "")).strip()
    heading_level    = int(rule.get("heading_level", 1))
    body_render      = str(rule.get("body_render", "key_value")).strip()
    skip_empty       = bool(rule.get("skip_empty_heading", True))
    body_cfg         = config.get("body", {})
    table_cfg        = config.get("table", {})

    headers, data = _read_worksheet_data(ws, rule)
    if not headers:
        logging.warning("Sheet has no header row — skipping.")
        return

    heading_col_idx = _find_column_index(headers, heading_col_name)
    if heading_col_idx is None and heading_col_name:
        logging.warning("heading_column '%s' not found in %s — will use row index as heading.",
                         heading_col_name, headers)

    body_cols = _resolve_columns(headers, rule, heading_col_idx)

    for row_num, row in enumerate(data, start=1):
        # Determine heading text
        if heading_col_idx is not None and heading_col_idx < len(row):
            heading_text = _cell_to_str(row[heading_col_idx])
        else:
            heading_text = f"Row {row_num}"

        if skip_empty and not heading_text:
            continue

        # Skip rows where all cells are empty (e.g. blank separator rows)
        if not any(row[i] is not None and _cell_to_str(row[i]) for i, _ in body_cols):
            if not heading_text:
                continue

        # Add the heading
        doc.add_heading(heading_text, level=heading_level)

        # Render body content
        if body_render == "key_value":
            for col_idx, display_name in body_cols:
                val_str = _cell_to_str(row[col_idx]) if col_idx < len(row) else ""
                if val_str:
                    _add_key_value_para(doc, display_name, val_str, body_cfg)

        elif body_render == "table_row":
            # 2-column table: Field | Value
            visible = [(col_idx, name, _cell_to_str(row[col_idx]) if col_idx < len(row) else "")
                       for col_idx, name in body_cols]
            visible = [(ci, n, v) for ci, n, v in visible if v]
            if visible:
                tbl = doc.add_table(rows=len(visible) + 1, cols=2)
                # Header row
                tbl.rows[0].cells[0].text = "Field"
                tbl.rows[0].cells[1].text = "Value"
                for r_idx, (_, name, val) in enumerate(visible, start=1):
                    tbl.rows[r_idx].cells[0].text = name
                    tbl.rows[r_idx].cells[1].text = val
                _style_table(tbl, table_cfg)
                if bool(rule.get("repeat_header", True)):
                    _set_header_row_repeat(tbl.rows[0])
                if bool(table_cfg.get("fit_to_page", True)):
                    _fit_columns_to_page(tbl, doc, float(table_cfg.get("max_col_width", 3.0)))
                doc.add_paragraph()  # spacing after table

        elif body_render == "prose":
            for col_idx, _ in body_cols:
                val_str = _cell_to_str(row[col_idx]) if col_idx < len(row) else ""
                if val_str:
                    doc.add_paragraph(val_str)


def render_table(doc: Document, ws, rule: dict, config: dict) -> None:
    """Render the entire sheet as a single styled Word table."""
    table_cfg = config.get("table", {})

    headers, data = _read_worksheet_data(ws, rule)
    if not headers:
        logging.warning("Sheet has no data — skipping table render.")
        return

    # Determine visible columns
    body_cols = _resolve_columns(headers, rule, None)  # None = no special heading column
    if not body_cols:
        body_cols = list(enumerate(headers))

    # Include heading column if it's set (for table mode, heading_column is just a normal col)
    heading_col = str(rule.get("heading_column", "")).strip()
    heading_idx = _find_column_index(headers, heading_col) if heading_col else None

    if heading_idx is not None:
        # Prepend the heading column
        display_name = (rule.get("rename_columns") or {}).get(headers[heading_idx], headers[heading_idx])
        body_cols = [(heading_idx, display_name)] + [c for c in body_cols if c[0] != heading_idx]

    skip_empty = bool(rule.get("skip_empty_heading", True))
    data_rows = []
    for row in data:
        if skip_empty and not any(_cell_to_str(row[ci]) if ci < len(row) else "" for ci, _ in body_cols):
            continue
        data_rows.append(row)

    if not data_rows:
        doc.add_paragraph("(No data)")
        return

    tbl = doc.add_table(rows=len(data_rows) + 1, cols=len(body_cols))

    # Write header row
    for j, (_, display_name) in enumerate(body_cols):
        tbl.rows[0].cells[j].text = display_name

    # Write data rows
    for i, row in enumerate(data_rows, start=1):
        for j, (col_idx, _) in enumerate(body_cols):
            tbl.rows[i].cells[j].text = _cell_to_str(row[col_idx]) if col_idx < len(row) else ""

    _style_table(tbl, table_cfg)
    if bool(rule.get("repeat_header", True)):
        _set_header_row_repeat(tbl.rows[0])
    if bool(table_cfg.get("fit_to_page", True)):
        _fit_columns_to_page(tbl, doc, float(table_cfg.get("max_col_width", 3.0)))
    doc.add_paragraph()  # spacing after table


def render_key_value(doc: Document, ws, rule: dict, config: dict) -> None:
    """
    Two-column sheet: column A = label (bold), column B = value.
    Useful for settings or metadata sheets.
    """
    body_cfg       = config.get("body", {})
    header_row_idx = int(rule.get("header_row", 1))
    data_start_idx = int(rule.get("data_start_row", 2))

    for row in ws.iter_rows(min_row=data_start_idx, values_only=True):
        label = _cell_to_str(row[0]) if len(row) > 0 else ""
        value = _cell_to_str(row[1]) if len(row) > 1 else ""
        if not label and not value:
            continue
        if label:
            _add_key_value_para(doc, label, value, body_cfg)
        else:
            doc.add_paragraph(value)


def render_list(doc: Document, ws, rule: dict, config: dict) -> None:
    """Single-column sheet: values become a bulleted list."""
    data_start_idx = int(rule.get("data_start_row", 2))

    for row in ws.iter_rows(min_row=data_start_idx, values_only=True):
        val = _cell_to_str(row[0]) if row else ""
        if val:
            doc.add_paragraph(val, style="List Bullet")


def render_prose(doc: Document, ws, rule: dict, config: dict) -> None:
    """Column A values become plain body paragraphs (no labels)."""
    data_start_idx = int(rule.get("data_start_row", 2))

    for row in ws.iter_rows(min_row=data_start_idx, values_only=True):
        val = _cell_to_str(row[0]) if row else ""
        if val:
            doc.add_paragraph(val)


def render_sheet(doc: Document, ws, rule: dict, config: dict, sheet_name: str) -> None:
    """
    Dispatch rendering for a single worksheet based on rule['render_as'].
    Optionally adds a section heading above the content.
    """
    render_as = str(rule.get("render_as", "table")).strip().lower()

    if render_as == "skip":
        logging.info("Skipping sheet '%s' (render_as: skip).", sheet_name)
        return

    # Determine whether to add a sheet-level heading
    add_heading_override = rule.get("add_section_heading")
    if add_heading_override is None:
        # Auto: row_as_heading mode provides its own headings per row
        add_section_heading = (render_as != "row_as_heading")
    else:
        add_section_heading = bool(add_heading_override)

    if add_section_heading:
        level = int(rule.get("section_heading_level", 1))
        doc.add_heading(sheet_name, level=level)

    logging.info("Rendering sheet '%s' as '%s'.", sheet_name, render_as)

    if render_as == "row_as_heading":
        render_row_as_heading(doc, ws, rule, config)
    elif render_as == "table":
        render_table(doc, ws, rule, config)
    elif render_as == "key_value":
        render_key_value(doc, ws, rule, config)
    elif render_as == "list":
        render_list(doc, ws, rule, config)
    elif render_as == "prose":
        render_prose(doc, ws, rule, config)
    else:
        logging.warning("Unknown render_as '%s' for sheet '%s' — skipping.", render_as, sheet_name)


# ── Document builder ──────────────────────────────────────────────────────────

def _build_document(config: dict, token_map: dict) -> Document:
    """Create a fresh Document with page layout, heading styles, and body defaults applied."""
    doc = Document()

    # Remove default blank paragraph
    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)

    page_cfg = config.get("page", {})
    _apply_page_layout(doc, page_cfg)
    _apply_body_defaults(doc, config.get("body", {}))
    _apply_core_properties(doc, config.get("identity", {}), token_map)

    # Apply H1 – H4 styles
    headings_cfg = config.get("headings", {})
    for level_key, style_name in [("h1", "Heading 1"), ("h2", "Heading 2"),
                                    ("h3", "Heading 3"), ("h4", "Heading 4")]:
        if level_key in headings_cfg:
            _apply_heading_style(doc, style_name, headings_cfg[level_key])

    return doc


# ── File conversion ───────────────────────────────────────────────────────────

def convert_file(xlsx_path: str, config: dict, output_dir: str) -> str:
    """
    Convert a single Excel file to a Word document.
    Returns the path to the saved .docx.
    """
    xlsx_path = os.path.abspath(xlsx_path)
    stem      = Path(xlsx_path).stem
    filename  = str(config["output"].get("filename_template", "{input_stem}.docx"))
    filename  = filename.replace("{input_stem}", stem)
    out_path  = os.path.join(output_dir, filename)

    logging.info("Converting: %s → %s", xlsx_path, out_path)

    # Build token map (may include the file stem as document_title fallback)
    identity = copy.deepcopy(config.get("identity", {}))
    if not identity.get("document_title"):
        identity["document_title"] = stem.replace("_", " ").replace("-", " ")
    token_map = _build_token_map(identity)

    doc = _build_document(config, token_map)

    # Cover page
    _build_cover_page(doc, config.get("cover_page", {}), token_map)

    # TOC
    if bool(config.get("toc", {}).get("enabled", False)):
        _insert_toc(doc, config.get("toc", {}))

    # Open workbook
    try:
        wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    except Exception as exc:
        logging.error("Could not open '%s': %s", xlsx_path, exc)
        raise

    rules          = config.get("sheets", [])
    include_unmatch = bool(config.get("include_unmatched", True))
    fallback_rule   = _deep_merge(DEFAULT_SHEET_RULE, {"render_as": "table"})

    for sheet_name in wb.sheetnames:
        rule = match_sheet_rule(sheet_name, rules)
        if rule is None:
            if include_unmatch:
                rule = fallback_rule
            else:
                logging.info("Sheet '%s' not matched by any rule — skipping.", sheet_name)
                continue

        ws = wb[sheet_name]
        render_sheet(doc, ws, rule, config, sheet_name)

    wb.close()

    # Header / footer
    _apply_header_footer(doc, config, token_map)

    os.makedirs(output_dir, exist_ok=True)
    doc.save(out_path)
    logging.info("Saved: %s", out_path)
    return out_path


def resolve_input_files(config: dict, cli_input: str | None) -> list[str]:
    """
    Resolve the list of .xlsx files to convert.
    Priority: --input CLI arg > config.input.files > config.input.directory scan.
    """
    exclude = config.get("input", {}).get("exclude_patterns", ["~$*"])

    def _is_excluded(name: str) -> bool:
        return any(fnmatch.fnmatch(name, pat) for pat in exclude)

    # CLI --input override
    if cli_input:
        p = os.path.abspath(cli_input)
        if os.path.isfile(p):
            return [p]
        if os.path.isdir(p):
            files = [os.path.join(p, f) for f in os.listdir(p)
                     if f.lower().endswith(".xlsx") and not _is_excluded(f)]
            files.sort()
            return files
        logging.error("--input path not found: %s", cli_input)
        return []

    # Explicit files list from config
    explicit = config.get("input", {}).get("files", []) or []
    if explicit:
        return [os.path.abspath(str(f)) for f in explicit]

    # Directory scan
    directory = config.get("input", {}).get("directory", "")
    if not directory:
        logging.error("No input specified. Set input.directory in config or use --input.")
        return []

    directory = os.path.abspath(directory)
    if not os.path.isdir(directory):
        logging.error("Input directory not found: %s", directory)
        return []

    files = [os.path.join(directory, f) for f in os.listdir(directory)
             if f.lower().endswith(".xlsx") and not _is_excluded(f)]
    files.sort()
    return files


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    parser = argparse.ArgumentParser(
        description="Convert Excel files to formatted Word documents.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python xlsx2docx.py
  python xlsx2docx.py --input controls.xlsx
  python xlsx2docx.py --input ../../input/ --output ../../output/xlsx2docx/
  python xlsx2docx.py --config my_config.yaml --input gap_analysis.xlsx
""",
    )
    parser.add_argument(
        "--config", "-c",
        default=None,
        help="Path to xlsx2docx_config.yaml (default: xlsx2docx_config.yaml in script dir)",
    )
    parser.add_argument(
        "--input", "-i",
        default=None,
        help="Path to a single .xlsx file or directory to scan (overrides config)",
    )
    parser.add_argument(
        "--output", "-o",
        default=None,
        help="Output directory for .docx files (overrides config)",
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Enable verbose (DEBUG) logging",
    )
    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # Resolve config path
    if args.config:
        config_path = os.path.abspath(args.config)
    else:
        config_path = os.path.join(os.path.dirname(__file__), "xlsx2docx_config.yaml")

    config = load_config(config_path)

    # Override output directory from CLI
    output_dir = args.output or config.get("output", {}).get("directory", "output/xlsx2docx/")
    output_dir = os.path.abspath(output_dir)

    # Resolve input files (paths relative to config file's directory)
    original_dir = os.getcwd()
    os.chdir(os.path.dirname(config_path))  # so relative paths in config resolve correctly
    try:
        input_files = resolve_input_files(config, args.input)
    finally:
        os.chdir(original_dir)

    if not input_files:
        print("No .xlsx files found. Check --input or input.directory in config.", file=sys.stderr)
        sys.exit(1)

    print(f"Converting {len(input_files)} file(s) → {output_dir}")

    errors = 0
    for xlsx_path in input_files:
        try:
            out_path = convert_file(xlsx_path, config, output_dir)
            print(f"  ✓ {os.path.basename(xlsx_path)} → {os.path.relpath(out_path)}")
        except Exception as exc:
            logging.error("Failed to convert '%s': %s", xlsx_path, exc, exc_info=args.verbose)
            print(f"  ✗ {os.path.basename(xlsx_path)} — ERROR: {exc}", file=sys.stderr)
            errors += 1

    if errors:
        print(f"\n{errors} file(s) failed.")
        sys.exit(1)
    else:
        print(f"\nDone. {len(input_files)} file(s) written to: {output_dir}")


if __name__ == "__main__":
    main()
