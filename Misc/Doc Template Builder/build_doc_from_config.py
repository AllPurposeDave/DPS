#!/usr/bin/env python3
"""
build_doc_from_config.py — Build a Word doc template from doc_template_config.xlsx

Reads the Excel config produced by create_config_excel.py, applies heading styles,
document identity, standard sections, headers/footers, table styles, and cover page
to a fresh .docx optimized for the DPS pipeline.

Usage:
  python build_doc_from_config.py
  python build_doc_from_config.py --config doc_template_config.xlsx -o my_template.docx
"""

from __future__ import annotations

import argparse
import copy
import logging
import os
import re
import sys
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook


# ── Config parsing ───────────────────────────────────────────────────────────

def _coerce_value(val: Any) -> Any:
    """Convert an Excel cell value to a Python type."""
    if val is None:
        return None
    if isinstance(val, bool):
        return val
    if isinstance(val, (int, float)):
        return val
    s = str(val).strip()
    if s.upper() == "TRUE":
        return True
    if s.upper() == "FALSE":
        return False
    try:
        as_int = int(s)
        if str(as_int) == s:
            return as_int
    except ValueError:
        pass
    try:
        return float(s)
    except ValueError:
        pass
    return s


def _parse_settings_sheet(ws) -> dict:
    """
    Read a Setting | Value | Description sheet.
    Skips rows where column A is blank or starts with '#'.
    Returns {setting_key: coerced_value}.
    """
    result = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        key = row[0] if row else None
        if not key:
            continue
        key = str(key).strip()
        if key.startswith("#"):
            continue
        val = _coerce_value(row[1] if len(row) > 1 else None)
        if val is not None:
            result[key] = val
    return result


def _parse_standard_sections(ws) -> list[dict]:
    """
    Parse the Standard Sections sheet (4-column: Section | Include | Heading Text | Placeholder Body).
    Returns list of dicts with keys: section, include, heading_text, placeholder.
    """
    sections = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row or not row[0]:
            continue
        section_name = str(row[0]).strip()
        if section_name.startswith("#"):
            continue
        include = _coerce_value(row[1] if len(row) > 1 else True)
        heading_text = str(row[2]).strip() if len(row) > 2 and row[2] else section_name
        placeholder = str(row[3]).strip() if len(row) > 3 and row[3] else ""
        sections.append({
            "section": section_name,
            "include": bool(include) if include is not None else True,
            "heading_text": heading_text,
            "placeholder": placeholder,
        })
    return sections


def _load_config(xlsx_path: str) -> dict:
    """
    Load all config sheets from the workbook.

    Returns dict with keys: document, headings, identity, page_layout,
    header_footer, standard_sections, table_styles, cover_page.
    """
    wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    config: dict = {}

    # 3-column settings sheets
    settings_sheets = {
        "Document": "document",
        "Heading Styles": "_headings_flat",
        "Document Identity": "identity",
        "Page Layout": "page_layout",
        "Header Footer": "header_footer",
        "Table Styles": "table_styles",
        "Cover Page": "cover_page",
    }
    for sheet_name, config_key in settings_sheets.items():
        if sheet_name in wb.sheetnames:
            config[config_key] = _parse_settings_sheet(wb[sheet_name])
        else:
            config[config_key] = {}

    # Parse headings into nested dict (h1.font_name -> headings["h1"]["font_name"])
    flat = config.pop("_headings_flat", {})
    headings: dict = {}
    for dotkey, val in flat.items():
        level, _, field = dotkey.partition(".")
        if level and field:
            headings.setdefault(level, {})[field] = val
    config["headings"] = headings

    # Standard Sections (4-column)
    if "Standard Sections" in wb.sheetnames:
        config["standard_sections"] = _parse_standard_sections(wb["Standard Sections"])
    else:
        config["standard_sections"] = []

    wb.close()
    return config


# ── Token resolution ─────────────────────────────────────────────────────────

def _build_token_map(identity: dict) -> dict[str, str]:
    """Build a token replacement map from Document Identity values."""
    return {
        "{document_title}": str(identity.get("document_title", "")),
        "{doc_id}": str(identity.get("document_id", "")),
        "{classification}": str(identity.get("classification", "")),
        "{organization_name}": str(identity.get("organization_name", "")),
        "{effective_date}": str(identity.get("effective_date", "")),
        "{department}": str(identity.get("department", "")),
        "{version}": str(identity.get("version", "")),
        "{document_owner}": str(identity.get("document_owner", "")),
    }


def _resolve_tokens(text: str, token_map: dict[str, str]) -> str:
    """Replace {token} placeholders with values from identity config."""
    if not text:
        return text
    for token, value in token_map.items():
        text = text.replace(token, value)
    return text


# ── Color/style helpers ──────────────────────────────────────────────────────

def _parse_hex_color(hex_str: Any) -> RGBColor:
    """Parse a 6-digit hex string (e.g. '2F5496') to RGBColor. Falls back to DPS blue."""
    try:
        s = str(hex_str).strip().lstrip("#")
        return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
    except (ValueError, IndexError, TypeError):
        logging.warning("Could not parse color '%s' — using default blue.", hex_str)
        return RGBColor(0x2F, 0x54, 0x96)


def _get_alignment(value: Any) -> WD_ALIGN_PARAGRAPH | None:
    """Convert alignment string to WD_ALIGN_PARAGRAPH enum."""
    s = str(value).strip().lower() if value else ""
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(s)


# ── Style application ────────────────────────────────────────────────────────

def _apply_heading_style(doc: Document, style_name: str, hcfg: dict) -> None:
    """Modify a built-in heading style in doc using values from hcfg."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        logging.warning("Style '%s' not found in document — skipping.", style_name)
        return

    font = style.font
    if "font_name" in hcfg:
        font.name = str(hcfg["font_name"])
    if "font_size" in hcfg:
        font.size = Pt(hcfg["font_size"])
    if "bold" in hcfg:
        font.bold = bool(hcfg["bold"])
    if "italic" in hcfg:
        font.italic = bool(hcfg["italic"])
    if "underline" in hcfg:
        font.underline = bool(hcfg["underline"])
    if "all_caps" in hcfg:
        font.all_caps = bool(hcfg["all_caps"])
    if "color" in hcfg:
        font.color.rgb = _parse_hex_color(hcfg["color"])

    pf = style.paragraph_format
    if "space_before" in hcfg:
        pf.space_before = Pt(hcfg["space_before"])
    if "space_after" in hcfg:
        pf.space_after = Pt(hcfg["space_after"])
    if "keep_with_next" in hcfg:
        pf.keep_with_next = bool(hcfg["keep_with_next"])
    if "alignment" in hcfg:
        align = _get_alignment(hcfg["alignment"])
        if align is not None:
            pf.alignment = align

    logging.info("Applied style '%s': %s", style_name, hcfg)


def _set_page_margins(doc: Document, doc_cfg: dict) -> None:
    """Set page margins on the first section from config values (in inches)."""
    section = doc.sections[0]
    if "margin_top" in doc_cfg:
        section.top_margin = Inches(doc_cfg["margin_top"])
    if "margin_bottom" in doc_cfg:
        section.bottom_margin = Inches(doc_cfg["margin_bottom"])
    if "margin_left" in doc_cfg:
        section.left_margin = Inches(doc_cfg["margin_left"])
    if "margin_right" in doc_cfg:
        section.right_margin = Inches(doc_cfg["margin_right"])


def _apply_page_layout(doc: Document, layout_cfg: dict) -> None:
    """Apply page size and orientation."""
    section = doc.sections[0]

    page_size = str(layout_cfg.get("page_size", "Letter")).strip()
    if page_size.upper() == "A4":
        width, height = Inches(8.27), Inches(11.69)
    else:  # Letter
        width, height = Inches(8.5), Inches(11.0)

    orientation = str(layout_cfg.get("orientation", "Portrait")).strip()
    if orientation.lower() == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = height
        section.page_height = width
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = width
        section.page_height = height


def _apply_body_defaults(doc: Document, doc_cfg: dict) -> None:
    """Apply body font and paragraph settings to the Normal style."""
    try:
        normal = doc.styles["Normal"]
    except KeyError:
        logging.warning("'Normal' style not found — body defaults not applied.")
        return

    if "body_font_name" in doc_cfg:
        normal.font.name = str(doc_cfg["body_font_name"])
    if "body_font_size" in doc_cfg:
        normal.font.size = Pt(doc_cfg["body_font_size"])

    pf = normal.paragraph_format
    if "line_spacing" in doc_cfg:
        pf.line_spacing = float(doc_cfg["line_spacing"])
    if "paragraph_space_before" in doc_cfg:
        pf.space_before = Pt(doc_cfg["paragraph_space_before"])
    if "paragraph_space_after" in doc_cfg:
        pf.space_after = Pt(doc_cfg["paragraph_space_after"])
    if "first_line_indent" in doc_cfg:
        indent = float(doc_cfg["first_line_indent"])
        if indent > 0:
            pf.first_line_indent = Inches(indent)
    if "text_alignment" in doc_cfg:
        align = _get_alignment(doc_cfg["text_alignment"])
        if align is not None:
            pf.alignment = align


def _apply_core_properties(doc: Document, identity: dict) -> None:
    """Write Document Identity values into Word core_properties for pipeline extraction."""
    props = doc.core_properties
    if identity.get("document_title"):
        props.title = str(identity["document_title"])
    if identity.get("document_owner"):
        props.author = str(identity["document_owner"])
    if identity.get("classification"):
        props.subject = str(identity["classification"])
    # Pack version + status + doc_id into comments for metadata extraction
    parts = []
    if identity.get("document_id"):
        parts.append(f"Doc ID: {identity['document_id']}")
    if identity.get("version"):
        parts.append(f"Version: {identity['version']}")
    if identity.get("document_status"):
        parts.append(f"Status: {identity['document_status']}")
    if parts:
        props.comments = " | ".join(parts)
    if identity.get("organization_name"):
        props.category = str(identity["organization_name"])


# ── Header/Footer builder ───────────────────────────────────────────────────

def _add_page_number_field(paragraph):
    """Insert a Word PAGE field code into a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3._element.append(fld_char_end)


def _add_num_pages_field(paragraph):
    """Insert a Word NUMPAGES field code into a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " NUMPAGES "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3._element.append(fld_char_end)


def _build_hf_paragraph(hf, text: str, alignment, font_name: str, font_size: int,
                         color: str, token_map: dict):
    """Build a header/footer paragraph with token resolution and page field codes."""
    resolved = _resolve_tokens(text, token_map)
    paragraph = hf.paragraphs[0] if hf.paragraphs else hf.add_paragraph()
    paragraph.alignment = alignment

    # Split on {page} and {pages} to insert field codes
    parts = re.split(r"(\{page\}|\{pages\})", resolved)
    for part in parts:
        if part == "{page}":
            _add_page_number_field(paragraph)
        elif part == "{pages}":
            _add_num_pages_field(paragraph)
        elif part:
            run = paragraph.add_run(part)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = _parse_hex_color(color)

    return paragraph


def _build_tab_stop_hf(hf, left_text: str, center_text: str, right_text: str,
                        font_name: str, font_size: int, color: str, token_map: dict,
                        page_width_inches: float = 6.0):
    """Build a header/footer with left/center/right text using tab stops."""
    paragraph = hf.paragraphs[0] if hf.paragraphs else hf.add_paragraph()
    # Clear existing runs
    for run in paragraph.runs:
        run.text = ""

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add tab stops for center and right alignment
    pf = paragraph.paragraph_format
    tab_stops = pf.tab_stops
    center_pos = Inches(page_width_inches / 2)
    right_pos = Inches(page_width_inches)

    from docx.enum.text import WD_TAB_ALIGNMENT
    tab_stops.add_tab_stop(center_pos, WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(right_pos, WD_TAB_ALIGNMENT.RIGHT)

    def _add_text_or_field(text_val: str):
        resolved = _resolve_tokens(text_val, token_map)
        parts = re.split(r"(\{page\}|\{pages\})", resolved)
        for part in parts:
            if part == "{page}":
                _add_page_number_field(paragraph)
            elif part == "{pages}":
                _add_num_pages_field(paragraph)
            elif part:
                run = paragraph.add_run(part)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.color.rgb = _parse_hex_color(color)

    # Left text
    if left_text:
        _add_text_or_field(left_text)

    # Tab to center
    run = paragraph.add_run("\t")
    run.font.size = Pt(font_size)
    if center_text:
        _add_text_or_field(center_text)

    # Tab to right
    run = paragraph.add_run("\t")
    run.font.size = Pt(font_size)
    if right_text:
        _add_text_or_field(right_text)


def _apply_header_footer(doc: Document, hf_cfg: dict, token_map: dict) -> None:
    """Apply header and footer to all sections (except first page if configured)."""
    if not hf_cfg:
        return

    different_first = bool(hf_cfg.get("different_first_page", True))

    for section in doc.sections:
        section.different_first_page_header_footer = different_first

        # Calculate usable width for tab stops
        usable_width = (section.page_width - section.left_margin - section.right_margin)
        page_width_inches = usable_width / 914400  # EMU to inches

        # Header
        header = section.header
        header.is_linked_to_previous = False
        _build_tab_stop_hf(
            header,
            left_text=str(hf_cfg.get("header_left", "")),
            center_text=str(hf_cfg.get("header_center", "")),
            right_text=str(hf_cfg.get("header_right", "")),
            font_name=str(hf_cfg.get("header_font_name", "Arial")),
            font_size=int(hf_cfg.get("header_font_size", 8)),
            color=str(hf_cfg.get("header_color", "666666")),
            token_map=token_map,
            page_width_inches=page_width_inches,
        )

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        _build_tab_stop_hf(
            footer,
            left_text=str(hf_cfg.get("footer_left", "")),
            center_text=str(hf_cfg.get("footer_center", "")),
            right_text=str(hf_cfg.get("footer_right", "")),
            font_name=str(hf_cfg.get("footer_font_name", "Arial")),
            font_size=int(hf_cfg.get("footer_font_size", 8)),
            color=str(hf_cfg.get("footer_color", "666666")),
            token_map=token_map,
            page_width_inches=page_width_inches,
        )


# ── Table builder ────────────────────────────────────────────────────────────

def _style_table(table, table_cfg: dict) -> None:
    """Apply table styling from config."""
    header_font = str(table_cfg.get("table_header_font_name", "Arial"))
    header_size = int(table_cfg.get("table_header_font_size", 10))
    header_bold = bool(table_cfg.get("table_header_bold", True))
    header_font_color = _parse_hex_color(table_cfg.get("table_header_font_color", "FFFFFF"))
    header_fill = str(table_cfg.get("table_header_fill_color", "2F5496"))

    body_font = str(table_cfg.get("table_body_font_name", "Calibri"))
    body_size = int(table_cfg.get("table_body_font_size", 10))
    banding = bool(table_cfg.get("table_row_banding", True))
    banding_color = str(table_cfg.get("table_banding_color", "D6E4F0"))

    border_style = str(table_cfg.get("table_border_style", "single"))
    border_color = str(table_cfg.get("table_border_color", "CCCCCC"))

    # Style header row
    if table.rows:
        for cell in table.rows[0].cells:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), header_fill)
            shading.set(qn("w:val"), "clear")
            cell._element.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = header_font
                    run.font.size = Pt(header_size)
                    run.font.bold = header_bold
                    run.font.color.rgb = header_font_color

    # Style body rows
    for i, row in enumerate(table.rows[1:], start=1):
        for cell in row.cells:
            if banding and i % 2 == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), banding_color)
                shading.set(qn("w:val"), "clear")
                cell._element.get_or_add_tcPr().append(shading)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = body_font
                    run.font.size = Pt(body_size)

    # Apply borders
    if border_style != "none":
        tbl = table._element
        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), border_color)
            borders.append(element)
        tbl_pr.append(borders)

    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def _insert_revision_history_table(doc: Document, table_cfg: dict) -> None:
    """Insert a revision history table with styled headers."""
    headers = ["Version", "Date", "Author", "Description"]
    table = doc.add_table(rows=2, cols=len(headers))

    # Write headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header

    # Write placeholder row
    placeholders = ["1.0", "", "", "Initial release"]
    for i, val in enumerate(placeholders):
        table.rows[1].cells[i].text = val

    _style_table(table, table_cfg)


# ── Cover page builder ───────────────────────────────────────────────────────

def _build_cover_page(doc: Document, cover_cfg: dict, token_map: dict) -> None:
    """Build a cover/title page and add a section break after it."""
    if not bool(cover_cfg.get("enable_cover_page", False)):
        return

    # Classification banner at top
    classification = _resolve_tokens(
        str(cover_cfg.get("cover_classification", "{classification}")), token_map
    )
    if classification:
        banner = doc.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = banner.add_run(classification)
        run.font.size = Pt(14)
        run.font.bold = True
        color_str = str(cover_cfg.get("cover_classification_color", "FF0000"))
        run.font.color.rgb = _parse_hex_color(color_str)
        banner.paragraph_format.space_after = Pt(36)

    # Logo placeholder
    if bool(cover_cfg.get("cover_logo_placeholder", False)):
        logo = doc.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = logo.add_run("[Organization Logo]")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        logo.paragraph_format.space_after = Pt(48)

    # Spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(72)

    # Title
    title_text = _resolve_tokens(
        str(cover_cfg.get("cover_title", "{document_title}")), token_map
    )
    if title_text:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title_text)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = _parse_hex_color("2F5496")
        title_para.paragraph_format.space_after = Pt(12)

    # Subtitle
    subtitle_text = _resolve_tokens(
        str(cover_cfg.get("cover_subtitle", "{organization_name}")), token_map
    )
    if subtitle_text:
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle_para.add_run(subtitle_text)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        subtitle_para.paragraph_format.space_after = Pt(48)

    # Prepared By / Approved By
    prepared = str(cover_cfg.get("cover_prepared_by", ""))
    approved = str(cover_cfg.get("cover_approved_by", ""))
    if prepared:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Prepared By: {prepared}")
        run.font.size = Pt(11)
    if approved:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Approved By: {approved}")
        run.font.size = Pt(11)

    # Date
    date_text = _resolve_tokens(
        str(cover_cfg.get("cover_date", "{effective_date}")), token_map
    )
    if date_text:
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(date_text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        date_para.paragraph_format.space_before = Pt(24)

    # Bottom classification banner
    if classification:
        banner2 = doc.add_paragraph()
        banner2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = banner2.add_run(classification)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = _parse_hex_color(
            str(cover_cfg.get("cover_classification_color", "FF0000"))
        )
        banner2.paragraph_format.space_before = Pt(72)

    # Add section break (new page) after cover
    from docx.enum.section import WD_SECTION_START
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    # Copy margins from first section
    first = doc.sections[0]
    new_section.top_margin = first.top_margin
    new_section.bottom_margin = first.bottom_margin
    new_section.left_margin = first.left_margin
    new_section.right_margin = first.right_margin
    new_section.page_width = first.page_width
    new_section.page_height = first.page_height
    new_section.orientation = first.orientation


# ── Standard sections builder ────────────────────────────────────────────────

def _build_standard_sections(doc: Document, sections: list[dict], table_cfg: dict) -> None:
    """Insert standard sections as H1 headings with placeholder body text."""
    for sec in sections:
        if not sec.get("include", True):
            continue

        heading_text = sec.get("heading_text", sec["section"])
        placeholder = sec.get("placeholder", "")

        # Add H1 heading (uses built-in Heading 1 style)
        doc.add_heading(heading_text, level=1)

        # Special handling for Revision History — insert a table
        if sec["section"] == "Revision History":
            _insert_revision_history_table(doc, table_cfg)
        elif placeholder:
            p = doc.add_paragraph(placeholder)


# ── Document builder ─────────────────────────────────────────────────────────

def build_doc_from_config(config_path: str, output_path: str | None) -> None:
    logging.info("Loading config: %s", config_path)
    config = _load_config(config_path)

    doc_cfg = config.get("document", {})
    headings_cfg = config.get("headings", {})
    identity_cfg = config.get("identity", {})
    layout_cfg = config.get("page_layout", {})
    hf_cfg = config.get("header_footer", {})
    sections_cfg = config.get("standard_sections", [])
    table_cfg = config.get("table_styles", {})
    cover_cfg = config.get("cover_page", {})

    # Resolve output path: CLI arg > identity doc_id pattern > config setting > default
    if not output_path:
        doc_id = str(identity_cfg.get("document_id", ""))
        doc_title = str(identity_cfg.get("document_title", ""))
        if doc_id and doc_title:
            # Build filename: Title_With_Underscores_DOC-ID.docx
            safe_title = re.sub(r"[^\w\s-]", "", doc_title).strip().replace(" ", "_")
            output_path = f"{safe_title}_{doc_id}.docx"
        elif doc_id:
            output_path = f"{doc_id}.docx"
        else:
            output_path = str(doc_cfg.get("output_filename", "doc_template.docx"))

    # Build token map from identity
    token_map = _build_token_map(identity_cfg)

    doc = Document()

    # Phase 1: Document-level settings
    _set_page_margins(doc, doc_cfg)
    _apply_page_layout(doc, layout_cfg)
    _apply_body_defaults(doc, doc_cfg)
    _apply_core_properties(doc, identity_cfg)

    # Phase 1: Apply heading styles (H1–H4)
    level_map = [
        ("h1", "Heading 1", 1),
        ("h2", "Heading 2", 2),
        ("h3", "Heading 3", 3),
        ("h4", "Heading 4", 4),
    ]
    for level_key, style_name, _ in level_map:
        if level_key in headings_cfg:
            _apply_heading_style(doc, style_name, headings_cfg[level_key])

    # Remove the default blank paragraph that Document() creates
    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)

    # Phase 3: Cover page (before body content)
    _build_cover_page(doc, cover_cfg, token_map)

    # Phase 2: Standard sections (or fallback to demo headings)
    if sections_cfg:
        _build_standard_sections(doc, sections_cfg, table_cfg)
    else:
        # Fallback: insert one placeholder paragraph per heading level
        for level_key, _, heading_level in level_map:
            default_text = {1: "Section Title", 2: "Subsection Title",
                           3: "Sub-subsection Title", 4: "Detail Heading"}.get(heading_level, "Heading")
            placeholder = headings_cfg.get(level_key, {}).get("placeholder_text", default_text)
            doc.add_heading(str(placeholder), level=heading_level)

    # Phase 3: Header/footer
    _apply_header_footer(doc, hf_cfg, token_map)

    doc.save(output_path)
    logging.info("Saved: %s", output_path)
    print(f"  Document template saved: {output_path}")


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    parser = argparse.ArgumentParser(
        description="Build a Word doc template from doc_template_config.xlsx"
    )
    parser.add_argument(
        "--config", "-c",
        default="doc_template_config.xlsx",
        help="Path to the Excel config workbook (default: doc_template_config.xlsx)",
    )
    parser.add_argument(
        "-o", "--output",
        default=None,
        help="Output .docx path — overrides output_filename in config",
    )
    args = parser.parse_args()

    if not os.path.isfile(args.config):
        print(f"ERROR: Config file not found: {args.config}", file=sys.stderr)
        print("  Run create_config_excel.py first to generate the config workbook.", file=sys.stderr)
        sys.exit(1)

    try:
        build_doc_from_config(args.config, args.output)
    except Exception as exc:
        logging.error("Failed: %s", exc, exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()
