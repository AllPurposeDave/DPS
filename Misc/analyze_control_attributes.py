"""
Control Attribute Analyzer
==========================

Scans .docx documents for all instances of control IDs (pattern: ABC##.###)
and extracts text attributes for each match. Useful for identifying false
positives and understanding which attributes distinguish real controls.

Features:
    - Finds ALL control ID matches (including false positives)
    - Captures text attributes: bold, style, source (text/table), context
    - Outputs detailed CSV for analysis and filtering
    - Tracks paragraph position and surrounding context

Usage:
    python Misc/analyze_control_attributes.py
    python Misc/analyze_control_attributes.py --config dps_config.yaml
    python Misc/analyze_control_attributes.py ./input ./output

Output:
    control_attributes_analysis.csv — one row per control match with attributes
"""

import re
import csv
import os
import sys
from pathlib import Path
from dataclasses import dataclass, asdict

# Add scripts directory to path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))

from docx import Document

from shared_utils import (
    ensure_output_dir,
    get_input_dir,
    get_output_dir,
    iter_docx_files,
    load_config,
    setup_argparse,
)


@dataclass
class ControlAttribute:
    """Represents a single control ID instance with all its attributes."""
    source_file: str = ""
    paragraph_index: int = 0
    control_id: str = ""
    match_position: int = 0  # character position in paragraph
    paragraph_text: str = ""
    context_before: str = ""  # 50 chars before match
    context_after: str = ""   # 50 chars after match
    # Formatting attributes
    paragraph_bold: bool = False
    paragraph_italic: bool = False
    paragraph_underline: bool = False
    font_size_pt: str = ""  # in points (e.g., "12")
    font_name: str = ""
    paragraph_style: str = ""
    is_heading_style: bool = False
    # Context attributes
    paragraph_source: str = ""  # Text or Table
    sentence_contains_period: bool = False  # is match in a sentence?
    line_number_in_document: int = 0


def extract_paragraphs_with_position(filepath):
    """Extract all paragraphs with full metadata and position tracking."""
    document = Document(str(filepath))
    paragraphs = []
    position = 0

    for para_idx, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        # Extract first run formatting
        first_run_bold = False
        first_run_italic = False
        first_run_underline = False
        font_size = ""
        font_name = ""

        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_bold = bool(first_run.bold)
            first_run_italic = bool(first_run.italic)
            first_run_underline = bool(first_run.underline)

            if first_run.font.size:
                font_size = str(first_run.font.size.pt)
            if first_run.font.name:
                font_name = first_run.font.name

        style_name = paragraph.style.name if paragraph.style else ""
        is_heading = "Heading" in style_name

        paragraphs.append({
            "index": para_idx,
            "text": text,
            "bold": first_run_bold,
            "italic": first_run_italic,
            "underline": first_run_underline,
            "font_size_pt": font_size,
            "font_name": font_name,
            "style": style_name,
            "is_heading": is_heading,
            "source": "Text",
            "line_number": position,
        })
        position += 1

    # Extract from table cells
    for table_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if not text:
                    continue

                # Try to get formatting from first run in cell
                cell_bold = False
                cell_italic = False
                cell_underline = False
                cell_font_size = ""
                cell_font_name = ""

                for para in cell.paragraphs:
                    if para.runs:
                        first_run = para.runs[0]
                        cell_bold = bool(first_run.bold)
                        cell_italic = bool(first_run.italic)
                        cell_underline = bool(first_run.underline)
                        if first_run.font.size:
                            cell_font_size = str(first_run.font.size.pt)
                        if first_run.font.name:
                            cell_font_name = first_run.font.name
                        break

                paragraphs.append({
                    "index": len(paragraphs),
                    "text": text,
                    "bold": cell_bold,
                    "italic": cell_italic,
                    "underline": cell_underline,
                    "font_size_pt": cell_font_size,
                    "font_name": cell_font_name,
                    "style": "",
                    "is_heading": False,
                    "source": f"Table[{table_idx}][{row_idx}][{cell_idx}]",
                    "line_number": position,
                })
                position += 1

    return paragraphs


def analyze_match_context(paragraph_text, match_start, match_end, control_id):
    """Extract context around the match and analyze sentence structure."""
    context_before = paragraph_text[max(0, match_start - 50):match_start]
    context_after = paragraph_text[match_end:min(len(paragraph_text), match_end + 50)]

    # Check if match is in a sentence (surrounded by sentence boundaries)
    text_before = paragraph_text[:match_start]
    sentence_has_period = "." in text_before[max(0, match_start - 100):]

    return {
        "context_before": context_before.strip(),
        "context_after": context_after.strip(),
        "sentence_contains_period": sentence_has_period,
    }


def find_all_control_matches(paragraphs, pattern):
    """Find all control ID matches in all paragraphs."""
    matches = []

    for para_dict in paragraphs:
        text = para_dict["text"]
        para_idx = para_dict["index"]

        # Find all matches in this paragraph
        for match in pattern.finditer(text):
            control_id = match.group(0)
            match_start = match.start()
            match_end = match.end()

            context = analyze_match_context(text, match_start, match_end, control_id)

            attr = ControlAttribute(
                source_file=para_dict.get("source_file", ""),
                paragraph_index=para_idx,
                control_id=control_id,
                match_position=match_start,
                paragraph_text=text,
                context_before=context["context_before"],
                context_after=context["context_after"],
                paragraph_bold=para_dict["bold"],
                paragraph_italic=para_dict["italic"],
                paragraph_underline=para_dict["underline"],
                font_size_pt=para_dict["font_size_pt"],
                font_name=para_dict["font_name"],
                paragraph_style=para_dict["style"],
                is_heading_style=para_dict["is_heading"],
                paragraph_source=para_dict["source"],
                sentence_contains_period=context["sentence_contains_period"],
                line_number_in_document=para_dict["line_number"],
            )
            matches.append(attr)

    return matches


def process_single_document(filepath, pattern):
    """Run the full analysis on a single .docx file."""
    filepath = Path(filepath)
    source_file = filepath.name
    paragraphs = extract_paragraphs_with_position(filepath)

    if not paragraphs:
        print(f"  WARNING: {source_file} has 0 paragraphs, skipping.")
        return []

    # Add source_file to paragraph metadata
    for para in paragraphs:
        para["source_file"] = source_file

    matches = find_all_control_matches(paragraphs, pattern)
    print(f"  Control ID instances found: {len(matches)}")

    return matches


def main():
    parser = setup_argparse("Analyze control ID attributes for false positive detection")
    args = parser.parse_args()

    config = load_config(args.config)
    ctrl_cfg = config.get("control_extraction", {})

    input_dir = get_input_dir(config, args.input_dir)
    output_dir = get_output_dir(config, "controls", args.output_dir)
    ensure_output_dir(output_dir)

    # Build control ID pattern from config (same as extract_controls.py)
    id_patterns = ctrl_cfg.get("control_id_patterns", None)
    if id_patterns is None:
        legacy = ctrl_cfg.get("control_id_pattern", r'\b[A-Z]{2,4}[-.]?\d{1,3}[-.]\d{2,4}\b')
        id_patterns = [legacy]

    combined_pattern = "|".join(f"({p})" for p in id_patterns)
    control_id_regex = re.compile(combined_pattern)

    output_csv = os.path.join(output_dir, "control_attributes_analysis.csv")

    docx_files = iter_docx_files(input_dir, config)
    if not docx_files:
        print(f"No .docx files found in {input_dir}")
        return

    all_matches = []
    total_instances = 0
    files_processed = 0

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        print(f"  Processing: {filename}")

        try:
            matches = process_single_document(filepath, control_id_regex)
            all_matches.extend(matches)
            total_instances += len(matches)
            files_processed += 1

        except Exception as error:
            print(f"  ERROR processing {filename}: {error}")

    # Write all matches to CSV
    if all_matches:
        with open(output_csv, "w", newline="", encoding="utf-8") as csv_file:
            columns = [
                "source_file", "paragraph_index", "line_number_in_document",
                "control_id", "match_position",
                "paragraph_text", "context_before", "context_after",
                "paragraph_bold", "paragraph_italic", "paragraph_underline",
                "font_size_pt", "font_name",
                "paragraph_style", "is_heading_style", "paragraph_source",
                "sentence_contains_period",
            ]
            writer = csv.DictWriter(csv_file, fieldnames=columns, quoting=csv.QUOTE_ALL)
            writer.writeheader()
            for match in all_matches:
                writer.writerow(asdict(match))

    print("\n" + "=" * 60)
    print("CONTROL ATTRIBUTE ANALYSIS SUMMARY")
    print("=" * 60)
    print(f"Files processed:     {files_processed}")
    print(f"Control instances:   {total_instances}")
    print(f"Output: {output_csv}")
    print("\nUse this CSV to:")
    print("  1. Identify patterns in false positives (table cells, sentences, etc.)")
    print("  2. Filter by attributes (paragraph_bold, paragraph_source, etc.)")
    print("  3. Improve whitelist/blacklist in config")
    print("  4. Refine control_id_pattern if needed")


if __name__ == "__main__":
    main()
