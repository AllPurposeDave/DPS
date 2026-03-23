"""
Word Counter — DOCX Word Count to CSV

Scans a folder for .docx files, counts the words in each document,
and writes a CSV with per-document and total word counts.

Usage:
    python word_counter.py <input_dir> [output_csv]

Arguments:
    input_dir   — Directory containing .docx files to count
    output_csv  — Path for the output CSV (default: word_counts.csv in input_dir)

Output:
    CSV with columns: filename, word_count
    Final row is a TOTAL across all documents.

REQUIREMENTS:
    pip install python-docx
    Python 3.10 or later

FAILURE POINT: Word temp files (~$*.docx) are automatically skipped.
FAILURE POINT: Corrupted or password-protected .docx files are logged and skipped.
"""

import argparse
import csv
import glob
import os
import sys

from docx import Document


def iter_docx_files(input_dir: str) -> list[str]:
    """
    Return sorted list of .docx file paths in input_dir.
    Skips Word temp files (~$*.docx). Top-level only — no sub-folders.
    """
    pattern = os.path.join(input_dir, "*.docx")
    files = glob.glob(pattern)
    files = [f for f in files if not os.path.basename(f).startswith("~$")]
    files.sort()
    return files


def count_words(doc_path: str) -> int:
    """
    Count total words in a .docx file.
    Includes paragraph text and table cell text.
    """
    doc = Document(doc_path)
    total = 0

    # Count words in paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            total += len(text.split())

    # Count words in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    total += len(text.split())

    return total


def main():
    parser = argparse.ArgumentParser(
        description="Count words in all .docx files in a folder and output a CSV."
    )
    parser.add_argument(
        "input_dir",
        help="Directory containing .docx files to count",
    )
    parser.add_argument(
        "output_csv",
        nargs="?",
        default=None,
        help="Output CSV path (default: word_counts.csv in input_dir)",
    )
    args = parser.parse_args()

    input_dir = args.input_dir
    if not os.path.isdir(input_dir):
        print(f"ERROR: '{input_dir}' is not a valid directory.")
        sys.exit(1)

    output_csv = args.output_csv or os.path.join(input_dir, "word_counts.csv")

    docx_files = iter_docx_files(input_dir)
    if not docx_files:
        print(f"No .docx files found in '{input_dir}'.")
        sys.exit(1)

    print(f"Found {len(docx_files)} .docx file(s) in '{input_dir}'.\n")

    results: list[tuple[str, int]] = []
    grand_total = 0

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        try:
            wc = count_words(filepath)
            results.append((filename, wc))
            grand_total += wc
            print(f"  {filename:.<60s} {wc:>8,} words")
        except Exception as e:
            print(f"  ERROR processing {filename}: {e}")
            results.append((filename, 0))

    print(f"\n  {'TOTAL':.<60s} {grand_total:>8,} words")

    # Write CSV
    os.makedirs(os.path.dirname(output_csv) or ".", exist_ok=True)
    with open(output_csv, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["filename", "word_count"])
        for filename, wc in results:
            writer.writerow([filename, wc])
        writer.writerow(["TOTAL", grand_total])

    print(f"\nCSV written to: {output_csv}")


if __name__ == "__main__":
    main()
