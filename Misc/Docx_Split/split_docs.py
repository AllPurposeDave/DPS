"""
Docx_Split - split large Word documents into smaller URL-annotated files.
=========================================================================

Behaviour
---------
For every .docx file in the configured input directory the script will:

  1. Look up the file in Doc_URL.xlsx by its filename stem (no extension).
  2. Remove every section whose heading appears in that row's
     "Delete_Headings" column, including its subsections. Matching is
     case-insensitive and exact on the trimmed heading text.
  3. Pack the remaining content greedily into one or more output .docx
     files, each at or below max_char_count characters. Splits occur only
     at Heading 1, 2, or 3 boundaries so that sections are never broken
     mid-body.
  4. Write a "Published URL: <url>" paragraph at the top of each output.
     When url_repeat_in_page_header is enabled, the URL is additionally
     placed in the document page header so that it appears on every
     rendered page.
  5. Name each output "<source_stem> - <first_heading_in_chunk>.docx".

Usage
-----
  1. pip install -r requirements.txt
  2. Place source .docx files in ./input.
  3. Populate ./input/Doc_URL.xlsx with the URL and optional deletion list
     for each source document.
  4. python split_docs.py

Runtime behaviour is defined in ./config.yaml. An alternate configuration
file may be provided via --config <path>.


"""

from __future__ import annotations

import argparse
import copy
import difflib
import re
import shutil
import sys
import traceback
import unicodedata
from pathlib import Path

import yaml
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook


# ----------------------------------------------------------------------------
# Config loading
# ----------------------------------------------------------------------------

DEFAULT_CONFIG = {
    "input_dir": "input",
    "output_dir": "output",
    "url_xlsx": "input/Doc_URL.xlsx",
    "url_xlsx_sheet": "URL Mapping",
    "doc_name_column": "Document_Name",
    "url_column": "URL",
    "delete_headings_column": "Delete_Headings",
    "max_char_count": 36000,
    "heading_levels": [1, 2, 3],
    "url_preamble_enabled": True,
    "url_preamble_label": "Published URL:",
    "url_repeat_in_page_header": False,
    "flatten_2_col_table": False,
    "flatten_2_col_separator": ": ",
    "flatten_2_col_heading_level": 2,
    "flatten_2_col_heading_default": "Definitions",
    "filename_suffix_max_len": 50,
}


def load_config(config_path: Path) -> dict:
    """Read config.yaml and merge on top of DEFAULT_CONFIG so missing keys fall back."""
    cfg = dict(DEFAULT_CONFIG)
    if config_path.exists():
        with config_path.open("r", encoding="utf-8") as f:
            user_cfg = yaml.safe_load(f) or {}
        cfg.update(user_cfg)
    else:
        print(f"  (no config.yaml found at {config_path}; using built-in defaults)")

    # Coerce numeric fields. YAML values quoted as strings would otherwise
    # break comparisons silently (e.g. heading_levels: ["1","2"] matches nothing
    # because heading_level() returns an int).
    def _coerce_int(key: str) -> None:
        try:
            cfg[key] = int(cfg[key])
        except (TypeError, ValueError) as exc:
            raise ValueError(f"config '{key}' must be an integer, got: {cfg[key]!r}") from exc

    _coerce_int("max_char_count")
    _coerce_int("flatten_2_col_heading_level")
    _coerce_int("filename_suffix_max_len")
    try:
        cfg["heading_levels"] = [int(v) for v in cfg.get("heading_levels", [1, 2, 3])]
    except (TypeError, ValueError) as exc:
        raise ValueError(
            f"config 'heading_levels' must be a list of integers, got: {cfg.get('heading_levels')!r}"
        ) from exc
    return cfg


def resolve_path(base: Path, path_str: str) -> Path:
    """Resolve a config path relative to the config file's folder (absolute paths pass through)."""
    p = Path(path_str)
    return p if p.is_absolute() else (base / p).resolve()


# ----------------------------------------------------------------------------
# Doc_URL.xlsx lookup
# ----------------------------------------------------------------------------

def load_doc_index(xlsx_path: Path, sheet: str, name_col: str, url_col: str,
                   delete_col: str) -> dict:
    """
    Read the URL/deletion spreadsheet and return a lookup dict.

    Keys are canonicalised document stems (filename without .docx, then run
    through canonicalize). Values are:
        {"url": str, "delete_headings": [str, ...], "raw_name": str}

    `raw_name` preserves the original spelling from the spreadsheet for use
    in diagnostic messages.

    Missing columns degrade gracefully: if Delete_Headings is absent the
    list is always []. If the workbook is currently open in Excel the
    function refuses to proceed, to avoid reading the last-saved copy while
    unsaved edits remain in Excel.
    """
    if not xlsx_path.exists():
        print(f"  WARNING: URL lookup file not found at {xlsx_path}")
        print("           Outputs will have no PublishedURL preamble.")
        return {}

    lock = excel_lock_file(xlsx_path)
    if lock is not None:
        print(f"  ERROR: {xlsx_path.name} appears to be open in Excel (lock file '{lock.name}' is present).")
        print("         The script reads only the saved copy, so unsaved changes in Excel")
        print("         would be silently ignored. Close the workbook and re-run.")
        return {}

    try:
        wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    except Exception as exc:
        print(f"  WARNING: could not open {xlsx_path.name}: {exc}")
        return {}

    try:
        if sheet not in wb.sheetnames:
            print(f"  WARNING: sheet '{sheet}' not found in {xlsx_path.name}.")
            print(f"           Available sheets: {', '.join(wb.sheetnames)}")
            return {}

        ws = wb[sheet]
        rows = ws.iter_rows(values_only=True)

        # Header row: map column name -> column index.
        try:
            header = next(rows)
        except StopIteration:
            return {}
        # Canonicalise header cells so that hidden NBSP / casing differences
        # in the column label don't silently drop a column.
        header_map = {canonicalize(h): i for i, h in enumerate(header) if h is not None}

        name_key = canonicalize(name_col)
        url_key = canonicalize(url_col)
        delete_key = canonicalize(delete_col)

        if name_key not in header_map:
            print(f"  WARNING: column '{name_col}' not found in sheet '{sheet}'.")
            print(f"           Columns present: {', '.join(str(h) for h in header if h is not None)}")
            return {}
        if url_key not in header_map:
            print(f"  WARNING: column '{url_col}' not found in sheet '{sheet}'.")

        name_idx = header_map[name_key]
        url_idx = header_map.get(url_key)
        delete_idx = header_map.get(delete_key)

        index: dict = {}
        for row in rows:
            name = row[name_idx] if name_idx < len(row) else None
            if name is None or str(name).strip() == "":
                continue
            raw_name = str(name).strip()
            stem = canonicalize(raw_name)
            url = str(row[url_idx]).strip() if url_idx is not None and url_idx < len(row) and row[url_idx] else ""
            delete_raw = row[delete_idx] if delete_idx is not None and delete_idx < len(row) else None
            delete_list: list[str] = []
            if delete_raw:
                delete_list = [h.strip() for h in str(delete_raw).split(",") if h.strip()]
            if stem in index:
                print(f"  WARNING: duplicate Document_Name '{raw_name}' in {xlsx_path.name}; last occurrence wins.")
            index[stem] = {"url": url, "delete_headings": delete_list, "raw_name": raw_name}

        return index
    finally:
        wb.close()


# ----------------------------------------------------------------------------
# DOCX helpers
# ----------------------------------------------------------------------------


_HEADING_STYLE_RE = re.compile(r"^Heading (\d)$")
_WHITESPACE_RE = re.compile(r"\s+")
_ZERO_WIDTH_RE = re.compile(r"[​‌‍⁠﻿]")
_DASH_CHARS = ("‐", "‑", "‒", "–", "—", "―", "−")
_SINGLE_QUOTE_CHARS = ("‘", "’", "‚", "‛")
_DOUBLE_QUOTE_CHARS = ("“", "”", "„", "‟")


def canonicalize(text) -> str:
    """Normalise text for comparison between Excel cells and Word paragraphs.

    Excel and Word frequently introduce invisible differences that defeat
    naive `==` / `.lower()` matching: non-breaking spaces (U+00A0), zero-width
    joiners, NFD vs NFC composition of accented characters, em/en dashes in
    place of hyphens, and smart quotes in place of ASCII quotes. This helper
    collapses all of those to a single canonical form so that what looks
    identical to a human reader actually compares equal.
    """
    if text is None:
        return ""
    t = unicodedata.normalize("NFC", str(text))
    t = t.replace(" ", " ")
    t = _ZERO_WIDTH_RE.sub("", t)
    for ch in _DASH_CHARS:
        t = t.replace(ch, "-")
    for ch in _SINGLE_QUOTE_CHARS:
        t = t.replace(ch, "'")
    for ch in _DOUBLE_QUOTE_CHARS:
        t = t.replace(ch, '"')
    t = _WHITESPACE_RE.sub(" ", t).strip().lower()
    return t


def excel_lock_file(xlsx_path: Path) -> Path | None:
    """Return Excel's lock file path for this workbook if present, else None.

    When Excel (or Word, or LibreOffice) has a workbook open, it writes a
    hidden sibling file named '~$<workbook>.xlsx' alongside the original.
    On macOS the operating system does not enforce file locks, so the script
    would otherwise silently read the last-saved state and miss any unsaved
    edits. Detecting the lock file lets us halt with a clear message instead.
    """
    lock = xlsx_path.parent / f"~${xlsx_path.name}"
    return lock if lock.exists() else None


def heading_level(style) -> int | None:
    """Return 1-9 for a built-in 'Heading N' style, else None."""
    if style is None:
        return None
    name = style.name if hasattr(style, "name") else str(style)
    m = _HEADING_STYLE_RE.match(name)
    return int(m.group(1)) if m else None


def build_element_sequence(doc) -> list[tuple]:
    """
    Walk the document body in natural order and return a list of
    ("paragraph", Paragraph) or ("table", Table) tuples.

    Preserves the original interleaving of paragraphs and tables. The walker
    descends into Structured Document Tag wrappers (w:sdt) so that content
    held inside Word content controls is included in the output. Any other
    top-level element (bookmarks, tracked-change markers, section properties)
    is ignored — they survive verbatim through the clone-and-repopulate
    write path because the output file is cloned from the source.
    """
    elements: list[tuple] = []

    def walk(parent_xml) -> None:
        for child in parent_xml:
            tag = child.tag
            if tag == qn("w:p"):
                elements.append(("paragraph", Paragraph(child, doc)))
            elif tag == qn("w:tbl"):
                elements.append(("table", Table(child, doc)))
            elif tag == qn("w:sdt"):
                content = child.find(qn("w:sdtContent"))
                if content is not None:
                    walk(content)

    walk(doc.element.body)
    return elements


def _append_before_sectpr(body, new_elem) -> None:
    """
    Insert new_elem immediately before the body-level w:sectPr, if present.

    The OOXML schema requires w:sectPr (if body-level) to be the final child
    of w:body. Appending after it produces a technically invalid document
    and can cause Word to emit repair warnings on open.
    """
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(new_elem)
    else:
        body.append(new_elem)


def copy_paragraph(source_para, target_doc) -> None:
    """Deep-copy a paragraph into target_doc.body. Preserves runs, formatting, hyperlinks."""
    _append_before_sectpr(target_doc.element.body, copy.deepcopy(source_para._element))


def copy_table(source_table, target_doc) -> None:
    """Deep-copy a table into target_doc.body. Preserves rows, cells, merges, formatting."""
    _append_before_sectpr(target_doc.element.body, copy.deepcopy(source_table._element))


def count_chars(elements) -> int:
    """Sum character counts across paragraphs and table cells. Matches section_splitter's accounting."""
    total = 0
    for elem_type, elem_obj in elements:
        if elem_type == "paragraph":
            total += len(elem_obj.text)
        elif elem_type == "table":
            for row in elem_obj.rows:
                for cell in row.cells:
                    total += len(cell.text)
    return total


# ----------------------------------------------------------------------------
# Section deletion
# ----------------------------------------------------------------------------

def remove_delete_headings(elements: list,
                           headings_to_delete: list[str]) -> tuple[list, list[str]]:
    """
    Remove every section whose heading matches an entry in headings_to_delete.

    Matching is performed on the canonicalised heading text (see
    canonicalize), so non-breaking spaces, smart quotes, em/en dashes and
    minor whitespace differences between Excel and Word do not defeat the
    match. Subsections are included: once a matching heading is found at
    level L, every following element is dropped until another heading at
    level <= L is encountered (or the document ends).

    Returns (surviving_elements, unmatched_headings). Any Delete_Headings
    entry that was never found in the document is returned so the caller
    can surface it as a warning.
    """
    if not headings_to_delete:
        return elements, []

    # Map canonical form -> original spelling, so diagnostic messages echo
    # what the user typed rather than the normalised lowercase form.
    target_map: dict[str, str] = {}
    for raw in headings_to_delete:
        if raw and raw.strip():
            target_map[canonicalize(raw)] = raw.strip()

    matched: set[str] = set()
    result: list = []
    i = 0
    while i < len(elements):
        elem_type, elem_obj = elements[i]
        if elem_type == "paragraph":
            lvl = heading_level(elem_obj.style)
            if lvl is not None:
                canon = canonicalize(elem_obj.text)
                if canon in target_map:
                    matched.add(canon)
                    # Skip this heading and everything beneath it.
                    i += 1
                    while i < len(elements):
                        et, eo = elements[i]
                        if et == "paragraph":
                            nxt_lvl = heading_level(eo.style)
                            if nxt_lvl is not None and nxt_lvl <= lvl:
                                break  # reached a sibling or parent heading
                        i += 1
                    continue
        result.append(elements[i])
        i += 1

    unmatched = [target_map[k] for k in target_map if k not in matched]
    return result, unmatched


# ----------------------------------------------------------------------------
# 2-column table flattening
# ----------------------------------------------------------------------------
# When enabled, 2-column tables without merged cells are replaced by a heading
# paragraph + one "col1: col2" paragraph per row. If the first row looks like a
# header ("Term | Definition"), it's used as the heading text (joined with
# " and ") and dropped from the data rows. Otherwise the config default is
# used and all rows are kept.
#
# Tables with merged cells (gridSpan horizontal OR vMerge vertical) are left
# untouched — flattening them would reorder content or duplicate data. The
# script scans the table XML directly to detect any merge markers.


def _table_has_merged_cells(table) -> bool:
    """True if the table contains any horizontal or vertical cell merges."""
    tbl_elem = table._element
    if tbl_elem.findall(f".//{qn('w:gridSpan')}"):
        return True
    if tbl_elem.findall(f".//{qn('w:vMerge')}"):
        return True
    return False


def _is_flattenable_2col(table) -> bool:
    """Exactly two grid columns, two logical cells per row, no merges."""
    try:
        if len(table.columns) != 2:
            return False
        if _table_has_merged_cells(table):
            return False
        for row in table.rows:
            if len(row.cells) != 2:
                return False
        return True
    except Exception:
        return False


def _row_is_header(row) -> bool:
    """
    Heuristic: explicit <w:tblHeader> marker, OR both cells short with bold text.
    Returning True means we'll use this row as the heading and drop it from rows.
    """
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is not None and trPr.find(qn("w:tblHeader")) is not None:
        return True

    cells = row.cells
    if len(cells) != 2:
        return False
    c1_text = cells[0].text.strip()
    c2_text = cells[1].text.strip()
    if not c1_text or not c2_text:
        return False
    if len(c1_text) >= 60 or len(c2_text) >= 60:
        return False
    # Any bold run anywhere in either cell is a strong header signal.
    for cell in cells:
        for para in cell.paragraphs:
            for run in para.runs:
                if run.bold:
                    return True
    return False


def flatten_2_col_tables(elements: list, config: dict) -> tuple[list, int, int]:
    """
    Replace clean 2-column tables with (heading paragraph + 'col1: col2' paragraphs).

    Returns (new_elements, flattened_count, skipped_merged_count). Merged-cell
    tables are kept as-is and reported in skipped_merged_count.
    """
    if not config.get("flatten_2_col_table", False):
        return elements, 0, 0

    separator = config.get("flatten_2_col_separator", ": ")
    heading_level_n = int(config.get("flatten_2_col_heading_level", 2))
    heading_default = config.get("flatten_2_col_heading_default", "Definitions")
    heading_style_name = f"Heading {heading_level_n}"

    # A scratch document owns the synthetic paragraph XML. When copy_paragraph
    # deep-copies these elements into the target docx, styles resolve by NAME
    # against the target's style definitions — so "Heading 2" renders correctly
    # as long as the target doc (cloned from source) defines it, which real
    # Word documents always do.
    scratch = Document()

    def make_para(text: str, style_name: str | None = None):
        para = scratch.add_paragraph(text)
        if style_name:
            try:
                para.style = scratch.styles[style_name]
            except KeyError:
                pass
        return para

    new_elements: list = []
    flattened = 0
    skipped_merged = 0

    for elem_type, elem_obj in elements:
        if elem_type != "table":
            new_elements.append((elem_type, elem_obj))
            continue

        # Not 2 columns → leave untouched.
        try:
            if len(elem_obj.columns) != 2:
                new_elements.append((elem_type, elem_obj))
                continue
        except Exception:
            new_elements.append((elem_type, elem_obj))
            continue

        # 2 columns but has merges → skip flattening (per user request).
        if not _is_flattenable_2col(elem_obj):
            new_elements.append((elem_type, elem_obj))
            skipped_merged += 1
            continue

        rows = list(elem_obj.rows)
        if not rows:
            new_elements.append((elem_type, elem_obj))
            continue

        if _row_is_header(rows[0]):
            c1 = rows[0].cells[0].text.strip()
            c2 = rows[0].cells[1].text.strip()
            heading_text = f"{c1} and {c2}" if (c1 and c2) else heading_default
            data_rows = rows[1:]
        else:
            heading_text = heading_default
            data_rows = rows

        new_elements.append(("paragraph", make_para(heading_text, heading_style_name)))

        for row in data_rows:
            left = row.cells[0].text.strip()
            right = row.cells[1].text.strip()
            if not left and not right:
                continue
            new_elements.append(("paragraph", make_para(f"{left}{separator}{right}")))

        flattened += 1

    return new_elements, flattened, skipped_merged


# ----------------------------------------------------------------------------
# Greedy section packer
# ----------------------------------------------------------------------------

def pack_chunks(elements: list, max_chars: int,
                boundary_levels: list[int]) -> tuple[list[dict], list[str]]:
    """
    Greedy-pack elements into chunks of at most max_chars characters.

    Only headings at boundary_levels are eligible split points. Content
    between two adjacent eligible headings is an atomic "section" — we never
    split inside one. Sections are accumulated into the current chunk until
    adding the next section would exceed max_chars; then the chunk is closed.

    Returns (chunks, warnings). Each chunk is:
        {"elements": [...], "first_heading": str, "char_count": int}
    """
    warnings: list[str] = []
    boundary_levels_set = set(boundary_levels)

    # Find positions of eligible heading paragraphs.
    boundaries = [
        i for i, (t, o) in enumerate(elements)
        if t == "paragraph" and heading_level(o.style) in boundary_levels_set
    ]

    # No eligible boundaries: emit the entire doc as a single chunk.
    if not boundaries:
        chars = count_chars(elements)
        levels_label = "/".join(f"H{level}" for level in boundary_levels)
        warnings.append(
            f"no {levels_label} headings found; emitting the entire document as one chunk ({chars} chars)"
        )
        return ([{
            "elements": list(elements),
            "first_heading": "(Full Document)",
            "char_count": chars,
        }], warnings)

    # Slice into sections. Content before the first boundary is a
    # pseudo-section with no heading (called "(Preamble)").
    sections: list[dict] = []
    if boundaries[0] > 0:
        preamble = elements[:boundaries[0]]
        sections.append({
            "elements": preamble,
            "heading": "",
            "char_count": count_chars(preamble),
        })
    boundaries_ext = boundaries + [len(elements)]
    for k in range(len(boundaries)):
        start = boundaries_ext[k]
        end = boundaries_ext[k + 1]
        sec_elems = elements[start:end]
        heading_text = sec_elems[0][1].text.strip() if sec_elems else ""
        sections.append({
            "elements": sec_elems,
            "heading": heading_text,
            "char_count": count_chars(sec_elems),
        })

    # Greedy pack sections into chunks.
    chunks: list[dict] = []
    current = _empty_chunk()

    def flush():
        nonlocal current
        if current["elements"]:
            chunks.append(current)
            current = _empty_chunk()

    for sec in sections:
        # A single section bigger than the cap can't be packed with anything
        # else. Flush what we have and emit it solo, with a warning.
        if sec["char_count"] > max_chars:
            flush()
            chunks.append({
                "elements": sec["elements"],
                "first_heading": sec["heading"] or "(Preamble)",
                "char_count": sec["char_count"],
            })
            warnings.append(
                f"section '{sec['heading'] or '(Preamble)'}' is {sec['char_count']} chars, exceeds max_char_count={max_chars}"
            )
            continue

        # Will this section push the current chunk over the limit? If so,
        # flush and start a new chunk at this heading boundary.
        if current["elements"] and current["char_count"] + sec["char_count"] > max_chars:
            flush()

        if not current["elements"]:
            current["first_heading"] = sec["heading"] or "(Preamble)"
        current["elements"].extend(sec["elements"])
        current["char_count"] += sec["char_count"]

    flush()
    return chunks, warnings


def _empty_chunk() -> dict:
    return {"elements": [], "first_heading": "", "char_count": 0}


# ----------------------------------------------------------------------------
# Output writing
# ----------------------------------------------------------------------------

def create_sub_docx(source_path: Path, chunk: dict, output_path: Path,
                    published_url: str, config: dict) -> None:
    """
    Write one split chunk as a new .docx.

    We clone the SOURCE file, empty its body, then repopulate it. Cloning
    preserves styles, themes, numbering, and relationships so Word opens the
    output without "unreadable content" warnings — you can't get the same
    result by starting from Document() and copying XML in.
    """
    shutil.copy2(source_path, output_path)
    sub_doc = Document(str(output_path))

    # Clear body children except sectPr (page-layout section properties must stay).
    body = sub_doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    # Visible preamble paragraph at the top of the body.
    if config.get("url_preamble_enabled", True) and published_url:
        label = config.get("url_preamble_label", "Published URL:")
        sub_doc.add_paragraph(f"{label} {published_url}")

    # Optional: write URL into the page header so every rendered page shows it.
    # Useful for RAG pipelines that chunk per-page and would otherwise lose the
    # URL after page 1.
    if config.get("url_repeat_in_page_header", False) and published_url:
        label = config.get("url_preamble_label", "Published URL:")
        header = sub_doc.sections[0].header
        header.is_linked_to_previous = False
        if header.paragraphs and not header.paragraphs[0].text:
            header.paragraphs[0].text = f"{label} {published_url}"
        else:
            header.add_paragraph(f"{label} {published_url}")

    # Repopulate with the chunk's content.
    for elem_type, elem_obj in chunk["elements"]:
        if elem_type == "paragraph":
            copy_paragraph(elem_obj, sub_doc)
        elif elem_type == "table":
            copy_table(elem_obj, sub_doc)

    sub_doc.save(str(output_path))


# ----------------------------------------------------------------------------
# Filename building
# ----------------------------------------------------------------------------

_FILENAME_UNSAFE_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def sanitize_filename_piece(text: str, max_len: int) -> str:
    """Strip filesystem-unsafe characters, collapse whitespace, and truncate.

    Trailing dots and spaces are stripped because Windows silently rejects
    them at the end of filenames, which would cause downstream copies or
    uploads to fail without a clear error.
    """
    cleaned = _FILENAME_UNSAFE_RE.sub("", text)
    cleaned = _WHITESPACE_RE.sub(" ", cleaned).strip()
    cleaned = cleaned.rstrip(". ")
    if len(cleaned) > max_len:
        cleaned = cleaned[:max_len].rstrip(". ")
    return cleaned or "untitled"


def build_output_name(source_stem: str, first_heading: str, max_len: int,
                      used_names: set) -> str:
    """Build '<source_stem> - <first_heading>.docx' with case-insensitive dedup.

    used_names is expected to hold lowercase entries: macOS and Windows
    filesystems default to case-insensitive behaviour, so "Foo.docx" and
    "foo.docx" would collide and the second write would overwrite the first.
    """
    safe_heading = sanitize_filename_piece(first_heading, max_len)
    base = f"{source_stem} - {safe_heading}"
    name = f"{base}.docx"
    n = 2
    while name.lower() in used_names:
        name = f"{base} ({n}).docx"
        n += 1
    used_names.add(name.lower())
    return name


# ----------------------------------------------------------------------------
# Main orchestration
# ----------------------------------------------------------------------------

def process_document(source_path: Path, output_dir: Path, doc_entry: dict,
                     config: dict, used_names: set) -> tuple[int, list[str]]:
    """Process one source docx. Returns (num_chunks_written, warning_messages)."""
    doc = Document(str(source_path))
    elements = build_element_sequence(doc)

    delete_headings = doc_entry.get("delete_headings", [])
    published_url = doc_entry.get("url", "")

    elements, unmatched_headings = remove_delete_headings(elements, delete_headings)

    elements, flattened_n, skipped_merged_n = flatten_2_col_tables(elements, config)

    chunks, warnings = pack_chunks(
        elements,
        max_chars=config["max_char_count"],
        boundary_levels=config.get("heading_levels", [1, 2, 3]),
    )

    if unmatched_headings:
        warnings.append(
            "Delete_Headings entries not found in document: "
            + "; ".join(f"'{h}'" for h in unmatched_headings)
        )
    if flattened_n:
        warnings.append(f"flattened {flattened_n} two-column table(s)")
    if skipped_merged_n:
        warnings.append(f"skipped {skipped_merged_n} two-column table(s) with merged cells (kept as-is)")

    source_stem = source_path.stem
    for chunk in chunks:
        out_name = build_output_name(
            source_stem,
            chunk["first_heading"],
            config.get("filename_suffix_max_len", 50),
            used_names,
        )
        out_path = output_dir / out_name
        create_sub_docx(source_path, chunk, out_path, published_url, config)

    return len(chunks), warnings


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Split large .docx files into smaller URL-stamped chunks.",
    )
    parser.add_argument(
        "--config",
        type=str,
        default=None,
        help="Path to config.yaml (default: ./config.yaml next to this script).",
    )
    args = parser.parse_args()

    script_dir = Path(__file__).parent.resolve()
    config_path = Path(args.config).resolve() if args.config else (script_dir / "config.yaml")
    config = load_config(config_path)
    base = config_path.parent

    input_dir = resolve_path(base, config["input_dir"])
    output_dir = resolve_path(base, config["output_dir"])
    xlsx_path = resolve_path(base, config["url_xlsx"])

    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Input dir:   {input_dir}")
    print(f"Output dir:  {output_dir}")
    print(f"URL lookup:  {xlsx_path}")
    print(f"Max chars:   {config['max_char_count']:,}")
    print("-" * 60)

    if not input_dir.exists():
        print(f"ERROR: input_dir does not exist: {input_dir}")
        return 1

    doc_index = load_doc_index(
        xlsx_path,
        config["url_xlsx_sheet"],
        config["doc_name_column"],
        config["url_column"],
        config["delete_headings_column"],
    )

    # Skip Word lock files (~$*.docx) and hidden files.
    docx_files = sorted(
        p for p in input_dir.glob("*.docx")
        if not p.name.startswith("~$") and not p.name.startswith(".")
    )

    if not docx_files:
        print(f"No .docx files found in {input_dir}")
        return 0

    used_names: set = set()
    total_chunks = 0
    files_ok = 0
    files_failed = 0
    all_warnings: list[str] = []

    for src in docx_files:
        stem_key = canonicalize(src.stem)
        entry = doc_index.get(stem_key, {"url": "", "delete_headings": []})
        if stem_key not in doc_index:
            close = difflib.get_close_matches(stem_key, doc_index.keys(), n=3, cutoff=0.6)
            if close:
                hints = ", ".join(doc_index[c]["raw_name"] for c in close)
                all_warnings.append(
                    f"{src.name}: stem not found in {xlsx_path.name} — no URL will be stamped. "
                    f"Close matches in spreadsheet: {hints}"
                )
            else:
                all_warnings.append(
                    f"{src.name}: stem '{src.stem}' not found in {xlsx_path.name} — no URL will be stamped"
                )

        try:
            n_chunks, warnings = process_document(src, output_dir, entry, config, used_names)
            total_chunks += n_chunks
            files_ok += 1
            for w in warnings:
                all_warnings.append(f"{src.name}: {w}")
            url_note = "" if entry.get("url") else "  (no URL)"
            del_note = f"  (-{len(entry.get('delete_headings') or [])} deletions)" if entry.get("delete_headings") else ""
            print(f"  {src.name}: {n_chunks} chunk(s){del_note}{url_note}")
        except Exception as e:
            files_failed += 1
            print(f"  ERROR processing {src.name}: {e}")
            traceback.print_exc()

    print("-" * 60)
    print(f"Processed: {files_ok}    Failed: {files_failed}    Chunks written: {total_chunks}")
    if all_warnings:
        print(f"Warnings ({len(all_warnings)}):")
        for w in all_warnings:
            print(f"  - {w}")
    return 0 if files_failed == 0 else 2


if __name__ == "__main__":
    sys.exit(main())
